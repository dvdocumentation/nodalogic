from __future__ import annotations

import json
import os
import ast
import io
import pickle
import sqlite3
import base64
import uuid
import re
import socket
import subprocess
import tempfile
import math
import time
import threading
import traceback
import sys
from datetime import datetime, timezone, timedelta
from typing import Any, Dict, List, Optional, Tuple
from pathlib import Path
from urllib.parse import urlparse, unquote

import requests
from flask import Blueprint, abort, flash, jsonify, redirect, render_template, request, url_for, after_this_request, send_from_directory, send_file, Response, make_response, g, has_request_context, current_app
from markupsafe import escape
from html import unescape
from flask_login import current_user, login_required
from functools import wraps
from contextlib import contextmanager
from werkzeug.security import check_password_hash

from .nodalayout import render_nodalayout_html, DEFAULT_NL_CSS
from llm_credentials import chat_completion as _shared_chat_completion, message_content as _shared_message_content
from . import models
from . import ngenie_core
from . import ngenie_skill_registry
import nodes as _nodes_mod
import background_jobs as _background_jobs
import hashlib
import inspect
import mimetypes
from io import BytesIO
from sqlalchemy import select, or_, and_, func
from sqlalchemy.orm import selectinload
from jinja2.sandbox import SandboxedEnvironment
from jinja2 import select_autoescape

try:
    import qrcode
except Exception:  # optional dependency
    qrcode = None

_REPORTLAB_IMPORT_ERROR = None
try:
    from reportlab.graphics.barcode import createBarcodeDrawing, getCodes as _reportlab_barcode_codes
    from reportlab.graphics.barcode.code128 import stop as _REPORTLAB_CODE128_STOP
    from reportlab.graphics.shapes import Drawing as _ReportLabDrawing
    _ReportLabCode128Widget = _reportlab_barcode_codes().get("Code128")
except Exception as _reportlab_import_error:  # optional dependency
    _REPORTLAB_IMPORT_ERROR = _reportlab_import_error
    createBarcodeDrawing = None
    _REPORTLAB_CODE128_STOP = None
    _ReportLabDrawing = None
    _ReportLabCode128Widget = None


if _ReportLabCode128Widget is not None:
    class _Utf8BarcodeCode128(_ReportLabCode128Widget):
        """Code 128 widget that encodes Unicode as UTF-8 bytes via FNC4.

        ReportLab's stock Code128 validator accepts only 7-bit ASCII. Code 128
        itself can carry extended byte values with FNC4, so this widget keeps
        the original text and encodes its UTF-8 byte sequence.
        """

        def validate(self):
            self.valid = 1
            self.validated = str(self.value or "")
            self._utf8_bytes = self.validated.encode("utf-8")
            return self.validated

        def encode(self):
            data = getattr(self, "_utf8_bytes", str(self.value or "").encode("utf-8"))
            current_set = "B"
            encoded = [104]  # START_B

            for byte_value in data:
                extended = byte_value >= 128
                base_value = byte_value - 128 if extended else byte_value
                target_set = "A" if base_value < 32 else "B"

                if target_set != current_set:
                    encoded.append(101 if current_set == "B" else 100)
                    current_set = target_set

                if extended:
                    # FNC4 is code 101 in set A and code 100 in set B.
                    encoded.append(101 if current_set == "A" else 100)

                if current_set == "A" and base_value < 32:
                    codeword = base_value + 64
                else:
                    codeword = base_value - 32

                if codeword < 0 or codeword > 95:
                    raise ValueError(f"Cannot encode byte {byte_value} in Code128")
                encoded.append(codeword)

            checksum = encoded[0]
            for position, codeword in enumerate(encoded[1:], 1):
                checksum += position * codeword
            self.encoded = encoded + [checksum % 103, _REPORTLAB_CODE128_STOP]
            return self.encoded
else:
    _Utf8BarcodeCode128 = None


class _PrintAttrDict(dict):
    """Dictionary wrapper for PrintForm templates.

    Dot-access in Jinja should address _data keys, including names that collide
    with dict methods such as `items`, `keys`, or `values`.
    """
    def __getattribute__(self, name):
        if name.startswith('__'):
            return dict.__getattribute__(self, name)
        try:
            return dict.__getitem__(self, name)
        except KeyError:
            return dict.__getattribute__(self, name)


def _print_attr_tree(value: Any) -> Any:
    if isinstance(value, dict):
        return _PrintAttrDict({k: _print_attr_tree(v) for k, v in value.items()})
    if isinstance(value, list):
        return [_print_attr_tree(v) for v in value]
    if isinstance(value, tuple):
        return tuple(_print_attr_tree(v) for v in value)
    return value


class _PrintSandboxedEnvironment(SandboxedEnvironment):
    def is_safe_attribute(self, obj: Any, attr: str, value: Any) -> bool:
        if isinstance(obj, _PrintAttrDict) and attr in obj and not str(attr).startswith('__'):
            return True
        return super().is_safe_attribute(obj, attr, value)



def string_to_color(text: str) -> str:
    """Stable color for a tag string."""
    hash_object = hashlib.md5(str(text or '').encode('utf-8'))
    return f"#{hash_object.hexdigest()[:6]}"


def _tag_text_color(bg: str) -> str:
    bg = str(bg or '').strip()
    if not re.match(r'^#([0-9a-fA-F]{6})$', bg):
        return '#000000'
    r = int(bg[1:3], 16)
    g = int(bg[3:5], 16)
    b = int(bg[5:7], 16)
    # WCAG relative luminance approximation good enough for black/white choice.
    luminance = (0.299 * r + 0.587 * g + 0.114 * b)
    return '#000000' if luminance > 150 else '#FFFFFF'


def _normalize_node_tags(data: Optional[Dict[str, Any]]) -> List[Dict[str, str]]:
    raw = (data or {}).get('_tags') if isinstance(data, dict) else None
    if not isinstance(raw, list):
        return []
    out: List[Dict[str, str]] = []
    seen = set()
    for item in raw:
        tag_id = ''
        color = ''
        if isinstance(item, str):
            tag_id = item.strip()
            color = string_to_color(tag_id) if tag_id else ''
        elif isinstance(item, dict):
            tag_id = str(item.get('id') or item.get('name') or item.get('tag') or '').strip()
            color = str(item.get('color') or '').strip()
            if tag_id and not re.match(r'^#([0-9a-fA-F]{6})$', color):
                color = string_to_color(tag_id)
        if not tag_id or tag_id in seen:
            continue
        seen.add(tag_id)
        out.append({'id': tag_id, 'color': color, 'text_color': _tag_text_color(color)})
    return out


def _render_tags_html(data: Optional[Dict[str, Any]]) -> str:
    tags = _normalize_node_tags(data)
    if not tags:
        return ''
    parts = []
    for t in tags:
        tid = str(t.get('id') or '')
        bg = str(t.get('color') or string_to_color(tid))
        fg = str(t.get('text_color') or _tag_text_color(bg))
        parts.append(
            f'<span class="nl-tag-badge" data-nl-tag="{escape(tid)}" '
            f'style="display:inline-flex;align-items:center;border-radius:999px;padding:2px 8px;font-size:12px;line-height:1.4;background:{escape(bg)};color:{escape(fg)};">'
            f'{escape(tid)}</span>'
        )
    return '<div class="nl-tag-cloud d-flex gap-1 flex-wrap mt-2">' + ''.join(parts) + '</div>'


def _cover_with_tags(html: str, data: Optional[Dict[str, Any]], enabled: bool = False) -> str:
    # _tags are now rendered whenever they exist.  The enabled argument is kept
    # only for backward compatibility with older call sites / configs.
    base = str(html or '')
    tags_html = _render_tags_html(data)
    if not tags_html:
        return base
    return base + tags_html


def _tag_ids(data: Optional[Dict[str, Any]]) -> List[str]:
    return [str(t.get('id') or '') for t in _normalize_node_tags(data) if str(t.get('id') or '')]

import __main__ as main


# client_bp is registered by server/app.py as url_prefix="/client"
client_bp = Blueprint(
    "client",
    __name__,
    url_prefix="/client",
)

APP_TITLE = "NodaLogic Client"
DEFAULT_LIMIT_PER_CLASS = 50
AUTO_REFRESH_SECONDS = 10
RAW_NODES_SECTION_CODE = "__received_nodes__"
RAW_NODES_SECTION_NAME = "Received Nodes"
DASHBOARD_SECTION_CODE = "__dashboard__"
DASHBOARD_SECTION_NAME = "Dashboard"

PROJECTION_CLASS_TYPE = "projection"
PROJECTION_KANBAN_TYPE = "kanban_projection"
PROJECTION_DIAGRAM_TYPE = "diagram_projection"
PROJECTION_SCHEDULE_TYPE = "schedule_projection"
PROJECTION_GANTT_TYPE = "gantt_projection"
PROJECTION_NODE_LIST_TYPE = "node_list_projection"
PROJECTION_NODES_LIST_TYPE = "nodes_list_projection"
PROJECTION_HTML_TYPE = "html_projection"
# Projection nodes are reports. They may receive a transient list of object UIDs
# from onRunProjection so the browser can render immediately, but that list must
# not become saved projection state. Object positions/statuses live on the objects
# themselves in _projection_values[projection_uid].
PROJECTION_TRANSIENT_SAVE_FIELDS = {"_projection_objects", "_projection_html"}
PRINT_FORM_CLASS_TYPE = "print_form"
PRINT_FORM_TEMPLATE_HTML_JINJA = "html_jinja"
SINGLETON_CLASS_TYPES = {"custom_process", PROJECTION_CLASS_TYPE}

def _class_type_value(cls_or_type: Any) -> str:
    if isinstance(cls_or_type, dict):
        return str(cls_or_type.get("class_type") or "data_node").strip()
    return str(cls_or_type or "data_node").strip()

def _is_singleton_class_type(cls_or_type: Any) -> bool:
    return _class_type_value(cls_or_type) in SINGLETON_CLASS_TYPES

def _is_projection_class_type(cls_or_type: Any) -> bool:
    return _class_type_value(cls_or_type) == PROJECTION_CLASS_TYPE

def _is_node_list_projection_type(value: Any) -> bool:
    return str(value or "").strip() in {PROJECTION_NODE_LIST_TYPE, PROJECTION_NODES_LIST_TYPE, "node_list", "list_projection"}

def _is_html_projection_type(value: Any) -> bool:
    return str(value or "").strip() in {PROJECTION_HTML_TYPE, "html", "HTML"}

def _is_print_form_class_type(cls_or_type: Any) -> bool:
    return _class_type_value(cls_or_type) == PRINT_FORM_CLASS_TYPE


def _is_probably_print_template_base64(value: Any) -> bool:
    s = str(value or "").strip()
    if not s or len(s) % 4:
        return False
    try:
        raw = base64.b64decode(s.encode("ascii"), validate=True)
        text = raw.decode("utf-8")
    except Exception:
        return False
    return "\x00" not in text


def _decode_print_html_template(value: Any) -> str:
    s = str(value or "")
    if _is_probably_print_template_base64(s):
        try:
            return base64.b64decode(s.strip().encode("ascii"), validate=True).decode("utf-8")
        except Exception:
            return s
    return s


def _normalize_print_html_templates_in_config(cfg: Dict[str, Any]) -> Dict[str, Any]:
    """Decode PrintForm HTML templates for the web client runtime.

    Public/editor API exports print_html_template as base64 to make JSON import/export
    safe. The client renderer needs the original HTML/Jinja text.
    """
    if not isinstance(cfg, dict):
        return cfg
    for c in (cfg.get("classes") or []):
        if not isinstance(c, dict):
            continue
        if "print_html_template" in c:
            c["print_html_template"] = _decode_print_html_template(c.get("print_html_template") or "")
    return cfg

# Small per-request-ish memoization. These helpers are called many times while
# rendering Received Nodes; without memoization they repeatedly query device/ack
# tables and can burn CPU on large servers. Values are safe to reuse only for
# the current Flask request/user, so the cache key includes the current user id.
_CURRENT_USER_KEYS_CACHE: Dict[Any, List[str]] = {}
_CURRENT_USER_GROUP_IDS_CACHE: Dict[Tuple[Any, Tuple[str, ...]], set] = {}


def _guess_image_mimetype_from_url(value: str) -> str:
    parsed_path = ""
    try:
        parsed_path = urlparse(str(value or "")).path or ""
    except Exception:
        parsed_path = str(value or "")
    mimetype, _ = mimetypes.guess_type(parsed_path)
    if mimetype and str(mimetype).startswith("image/"):
        return mimetype
    return "image/jpeg"


def _is_cacheable_chat_image_url(value: str) -> bool:
    raw = str(value or "").strip()
    if not raw:
        return False
    parsed = urlparse(raw)
    if parsed.scheme not in ("http", "https"):
        return False
    s3_key_from_public_url = getattr(main, "_s3_key_from_public_url", None)
    try:
        if callable(s3_key_from_public_url) and s3_key_from_public_url(raw):
            return True
    except Exception:
        pass
    # Keep this permissive enough for older/mobile messages that already stored
    # a public image_url outside the configured S3 endpoint, while still only
    # proxying explicit web URLs.
    return True

def _received_nodes_section() -> Dict[str, str]:
    return {"code": RAW_NODES_SECTION_CODE, "name": RAW_NODES_SECTION_NAME}


def _dashboard_section() -> Dict[str, str]:
    return {"code": DASHBOARD_SECTION_CODE, "name": DASHBOARD_SECTION_NAME}


def _class_dashboard_enabled(cls: Dict[str, Any]) -> bool:
    if not isinstance(cls, dict):
        return False
    return bool(cls.get("dashboard_enabled") or cls.get("dashboardEnabled"))


def _has_dashboard_classes(repos: List[models.Repo], db) -> bool:
    try:
        for r in (repos or []):
            parsed = get_parsed_config(r, db)
            if not parsed:
                continue
            cfg = parsed.get("cfg") or {}
            for c in (cfg.get("classes") or []):
                if bool(c.get("hidden")) or _is_print_form_class_type(c):
                    continue
                cn = str(c.get("name") or "").strip()
                if not cn or not _class_dashboard_enabled(c):
                    continue
                if _client_user_can_access_class(r.config_uid, cn):
                    return True
    except Exception as e:
        print("dashboard availability check failed:", e)
    return False


def _with_received_nodes_section(sections: List[Dict[str, str]], repos: Optional[List[models.Repo]] = None) -> List[Dict[str, str]]:
    out = []
    if repos is not None and _has_dashboard_classes(repos, models.db):
        out.append(_dashboard_section())
    out.append(_received_nodes_section())
    for item in (sections or []):
        code = item.get("code") or ""
        if code not in (RAW_NODES_SECTION_CODE, DASHBOARD_SECTION_CODE):
            out.append(item)
    return out


def _default_section_code(sections: List[Dict[str, str]]) -> str:
    """Return the first regular section; keep Received Nodes pinned but not default."""
    for item in (sections or []):
        code = item.get("code") or ""
        if code != RAW_NODES_SECTION_CODE:
            return code
    return (sections[0].get("code") if sections else "") or ""


def _server_model(name: str):
    return getattr(main, name, None)


def _extract_raw_node_class_name(value) -> str:
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, dict):
        for key in ("full_name", "code", "uid", "id", "name", "class_name", "_name"):
            item = value.get(key)
            if isinstance(item, str) and item.strip():
                return item.strip()
    return str(value or "").strip()


def _raw_node_payload(obj) -> Dict[str, Any]:
    payload = getattr(obj, "payload_json", None)
    return payload if isinstance(payload, dict) else {}


def _extract_raw_node_class_json(payload: Dict[str, Any]) -> Dict[str, Any]:
    """Return embedded class JSON from a raw-node payload, if present.

    Android/server raw-node messages can carry either:
      * _class: "ClassName"              -> class is resolved from a client repo/config
      * _class: { ... full class json ... } -> class travels with the raw-node

    The web client must not render raw-nodes as plain _data only; layouts and
    events must see the real class JSON, exactly like PythonScript event flow
    does in app.py.
    """
    payload = payload if isinstance(payload, dict) else {}

    # Prefer the server-side helper when it exists, so web-client and runtime
    # event dispatch keep the same accepted aliases (_class/class/class_json...).
    helper = getattr(main, "_extract_class_json_from_node_json", None)
    if callable(helper):
        try:
            obj = helper(payload)
            if isinstance(obj, dict) and obj:
                return obj
        except Exception:
            pass

    for key in ("_class", "class", "class_json", "_class_json", "schema", "node_class"):
        value = payload.get(key)
        if isinstance(value, dict) and value:
            return value
    return {}


def _raw_node_data_from_payload(payload: Dict[str, Any]) -> Dict[str, Any]:
    """Normalize node data without losing top-level raw-node metadata.

    The form still edits/renders _data, but class/layout/event resolution uses
    the full payload. If an older/raw sender did not wrap fields in _data, use
    the payload fields except structural wrappers.
    """
    payload = payload if isinstance(payload, dict) else {}
    data = payload.get("_data")
    if isinstance(data, dict):
        return dict(data or {})

    skip = {
        "_data", "data", "payload",
        "_class", "class", "class_json", "_class_json", "schema", "node_class",
        "_download_url", "download_url", "raw_node_url", "node_url", "thread_ref",
    }
    return {k: v for k, v in payload.items() if k not in skip}


def _raw_node_download_url(raw_node_id: str) -> str:
    raw_node_id = str(raw_node_id or "").strip()
    explicit = ""
    try:
        explicit = str(request.url_root.rstrip("/")) + f"/api/raw-node/{raw_node_id}"
    except Exception:
        explicit = f"/api/raw-node/{raw_node_id}"
    return explicit


def _raw_node_download_ref(payload: Dict[str, Any], raw_node_id: str) -> str:
    payload = payload if isinstance(payload, dict) else {}
    for key in ("_download_url", "download_url", "raw_node_url", "node_url", "thread_ref"):
        value = str(payload.get(key) or "").strip()
        if value:
            return value
    return _raw_node_download_url(raw_node_id)


def _class_name_from_embedded_class(raw_class: Dict[str, Any], fallback: str = "", payload: Optional[Dict[str, Any]] = None) -> str:
    helper = getattr(main, "_extract_class_name_from_class_json", None)
    if callable(helper):
        try:
            value = helper(raw_class, fallback_class_name=fallback, node_json=payload or {})
            if isinstance(value, str) and value.strip():
                return value.strip()
        except Exception:
            pass
    value = _extract_raw_node_class_name(raw_class)
    return value or str(fallback or "").strip()


def _raw_node_identity(payload: Dict[str, Any], fallback_node_id: str = "") -> Tuple[str, str, Dict[str, Any]]:
    data = _raw_node_data_from_payload(payload)
    embedded_class = _extract_raw_node_class_json(payload)
    class_name = _class_name_from_embedded_class(
        embedded_class,
        fallback=_extract_raw_node_class_name(payload.get("_class") or payload.get("class_name") or data.get("_class")),
        payload=payload,
    )
    node_id = str(payload.get("_id") or payload.get("node_id") or payload.get("node_uid") or data.get("_id") or fallback_node_id or "").strip()
    return class_name, node_id, dict(data or {})


def _merge_raw_class_into_parsed(base_parsed: Optional[Dict[str, Any]], class_name: str, class_obj: Dict[str, Any], payload: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    """Build a parsed-config-like object where embedded raw class wins.

    render_nodalayout_html helpers expect the same parsed structure used by
    normal /section and /node pages. This lets raw-node classes reuse
    CommonLayouts/NodeInput rendering where possible.
    """
    parsed = dict(base_parsed or {})
    cfg = dict((parsed.get("cfg") or {}) if isinstance(parsed.get("cfg"), dict) else {})
    payload = payload if isinstance(payload, dict) else {}

    for key in ("config", "_config", "configuration", "cfg"):
        value = payload.get(key)
        if isinstance(value, dict):
            # Embedded config may contain CommonLayouts or other class-level
            # references; prefer it only for keys that are absent in the repo cfg.
            for k, v in value.items():
                cfg.setdefault(k, v)

    if isinstance(class_obj, dict):
        for key in ("CommonLayouts", "common_layouts"):
            value = class_obj.get(key)
            if isinstance(value, list) and value and "CommonLayouts" not in cfg:
                cfg["CommonLayouts"] = value

    classes = dict((parsed.get("classes") or {}) if isinstance(parsed.get("classes"), dict) else {})
    if class_name and isinstance(class_obj, dict) and class_obj:
        classes[class_name] = class_obj

    parsed["cfg"] = cfg
    parsed["classes"] = classes
    parsed.setdefault("sections", [])
    parsed.setdefault("classes_by_section", {})
    parsed.setdefault("rooms", {})
    return parsed


def _resolve_raw_node_class(payload: Dict[str, Any], class_name: str, preferred_repo=None):
    """Resolve class for a raw-node.

    Returns (repo, parsed, class_obj). If payload carries embedded class JSON,
    that JSON takes precedence; otherwise class_name is resolved from the
    user's configured repositories.
    """
    embedded_class = _extract_raw_node_class_json(payload)
    embedded_name = _class_name_from_embedded_class(embedded_class, fallback=class_name, payload=payload) if embedded_class else ""
    effective_name = embedded_name or str(class_name or "").strip()

    repo = None
    parsed = None

    if embedded_class:
        repo, parsed, _repo_cls = _find_repo_for_raw_node(effective_name, payload)
        if repo is None and preferred_repo is not None:
            repo = preferred_repo
            parsed = get_parsed_config(repo, models.db) or {}
        if repo is None:
            repos = models.Repo.query.filter_by(user_id=_ngenie_effective_user_id()).order_by(models.Repo.id.asc()).all()
            repo = repos[0] if repos else None
            parsed = get_parsed_config(repo, models.db) if repo else {}
        parsed = _merge_raw_class_into_parsed(parsed or {}, effective_name, embedded_class, payload=payload)
        return repo, parsed, embedded_class

    repo, parsed, cls = _find_repo_for_raw_node(effective_name, payload)
    if repo is None and preferred_repo is not None:
        repo = preferred_repo
        parsed = get_parsed_config(repo, models.db) or {}
        cls = ((parsed or {}).get("classes") or {}).get(effective_name) or cls
    return repo, parsed, cls or {}


def _raw_node_search_text(obj, payload: Dict[str, Any], class_name: str, node_id: str, data: Dict[str, Any]) -> str:
    parts = [
        str(getattr(obj, "node_id", "") or ""),
        str(node_id or ""),
        str(class_name or ""),
        str(getattr(obj, "content_type", "") or ""),
    ]
    try:
        parts.append(json.dumps(data, ensure_ascii=False, default=str))
    except Exception:
        parts.append(str(data))
    try:
        parts.append(json.dumps(payload, ensure_ascii=False, default=str))
    except Exception:
        parts.append(str(payload))
    return "\n".join(parts).lower()


_RAW_NODE_URL_RE = re.compile(r"(?:^|[\s\"'(<])(?:https?://[^\s\"'<>]+)?/(?:api/)?raw-node/([^\s\"'<>?#]+)", re.UNICODE)


def _extract_raw_node_id_from_url(value: str) -> str:
    """Extract the DB RawNode.node_id from absolute/relative raw-node links."""
    raw = str(value or "").strip()
    if not raw:
        return ""
    for marker in ("/api/raw-node/", "api/raw-node/", "/raw-node/", "raw-node/"):
        if marker in raw:
            tail = raw.rsplit(marker, 1)[-1]
            return tail.split("?", 1)[0].split("#", 1)[0].strip().strip('"\'<>')
    return ""


def _extract_raw_node_ids_from_message_payload(value, *, deep: bool = False, _depth: int = 0) -> set:
    """Extract raw-node ids from known node-message shapes.

    Keep the default path intentionally shallow. The previous recursive scanner
    walked arbitrary JSON blobs for every request; on production message tables
    that can pin a CPU core. Received Nodes only needs the fields that the
    server/mobile clients actually use: type=node + node_id/node_uid and the
    raw-node URL fields.
    """
    ids = set()

    def add(item):
        item = str(item or "").strip()
        if item:
            ids.add(item)

    if value is None:
        return ids

    if isinstance(value, str):
        raw = value.strip()
        if not raw:
            return ids
        add(_extract_raw_node_id_from_url(raw))
        for match in _RAW_NODE_URL_RE.finditer(" " + raw):
            add(match.group(1))
        if deep and _depth < 2 and raw[:1] in ("{", "["):
            try:
                ids.update(_extract_raw_node_ids_from_message_payload(json.loads(raw), deep=False, _depth=_depth + 1))
            except Exception:
                pass
        return {x for x in ids if x}

    if isinstance(value, (list, tuple)):
        for item in list(value)[:50]:
            ids.update(_extract_raw_node_ids_from_message_payload(item, deep=False, _depth=_depth + 1))
        return {x for x in ids if x}

    if not isinstance(value, dict):
        return ids

    obj = value
    for key in ("download_url", "_download_url", "raw_node_url", "node_url", "thread_ref"):
        add(_extract_raw_node_id_from_url(obj.get(key)))

    ptype = str(obj.get("type") or obj.get("message_type") or "").strip().lower()
    has_raw_url = any(_extract_raw_node_id_from_url(obj.get(k)) for k in ("download_url", "_download_url", "raw_node_url", "node_url", "thread_ref"))
    if ptype in {"node", "node_download", "raw_node", "raw-node", "node_message"} or has_raw_url:
        for key in ("node_id", "node_uid", "raw_node_id", "_id"):
            add(obj.get(key))

    # Common one-level wrappers only. Do not recurse over every payload value.
    for key in ("data", "payload", "node"):
        nested = obj.get(key)
        if isinstance(nested, dict):
            ids.update(_extract_raw_node_ids_from_message_payload(nested, deep=False, _depth=_depth + 1))

    # Batch payloads are explicit and bounded.
    for key in ("items", "nodes"):
        nested = obj.get(key)
        if isinstance(nested, list):
            for item in nested[:50]:
                ids.update(_extract_raw_node_ids_from_message_payload(item, deep=False, _depth=_depth + 1))

    for key in ("items_json", "nodes_json", "payload_json", "data_json"):
        raw = obj.get(key)
        if isinstance(raw, str) and raw.strip()[:1] in ("{", "["):
            try:
                ids.update(_extract_raw_node_ids_from_message_payload(json.loads(raw), deep=False, _depth=_depth + 1))
            except Exception:
                pass

    return {x for x in ids if x}


def _dt_sort_value(value) -> float:
    if value is None:
        return 0.0
    try:
        if getattr(value, "tzinfo", None) is None:
            value = value.replace(tzinfo=timezone.utc)
        return float(value.timestamp())
    except Exception:
        try:
            return float(datetime.fromisoformat(str(value)).timestamp())
        except Exception:
            return 0.0


def _current_user_cache_key():
    try:
        return getattr(current_user, "id", None) or getattr(current_user, "email", None) or id(current_user)
    except Exception:
        return None


def _current_user_keys() -> List[str]:
    """Return message aliases known for the logged-in web/API user.

    This is intentionally cached for the duration of rendering because Received
    Nodes calls it from several helpers. It also includes device_uid values so
    direct device-targeted Android deliveries can be matched by indexed
    outgoing_message_log.target_id instead of scanning JSON payloads globally.
    """
    cache_key = _current_user_cache_key()
    if cache_key in _CURRENT_USER_KEYS_CACHE:
        return list(_CURRENT_USER_KEYS_CACHE.get(cache_key) or [])

    keys = []

    def add(value):
        value = str(value or "").strip()
        if value:
            keys.append(value)

    try:
        add(getattr(current_user, "email", ""))
    except Exception:
        pass
    try:
        add(getattr(current_user, "id", ""))
    except Exception:
        pass
    try:
        add(getattr(current_user, "config_display_name", ""))
    except Exception:
        pass

    user_id = None
    try:
        user_id = getattr(current_user, "id", None)
    except Exception:
        user_id = None

    RoomDevice = _server_model("RoomDevice")
    UserDevice = _server_model("UserDevice")
    OutgoingMessageDeviceAck = _server_model("OutgoingMessageDeviceAck")
    device_uids = set()

    if user_id is not None:
        try:
            if RoomDevice is not None:
                for rd in RoomDevice.query.filter_by(user_id=user_id).all():
                    add(getattr(rd, "user_key", ""))
                    du = str(getattr(rd, "device_uid", "") or "").strip()
                    if du:
                        device_uids.add(du)
                        add(du)
                    extra = getattr(rd, "extra_json", None)
                    if isinstance(extra, dict):
                        for k in ("user_key", "target_user", "recipient", "to_user"):
                            add(extra.get(k))
        except Exception:
            pass
        try:
            if UserDevice is not None:
                for ud in UserDevice.query.filter_by(user_id=user_id).all():
                    du = str(getattr(ud, "device_uid", "") or "").strip()
                    if du:
                        device_uids.add(du)
                        add(du)
                    add(getattr(ud, "android_id", ""))
                    extra = getattr(ud, "extra_json", None)
                    if isinstance(extra, dict):
                        for k in ("user_key", "target_user", "recipient", "to_user"):
                            add(extra.get(k))
        except Exception:
            pass

    # Keep this bounded and only once per request/user. It is a compatibility
    # fallback for old Android rows, not the primary lookup path.
    if device_uids and OutgoingMessageDeviceAck is not None:
        try:
            for ack in OutgoingMessageDeviceAck.query.filter(OutgoingMessageDeviceAck.device_uid.in_(list(device_uids))).order_by(OutgoingMessageDeviceAck.id.desc()).limit(200).all():
                add(getattr(ack, "user_key", ""))
                add(getattr(ack, "ack_by", ""))
                ack_payload = getattr(ack, "ack_payload", None)
                if isinstance(ack_payload, dict):
                    for k in ("user_key", "ack_user", "target_user", "recipient", "to_user"):
                        add(ack_payload.get(k))
        except Exception:
            pass

    seen = set()
    out = []
    for key in keys:
        low = str(key or "").strip().lower()
        if not low or low in seen:
            continue
        seen.add(low)
        out.append(str(key).strip())

    _CURRENT_USER_KEYS_CACHE[cache_key] = list(out)
    return out


def _current_user_group_ids(user_keys: Optional[List[str]] = None) -> set:
    MessageGroupMember = _server_model("MessageGroupMember")
    keys = user_keys if user_keys is not None else _current_user_keys()
    lows_tuple = tuple(sorted({str(k or "").strip().lower() for k in keys if str(k or "").strip()}))
    cache_key = (_current_user_cache_key(), lows_tuple)
    if cache_key in _CURRENT_USER_GROUP_IDS_CACHE:
        return set(_CURRENT_USER_GROUP_IDS_CACHE.get(cache_key) or set())
    if MessageGroupMember is None or not lows_tuple:
        return set()
    try:
        rows = MessageGroupMember.query.filter(func.lower(MessageGroupMember.user_key).in_(list(lows_tuple))).all()
        group_ids = {str(r.group_id or "").strip() for r in rows if str(r.group_id or "").strip()}
    except Exception:
        try:
            rows = MessageGroupMember.query.all()
            lows = set(lows_tuple)
            group_ids = {
                str(r.group_id or "").strip()
                for r in rows
                if str(getattr(r, "user_key", "") or "").strip().lower() in lows and str(r.group_id or "").strip()
            }
        except Exception:
            group_ids = set()
    _CURRENT_USER_GROUP_IDS_CACHE[cache_key] = set(group_ids)
    return group_ids


def _raw_node_payload_has_current_user_hint(payload: Dict[str, Any], include_sender: bool = True, *, user_keys: Optional[List[str]] = None, group_ids: Optional[set] = None) -> bool:
    payload = payload if isinstance(payload, dict) else {}
    keys = user_keys if user_keys is not None else _current_user_keys()
    lows = {k.lower() for k in keys}
    groups = group_ids if group_ids is not None else _current_user_group_ids(keys)
    if not lows:
        return False

    hints = payload.get("_node_message_targets")
    if isinstance(hints, list):
        for hint in hints:
            if not isinstance(hint, dict):
                continue
            target_type = str(hint.get("target_type") or "").strip().lower()
            target_id = str(hint.get("target_id") or hint.get("user_key") or hint.get("group_id") or "").strip()
            if target_type == "user" and target_id.lower() in lows:
                return True
            if target_type == "group" and target_id in groups:
                return True
            if include_sender and str(hint.get("sender_user") or "").strip().lower() in lows:
                return True

    if include_sender and str(payload.get("sender_user") or "").strip().lower() in lows:
        return True

    for key in ("target_user", "user_key", "target_key", "target_id", "recipient", "recipient_user", "to", "to_user", "peer", "peer_user", "receiver"):
        value = str(payload.get(key) or "").strip()
        if value and value.lower() in lows:
            return True

    group_id = str(payload.get("group_id") or payload.get("discussion_group_id") or "").strip()
    if group_id and group_id in groups:
        return True

    # Only known wrappers, not arbitrary recursion.
    for key in ("data", "_data", "payload", "node"):
        nested = payload.get(key)
        if isinstance(nested, dict) and _raw_node_payload_has_current_user_hint(nested, include_sender=include_sender, user_keys=keys, group_ids=groups):
            return True

    return False


def _message_row_visible_to_current_user(row, *, user_keys: Optional[List[str]] = None, group_ids: Optional[set] = None) -> bool:
    keys = user_keys if user_keys is not None else _current_user_keys()
    lows = {k.lower() for k in keys}
    if not lows:
        return False
    groups = group_ids if group_ids is not None else _current_user_group_ids(keys)
    payload = getattr(row, "payload_json", None)
    payload = payload if isinstance(payload, dict) else {}

    sender = str(getattr(row, "sender_user", "") or payload.get("sender_user") or "").strip().lower()
    if sender and sender in lows:
        return True

    target_type = str(getattr(row, "target_type", "") or "").strip().lower()
    target_id = str(getattr(row, "target_id", "") or "").strip()
    if target_type in {"user", "device"} and target_id.lower() in lows:
        return True
    if target_type == "group" and target_id in groups:
        return True

    for key in ("user_key", "target_key", "target_user", "target_id", "recipient", "recipient_user", "to", "to_user", "peer", "peer_user", "receiver"):
        value = str(payload.get(key) or "").strip().lower()
        if value and value in lows:
            return True
    group_id = str(payload.get("group_id") or payload.get("discussion_group_id") or "").strip()
    if group_id and group_id in groups:
        return True
    return False


def _query_current_user_message_rows(limit: int = 1000):
    """Return only rows that are plausibly related to current user, using DB indexes.

    No network calls and no full-table JSON scan. A small recent fallback catches
    older direct-device payloads where only payload.user_key carries the user.
    """
    OutgoingMessageLog = _server_model("OutgoingMessageLog")
    if OutgoingMessageLog is None:
        return []

    keys = _current_user_keys()
    groups = _current_user_group_ids(keys)
    clauses = []
    if keys:
        clauses.append(OutgoingMessageLog.sender_user.in_(keys))
        clauses.append(and_(OutgoingMessageLog.target_type.in_(("user", "device")), OutgoingMessageLog.target_id.in_(keys)))
    if groups:
        clauses.append(and_(OutgoingMessageLog.target_type == "group", OutgoingMessageLog.target_id.in_(list(groups))))

    rows_by_id = {}
    try:
        if clauses:
            rows = OutgoingMessageLog.query.filter(or_(*clauses)).order_by(
                OutgoingMessageLog.id.desc(),
            ).limit(limit).all()
            for row in rows:
                rows_by_id[getattr(row, "id", id(row))] = row
    except Exception:
        pass

    # Bounded compatibility fallback: catches Android rows whose route target
    # is not one of the web user's aliases but payload.user_key/peer_user is.
    # Order by primary key only; created_at is not guaranteed to be indexed.
    try:
        fallback_limit = min(1000, max(100, int(limit)))
        rows = OutgoingMessageLog.query.order_by(
            OutgoingMessageLog.id.desc(),
        ).limit(fallback_limit).all()
        for row in rows:
            if _message_row_visible_to_current_user(row, user_keys=keys, group_ids=groups):
                rows_by_id[getattr(row, "id", id(row))] = row
    except Exception:
        pass

    out = list(rows_by_id.values())
    out.sort(key=lambda r: (_dt_sort_value(getattr(r, "created_at", None)), getattr(r, "id", 0) or 0), reverse=True)
    return out[:limit]


def _raw_node_ids_from_message_history() -> set:
    """Raw node ids delivered to/currently sent by the current user.

    This is based on indexed OutgoingMessageLog columns plus a tiny recent
    fallback. It never downloads /api/raw-node URLs and never loops over RawNode.
    """
    ids = []
    seen = set()
    for row in _query_current_user_message_rows(limit=1000):
        payload = getattr(row, "payload_json", None)
        for raw_id in _extract_raw_node_ids_from_message_payload(payload):
            if raw_id and raw_id not in seen:
                seen.add(raw_id)
                ids.append(raw_id)
    return set(ids)


def _current_user_can_access_raw_node(raw_node_id: str, obj=None, include_sender: bool = True) -> bool:
    raw_node_id = str(raw_node_id or "").strip()
    if not raw_node_id or not getattr(current_user, "is_authenticated", False):
        return False

    RawNode = _server_model("RawNode")
    if obj is None and RawNode is not None:
        try:
            obj = RawNode.query.filter_by(node_id=raw_node_id).first()
        except Exception:
            obj = None

    try:
        if obj is not None and getattr(obj, "owner_user_id", None) == getattr(current_user, "id", None):
            return True
    except Exception:
        pass

    keys = _current_user_keys()
    groups = _current_user_group_ids(keys)
    payload = _raw_node_payload(obj) if obj is not None else {}
    if _raw_node_payload_has_current_user_hint(payload, include_sender=include_sender, user_keys=keys, group_ids=groups):
        return True

    return raw_node_id in _raw_node_ids_from_message_history()


def _message_dict_visible_to_current_user(msg: Dict[str, Any]) -> bool:
    msg = msg if isinstance(msg, dict) else {}
    keys = _current_user_keys()
    lows = {k.lower() for k in keys}
    if not lows:
        return False
    group_ids = _current_user_group_ids(keys)

    data = msg.get("data") if isinstance(msg.get("data"), dict) else {}
    sender = str(msg.get("sender_user") or data.get("sender_user") or "").strip().lower()
    if sender and sender in lows:
        return True
    target_type = str(msg.get("target_type") or data.get("target_type") or "").strip().lower()
    target_id = str(msg.get("target_id") or data.get("target_id") or "").strip()
    if target_type in {"user", "device"} and target_id.lower() in lows:
        return True
    if target_type == "group" and target_id in group_ids:
        return True
    group_id = str(msg.get("group_id") or data.get("group_id") or "").strip()
    if group_id and group_id in group_ids:
        return True
    user_key = str(msg.get("user_key") or data.get("user_key") or "").strip().lower()
    if user_key and user_key in lows:
        return True
    return False


def _current_user_can_access_node_discussion(node_id: str) -> bool:
    node_id = str(node_id or "").strip()
    if not node_id:
        return False
    if _current_user_can_access_raw_node(node_id):
        return True
    NodeDiscussionMessage = _server_model("NodeDiscussionMessage")
    if NodeDiscussionMessage is not None:
        try:
            keys = _current_user_keys()
            groups = _current_user_group_ids(keys)
            rows = NodeDiscussionMessage.query.filter_by(node_id=node_id).order_by(NodeDiscussionMessage.id.desc()).limit(200).all()
            if any(_message_row_visible_to_current_user(r, user_keys=keys, group_ids=groups) for r in rows):
                return True
        except Exception:
            pass
    return False

def _find_repo_for_raw_node(class_name: str, payload: Optional[Dict[str, Any]] = None):
    repos = models.Repo.query.filter_by(user_id=_ngenie_effective_user_id()).all()
    if not repos:
        return None, None, None

    payload = payload if isinstance(payload, dict) else {}
    class_obj = payload.get("_class")
    possible_cfg_uids = []
    if isinstance(class_obj, dict):
        for key in ("config_uid", "configuration_uid", "config", "repo_uid"):
            val = str(class_obj.get(key) or "").strip()
            if val:
                possible_cfg_uids.append(val)

    def repo_score(repo):
        score = 0
        if possible_cfg_uids and str(repo.config_uid or "") in possible_cfg_uids:
            score += 10
        return score

    candidates = sorted(repos, key=repo_score, reverse=True)
    first_parsed = None
    first_repo = candidates[0] if candidates else None
    for repo in candidates:
        parsed = get_parsed_config(repo, models.db)
        if first_parsed is None:
            first_parsed = parsed
        cls = ((parsed or {}).get("classes") or {}).get(class_name) if class_name else None
        if cls:
            return repo, parsed, cls
    return first_repo, first_parsed, None


def _build_raw_node_items(q: str = "") -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
    """Build Received Nodes directly from local DB rows.

    Important: do not call /api/raw-node URLs from the web server. Those links
    are hosted by this same Flask app, and the source of truth is RawNode. The
    list is scoped by OutgoingMessageLog/RawNode metadata, then RawNode rows are
    fetched by node_id with indexed queries.
    """
    RawNode = _server_model("RawNode")
    if RawNode is None:
        return [], {"classes_ui": [], "table_headers": ["Created", "Updated"], "filter_indexes": []}

    delivered_ids = _raw_node_ids_from_message_history()
    user_id = getattr(current_user, "id", None)
    rows_by_id = {}

    # Nodes received/sent through messages: fetch by indexed RawNode.node_id.
    if delivered_ids:
        try:
            for obj in RawNode.query.filter(RawNode.node_id.in_(list(delivered_ids))).all():
                raw_id = str(getattr(obj, "node_id", "") or "").strip()
                if raw_id:
                    rows_by_id[raw_id] = obj
        except Exception:
            pass

    # Nodes uploaded by this web/API user: indexed owner_user_id.
    if user_id is not None:
        try:
            q_owner = RawNode.query.filter_by(owner_user_id=user_id)
            try:
                owner_rows = q_owner.order_by(RawNode.updated_at.desc(), RawNode.created_at.desc(), RawNode.id.desc()).limit(500).all()
            except Exception:
                owner_rows = q_owner.order_by(RawNode.id.desc()).limit(500).all()
            for obj in owner_rows:
                raw_id = str(getattr(obj, "node_id", "") or "").strip()
                if raw_id:
                    rows_by_id[raw_id] = obj
        except Exception:
            pass

    rows = list(rows_by_id.values())
    rows.sort(
        key=lambda obj: (
            _dt_sort_value(getattr(obj, "updated_at", None) or getattr(obj, "created_at", None)),
            getattr(obj, "id", 0) or 0,
        ),
        reverse=True,
    )
    rows = rows[:500]

    # Resolve repositories/classes once as much as possible.
    items = []
    q_low = str(q or "").strip().lower()
    repo_cache: Dict[str, Tuple[Any, Any, Any]] = {}

    for obj in rows:
        raw_id = str(getattr(obj, "node_id", "") or "").strip()
        if not raw_id:
            continue

        payload = _raw_node_payload(obj)
        class_name, payload_node_id, data = _raw_node_identity(payload, raw_id)
        node_id = payload_node_id or raw_id
        data.setdefault("_id", node_id)
        if class_name:
            data.setdefault("_class", class_name)
        data.setdefault("_raw_node_id", raw_id)
        data.setdefault("_download_url", _raw_node_download_ref(payload, raw_id))

        if q_low and q_low not in _raw_node_search_text(obj, payload, class_name, node_id, data):
            continue

        embedded_key = "embedded:" + str(id(payload.get("_class"))) if isinstance(payload.get("_class"), dict) else ""
        cache_key = embedded_key or class_name or "raw-node"
        if cache_key in repo_cache:
            repo, parsed, cls = repo_cache[cache_key]
        else:
            repo, parsed, cls = _resolve_raw_node_class(payload, class_name)
            repo_cache[cache_key] = (repo, parsed, cls)

        if isinstance(cls, dict):
            resolved_name = _class_name_from_embedded_class(cls, fallback=class_name, payload=payload)
            if resolved_name:
                class_name = resolved_name
                data.setdefault("_class", class_name)
        repo_id = getattr(repo, "id", 0) or 0
        repo_name = getattr(repo, "display_name", "") or getattr(repo, "name", "") or RAW_NODES_SECTION_NAME
        display_image_html = ""
        tv = {
            "Created": getattr(obj, "created_at", None).isoformat() if getattr(obj, "created_at", None) else "",
            "Updated": getattr(obj, "updated_at", None).isoformat() if getattr(obj, "updated_at", None) else "",
        }
        try:
            if repo and parsed and cls:
                cover_layout = cls.get("cover_image")
                cover_web_layout = cls.get("display_image_web") or ""
                layout_to_use = cover_web_layout if str(cover_web_layout or "").strip() else cover_layout
                if layout_to_use is not None:
                    layout_to_use = resolve_common_layout(parsed, layout_to_use)
                    _fill_nodeinput_views(repo, parsed, layout_to_use, data)
                    display_image_html = _cover_with_tags(_wrap_client_tpl_html(str(render_nodalayout_html(
                        layout_to_use,
                        data,
                        assets_base_dir=_userfiles_dir_for_repo(repo),
                        context=_nl_context(repo, class_name=class_name, node_id=node_id),
                    ) or ""), data), data, bool((cls or {}).get("show_tag_cloud")))
        except Exception:
            display_image_html = ""
        if not display_image_html:
            display_image_html = _render_tags_html(data)

        items.append({
            "repo": repo_name,
            "repo_id": repo_id,
            "class": class_name or "raw-node",
            "id": node_id,
            "raw_node_id": raw_id,
            "data": data,
            "class_obj": cls or {},
            "is_raw_node": True,
            "is_custom_process": True,
            "display_image_html": display_image_html,
            "table_values": tv,
            "use_standard_commands": False,
            "repo_uid": getattr(repo, "config_uid", "") or "",
        })

    return items, {
        "classes_ui": [],
        "table_headers": ["Created", "Updated"],
        "start_menu_cmds_ui": [],
        "filter_indexes": [],
    }


_CLASS_VIEW_RE = re.compile(r"\{([\w.]+)\}", re.UNICODE)

def _render_class_record_view(parsed: Optional[Dict[str, Any]], class_name: str, node_id: str, data: Optional[Dict[str, Any]]) -> str:
    """Render class-level record view template using node data."""
    data = data if isinstance(data, dict) else {}

    if isinstance(data.get("_view"), str) and data.get("_view", "").strip():
        return data.get("_view", "").strip()

    cls_cfg: Dict[str, Any] = {}
    try:
        cls_cfg = ((parsed or {}).get("classes") or {}).get(class_name) or {}
    except Exception:
        cls_cfg = {}

    tpl = str(cls_cfg.get("record_view") or "").strip()
    if tpl:
        def repl(m: re.Match) -> str:
            key = m.group(1)
            val = data.get(key)
            return "" if val is None else str(val)

        rendered = _CLASS_VIEW_RE.sub(repl, tpl).strip()
        if rendered:
            return rendered

    return str(node_id or "")

# In-memory parsed configuration cache. Local configurations are cached per
# configuration revision *and* ACL signature. The scheduler and simultaneous
# web users must never evict each other's ACL-specific snapshots.
CONFIG_MEM: Dict[Any, Dict[str, Any]] = {}
CONFIG_MEM_LOCK = threading.RLock()


def _invalidate_repo_config_mem(repo_id: Any) -> None:
    """Drop every parsed-config cache entry belonging to one repository."""
    try:
        rid = int(repo_id)
    except Exception:
        return
    with CONFIG_MEM_LOCK:
        for key in list(CONFIG_MEM.keys()):
            same_repo = key == rid or (
                isinstance(key, tuple)
                and len(key) > 0
                and key[0] == rid
            )
            if same_repo:
                CONFIG_MEM.pop(key, None)


def _client_repo_actor(repo):
    """Return the repository owner/effective user without relying on current_user.

    Background scheduler threads have an application context but no authenticated
    request. Falling back to ``current_user`` there produced an anonymous ACL
    snapshot and invalidated the owner's hot cache.
    """
    try:
        user_id = int(getattr(repo, "user_id", 0) or 0)
    except Exception:
        user_id = 0
    if not user_id:
        return None
    User = _client_root_model("User")
    if User is None:
        return None
    try:
        return models.db.session.get(User, user_id)
    except Exception:
        try:
            return User.query.filter_by(id=user_id).first()
        except Exception:
            return None


# in-memory cache for exec()'ed server handlers modules (per config_uid)
SERVER_HANDLERS_MEM: Dict[str, Dict[str, Any]] = {}
_SERVER_HANDLERS_NS_MEM: Dict[str, Dict[str, Any]] = {}
_SERVER_NODE_CLASS_MEM: Dict[Tuple[str, str, str], Any] = {}
_SERVER_HANDLERS_LOCK = threading.RLock()


# -------- client settings (stored in client.sqlite) --------

def _split_dataset_item_uid(ds_name: str, item_uid: str) -> Tuple[str, str]:
    """Return (dataset_name, item_id) from DatasetLink value.

    Dataset links are self-describing: ``Goods$123`` already contains the
    dataset name, so cover layouts do not need an additional ``dataset`` field.
    If ``ds_name`` is explicitly provided, it is kept as fallback/override for
    old layouts that store only ``123``.
    """
    ds = str(ds_name or "").strip()
    uid = str(item_uid or "").strip()
    if "$" in uid:
        left, right = uid.split("$", 1)
        if not ds:
            ds = left.strip()
        return ds, right.strip()
    return ds, uid


def _render_template_fields(template: str, data: Dict[str, Any]) -> str:
    """Small {field} renderer for dataset/record views."""
    template = str(template or "")
    if not template:
        return ""

    def repl(match: re.Match) -> str:
        field_name = match.group(1)
        value = data.get(field_name, "")
        return str(value) if value is not None else ""

    return _CLASS_VIEW_RE.sub(repl, template).strip()


def _get_dataset_item_direct(config_uid: str, ds_name: str, item_id: str) -> Optional[Dict[str, Any]]:
    """Directly get dataset item from database without HTTP."""
    try:
        ds_name, item_id = _split_dataset_item_uid(ds_name, item_id)
        if not ds_name or not item_id:
            return None

        Configuration = main.Configuration
        Dataset = main.Dataset
        DatasetItem = main.DatasetItem

        cfg = Configuration.query.filter_by(uid=config_uid).first() if config_uid else None
        ds = None

        # Сначала ищем в текущей конфигурации репозитория.
        if cfg:
            ds = Dataset.query.filter_by(config_id=cfg.id, name=ds_name).first()

        # В ссылке DatasetLink/DatasetInput хранится только Dataset$Id, без config_uid.
        # Поэтому, если dataset не найден в текущей конфе, пробуем найти такой dataset
        # в конфигурациях репозиториев текущего пользователя. Это особенно важно для
        # node_form и динамических layout-ов, где layout может прийти из другого контекста.
        if not ds:
            try:
                repos = models.Repo.query.filter_by(user_id=_ngenie_effective_user_id()).all() if current_user.is_authenticated else []
                for r in repos:
                    cfg_uid = str(getattr(r, "config_uid", "") or "").strip()
                    if not cfg_uid:
                        continue
                    c2 = Configuration.query.filter_by(uid=cfg_uid).first()
                    if not c2:
                        continue
                    ds2 = Dataset.query.filter_by(config_id=c2.id, name=ds_name).first()
                    if ds2:
                        cfg = c2
                        ds = ds2
                        break
            except Exception:
                pass

        if not ds:
            return None

        item = DatasetItem.query.filter_by(dataset_id=ds.id, item_id=item_id).first()
        if not item:
            return None

        data = item.data or {}
        if not isinstance(data, dict):
            data = {}

        view = str(data.get("_view") or "").strip()
        if not view:
            view = _render_template_fields(ds.view_template or "", data)
        if not view:
            view = str(data.get("title") or data.get("name") or item_id)

        return {
            "_id": item_id,
            "_view": view,
            "_data": data,
            "_dataset": ds_name,
        }
    except Exception as e:
        print(f"Error getting dataset item: {e}")
        return None

def _get_setting(key: str, default: str = "") -> str:
    """Get per-user client setting."""
    if not current_user.is_authenticated:
        return default
    row = models.ClientSetting.query.filter_by(user_id=current_user.id, key=key).first()
    if not row:
        return default
    return (row.value or "")


def _set_setting(key: str, value: str) -> None:
    if not current_user.is_authenticated:
        return
    row = models.ClientSetting.query.filter_by(user_id=current_user.id, key=key).first()
    if not row:
        row = models.ClientSetting(user_id=current_user.id, key=key, value=value or "")
        models.db.session.add(row)
    else:
        row.value = value or ""
    models.db.session.commit()


def _client_is_admin() -> bool:
    """Client-side admin switch: currently mapped to designer/admin users."""
    try:
        if not current_user.is_authenticated:
            return False
        for attr in ("is_admin", "admin", "can_designer"):
            val = getattr(current_user, attr, False)
            if callable(val):
                try:
                    val = val()
                except Exception:
                    val = False
            if bool(val):
                return True
        return False
    except Exception:
        return False


def _client_show_node_json() -> bool:
    return _client_is_admin() and str(_get_setting("show_node_json", "0")).lower() in ("1", "true", "yes", "on")


def _client_ngenie_enabled() -> bool:
    """Whether the nGenie web UI is visible for the current user.

    This setting controls only the interactive web entry points. Runtime calls
    from Python/NodaScript (``ngenie(...)``) remain available for handlers and
    timers even when the UI is hidden. The default is enabled to preserve the
    existing behaviour.
    """
    return str(_get_setting("ngenie_enabled", "1")).lower() not in ("0", "false", "no", "off")


def _client_root_app_module():
    return sys.modules.get("app") or main


def _client_runtime_system_user_payload() -> Dict[str, Any]:
    """Logged-in ``_System/_User`` payload for web-client Node handlers.

    For browser requests Flask-Login's ``current_user`` is authoritative.  A
    value in ``g.api_user`` can represent an API/repository transport account
    and must not override the account that actually opened the web client.
    ``g.api_user`` is only a fallback for non-Flask-Login request contexts.
    """
    if not has_request_context():
        return {}

    actor = None
    try:
        if getattr(current_user, "is_authenticated", False):
            actor = current_user
    except Exception:
        actor = None

    if actor is None:
        try:
            actor = getattr(g, "api_user", None)
        except Exception:
            actor = None

    root = _client_root_app_module()
    resolver = getattr(root, "_system_user_payload", None)
    if callable(resolver):
        try:
            payload = resolver(actor)
            return payload if isinstance(payload, dict) else {}
        except Exception:
            pass
    return {}


def _client_remote_system_user_headers() -> Dict[str, str]:
    """Preserve the interactive user when Repo HTTP uses owner credentials."""
    payload = _client_runtime_system_user_payload()
    data = payload.get("_data") if isinstance(payload.get("_data"), dict) else {}
    user_uid = str(payload.get("_id") or data.get("_id") or "").strip()
    user_login = str(data.get("login") or data.get("email") or payload.get("login") or payload.get("email") or "").strip()
    headers: Dict[str, str] = {}
    if user_uid:
        headers["X-System-User-Id"] = user_uid
    if user_login:
        headers["X-System-User-Login"] = user_login
    return headers


def _client_system_user_view(uid: str) -> str:
    root = _client_root_app_module()
    resolver = getattr(root, "resolve_system_user_view_global", None)
    if not callable(resolver):
        return ""
    try:
        value = resolver(uid, default="")
        return str(value or "").strip()
    except Exception:
        return ""


def _client_system_user_picker_items(q: str = "", limit: int = 80) -> List[Dict[str, Any]]:
    """NodeInput exception for the reserved ``_System/_User`` class only."""
    root = _client_root_app_module()
    owner_resolver = getattr(root, "_system_owner_id_for_user", None)
    ensure_config = getattr(root, "_ensure_system_config_for_owner", None)
    users_query = getattr(root, "_owner_scope_users_query", None)
    payload_builder = getattr(root, "_system_user_node_payload_from_user", None)
    if not all(callable(fn) for fn in (owner_resolver, ensure_config, users_query, payload_builder)):
        return []

    try:
        owner_id = owner_resolver(current_user)
        system_cfg = ensure_config(owner_id, sync_users=True)
        if not owner_id or system_cfg is None:
            return []
        db_obj = getattr(root, "db", None)
        user_model = getattr(root, "User", None)
        if db_obj is None or user_model is None:
            return []
        sql_users = db_obj.session.execute(
            users_query(owner_id).order_by(user_model.email)
        ).scalars().all()
    except Exception:
        return []

    needle = str(q or "").strip().casefold()
    hard_limit = max(1, min(int(limit or 80), 200))
    items: List[Dict[str, Any]] = []
    for sql_user in sql_users:
        try:
            payload = payload_builder(sql_user, system_cfg.uid, owner_id=owner_id)
            data = payload.get("_data") if isinstance(payload, dict) else None
            if not isinstance(data, dict):
                continue
            uid = str(data.get("_id") or "").strip()
            login = str(data.get("login") or data.get("email") or "").strip()
            name = str(data.get("name") or login or uid).strip()
            if needle and needle not in f"{login} {name} {uid}".casefold():
                continue
            items.append({
                "uid": uid,
                "_id": str(payload.get("_id") or ""),
                "_class": "_User",
                "_view": login or name,
                "cover_html": "",
                "data": data,
                "repo_id": None,
                "repo_uid": str(system_cfg.uid or ""),
            })
            if len(items) >= hard_limit:
                break
        except Exception:
            continue
    return items


def _client_repo_is_local(repo: Optional[models.Repo]) -> bool:
    """True only when this repo has a live configuration in the current DB.

    Historically an empty ``base_url`` was treated as proof that a repository
    was local.  Imported/cached repositories can also have an empty base URL,
    and old ``/repos/add`` code even discarded the remote host.  That made every
    request try ``fetch_config_from_local_db()`` first and print a permanent
    ``Configuration ... not found`` warning before falling back to RepoConfig.
    """
    if repo is None:
        return False
    try:
        base_url = (getattr(repo, "base_url", "") or "").strip().rstrip("/")
        current = (request.host_url or "").rstrip("/")
        if base_url and base_url != current:
            return False
        config_uid = str(getattr(repo, "config_uid", "") or "").strip()
        return bool(config_uid and _client_cfg_by_uid(config_uid) is not None)
    except Exception:
        return False


def _client_root_model(name: str):
    root_mod = sys.modules.get("models")
    obj = getattr(main, name, None)
    if obj is not None:
        return obj
    return getattr(root_mod, name, None) if root_mod is not None else None


def _client_acl_cache() -> Dict[Any, Any]:
    try:
        cache = getattr(g, '_client_acl_cache', None)
        if cache is None:
            cache = {}
            g._client_acl_cache = cache
        return cache
    except Exception:
        return {}


def _client_cfg_by_uid(config_uid: str):
    Configuration = _client_root_model("Configuration")
    if not Configuration:
        return None
    uid = str(config_uid or "")
    cache = _client_acl_cache()
    key = ('config', uid)
    if key in cache:
        return cache[key]
    try:
        cfg = models.db.session.execute(select(Configuration).where(Configuration.uid == uid)).scalar_one_or_none()
    except Exception:
        try:
            cfg = Configuration.query.filter_by(uid=uid).first()
        except Exception:
            cfg = None
    cache[key] = cfg
    return cfg


def _client_user_is_config_scope_owner(user, cfg) -> bool:
    if not user or not cfg:
        return False
    user_id = getattr(user, "id", None)
    cache = _client_acl_cache()
    key = ('scope_owner', user_id, getattr(cfg, 'id', None))
    if key in cache:
        return bool(cache[key])
    try:
        if int(getattr(cfg, "user_id", 0) or 0) == int(user_id or 0):
            cache[key] = True
            return True
    except Exception:
        if getattr(cfg, "user_id", None) == user_id:
            cache[key] = True
            return True
    result = False
    try:
        User = _client_root_model("User")
        cfg_owner_id = int(getattr(cfg, "user_id", 0) or 0)
        owner_key = ('sql_user', cfg_owner_id)
        cfg_owner = cache.get(owner_key)
        if owner_key not in cache and User is not None:
            cfg_owner = models.db.session.get(User, cfg_owner_id)
            cache[owner_key] = cfg_owner
        result = bool(cfg_owner and getattr(cfg_owner, "parent_user_id", None) is not None and int(cfg_owner.parent_user_id) == int(user_id))
    except Exception:
        try:
            cfg_owner = getattr(cfg, "user", None)
            result = bool(cfg_owner and getattr(cfg_owner, "parent_user_id", None) == user_id)
        except Exception:
            result = False
    cache[key] = result
    return result


def _client_assigned_profile_rows(user, config_id=None, class_name=None):
    if not user:
        return []
    UserProfile = _client_root_model("UserProfile")
    UserProfileRole = _client_root_model("UserProfileRole")
    UserProfileClassAccess = _client_root_model("UserProfileClassAccess")
    if not (UserProfile and UserProfileRole and UserProfileClassAccess):
        return []
    user_id = int(getattr(user, "id", 0) or 0)
    cache = _client_acl_cache()
    all_key = ('profile_rows_all', user_id)
    if all_key not in cache:
        try:
            q = select(UserProfileClassAccess).join(UserProfile, UserProfileClassAccess.profile_id == UserProfile.id).join(
                UserProfileRole, UserProfileRole.profile_id == UserProfile.id
            ).where(UserProfileRole.user_id == user_id)
            cache[all_key] = models.db.session.execute(q).scalars().all()
        except Exception as e:
            print("client profile rows error:", e)
            cache[all_key] = []
    rows = cache.get(all_key) or []
    if config_id is not None:
        rows = [r for r in rows if int(getattr(r, 'config_id', 0) or 0) == int(config_id)]
    if class_name is not None:
        rows = [r for r in rows if str(getattr(r, 'class_name', '') or '') == str(class_name)]
    return rows


def _client_acl_signature(user=None) -> str:
    user = user or current_user
    uid = int(getattr(user, "id", 0) or 0)
    if not uid:
        return "anonymous"
    # Profiles never restrict the main account.  Avoid even loading role rows
    # when computing the live-config cache key for the owner fast path.
    if getattr(user, "parent_user_id", None) in (None, 0, ""):
        return f"owner:{uid}"
    UserProfileRole = _client_root_model("UserProfileRole")
    UserProfileClassAccess = _client_root_model("UserProfileClassAccess")
    if not (UserProfileRole and UserProfileClassAccess):
        return f"u:{uid}"
    try:
        rows = _client_assigned_profile_rows(user, None, None)
        role_ids = sorted({str(getattr(r, "profile_id", "")) for r in rows})
        max_upd = ""
        for r in rows:
            v = getattr(r, "updated_at", None)
            sv = v.isoformat() if hasattr(v, "isoformat") else str(v or "")
            if sv > max_upd:
                max_upd = sv
        return f"u:{uid}|p:{','.join(role_ids)}|n:{len(rows)}|m:{max_upd}"
    except Exception:
        return f"u:{uid}|t:{int(time.time())}"


def _client_user_can_access_config(config_uid: str, user=None) -> bool:
    user = user or current_user
    uid = str(config_uid or "")
    user_id = getattr(user, 'id', None)
    cache = _client_acl_cache()
    key = ('config_access', user_id, uid)
    if key in cache:
        return bool(cache[key])
    cfg = _client_cfg_by_uid(uid)
    if not cfg or not user:
        cache[key] = False
        return False
    try:
        if bool(getattr(cfg, "is_system", False)):
            root = _client_root_app_module()
            result = bool(hasattr(root, "user_can_access_config") and root.user_can_access_config(user, uid))
            cache[key] = result
            return result
        if _client_user_is_config_scope_owner(user, cfg):
            cache[key] = True
            return True
        UserConfigAccess = _client_root_model("UserConfigAccess")
        if not UserConfigAccess:
            cache[key] = False
            return False
        result = bool(models.db.session.execute(select(UserConfigAccess.id).where(
            UserConfigAccess.user_id == int(user_id or 0),
            UserConfigAccess.config_id == int(getattr(cfg, "id", 0) or 0),
        )).scalar_one_or_none())
        cache[key] = result
        return result
    except Exception as e:
        print("client config ACL error:", e)
        cache[key] = False
        return False


def _client_user_can_access_class(config_uid: str, class_name: str, user=None) -> bool:
    user = user or current_user
    uid = str(config_uid or "")
    class_name = str(class_name or "").strip()
    user_id = getattr(user, 'id', None)
    cache = _client_acl_cache()
    key = ('class_access', user_id, uid, class_name)
    if key in cache:
        return bool(cache[key])
    cfg = _client_cfg_by_uid(uid)
    if not cfg or not user or not class_name:
        cache[key] = False
        return False
    try:
        if bool(getattr(cfg, "is_system", False)):
            result = class_name == "_User" and _client_user_can_access_config(uid, user=user)
            cache[key] = result
            return result
        if _client_user_is_config_scope_owner(user, cfg):
            cache[key] = True
            return True
        if not _client_user_can_access_config(uid, user=user):
            cache[key] = False
            return False

        all_profile_rows = _client_assigned_profile_rows(user, None, None)
        if not all_profile_rows:
            cache[key] = True
            return True
        result = any(
            int(getattr(r, "config_id", 0) or 0) == int(getattr(cfg, "id", 0) or 0)
            and str(getattr(r, "class_name", "") or "") == class_name
            and bool(getattr(r, "visible", False))
            for r in all_profile_rows
        )
        cache[key] = result
        return result
    except Exception as e:
        print("client class ACL error:", e)
        cache[key] = False
        return False


def _client_rls_get(source, path):
    cur = source or {}
    for part in str(path or "").split("."):
        if not part:
            continue
        if isinstance(cur, dict):
            cur = cur.get(part)
        else:
            return None
    return cur


def _client_rls_list(value):
    if value is None:
        return []
    if isinstance(value, (list, tuple, set)):
        return list(value)
    return [value]


def _client_rls_normalize_scalar(value, value_type="string"):
    value_type = str(value_type or "string").strip().lower()
    if isinstance(value, dict):
        value = value.get("_id") or value.get("id") or value.get("uid") or value.get("value") or ""
    if value_type in ("node", "nodelink", "nodeinput", "class"):
        text = str(value or "").strip()
        if "$" in text:
            text = text.split("$")[-1].strip()
        return text
    if value_type in ("number", "numeric", "float", "integer", "int"):
        try:
            return str(float(str(value).replace(",", ".")))
        except Exception:
            return str(value or "").strip()
    if value_type in ("boolean", "bool"):
        return "true" if str(value).strip().lower() in ("1", "true", "yes", "y", "да", "истина", "on") else "false"
    if value_type in ("date", "datetime"):
        return str(value or "").strip()[:10]
    return str(value or "").strip()


def _client_normalize_list(value):
    if value is None:
        return []
    if isinstance(value, (list, tuple, set)):
        return [str(x).strip() for x in value if str(x).strip()]
    text_value = str(value).replace("\n", ",").replace(";", ",")
    return [x.strip() for x in text_value.split(",") if x.strip()]


def _client_eval_rls_rules(data: dict, rules) -> bool:
    if not rules:
        return True
    if isinstance(rules, str):
        try:
            rules = json.loads(rules)
        except Exception:
            rules = []
    for row in (rules or []):
        if not isinstance(row, dict):
            continue
        field = row.get("field") or row.get("field_name") or row.get("name")
        op = str(row.get("op") or row.get("operator") or row.get("mode") or "in").strip().lower()
        value_type = str(row.get("value_type") or row.get("type") or "string").strip().lower()
        values = [_client_rls_normalize_scalar(x, value_type) for x in _client_normalize_list(row.get("values", row.get("value", row.get("list", []))))]
        actual_values = [_client_rls_normalize_scalar(x, value_type) for x in _client_rls_list(_client_rls_get(data, field))]
        actual = actual_values[0] if actual_values else ""
        expected = values[0] if values else ""
        if op in ("=", "eq", "equal", "equals"):
            match = (actual == expected)
        elif op in ("!=", "<>", "ne", "neq", "not_equal", "not equal"):
            match = (actual != expected)
        elif op in ("not", "not in", "not_in", "не", "не в списке", "exclude"):
            match = not any(x in values for x in actual_values)
        else:
            match = any(x in values for x in actual_values) if values else bool(actual_values)
        if not match:
            return False
    return True


def _client_eval_rls_handler(row, config_uid: str, class_name: str, node_id: str, data: dict):
    root = _client_root_app_module()
    fn = getattr(root, "_eval_rls_handler", None)
    if callable(fn):
        return fn(row, config_uid, class_name, node_id, data or {})
    return None


def _client_user_can_access_node(config_uid: str, class_name: str, node_id: str, data: Optional[Dict[str, Any]] = None, user=None) -> bool:
    user = user or current_user
    if not _client_user_can_access_class(config_uid, class_name, user=user):
        return False
    uid = str(config_uid or '')
    class_name = str(class_name or '')
    node_id = str(node_id or '')
    cfg = _client_cfg_by_uid(uid)
    if not cfg or _client_user_is_config_scope_owner(user, cfg):
        return True
    cache = _client_acl_cache()
    user_id = getattr(user, 'id', None)
    rows_key = ('visible_rls_rows', user_id, getattr(cfg, 'id', None), class_name)
    if rows_key not in cache:
        rows = [r for r in _client_assigned_profile_rows(user, getattr(cfg, "id", None), class_name) if bool(getattr(r, "visible", False))]
        cache[rows_key] = [r for r in rows if bool(getattr(r, "rls_enabled", False))]
    rls_rows = cache.get(rows_key) or []
    if not rls_rows:
        return True
    decision_key = ('node_access', user_id, uid, class_name, node_id)
    if decision_key in cache:
        return bool(cache[decision_key])
    node_data = data or {}
    any_allowed = False
    for row in rls_rows:
        try:
            handler_decision = _client_eval_rls_handler(row, uid, class_name, node_id, node_data)
        except Exception as e:
            handler_decision = False
            print("client RLS handler error:", e)
        listed = handler_decision if handler_decision is not None else _client_eval_rls_rules(node_data, getattr(row, "rls_rules_json", None) or [])
        mode = str(getattr(row, "rls_mode", "") or "allow").lower()
        allowed = (not listed) if mode in ("deny", "forbid", "exclude") else bool(listed)
        if mode in ("deny", "forbid", "exclude") and not allowed:
            cache[decision_key] = False
            return False
        if allowed:
            any_allowed = True
    cache[decision_key] = any_allowed
    return any_allowed


def _client_filter_nodes_for_acl(config_uid: str, class_name: str, items: List[Dict[str, Any]], user=None) -> List[Dict[str, Any]]:
    user = user or current_user
    if not _client_user_can_access_class(config_uid, class_name, user=user):
        return []
    cfg = _client_cfg_by_uid(str(config_uid or ''))
    if cfg and _client_user_is_config_scope_owner(user, cfg):
        return list(items or [])
    rows = [r for r in _client_assigned_profile_rows(user, getattr(cfg, 'id', None), class_name) if bool(getattr(r, 'visible', False))] if cfg else []
    if not any(bool(getattr(r, 'rls_enabled', False)) for r in rows):
        return list(items or [])
    out: List[Dict[str, Any]] = []
    for item in items or []:
        data = (item or {}).get("_data") or {}
        node_id = str((item or {}).get("_id") or data.get("_id") or "")
        if _client_user_can_access_node(config_uid, class_name, node_id, data if isinstance(data, dict) else {}, user=user):
            out.append(item)
    return out


def _client_filter_config_for_acl(cfg: Dict[str, Any], config_uid: str = "", user=None) -> Dict[str, Any]:
    """Remove classes unavailable to the current web-client user from cached local config JSON.

    Config-level access is checked when the repo is added.  Class/RLS access must also
    be applied at render time because old client_repo_config rows may still contain
    the full designer configuration.
    """
    if not isinstance(cfg, dict):
        return cfg
    user = user or current_user
    uid = str(config_uid or cfg.get("uid") or "").strip()
    classes = []
    for c in (cfg.get("classes") or []):
        if not isinstance(c, dict):
            continue
        name = str(c.get("name") or "").strip()
        if not name:
            continue
        if _client_user_can_access_class(uid, name, user=user):
            classes.append(c)
    filtered = dict(cfg)
    filtered["classes"] = classes
    return filtered


def _get_repo_by_config_uid_or_404(config_uid: str) -> models.Repo:
    config_uid = (config_uid or "").strip()
    repo = models.Repo.query.filter_by(config_uid=config_uid, user_id=current_user.id).first()
    if not repo:
        abort(404)
    return repo

def _exec_node_class(config: Any, class_name: str):

    ns = {}
    exec(config.nodes_server_handlers, ns)
    cls = ns.get(class_name)
    if not cls:
        raise RuntimeError(f"Node class '{class_name}' not found in handlers")
    return cls


def _is_local_repo(repo: models.Repo) -> bool:
    if not repo.base_url:
        return True
    return repo.base_url.rstrip("/") == request.host_url.rstrip("/")



# -------- NodaLayout context helpers (NodeChildren / node covers) --------

def _pick_node_title(data: Dict[str, Any]) -> str:
    """Best-effort human title for a node cover."""
    for k in ("title", "name", "caption", "label", "number", "code"):
        v = data.get(k)
        if isinstance(v, str) and v.strip():
            return v.strip()
    return ""

def _fetch_node_data_for_repo(repo: models.Repo, class_name: str, node_id: str) -> Dict[str, Any]:
    """Get _data for a node from local storage or remote base_url."""
    cfg_uid = repo.config_uid
    base_url = (repo.base_url or "").strip().rstrip("/")
    current = (request.host_url or "").rstrip("/")

    if not base_url or base_url == current:
        try:
            return _node_local_get_data(cfg_uid, class_name, node_id) or {}
        except Exception:
            return {}
    try:
        payload = _api_get_remote(repo, f"/api/config/{cfg_uid}/node/{class_name}/{node_id}")
        return (payload or {}).get("_data") or {}
    except Exception:
        return {}

def _wrap_client_tpl_html(html: str, data: dict) -> str:
    html = str(html or "").strip()
    if not html:
        return ""
    if 'data-nl-tpl-' not in html:
        return html
    try:
        payload = escape(json.dumps(data or {}, ensure_ascii=False), quote=True)
        return f'<div class="nl-cover-runtime" data-nl-cover-data="{payload}">{html}</div>'
    except Exception:
        return html

def _node_cover_html(repo: models.Repo, class_name: str, node_id: str, mode: str = "") -> str:

    # NOTE: `mode` is kept for backward/forward compatibility.
    # Some callers (e.g. Table rows) may pass mode="table" to request a more
    # compact look in the future. Currently we render the same cover.
    data = _fetch_node_data_for_repo(repo, class_name, node_id)
    assets_base_dir = _userfiles_dir_for_repo(repo)
    parsed = get_parsed_config(repo, models.db) or {}
    cls_cfg_for_tags = ((parsed.get("classes") or {}).get(class_name) or {}) if isinstance(parsed, dict) else {}
    show_tags = bool(cls_cfg_for_tags.get("show_tag_cloud"))
    nl_context = _nl_context(repo, class_name=class_name, node_id=node_id)

    try:
        cov = data.get("_cover") if isinstance(data, dict) else None
        if cov:
            if isinstance(cov, (dict, list)):
                html = str(render_nodalayout_html(cov, data, assets_base_dir=assets_base_dir, context=nl_context) or "").strip()
                if html:
                    return _cover_with_tags(_wrap_client_tpl_html(html, data), data, show_tags)
            elif isinstance(cov, str):
                s = cov.strip()
                # json layout as string
                if (s.startswith("[") or s.startswith("{")):
                    html = str(render_nodalayout_html(s, data, assets_base_dir=assets_base_dir, context=nl_context) or "").strip()
                    if html:
                        return _cover_with_tags(_wrap_client_tpl_html(html, data), data, show_tags)
                # plain image src
                pic_layout = [[{"type": "Picture", "value": s, "width": -1}]]
                html = str(render_nodalayout_html(pic_layout, data, assets_base_dir=assets_base_dir, context=nl_context) or "").strip()
                if html:
                    return _cover_with_tags(_wrap_client_tpl_html(html, data), data, show_tags)
    except Exception:
        pass

    
    try:
        cls = (parsed.get("classes") or {}).get(class_name) or {}

        cover_web_layout = (cls.get("display_image_web") or "").strip()
        cover_layout = cls.get("cover_image")  # может быть dict layout

        layout_to_use = None
        if cover_web_layout:
            layout_to_use = cover_web_layout
        elif cover_layout:
            layout_to_use = cover_layout

        if layout_to_use:
            _fill_nodeinput_views(repo, parsed, layout_to_use, data)
            html = str(render_nodalayout_html(layout_to_use, data, assets_base_dir=assets_base_dir, context=nl_context) or "").strip()
            if html:
                return _cover_with_tags(_wrap_client_tpl_html(html, data), data, show_tags)
    except Exception:
        pass

    
    title = _pick_node_title(data)
    subtitle = f"{class_name}/{node_id}"

    if title:
        return _cover_with_tags((
            f'<div class="card"><div class="card-body p-2">'
            f'<div class="fw-semibold">{escape(title)}</div>'
            f'<div class="text-muted small">{escape(subtitle)}</div>'
            f'</div></div>'
        ), data, show_tags)
    return _cover_with_tags((
        f'<div class="card"><div class="card-body p-2">'
        f'<div class="fw-semibold">{escape(subtitle)}</div>'
        f'</div></div>'
    ), data, show_tags)


def _node_children_tree(repo: models.Repo, class_name: str, node_id: str) -> List[Dict[str, Any]]:
    """Build recursive tree for NodeChildren renderer. Supports both old and new formats."""
    visited: set[tuple[str, str]] = set()
    
    def _parse_uid(s: str):
        s = str(s or "").strip()
        if "$" in s:
            parts = s.split("$")
            if len(parts) >= 3:
                # cfg$Class$Id
                return parts[-2], parts[-1]
            if len(parts) == 2:
                # Class$Id
                return parts[0], parts[1]
        return "", s
    
    def _get_children_from_node_data(data: Dict[str, Any]) -> List[Dict[str, str]]:
        """Extract children from node data in both formats"""
        children_data = data.get("_children") or []
        result = []
        
        # New format (dict)
        if isinstance(children_data, dict):
            for key, value in children_data.items():
                # key: "ClassName$nodeId", value: "config_uid$ClassName$nodeId"
                config_uid, child_class, child_id = None, None, None
                
                # Try to parse from key
                key_parts = key.split("$")
                if len(key_parts) == 2:
                    child_class_name, child_id = key_parts[0], key_parts[1]
                elif len(key_parts) == 3:
                    child_class_name, child_id = key_parts[1], key_parts[2]
                
                # If not successful, try from value
                if not child_class or not child_id:
                    value_parts = value.split("$")
                    if len(value_parts) >= 3:
                        child_class, child_id = value_parts[-2], value_parts[-1]
                
                if child_class and child_id:
                    result.append({
                        "class": child_class,
                        "id": child_id,
                        "uid": value
                    })
        
        # Old format (list)
        elif isinstance(children_data, list):
            for child in children_data:
                if isinstance(child, dict):
                    child_class = child.get("class") or child.get("_class")
                    child_id = child.get("id") or child.get("_id")
                    if child_class and child_id:
                        result.append({
                            "class": child_class,
                            "id": child_id,
                            "uid": child.get("uid")
                        })
        
        return result
    
    def build(cn: str, nid: str) -> List[Dict[str, Any]]:
        key = (cn, nid)
        if key in visited:
            return []
        visited.add(key)
        
        data = _fetch_node_data_for_repo(repo, cn, nid)
        children_list = _get_children_from_node_data(data)
        out: List[Dict[str, Any]] = []
        
        for child in children_list:
            cc = str(child.get("class") or "").strip()
            ci = str(child.get("id") or "").strip()
            
            if not cc or not ci:
                # Try to parse from uid
                uid = child.get("uid")
                if uid:
                    cc2, ci2 = _parse_uid(uid)
                    cc, ci = cc2, ci2
            
            if not cc or not ci:
                continue
            
            out.append({
                "class": cc,
                "id": ci,
                "cover_html": _node_cover_html(repo, cc, ci),
                "open_url": url_for("client.node_form_redirect", repo_id=repo.id, class_name=cc, node_id=ci),
                "children": build(cc, ci),
            })
        
        return out
    
    return build(class_name, node_id)

def _walk_layout_find_link_elements(layout_obj):
    """Yield link/input elements from layout (2d/1d/json str)."""
    import json
    if layout_obj is None:
        return
    if isinstance(layout_obj, str):
        try:
            layout_obj = json.loads(layout_obj)
        except Exception:
            return

    def walk(x):
        if isinstance(x, dict):
            t = x.get("type") or x.get("t")
            if t in ("NodeInput", "NodeLink", "DatasetInput", "DatasetField", "DatasetLink", "DataSetLink"):
                yield x
            # walk common nested places
            for k in ("layout", "tabs", "rows", "cols", "items", "children"):
                v = x.get(k)
                if isinstance(v, list):
                    for it in v:
                        yield from walk(it)
                elif isinstance(v, dict):
                    yield from walk(v)

            # Tabs keep Tab objects in ``value`` rather than ``layout``.
            # Do not walk every element value blindly: inputs/tables may store
            # ordinary data there. Only the container element is recursive.
            if str(t or "") == "Tabs":
                v = x.get("value")
                if isinstance(v, list):
                    for it in v:
                        yield from walk(it)
                elif isinstance(v, dict):
                    yield from walk(v)
        elif isinstance(x, list):
            for it in x:
                yield from walk(it)

    yield from walk(layout_obj)


def _walk_layout_find_nodeinputs(layout_obj):
    """Backward-compatible iterator name."""
    yield from _walk_layout_find_link_elements(layout_obj)

def _prepare_nodalayout(layout: Any) -> Any:
    """Parse a JSON layout once per class instead of once per rendered card."""
    if not isinstance(layout, str):
        return layout
    text = layout.strip()
    if not text or text[:1] not in ("[", "{"):
        return layout
    try:
        return json.loads(text)
    except Exception:
        return layout


def _node_storage_data_direct(
    config_uid: str,
    class_name: str,
    node_id: str,
    cache: Optional[Dict[Any, Any]] = None,
) -> Optional[Dict[str, Any]]:
    """Read one node payload directly from its SQLite store.

    Cover rendering only needs node data to calculate ``record_view``. Loading a
    generated Python class and executing handlers for that is unnecessary and
    was the main cost in sections containing several linked classes.
    """
    cfg_uid = str(config_uid or "").strip()
    cls_name = str(class_name or "").strip()
    raw_id = str(node_id or "").strip()
    if not cfg_uid or not cls_name or not raw_id:
        return None

    try:
        _, _, internal_id = _nodes_mod.parse_uid_any(raw_id)
    except Exception:
        internal_id = raw_id
    internal_id = str(internal_id or raw_id)
    cache_key = (cfg_uid, cls_name, internal_id)
    if isinstance(cache, dict) and cache_key in cache:
        value = cache.get(cache_key)
        return dict(value) if isinstance(value, dict) else None

    db_path = os.path.join("node_storage", f"{cls_name}_{cfg_uid}.sqlite")
    result: Optional[Dict[str, Any]] = None
    if os.path.exists(db_path):
        candidates = [internal_id, raw_id, f"{cfg_uid}${cls_name}${internal_id}"]
        seen = set()
        try:
            conn = sqlite3.connect(db_path)
            try:
                cur = conn.cursor()
                for candidate in candidates:
                    candidate = str(candidate or "")
                    if not candidate or candidate in seen:
                        continue
                    seen.add(candidate)
                    cur.execute("SELECT value FROM unnamed WHERE key = ?", (candidate,))
                    row = cur.fetchone()
                    if not row:
                        continue
                    try:
                        obj = pickle.loads(row[0])
                    except Exception:
                        obj = None
                    if isinstance(obj, dict):
                        data = obj.get("_data")
                        if isinstance(data, dict):
                            result = dict(data)
                        else:
                            result = dict(obj)
                        break
            finally:
                conn.close()
        except Exception:
            result = None

    if isinstance(cache, dict):
        cache[cache_key] = dict(result) if isinstance(result, dict) else None
    return result


def _node_storage_data_batch_direct(
    config_uid: str,
    class_name: str,
    raw_uids: List[str],
) -> Dict[str, Optional[Dict[str, Any]]]:
    """Read many nodes of one class with one SQLite connection.

    The return mapping is keyed by the exact UID supplied by the caller.  Node
    stores historically used either the internal id, the raw UID, or the fully
    qualified UID as the sqlite key, so all three aliases are queried in one
    batch.
    """
    cfg_uid = str(config_uid or "").strip()
    cls_name = str(class_name or "").strip()
    requested = [str(x or "").strip() for x in (raw_uids or []) if str(x or "").strip()]
    result: Dict[str, Optional[Dict[str, Any]]] = {uid: None for uid in requested}
    if not cfg_uid or not cls_name or not requested:
        return result

    db_path = os.path.join("node_storage", f"{cls_name}_{cfg_uid}.sqlite")
    if not os.path.exists(db_path):
        return result

    candidate_to_uids: Dict[str, List[str]] = {}
    for raw_uid in requested:
        try:
            _, _, internal_id = _nodes_mod.parse_uid_any(raw_uid)
        except Exception:
            internal_id = raw_uid
        internal_id = str(internal_id or raw_uid).strip()
        for candidate in (internal_id, raw_uid, f"{cfg_uid}${cls_name}${internal_id}"):
            candidate = str(candidate or "").strip()
            if candidate:
                candidate_to_uids.setdefault(candidate, []).append(raw_uid)

    keys = list(candidate_to_uids)
    try:
        conn = sqlite3.connect(db_path)
        try:
            cur = conn.cursor()
            # Stay below SQLite's common 999 parameter limit.
            for pos in range(0, len(keys), 800):
                chunk = keys[pos:pos + 800]
                marks = ",".join("?" for _ in chunk)
                cur.execute(f"SELECT key, value FROM unnamed WHERE key IN ({marks})", chunk)
                for storage_key, payload in cur.fetchall():
                    try:
                        obj = pickle.loads(payload)
                    except Exception:
                        obj = None
                    if not isinstance(obj, dict):
                        continue
                    data = obj.get("_data") if isinstance(obj.get("_data"), dict) else obj
                    if not isinstance(data, dict):
                        continue
                    for raw_uid in candidate_to_uids.get(str(storage_key), []):
                        if result.get(raw_uid) is None:
                            result[raw_uid] = dict(data)
        finally:
            conn.close()
    except Exception:
        return result
    return result


def _fill_virtual_node_views_batch(
    repo,
    parsed: Dict[str, Any],
    layout: Any,
    node_data: Dict[str, Any],
    shared_cache: Optional[Dict[str, Any]] = None,
) -> None:
    """Fill missing NodeLink ``*_view`` values in all inline rows at once."""
    if not isinstance(node_data, dict):
        return
    layout_obj = _prepare_nodalayout(layout)
    cache = shared_cache if isinstance(shared_cache, dict) else {}
    parsed_by_config = cache.setdefault("parsed_by_config", {})
    parsed_by_config.setdefault(str(repo.config_uid or ""), parsed or {})

    tables: List[Dict[str, Any]] = []

    def walk(value: Any) -> None:
        if isinstance(value, list):
            for item in value:
                walk(item)
            return
        if not isinstance(value, dict):
            return
        if isinstance(value.get("virtual_node"), dict) and str(value.get("id") or "").strip():
            tables.append(value)
        for child in value.values():
            if isinstance(child, (list, dict)):
                walk(child)

    walk(layout_obj)
    if not tables:
        return

    pending: Dict[Tuple[str, str], List[Tuple[str, str, Dict[str, Any], str]]] = {}

    for table in tables:
        table_id = str(table.get("id") or "").strip()
        rows = node_data.get(table_id)
        if not table_id or not isinstance(rows, list):
            continue
        virtual = table.get("virtual_node") or {}
        paths = set()
        for link_el in _walk_layout_find_link_elements(virtual.get("layout")):
            t = str(link_el.get("type") or link_el.get("t") or "")
            if t not in ("NodeInput", "NodeLink"):
                continue
            raw_value = link_el.get("value")
            path = raw_value[1:] if isinstance(raw_value, str) and raw_value.startswith("@") else str(link_el.get("id") or "")
            path = str(path or "").strip()
            if path and not path.endswith("_view"):
                paths.add(path)
        for link_el in _walk_layout_find_link_elements(virtual.get("cover")):
            t = str(link_el.get("type") or link_el.get("t") or "")
            if t not in ("NodeInput", "NodeLink"):
                continue
            raw_value = link_el.get("value")
            path = raw_value[1:] if isinstance(raw_value, str) and raw_value.startswith("@") else str(link_el.get("id") or "")
            path = str(path or "").strip()
            if path and not path.endswith("_view"):
                paths.add(path)

        # Header-only links are uncommon but supported.  They are added only
        # when the actual row value looks like a Node UID below.
        header_items = table.get("table_header") or []
        if isinstance(header_items, str):
            try:
                parsed_headers = json.loads(header_items)
                header_items = parsed_headers if isinstance(parsed_headers, list) else []
            except Exception:
                header_items = []
        for col in header_items:
            if isinstance(col, dict):
                path = str(col.get("key") or col.get("id") or "").strip()
            else:
                parts = str(col or "").split("|")
                path = str(parts[1] if len(parts) > 1 else "").strip()
            if path and not path.endswith("_view"):
                paths.add(path)

        for row in rows:
            if not isinstance(row, dict):
                continue
            for path in paths:
                # Inline rows use simple field names in all generated configs.
                # Nested paths still work for the common dotted form.
                parts = path.split(".")
                parent = row
                for part in parts[:-1]:
                    parent = parent.get(part) if isinstance(parent, dict) else None
                    if not isinstance(parent, dict):
                        break
                if not isinstance(parent, dict):
                    continue
                leaf = parts[-1]
                raw_ref = parent.get(leaf)
                view_key = f"{leaf}_view"
                if parent.get(view_key) not in (None, ""):
                    continue
                if isinstance(raw_ref, dict):
                    inline_view = raw_ref.get("_view") or raw_ref.get("view") or raw_ref.get("name")
                    raw_ref = raw_ref.get("uid") or raw_ref.get("_uid") or raw_ref.get("_id") or raw_ref.get("id")
                    if inline_view not in (None, ""):
                        parent[view_key] = str(inline_view)
                        continue
                raw_uid = str(raw_ref or "").strip()
                if "$" not in raw_uid:
                    continue
                try:
                    uid_cfg, cls_name, internal_id = _nodes_mod.parse_uid_any(raw_uid)
                except Exception:
                    continue
                eff_cfg = str(uid_cfg or repo.config_uid or "").strip()
                cls_name = str(cls_name or "").strip()
                internal_id = str(internal_id or "").strip()
                if not eff_cfg or not cls_name or not internal_id:
                    continue
                pending.setdefault((eff_cfg, cls_name), []).append((raw_uid, internal_id, parent, view_key))

    def parsed_for(config_uid: str) -> Dict[str, Any]:
        if config_uid in parsed_by_config:
            return parsed_by_config.get(config_uid) or {}
        target_parsed: Dict[str, Any] = {}
        try:
            target_repo = repo if config_uid == str(repo.config_uid or "") else models.Repo.query.filter_by(
                user_id=repo.user_id, config_uid=config_uid
            ).first()
            if target_repo is not None:
                target_parsed = get_parsed_config(target_repo, models.db) or {}
            elif _client_user_can_access_config(config_uid, current_user):
                cfg = fetch_config_from_local_db(config_uid, user=current_user)
                target_parsed = build_parsed_config(cfg or {}) if cfg else {}
        except Exception:
            target_parsed = {}
        parsed_by_config[config_uid] = target_parsed
        return target_parsed

    for (config_uid, cls_name), entries in pending.items():
        # Exact platform exception for virtual-table NodeLink/NodeInput values.
        if cls_name == "_User":
            for uid, _, parent, view_key in entries:
                view = _client_system_user_view(uid)
                if view:
                    parent[view_key] = view
            continue

        unique_uids = list(dict.fromkeys(uid for uid, _, _, _ in entries))
        data_by_uid = _node_storage_data_batch_direct(config_uid, cls_name, unique_uids)
        parsed_cfg = parsed_for(config_uid)
        view_by_uid: Dict[str, str] = {}
        for uid, internal_id, _, _ in entries:
            if uid in view_by_uid:
                continue
            data = data_by_uid.get(uid)
            if not isinstance(data, dict):
                data = _node_storage_data_direct(config_uid, cls_name, internal_id, cache.setdefault("node_data", {}))
            if not isinstance(data, dict):
                continue
            view_by_uid[uid] = str(_render_class_record_view(parsed_cfg, cls_name, internal_id, data) or internal_id)
        for uid, _, parent, view_key in entries:
            view = view_by_uid.get(uid)
            if view:
                parent[view_key] = view


def _link_view_needs_resolution(existing_view: Any, raw_ref: str, dataset_class: str = "") -> bool:
    """Return True when a saved ``<field>_view`` is still only a node UID.

    Older clients sometimes persisted ``Class$id`` or ``config$Class$id`` in
    the companion view field.  It is not a human caption and must not suppress
    normal ``record_view`` resolution.
    """
    if existing_view in (None, ""):
        return True
    text = str(existing_view or "").strip()
    raw = str(raw_ref or "").strip()
    if not text or text == raw:
        return True
    if "$" not in text:
        return False
    try:
        _cfg, cls_name, internal_id = _nodes_mod.parse_uid_any(text)
    except Exception:
        return False
    cls_name = str(cls_name or "").strip()
    internal_id = str(internal_id or "").strip()
    if not cls_name or not internal_id:
        return False
    if re.fullmatch(r"[\w.:-]+", cls_name, re.UNICODE) is None or re.search(r"\s", internal_id):
        return False
    expected = str(dataset_class or "").strip()
    if expected and cls_name != expected:
        return False
    return True


def _fill_nodeinput_views(repo, parsed, layout, node_data, shared_cache: Optional[Dict[str, Any]] = None):
    """Pre-fill human-readable values for node/dataset links.

    ``shared_cache`` is request-scoped.  It makes repeated links such as the
    same warehouse or zone resolve once for the whole section, not once per
    card.
    """
    cache = shared_cache if isinstance(shared_cache, dict) else {}
    _fill_virtual_node_views_batch(repo, parsed or {}, layout, node_data, shared_cache=cache)
    node_cache = cache.setdefault("node_views", {})
    dataset_cache = cache.setdefault("dataset_views", {})
    node_data_cache = cache.setdefault("node_data", {})
    parsed_by_config = cache.setdefault("parsed_by_config", {})
    parsed_by_config.setdefault(str(repo.config_uid or ""), parsed or {})
    layout_links_cache = cache.setdefault("layout_links", {})

    def parsed_for_config(config_uid: str) -> Dict[str, Any]:
        cfg_uid = str(config_uid or repo.config_uid or "")
        if cfg_uid in parsed_by_config:
            return parsed_by_config.get(cfg_uid) or {}
        target_parsed: Dict[str, Any] = {}
        try:
            target_repo = repo if cfg_uid == str(repo.config_uid or "") else models.Repo.query.filter_by(
                user_id=repo.user_id, config_uid=cfg_uid
            ).first()
            if target_repo is not None:
                target_parsed = get_parsed_config(target_repo, models.db) or {}
        except Exception:
            target_parsed = {}
        parsed_by_config[cfg_uid] = target_parsed
        return target_parsed

    def raw_ref_for_el(el: dict) -> str:
        raw_val = el.get("value")
        if isinstance(raw_val, str) and raw_val.startswith("@"):
            return str(node_data.get(raw_val[1:], "") or "").strip()
        if isinstance(raw_val, str):
            return raw_val.strip()
        return ""

    layout_key = ("obj", id(layout)) if not isinstance(layout, str) else ("str", layout)
    elements = layout_links_cache.get(layout_key)
    if elements is None:
        elements = list(_walk_layout_find_link_elements(layout))
        layout_links_cache[layout_key] = elements

    for el in elements:
        t = str(el.get("type") or el.get("t") or "")
        lid = str(el.get("id") or "").strip()
        raw_ref = raw_ref_for_el(el)
        if not raw_ref:
            continue

        raw_val = el.get("value")
        field_name = raw_val[1:] if isinstance(raw_val, str) and raw_val.startswith("@") else lid
        if not field_name:
            continue
        view_key = f"{field_name}_view"
        existing_view = node_data.get(view_key)
        dataset_class = str(el.get("dataset") or "").strip() if isinstance(el.get("dataset"), str) else ""
        # Older saved nodes may contain a different UID shape in <field>_view
        # (for example Class$id while the value is config$Class$id).  Resolve
        # every UID-looking view instead of showing it as a caption.
        if not _link_view_needs_resolution(existing_view, raw_ref, dataset_class):
            continue

        if t in ("DatasetInput", "DatasetField", "DatasetLink", "DataSetLink"):
            ds_name, item_id = _split_dataset_item_uid(str(el.get("dataset") or "").strip(), raw_ref)
            if not ds_name or not item_id:
                continue
            ck = (str(repo.config_uid or ""), ds_name, item_id)
            if ck not in dataset_cache:
                try:
                    item = _get_dataset_item_direct(repo.config_uid, ds_name, item_id)
                    dataset_cache[ck] = str(item.get("_view") or item_id) if item else None
                except Exception:
                    dataset_cache[ck] = None
            if dataset_cache.get(ck) is not None:
                node_data[view_key] = dataset_cache[ck]
            continue

        if t in ("NodeInput", "NodeLink"):
            try:
                cfg_uid, cls_name, internal_id = _nodes_mod.parse_uid_any(raw_ref)
                # NodeInput knows the target class from dataset, so local/plain
                # ids are resolvable too.  This also repairs old data written
                # before full UIDs became mandatory.
                cls_name = str(cls_name or dataset_class or "").strip()
                if not cls_name or not internal_id:
                    continue
                eff_cfg = str(cfg_uid or repo.config_uid or "")
                ck = (eff_cfg, str(cls_name), str(internal_id))

                # Reserved users live only in the owner's hidden _System config.
                # Keep the ordinary NodeInput/NodeLink semantics and replace
                # only this exact class with the platform resolver.
                if cls_name == "_User":
                    if ck not in node_cache:
                        node_cache[ck] = _client_system_user_view(raw_ref) or None
                    if node_cache.get(ck) is not None:
                        node_data[view_key] = node_cache[ck]
                    continue

                if ck not in node_cache:
                    data = _node_storage_data_direct(eff_cfg, cls_name, internal_id, node_data_cache)
                    if data is None:
                        # Preserve compatibility with nonstandard storage backends.
                        try:
                            node_cls = _load_server_node_class(eff_cfg, cls_name)
                            node = node_cls.get(internal_id, eff_cfg)
                            data = node.get_data() if node else None
                        except Exception:
                            data = None
                    node_cache[ck] = (
                        _render_class_record_view(parsed_for_config(eff_cfg), cls_name, internal_id, data or {})
                        if isinstance(data, dict) else None
                    )
                if node_cache.get(ck) is not None:
                    node_data[view_key] = node_cache[ck]
            except Exception:
                pass


def _nl_context(repo: models.Repo, *, class_name: str, node_id: str, shared_cache: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    cache = shared_cache if isinstance(shared_cache, dict) else {}
    parsed_by_config: Dict[str, Dict[str, Any]] = cache.setdefault("parsed_by_config", {})
    node_view_cache: Dict[Any, Any] = cache.setdefault("node_views", {})
    dataset_view_cache: Dict[Any, Any] = cache.setdefault("dataset_views", {})
    node_data_cache: Dict[Any, Any] = cache.setdefault("node_data", {})

    def _parsed_for_config(config_uid: str) -> Dict[str, Any]:
        """Return the parsed class metadata needed to render record_view.

        NodeLink values may point to the current configuration or to another
        configuration available to the same web-client user. Keep this lookup
        local to the render context and cache it for all links in the layout.
        """
        cfg_uid = str(config_uid or repo.config_uid or "").strip()
        if cfg_uid in parsed_by_config:
            return parsed_by_config[cfg_uid]

        parsed_cfg: Dict[str, Any] = {}
        try:
            target_repo = repo
            if cfg_uid and cfg_uid != str(repo.config_uid or ""):
                target_repo = models.Repo.query.filter_by(
                    user_id=repo.user_id,
                    config_uid=cfg_uid,
                ).first()
            if target_repo is not None:
                parsed_cfg = get_parsed_config(target_repo, models.db) or {}
        except Exception:
            parsed_cfg = {}

        # A local cross-config link can exist even when that configuration was
        # not explicitly added as a client repository. Use the live Designer
        # config as a safe fallback so record_view still works.
        if not parsed_cfg and cfg_uid:
            try:
                # Never resolve a cross-tenant record_view merely because a
                # foreign configuration UID is present in stored data.
                if _client_user_can_access_config(cfg_uid, current_user):
                    cfg = fetch_config_from_local_db(cfg_uid, user=current_user)
                    parsed_cfg = build_parsed_config(cfg or {})
            except Exception:
                parsed_cfg = {}

        parsed_by_config[cfg_uid] = parsed_cfg
        return parsed_cfg

    def get_dataset_item_view(ds_name: str, item_uid: str) -> str:
        """Resolve DatasetLink value to display text.

        The normal value format is self-describing: ``DatasetName$ItemId``.
        ``ds_name`` is only a fallback for legacy layouts that store just ItemId.
        """
        raw = str(item_uid or "").strip()
        try:
            parsed_ds, item_id = _split_dataset_item_uid(ds_name, raw)
            if not parsed_ds or not item_id:
                return raw
            cache_key = (str(repo.config_uid or ""), parsed_ds, item_id)
            if cache_key not in dataset_view_cache:
                item_data = _get_dataset_item_direct(repo.config_uid, parsed_ds, item_id)
                dataset_view_cache[cache_key] = str(item_data.get("_view") or item_id) if item_data else item_id
            return str(dataset_view_cache.get(cache_key) or item_id)
        except Exception:
            return raw

    def get_node_view(node_uid: str) -> str:
        """Resolve NodeLink value to display text.

        Node links are also self-describing: ``config_uid$ClassName$Id``.
        For old ``ClassName$Id`` links we use the current repo config_uid.
        """
        uid = str(node_uid or "").strip()
        if not uid:
            return ""
        try:
            uid_cfg, cls_name, internal_id = _nodes_mod.parse_uid_any(uid)
            if not internal_id:
                return uid
            eff_cfg = uid_cfg or str(repo.config_uid)
            if not cls_name:
                return internal_id

            cache_key = (str(eff_cfg), str(cls_name), str(internal_id))
            if cache_key in node_view_cache:
                return str(node_view_cache.get(cache_key) or internal_id)

            # Exact platform exception: a NodeLink/NodeInput to _System/_User
            # is resolved by the system-user registry, not by a business repo.
            if cls_name == "_User":
                user_view = _client_system_user_view(uid)
                if user_view:
                    node_view_cache[cache_key] = user_view
                    return user_view

            d = _node_storage_data_direct(eff_cfg, cls_name, internal_id, node_data_cache)
            if d is None:
                try:
                    node_cls = _load_server_node_class(eff_cfg, cls_name)
                    n = node_cls.get(internal_id, eff_cfg)
                    d = n.get_data() if n else None
                except Exception:
                    d = None
            if not isinstance(d, dict):
                node_view_cache[cache_key] = internal_id
                return internal_id

            parsed_cfg = _parsed_for_config(eff_cfg)
            view = _render_class_record_view(parsed_cfg, cls_name, internal_id, d)
            node_view_cache[cache_key] = view
            return view
        except Exception:
            return uid

    def uid_resolve(uid: str):
        """Resolve global node uid to (class_name, internal_id) for link/table helpers."""
        raw = str(uid or "").strip()
        try:
            uid_cfg, cls_name, internal_id = _nodes_mod.parse_uid_any(raw)
            if not internal_id:
                return ("", "")
            return (str(cls_name or ""), str(internal_id or ""))
        except Exception:
            return ("", "")


    return {
        "target": {
            "repo_id": int(repo.id),
            "config_uid": str(repo.config_uid),
            "class_name": str(class_name),
            "node_id": str(node_id),
        },
        "is_admin": _client_is_admin(),
        "show_node_json": _client_show_node_json(),
        "codeframe_read_url": url_for("client.api_codeframe_read"),
        "codeframe_save_url": url_for("client.api_codeframe_save"),
        "node_url": lambda c, i: url_for("client.node_form_redirect", repo_id=repo.id, class_name=c, node_id=i),
        # For Table(nodes_source=True, table=True): allow renderer to fetch per-node data
        # to fill individual cells according to table_header.
        "node_data": lambda c, i: _fetch_node_data_for_repo(repo, c, i),
        "node_cover": lambda c, i: _node_cover_html(repo, c, i),
        "node_cover_table": lambda cls, nid: _node_cover_html(repo, cls, nid, mode="table"),
        "node_children_tree": lambda c, i: _node_children_tree(repo, c, i),
        "get_dataset_item_view": get_dataset_item_view,
        "get_node_view": get_node_view,
        "uid_resolve": uid_resolve,
    }


def parse_config_url(config_url: str) -> Tuple[str, str, str]:
    """Returns (base_url, config_uid, normalized_config_url)."""
    u = (config_url or "").strip()
    if not u:
        raise ValueError("empty url")
    parsed = urlparse(u)
    if not parsed.scheme or not parsed.netloc:
        raise ValueError("url must include scheme and host, e.g. https://...")

    path = parsed.path.rstrip("/")
    parts = [p for p in path.split("/") if p]

    uid = None
    for i in range(len(parts) - 2):
        if parts[i] == "api" and parts[i + 1] == "config":
            uid = parts[i + 2]
            break
    if not uid:
        raise ValueError("url path must contain /api/config/<uid>")

    base_url = f"{parsed.scheme}://{parsed.netloc}"
    normalized = base_url + "/api/config/" + uid
    return base_url, uid, normalized


def normalize_sections(cfg: Dict[str, Any]) -> List[Dict[str, Any]]:
    secs = cfg.get("sections") or []
    out = []
    for s in secs:
        if bool(s.get('hide_web_client') or s.get('hideWebClient')):
            continue
        out.append({
            "code": s.get("code") or "",
            "name": s.get("name") or (s.get("code") or "<no code>"),
            "commands": (s.get("commands") or "").strip(),
            "hide_mobile_client": bool(s.get("hide_mobile_client") or s.get("hideMobileClient")),
            "hide_web_client": bool(s.get("hide_web_client") or s.get("hideWebClient")),
        })
    return out



def class_section_code(cls: Dict[str, Any]) -> str:
    return str(cls.get("section_code") or cls.get("section") or "")


def build_parsed_config(cfg: Dict[str, Any]) -> Dict[str, Any]:
    sections = normalize_sections(cfg)
    classes = cfg.get("classes") or []
    classes_by_name: Dict[str, Dict[str, Any]] = {}
    classes_by_section: Dict[str, List[Dict[str, Any]]] = {}

    for c in classes:
        if bool(c.get('hide_web_client') or c.get('hideWebClient')):
            continue
        name = c.get("name")
        if not name:
            continue
        classes_by_name[name] = c
        sc = class_section_code(c)
        classes_by_section.setdefault(sc, []).append(c)

    # Rooms mapping: alias -> room_uid
    rooms_map: Dict[str, str] = {}
    try:
        for it in (cfg.get("rooms") or []):
            if not isinstance(it, dict):
                continue
            a = str(it.get("alias") or "").strip()
            rid = str(it.get("room_id") or "").strip()
            if a:
                rooms_map[a] = rid
    except Exception:
        rooms_map = {}

    return {
        "cfg": cfg,
        "sections": sections,
        "classes": classes_by_name,
        "classes_by_section": classes_by_section,
        "rooms": rooms_map,
    }


def get_parsed_config(repo: models.Repo, db, user=None) -> Optional[Dict[str, Any]]:
    """Return a parsed configuration snapshot with a stable hot cache.

    Local configurations still come directly from the Designer SQL models, so
    class/profile changes are visible without manually refreshing repositories.
    The expensive ORM-to-dict conversion is performed only once per
    ``(repo, config revision, ACL signature)`` rather than once per request.
    """
    is_local = _client_repo_is_local(repo)
    actor = user
    if actor is None:
        try:
            actor = current_user
            # A scheduler/app-context LocalProxy resolves to AnonymousUser. Use
            # the repository owner instead so background work shares the right
            # ACL snapshot and cannot evict the browser cache.
            if not bool(getattr(actor, "is_authenticated", False)):
                actor = _client_repo_actor(repo)
        except Exception:
            actor = _client_repo_actor(repo)

    if is_local:
        config_uid = str(getattr(repo, "config_uid", "") or "")
        acl_signature = _client_acl_signature(actor) if actor is not None else "anonymous"
        cfg_obj = _client_cfg_by_uid(config_uid)
        cfg_modified = getattr(cfg_obj, "last_modified", None) if cfg_obj is not None else None
        cfg_stamp = cfg_modified.isoformat() if hasattr(cfg_modified, "isoformat") else str(cfg_modified or "")
        live_stamp = f"live:{config_uid}:{cfg_stamp}:{acl_signature}"
        cache_key = (int(repo.id), cfg_stamp, acl_signature, "local")

        with CONFIG_MEM_LOCK:
            mem = CONFIG_MEM.get(cache_key)
        if mem and mem.get("stamp") == live_stamp:
            return mem

        try:
            cfg = fetch_config_from_local_db(config_uid, user=actor)
        except Exception as e:
            print("local config live fetch failed, falling back to client_repo_config:", e)
            row = db.session.query(models.RepoConfig).filter_by(repo_id=repo.id).first()
            if not row:
                return None
            try:
                cfg = json.loads(row.config_json)
            except Exception:
                return None

        cfg = _normalize_print_html_templates_in_config(cfg)
        cfg = _client_filter_config_for_acl(
            cfg,
            getattr(repo, "config_uid", "") or cfg.get("uid") or "",
            user=actor,
        )
        parsed = build_parsed_config(cfg)
        parsed["stamp"] = live_stamp

        with CONFIG_MEM_LOCK:
            # Remove only snapshots from an older SQL configuration revision.
            # Keep other ACL signatures for the same revision: owner, child
            # users and scheduler are allowed to coexist in the cache.
            for old_key in list(CONFIG_MEM.keys()):
                if not (isinstance(old_key, tuple) and len(old_key) >= 4):
                    continue
                if old_key[0] == int(repo.id) and old_key[3] == "local" and old_key[1] != cfg_stamp:
                    CONFIG_MEM.pop(old_key, None)
            CONFIG_MEM[cache_key] = parsed
        return parsed

    row = db.session.query(models.RepoConfig).filter_by(repo_id=repo.id).first()
    if not row:
        return None

    stamp = row.updated_at.isoformat() if row.updated_at else ""
    cache_key = (int(repo.id), stamp, "remote")
    with CONFIG_MEM_LOCK:
        mem = CONFIG_MEM.get(cache_key)
    if mem and mem.get("stamp") == stamp:
        return mem

    try:
        cfg = json.loads(row.config_json)
    except Exception:
        return None

    cfg = _normalize_print_html_templates_in_config(cfg)
    parsed = build_parsed_config(cfg)
    parsed["stamp"] = stamp
    with CONFIG_MEM_LOCK:
        CONFIG_MEM[cache_key] = parsed
    return parsed

def _layout_is_empty(layout_spec: Any) -> bool:
    """Return True when a layout value cannot render a form.

    Empty JSON layouts (``[]``/``{}``) are common Android placeholders. They
    must not override a valid ``init_screen_layout_web`` in the browser.
    """
    if layout_spec is None:
        return True
    if isinstance(layout_spec, (list, tuple, dict)):
        return len(layout_spec) == 0
    if isinstance(layout_spec, str):
        s = layout_spec.strip()
        if not s or s.lower() in {"null", "none"}:
            return True
        if s in {"[]", "{}"}:
            return True
        try:
            parsed_layout = json.loads(s)
            if isinstance(parsed_layout, (list, dict)) and len(parsed_layout) == 0:
                return True
        except Exception:
            pass
    return False


def _first_usable_layout(*layout_specs: Any) -> Any:
    for layout_spec in layout_specs:
        if not _layout_is_empty(layout_spec):
            return layout_spec
    return None


def resolve_common_layout(parsed: Optional[Dict[str, Any]], layout_spec: Any) -> Any:
    """Resolve '^layout_id' using cfg['CommonLayouts'].

    If not found or spec is not a '^' string, returns layout_spec unchanged.
    If spec is '^...' but not found, returns None (meaning: ignore).
    """
    if not isinstance(layout_spec, str):
        return layout_spec
    s = layout_spec.strip()
    if not s.startswith("^"):
        return layout_spec

    name = s[1:].strip()
    if not name:
        return None

    try:
        cfg = (parsed or {}).get("cfg") if isinstance(parsed, dict) else None
        items = (cfg or {}).get("CommonLayouts") if isinstance(cfg, dict) else None
        if not isinstance(items, list):
            return None
        for it in items:
            if isinstance(it, dict) and str(it.get("id") or "").strip() == name:
                return it.get("layout")
    except Exception:
        return None
    return None

def fetch_config_from_local_db(config_uid: str, user=None) -> Dict[str, Any]:
    
    
    Configuration = main.Configuration

    stmt = select(Configuration).where(Configuration.uid == config_uid)
    # The old live export lazily loaded methods, events and every actions list
    # class by class (classic N+1). A medium configuration could therefore issue
    # hundreds of SQL statements on a cold cache. Select-in loading keeps the
    # same SQL source of truth while reducing it to a small fixed query set.
    try:
        ConfigClass = main.ConfigClass
        ClassEvent = main.ClassEvent
        ConfigEvent = main.ConfigEvent
        ConfigTimer = main.ConfigTimer
        stmt = stmt.options(
            selectinload(Configuration.classes).selectinload(ConfigClass.methods),
            selectinload(Configuration.classes).selectinload(ConfigClass.event_objs).selectinload(ClassEvent.actions),
            selectinload(Configuration.sections),
            selectinload(Configuration.room_aliases),
            selectinload(Configuration.config_events).selectinload(ConfigEvent.actions),
            selectinload(Configuration.config_timers).selectinload(ConfigTimer.actions),
        )
    except Exception:
        # Compatibility with older/minimal model bundles: lazy loading still
        # works, only without the cold-start optimization.
        pass

    cfg_obj = models.db.session.execute(stmt).scalar_one_or_none()

    if not cfg_obj:
        raise ValueError(f"Configuration {config_uid} not found in DB")

    
    try:
        host_url = (request.host_url or "").rstrip("/")
    except Exception:
        host_url = ""
    url = (host_url + f"/api/config/{cfg_obj.uid}") if host_url else f"/api/config/{cfg_obj.uid}"

    actor = user
    if actor is None:
        try:
            actor = current_user
        except Exception:
            actor = None

    classes = []
    for c in cfg_obj.classes:
        # Local DB export for the web client must already be class-filtered.
        # Use the client ACL helper, not the root API helper, so the same
        # whitelist/RLS semantics are used by page render and section refresh.
        try:
            if bool(getattr(c, 'hide_web_client', False)):
                continue
            if not _client_user_can_access_class(cfg_obj.uid, c.name, user=actor):
                continue
        except Exception:
            continue
        classes.append({
            "name": c.name,
            "section": c.section,
            "section_code": c.section_code,
            "has_storage": c.has_storage,
            "display_name": c.display_name,
            "record_view": getattr(c, "record_view", "") or "",
            "cover_image": c.cover_image,
            
            "display_image_web": getattr(c, "display_image_web", "") or "",
            "display_image_table": getattr(c, "display_image_table", "") or "",
            "data_structure": getattr(c, "data_structure", "") or "",
            "ngenie_role": getattr(c, "ngenie_role", "") or "",
            "ngenie_prompt": getattr(c, "ngenie_prompt", "") or "",
            "show_tag_cloud": bool(getattr(c, "show_tag_cloud", False)),
            # Dashboard settings must be present in the live local-config export.
            # Local repos are read directly from the Designer DB on every request,
            # so omitting these fields makes the Dashboard tab invisible even
            # though the checkbox is saved on ConfigClass.
            "hide_mobile_client": bool(getattr(c, "hide_mobile_client", False)),
            "hide_web_client": bool(getattr(c, "hide_web_client", False)),
            "dashboard_enabled": bool(getattr(c, "dashboard_enabled", False)),
            "dashboard_width": str(getattr(c, "dashboard_width", "") or "100"),
            "dashboard_top": bool(getattr(c, "dashboard_top", False)),
            "mobile_print_enabled": bool(getattr(c, "mobile_print_enabled", False)),
            "commands": getattr(c, "commands", "") or "",
            "use_standard_commands": bool(getattr(c, "use_standard_commands", True)),
            "svg_commands": getattr(c, "svg_commands", "") or "",
            
            "migration_register_command": bool(getattr(c, "migration_register_command", False)),
            "migration_register_on_save": bool(getattr(c, "migration_register_on_save", False)),
            "migration_default_room_uid": getattr(c, "migration_default_room_uid", "") or "",
            "migration_default_room_alias": getattr(c, "migration_default_room_alias", "") or "",
            "indexes": getattr(c, "indexes_json", None) or [],
            "class_type": c.class_type,
            "projection_type": getattr(c, "projection_type", "") or "",
            "projection_kanban_columns": getattr(c, "projection_kanban_columns", "") or "",
            "print_template_type": getattr(c, "print_template_type", "") or "html_jinja",
            "print_target_classes": getattr(c, "print_target_classes", None) or [],
            "print_html_template": _decode_print_html_template(getattr(c, "print_html_template", "") or ""),
            "hidden": getattr(c, "hidden", False),
            "init_screen_layout": getattr(c, "init_screen_layout", "") or "",
            "init_screen_layout_web": getattr(c, "init_screen_layout_web", "") or "",
            "methods": [{
                "name": m.name,
                "source": m.source,
                "engine": m.engine,
                "code": m.code,
            } for m in (c.methods or [])],
            "events": [
                {
                    "event": e.event,
                    "listener": e.listener,
                    "actions": [
                        {
                            "action": a.action,
                            "source": a.source,
                            "server": a.server,
                            "method": a.method,
                            "methodText": getattr(a, "method_text", "") or "",
                            "postExecuteMethod": a.post_execute_method,
                            "postExecuteMethodText": getattr(a, "post_execute_method_text", "") or "",
                        }
                        for a in (e.actions or [])
                    ],
                }
                for e in (getattr(c, "event_objs", None) or [])
            ],
        })

    sections = []
    for s in (cfg_obj.sections or []):
        sections.append({
            "name": s.name,
            "code": s.code,
            "commands": s.commands,
            "hide_mobile_client": bool(getattr(s, "hide_mobile_client", False)),
            "hide_web_client": bool(getattr(s, "hide_web_client", False)),
        })

    common_events = []
    for e in (getattr(cfg_obj, "config_events", None) or []):
        common_events.append({
            "event": getattr(e, "event", "") or "",
            "listener": getattr(e, "listener", "") or "",
            "actions": e.actions_as_dicts() if hasattr(e, "actions_as_dicts") else [],
        })   

    return {
        "name": cfg_obj.name,
        "server_name": getattr(cfg_obj, "server_name", "") or "",
        "uid": cfg_obj.uid,
        "url": url,
        "content_uid": getattr(cfg_obj, "content_uid", "") or "",
        "nodes_handlers": getattr(cfg_obj, "nodes_handlers", None),
        "nodes_server_handlers": getattr(cfg_obj, "nodes_server_handlers", None),
        "ngenie_prompt": getattr(cfg_obj, "ngenie_prompt", "") or "",
        "version": getattr(cfg_obj, "version", "00.00.01") or "00.00.01",
        "last_modified": cfg_obj.last_modified.isoformat() if getattr(cfg_obj, "last_modified", None) else "",
        "provider": cfg_obj.vendor or "",
        "vendor": cfg_obj.vendor or "", 
        "display_name": cfg_obj.name or "",
        "classes": classes,
        "sections": sections,
        "rooms": [
            {"alias": a.alias, "room_id": a.room_uid}
            for a in (getattr(cfg_obj, "room_aliases", None) or [])
        ],
        "CommonEvents": common_events,
        "Timers": [t.to_dict() if hasattr(t, "to_dict") else {
            "id": getattr(t, "timer_id", "") or "",
            "timer_id": getattr(t, "timer_id", "") or "",
            "period_seconds": max(900 if (str(getattr(t, "runtime", "") or "server").strip().lower() == "server" or bool(getattr(t, "worker", False))) else 1, getattr(t, "period_seconds", 0) or 0),
            "active": bool(getattr(t, "active", False)),
            "worker": bool(getattr(t, "worker", False)),
            "runtime": str(getattr(t, "runtime", "") or "server").strip().lower(),
            "actions": t.actions_as_dicts() if hasattr(t, "actions_as_dicts") else [],
        } for t in (getattr(cfg_obj, "config_timers", None) or [])],
        "CommonLayouts": getattr(cfg_obj, "common_layouts", None) or getattr(cfg_obj, "CommonLayouts", None) or [],
    }



def _handlers_file_path(config_uid: str) -> str:
    
    base_dir = os.path.dirname(os.path.abspath(__file__))  # client_app/
   
    root = os.path.abspath(os.path.join(base_dir, ".."))
    return os.path.join(root, "Handlers", config_uid, "handlers.py")

def _decode_base64_text_maybe(value: Any) -> str:
    """Decode a base64 text blob; tolerate already-plain text."""
    raw = str(value or "")
    if not raw.strip():
        return ""
    try:
        return base64.b64decode(raw).decode("utf-8", errors="replace")
    except Exception:
        return raw


def _load_server_handlers_ns(config_uid: str, parsed_config: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    """Load server handlers for config_uid.

    Source priority:
      1) Handlers/<uid>/handlers.py for instant local edits;
      2) Configuration.nodes_server_handlers from the Designer DB;
      3) nodes_server_handlers/nodes_handlers from the cached repository config.

    The __file__ value is always the canonical Handlers/<uid>/handlers.py path,
    even when the code was read from DB/cache. This keeps existing helpers such
    as Node.get_all() able to resolve the current config UID from the handlers
    call stack.
    """
    config_uid = str(config_uid or "").strip()
    fp = _handlers_file_path(config_uid)

    code = None
    if os.path.isfile(fp):
        with open(fp, "r", encoding="utf-8") as f:
            code = f.read()

    if code is None:
        try:
            Configuration = getattr(main, "Configuration", None)
            if Configuration is not None:
                cfg = models.db.session.execute(
                    select(Configuration).where(Configuration.uid == config_uid)
                ).scalar_one_or_none()
                if cfg is not None and getattr(cfg, "nodes_server_handlers", None):
                    code = _decode_base64_text_maybe(cfg.nodes_server_handlers)
        except Exception:
            code = None

    if code is None and isinstance(parsed_config, dict):
        cfg_json = parsed_config.get("cfg") if isinstance(parsed_config.get("cfg"), dict) else parsed_config
        if isinstance(cfg_json, dict):
            # Server handlers are preferred. nodes_handlers is accepted as a
            # backward-compatible fallback for older cached configs.
            code = _decode_base64_text_maybe(
                cfg_json.get("nodes_server_handlers")
                or cfg_json.get("server_handlers")
                or cfg_json.get("nodes_handlers")
                or ""
            ) or None

    if code is None:
        raise ValueError(f"Handlers not found for config: {config_uid}")

    g: Dict[str, Any] = {
        "__name__": f"handlers_{config_uid}",
        "__file__": fp,
    }
    try:
        for _helper_name in ("CallSwarm",):
            _helper = getattr(_nodes_mod, _helper_name, None)
            if callable(_helper):
                g[_helper_name] = _helper
    except Exception:
        pass
    compiled = compile(code, fp, "exec")
    exec(compiled, g, g)

    try:
        for k, v in list(g.items()):
            if isinstance(v, type):
                try:
                    setattr(v, "_handlers_globals", g)
                except Exception:
                    pass
    except Exception:
        pass

    return g

def _handlers_file_path(config_uid: str) -> str:
    """
    Handlers/<uid>/handlers.py relative to project root.
    Adjust root if your folder layout differs.
    """
    base_dir = os.path.dirname(os.path.abspath(__file__))  # .../client_app
    root = os.path.abspath(os.path.join(base_dir, ".."))   # project root (usually)
    return os.path.join(root, "Handlers", config_uid, "handlers.py")


def _load_server_node_class(config_uid: str, class_name: str):
    """Load a generated server node class with source-aware module caching.

    Compiling and executing the whole generated ``handlers.py`` for every
    NodeLink on every card is extremely expensive.  A local handlers file is
    now cached by its nanosecond mtime and size, so an editor save is still
    picked up immediately while ordinary section rendering reuses one module.
    DB-backed handlers use the existing explicit invalidation paths.
    """
    config_uid = str(config_uid or "").strip()
    class_name = str(class_name or "").strip()
    if not config_uid or not class_name:
        raise ValueError("config_uid/class_name is empty")

    fp = _handlers_file_path(config_uid)
    code = None
    source_stamp = ""
    db_cfg = None

    if os.path.isfile(fp):
        st = os.stat(fp)
        source_stamp = f"file:{int(getattr(st, 'st_mtime_ns', int(st.st_mtime * 1_000_000_000)))}:{int(st.st_size)}"
    else:
        # Use the live SQL revision as the cache stamp. _client_cfg_by_uid is
        # request-cached, so repeated link resolutions do not repeat this query.
        try:
            db_cfg = _client_cfg_by_uid(config_uid)
        except Exception:
            db_cfg = None
        modified = getattr(db_cfg, "last_modified", None) if db_cfg is not None else None
        modified_stamp = modified.isoformat() if hasattr(modified, "isoformat") else str(modified or "")
        blob_size = len(str(getattr(db_cfg, "nodes_server_handlers", "") or "")) if db_cfg is not None else 0
        source_stamp = f"db:{modified_stamp}:{blob_size}"

    class_key = (config_uid, class_name, source_stamp)
    with _SERVER_HANDLERS_LOCK:
        cached_class = _SERVER_NODE_CLASS_MEM.get(class_key)
        if cached_class is not None:
            return cached_class

        ns_entry = _SERVER_HANDLERS_NS_MEM.get(config_uid)
        if isinstance(ns_entry, dict) and ns_entry.get("stamp") == source_stamp:
            g = ns_entry.get("globals") or {}
            cls = g.get(class_name)
            if cls is not None:
                _SERVER_NODE_CLASS_MEM[class_key] = cls
                return cls

        if os.path.isfile(fp):
            with open(fp, "r", encoding="utf-8") as f:
                code = f.read()
        else:
            cfg = db_cfg
            if cfg is None:
                import __main__ as main
                Configuration = main.Configuration
                cfg = models.db.session.execute(
                    select(Configuration).where(Configuration.uid == config_uid)
                ).scalar_one_or_none()
            if not cfg or not getattr(cfg, "nodes_server_handlers", None):
                raise ValueError(f"Server handlers not found for config (no file {fp} and no DB blob)")
            code = base64.b64decode(cfg.nodes_server_handlers).decode("utf-8", errors="replace")

        g: Dict[str, Any] = {
            "__name__": f"handlers_{config_uid}",
            "__file__": fp,
        }
        try:
            for helper_name in ("CallSwarm",):
                helper = getattr(_nodes_mod, helper_name, None)
                if callable(helper):
                    g[helper_name] = helper
        except Exception:
            pass

        compiled = compile(code, fp, "exec")
        exec(compiled, g, g)

        try:
            for value in list(g.values()):
                if isinstance(value, type):
                    try:
                        setattr(value, "_handlers_globals", g)
                    except Exception:
                        pass
        except Exception:
            pass

        cls = g.get(class_name)
        if cls is None:
            raise ValueError(f"Class {class_name} not found in server handlers")

        # Remove stale revisions of this configuration and keep one namespace.
        _SERVER_HANDLERS_NS_MEM[config_uid] = {"stamp": source_stamp, "globals": g}
        for key in list(_SERVER_NODE_CLASS_MEM.keys()):
            if key and key[0] == config_uid and key[2] != source_stamp:
                _SERVER_NODE_CLASS_MEM.pop(key, None)
        _SERVER_NODE_CLASS_MEM[class_key] = cls
        return cls




def fetch_config(config_url: str) -> Dict[str, Any]:
    resp = requests.get(config_url, timeout=20)
    resp.raise_for_status()
    return resp.json()


def build_global_sections(repos: List[models.Repo], db) -> List[Dict[str, str]]:
    seen: Dict[str, str] = {}
    has_empty = False

    for r in repos:
        parsed = get_parsed_config(r, db)
        if not parsed:
            continue
        cfg = parsed["cfg"]
        classes = [c for c in (cfg.get("classes") or []) if not bool(c.get("hidden")) and not _is_print_form_class_type(c)]
        section_codes_with_classes = {class_section_code(c) for c in classes}
        if "" in section_codes_with_classes:
            has_empty = True
        for s in normalize_sections(cfg):
            code = s["code"]
            # Do not show sections that contain no accessible classes.
            if code not in section_codes_with_classes:
                continue
            if code not in seen:
                seen[code] = s["name"]

    sections: List[Dict[str, str]] = []
    if has_empty:
        sections.append({"code": "", "name": "<...>"})
    for code in sorted(seen.keys()):
        if code == "":
            continue
        sections.append({"code": code, "name": seen[code]})
    return sections


def _node_id(node: Dict[str, Any]) -> str:
    return str(
        node.get("_id")
        or node.get("_Id")
        or (node.get("_data") or {}).get("_id")
        or (node.get("_data") or {}).get("_Id")
        or ""
    )


@contextmanager
def _nodes_config_context(config_uid: str, parsed_config: Optional[Dict[str, Any]] = None):
    """Expose the exact installed repo config to nodes.py for this lookup.

    The section UI is rendered from ``get_parsed_config()``.  Semantic search
    previously loaded index definitions independently from the Designer DB.
    For cached/remote repositories (and briefly after an index edit) that could
    mean that the filter was visible in the browser while nodes.py saw no such
    index or saw stale model/threshold settings.  Keep both paths on the same
    configuration snapshot.
    """
    cfg_var = parsed_var = None
    cfg_token = parsed_token = None
    try:
        cfg_var = getattr(_nodes_mod, "CURRENT_CONFIG_UID", None)
        parsed_var = getattr(_nodes_mod, "CURRENT_PARSED_CONFIG", None)
        if cfg_var is not None:
            cfg_token = cfg_var.set(str(config_uid or "").strip())
        if parsed_var is not None and parsed_config is not None:
            parsed_token = parsed_var.set(parsed_config)
        yield
    finally:
        try:
            if parsed_token is not None:
                parsed_var.reset(parsed_token)
        except Exception:
            pass
        try:
            if cfg_token is not None:
                cfg_var.reset(cfg_token)
        except Exception:
            pass


def _text_like_index_ids_for_class(config_uid: str, class_name: str, q: str, limit: int | None = None) -> Optional[List[str]]:
    q = (q or "").strip()
    if not q:
        return None
    try:
        node_cls = _load_server_node_class(config_uid, class_name)
        defs = node_cls._get_defined_indexes(config_uid) if hasattr(node_cls, "_get_defined_indexes") else []
    except Exception:
        return None

    idx_names: List[str] = []
    for idx in defs or []:
        if not isinstance(idx, dict):
            continue
        kind = str(idx.get("kind") or "hash_index").strip().lower()
        if kind not in ("text_index", "trigram_index", "text_index_full", "semantic", "semantic_index", "semanic_index"):
            continue
        name = str(idx.get("name") or "").strip()
        if name and name not in idx_names:
            idx_names.append(name)

    if not idx_names:
        return None

    out: List[str] = []
    seen = set()
    has_index_rows = False
    for name in idx_names:
        try:
            if hasattr(node_cls, "_defined_index_has_rows") and node_cls._defined_index_has_rows(name, config_uid):
                has_index_rows = True
            else:
                store = node_cls._defined_index_storage(name, config_uid)
                if list(store.keys()):
                    has_index_rows = True
        except Exception:
            pass
        try:
            ids = node_cls.find_ids_by_index(name, q, config_uid, limit=max(1, int(limit or 0) - len(out)) if limit else None)
        except Exception as exc:
            # Do not turn a model/index failure into an apparently successful
            # unfiltered result without leaving a useful server-side trace.
            print("section index search failed", config_uid, class_name, name, repr(exc))
            ids = []
        for nid in ids or []:
            sid = str(nid)
            if sid not in seen:
                seen.add(sid)
                out.append(sid)
                if limit and len(out) >= int(limit):
                    break
        if limit and len(out) >= int(limit):
            break
    return out if has_index_rows else None


def _nodes_storage_page(config_uid: str, class_name: str, *, offset: int, limit: int, q: str = "", index_name: str = "", index_value: str = "") -> List[Dict[str, Any]]:
    """Read nodes directly from the same storage as /api/.../page (no HTTP call)."""
    storage_key = f"{class_name}_{config_uid}"
    db_path = os.path.join("node_storage", f"{storage_key}.sqlite")
    if not os.path.exists(db_path):
        return []

    table = "unnamed"
    # Keep the exact user text for embedding models.  Lexical fallback matching
    # uses q_lower below, but semantic search must receive the same query text as
    # a direct SentenceTransformer call (including case for brands/model names).
    q_raw = (q or "").strip()
    q_lower = q_raw.lower()
    index_name = (index_name or "").strip()
    index_value = "" if index_value is None else str(index_value)

    def unpack(blob):
        try:
            return pickle.loads(blob)
        except Exception:
            return None

    conn = sqlite3.connect(db_path)
    try:
        cur = conn.cursor()

        def fetch_items_by_ids(ids: List[str]) -> List[Dict[str, Any]]:
            items: List[Dict[str, Any]] = []
            for nid in (ids or [])[offset: offset + limit]:
                try:
                    cur.execute(f"SELECT value FROM {table} WHERE key = ?", (str(nid),))
                    row = cur.fetchone()
                    if not row:
                        continue
                    obj = unpack(row[0])
                except Exception:
                    obj = None
                if obj is not None:
                    items.append(obj)
            return items

        if index_name and index_value != "":
            try:
                node_cls = _load_server_node_class(config_uid, class_name)
                ids = node_cls.find_ids_by_index(index_name, index_value, config_uid, limit=offset + limit)
            except Exception:
                ids = []
            return fetch_items_by_ids(ids)

        if q_raw:
            indexed_ids = _text_like_index_ids_for_class(config_uid, class_name, q_raw, limit=offset + limit)
            if indexed_ids is not None:
                return fetch_items_by_ids(indexed_ids)

        if not q_raw:
            cur.execute(
                f"SELECT value FROM {table} ORDER BY key LIMIT ? OFFSET ?",
                (limit, offset),
            )
            rows = cur.fetchall()
            items = []
            for (val_blob,) in rows:
                obj = unpack(val_blob)
                if obj is not None:
                    items.append(obj)
            return items

        # slow path: q scan
        cur.execute(f"SELECT value FROM {table}")
        rows = cur.fetchall()
        all_items = []
        for (val_blob,) in rows:
            obj = unpack(val_blob)
            if obj is not None:
                all_items.append(obj)

        def match(item: dict) -> bool:
            data = (item or {}).get("_data") or {}
            try:
                # prefer precomputed _search_index
                sidx = data.get("_search_index")
                if isinstance(sidx, str) and q_lower in sidx.lower():
                    return True
            except Exception:
                pass
            for v in data.values():
                try:
                    if q_lower in str(v).lower():
                        return True
                except Exception:
                    pass
            return False

        filtered = [it for it in all_items if match(it)]

        # mimic server sorting rule
        def sort_key(item: dict):
            d = (item or {}).get("_data") or {}
            if "_sort_string_desc" in d:
                return str(d.get("_sort_string_desc") or "")
            if "_sort_string" in d:
                return str(d.get("_sort_string") or "")
            return _node_id(item)

        any_desc = any("_sort_string_desc" in ((it or {}).get("_data") or {}) for it in filtered)
        filtered.sort(key=sort_key, reverse=any_desc)
        return filtered[offset: offset + limit]
    finally:
        conn.close()


def _auth_tuple(repo: models.Repo) -> Optional[tuple]:
    if repo.username:
        return (repo.username, repo.password)
    return None


def _api_get_remote(repo: models.Repo, path: str, *, params: Optional[dict] = None, timeout: int = 20) -> Any:
    url = repo.base_url.rstrip("/") + path
    resp = requests.get(url, params=params, auth=_auth_tuple(repo), headers=_client_remote_system_user_headers(), timeout=timeout)
    resp.raise_for_status()
    return resp.json()


def _fetch_nodes_for_class(repo: models.Repo, *, config_uid: str, class_name: str, q: str, limit: int, index_name: str = "", index_value: str = "", user=None, parsed_config: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
    """Fetch nodes either locally (same server) or remotely (repo.base_url override)."""
    # Default: no base_url override or points to this server -> read local storage.
    # If base_url is configured and does not match current host, do remote HTTP.
    base_url = (repo.base_url or "").strip().rstrip("/")
    current = ""
    if has_request_context():
        try:
            current = (request.host_url or "").rstrip("/")
        except Exception:
            current = ""
    storage_path = os.path.join("node_storage", f"{class_name}_{config_uid}.sqlite")
    is_local = os.path.exists(storage_path) or not base_url or (bool(current) and base_url == current)

    if is_local:
        actor = user if user is not None else _ngenie_effective_user()
        if actor is not None and not _client_user_can_access_class(config_uid, class_name, user=actor):
            return []
        with _nodes_config_context(config_uid, parsed_config):
            items = _nodes_storage_page(config_uid, class_name, offset=0, limit=limit, q=q, index_name=index_name, index_value=index_value)
        return _client_filter_nodes_for_acl(config_uid, class_name, items, user=actor) if actor is not None else items

    # Remote
    try:
        payload = _api_get_remote(
            repo,
            f"/api/config/{config_uid}/node/{class_name}/page",
            params=({"offset": 0, "limit": limit, "q": q} if q else {"offset": 0, "limit": limit}) | ({"index_name": index_name, "index_value": index_value} if index_name and index_value != "" else {}),
        )
        items = payload.get("items", [])
        return items if isinstance(items, list) else []
    except Exception:
        # fallback to full list
        try:
            all_nodes = _api_get_remote(repo, f"/api/config/{config_uid}/node/{class_name}")
        except Exception:
            return []

        if isinstance(all_nodes, dict):
            items = list(all_nodes.values())
        elif isinstance(all_nodes, list):
            items = all_nodes
        else:
            items = []

        ql = (q or "").strip().lower()
        if ql:
            def match(n: Dict[str, Any]) -> bool:
                data = n.get("_data") or {}
                if isinstance(data.get("_search_index"), str):
                    return ql in data["_search_index"].lower()
                try:
                    return ql in json.dumps(data, ensure_ascii=False).lower()
                except Exception:
                    return False
            items = [n for n in items if match(n)]

        return items[:limit]


def _api_get_remote(repo: models.Repo, path: str, *, params: Optional[dict] = None, timeout: int = 20) -> Any:
    url = repo.base_url.rstrip("/") + path
    resp = requests.get(url, params=params, auth=_auth_tuple(repo), headers=_client_remote_system_user_headers(), timeout=timeout)
    resp.raise_for_status()
    return resp.json()


def _api_post_remote(repo: models.Repo, path: str, *, json_data: Any = None, timeout: int = 20) -> Any:
    url = repo.base_url.rstrip("/") + path
    resp = requests.post(url, json=json_data, auth=_auth_tuple(repo), headers=_client_remote_system_user_headers(), timeout=timeout)
    resp.raise_for_status()
    if resp.content and resp.headers.get("content-type", "").lower().startswith("application/json"):
        return resp.json()
    if resp.content:
        try:
            return resp.json()
        except Exception:
            return resp.text
    return None


def _api_delete_remote(repo: models.Repo, path: str, *, json_data: Any = None, timeout: int = 20) -> Any:
    url = repo.base_url.rstrip("/") + path
    resp = requests.delete(url, json=json_data, auth=_auth_tuple(repo), headers=_client_remote_system_user_headers(), timeout=timeout)
    resp.raise_for_status()
    if resp.content and resp.headers.get("content-type", "").lower().startswith("application/json"):
        return resp.json()
    if resp.content:
        try:
            return resp.json()
        except Exception:
            return resp.text
    return None




# ---------- nGenie AI assistant ----------

NGENIE_DEEPSEEK_URL = "https://api.deepseek.com/v1/chat/completions"
NGENIE_PURPLE = "#6f42c1"


def _ngenie_root_models():
    import importlib
    return importlib.import_module("models")


def _ngenie_effective_user():
    try:
        u = getattr(g, "api_user", None)
        if u is not None:
            return u
    except Exception:
        pass
    try:
        if getattr(current_user, "is_authenticated", False):
            return current_user
    except Exception:
        pass

    # Server timers and server-side NodaScript run without a Flask login request.
    # set_runtime_context() stores the repository owner in CURRENT_SYSTEM_USER;
    # resolve that id back to the SQL User so config/class ACL and RLS use the
    # same actor as an interactive nGenie request.
    try:
        runtime_user = _nodes_mod.CURRENT_SYSTEM_USER.get()
        if isinstance(runtime_user, dict):
            runtime_id = int(runtime_user.get("id") or runtime_user.get("user_id") or 0)
        else:
            runtime_id = int(getattr(runtime_user, "id", 0) or 0)
        if runtime_id:
            root = _ngenie_root_models()
            User = getattr(root, "User", None)
            if User is not None:
                return models.db.session.get(User, runtime_id)
    except Exception:
        pass
    return None


def _ngenie_effective_user_id() -> int:
    u = _ngenie_effective_user()
    uid = int(getattr(u, "id", 0) or 0)
    if uid:
        return uid
    # Server timers run in an application context without flask-login's
    # current_user.  Their runtime context carries the repository owner.
    try:
        runtime_user = _nodes_mod.CURRENT_SYSTEM_USER.get()
        if isinstance(runtime_user, dict):
            return int(runtime_user.get("id") or runtime_user.get("user_id") or 0)
        return int(getattr(runtime_user, "id", 0) or 0)
    except Exception:
        return 0


def _ngenie_user_can_access_config(user, config_uid: str) -> bool:
    if not user or not str(config_uid or "").strip():
        return False
    root = _ngenie_root_models()
    cfg = root.Configuration.query.filter_by(uid=str(config_uid).strip()).first()
    if not cfg:
        return False
    if int(getattr(cfg, "user_id", 0) or 0) == int(getattr(user, "id", 0) or 0):
        return True
    try:
        return root.UserConfigAccess.query.filter_by(user_id=user.id, config_id=cfg.id).first() is not None
    except Exception:
        return False


def _ngenie_api_auth_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        user = None
        try:
            if getattr(current_user, "is_authenticated", False):
                user = current_user
        except Exception:
            user = None

        if user is None:
            token = ""
            bearer = request.headers.get("Authorization", "") or ""
            if bearer.lower().startswith("bearer "):
                token = bearer.split(" ", 1)[1].strip()
            token = token or (request.headers.get("X-API-Token", "") or "").strip()
            if token:
                try:
                    import app as _root_app
                    user = _root_app.check_api_token(token)
                except Exception:
                    user = None

        if user is None:
            auth = request.authorization
            if auth:
                try:
                    root = _ngenie_root_models()
                    cand = root.User.query.filter_by(email=auth.username).first()
                    if cand and check_password_hash(cand.password, auth.password):
                        user = cand
                except Exception:
                    user = None

        if not user:
            return jsonify({"ok": False, "error": "Unauthorized"}), 401
        if not bool(getattr(user, "can_api", False)) and not bool(getattr(current_user, "is_authenticated", False)):
            return jsonify({"ok": False, "error": "Forbidden"}), 403

        cfg_uid = kwargs.get("config_uid") or kwargs.get("uid")
        # For POST endpoints config_uid lives in JSON, not in the URL.
        try:
            if not cfg_uid and request.is_json:
                body = request.get_json(silent=True) or {}
                cfg_uid = body.get("config_uid") or ((body.get("state") or {}).get("config_uid") if isinstance(body.get("state"), dict) else "")
        except Exception:
            cfg_uid = ""
        if cfg_uid and not _ngenie_user_can_access_config(user, str(cfg_uid)):
            return jsonify({"ok": False, "error": "Forbidden"}), 403

        g.api_user = user
        return f(*args, **kwargs)
    return decorated


def _ngenie_split_top_level(text: Any, delimiter: str = ",") -> List[str]:
    raw = str(text or "")
    parts: List[str] = []
    buf: List[str] = []
    depth = 0
    quote = ""
    esc = False
    for ch in raw:
        if quote:
            buf.append(ch)
            if esc:
                esc = False
            elif ch == "\\":
                esc = True
            elif ch == quote:
                quote = ""
            continue
        if ch in ('"', "'"):
            quote = ch
            buf.append(ch)
            continue
        if ch in "([{" :
            depth += 1
        elif ch in ")]}" and depth > 0:
            depth -= 1
        if ch == delimiter and depth == 0:
            item = "".join(buf).strip()
            if item:
                parts.append(item)
            buf = []
        else:
            buf.append(ch)
    item = "".join(buf).strip()
    if item:
        parts.append(item)
    return parts


def _ngenie_split_once_top_level(text: Any, delimiter: str) -> Tuple[str, str]:
    raw = str(text or "")
    depth = 0
    quote = ""
    esc = False
    for i, ch in enumerate(raw):
        if quote:
            if esc:
                esc = False
            elif ch == "\\":
                esc = True
            elif ch == quote:
                quote = ""
            continue
        if ch in ('"', "'"):
            quote = ch
            continue
        if ch in "([{" :
            depth += 1
            continue
        if ch in ")]}" and depth > 0:
            depth -= 1
            continue
        if ch == delimiter and depth == 0:
            return raw[:i].strip(), raw[i + 1:].strip()
    return raw.strip(), ""


def _ngenie_unquote(value: Any) -> str:
    s = str(value or "").strip()
    if len(s) >= 2 and s[0] == s[-1] and s[0] in ('"', "'"):
        return s[1:-1]
    return s


def _ngenie_parse_type(type_text: Any) -> Dict[str, Any]:
    t = str(type_text or "").strip()

    # Wizard/DataStructure function types:
    # Node("Goods"), ChildNode("OrderPosition"), DataSet("goods") / Dataset("goods")
    m = re.match(r"^(Node|ChildNode|DataSet|Dataset)\s*\((.*)\)$", t, re.I)
    if m:
        kind = m.group(1).lower()
        if kind == "dataset":
            kind = "dataset"
        if kind == "dataset":
            kind = "dataset"
        if kind == "dataset":
            kind = "dataset"
        if kind == "dataset":
            kind = "dataset"

        # normalize DataSet/Dataset to dataset
        if kind in {"dataset", "dataset"}:
            kind = "dataset"

        return {
            "kind": kind,
            "target": _ngenie_unquote(m.group(2)),
            "raw": t,
        }

    # List/table syntax:
    # positions:[Node("ContainerLine")]
    # positions:[ChildNode("ContainerLine")]
    # lines:[Node("CommonLine")]
    if t.startswith("[") and t.endswith("]"):
        inner = t[1:-1].strip()
        inner_type = _ngenie_parse_type(inner)
        return {
            "kind": "list",
            "item_type": inner_type,
            "target": inner_type.get("target") or "",
            "relation": inner_type.get("kind") or "",
            "raw": t,
        }

    # Spinner/select syntax from Wizard
    m = re.match(r"^select\s*\((.*)\)$", t, re.I)
    if m:
        return {
            "kind": "select",
            "raw": t,
        }

    low = t.lower()
    if low in {"str", "string", "text", "textarea"}:
        kind = "string"
    elif low in {"int", "integer", "number", "float", "double", "decimal"}:
        kind = "number"
    elif low in {"bool", "boolean", "checkbox", "check", "switch", "галочка"}:
        kind = "boolean"
    elif low in {"date", "datetime", "time"}:
        kind = low
    else:
        kind = low or "string"

    return {"kind": kind, "raw": t}


def _ngenie_parse_field_spec(spec: Any) -> Optional[Dict[str, Any]]:
    left, typ = _ngenie_split_once_top_level(spec, ":")
    if not left:
        return None
    label, name = _ngenie_split_once_top_level(left, "|")
    if not name:
        name = re.sub(r"[^a-zA-Z0-9_]+", "_", label.strip()).strip("_") or label.strip()
    name = str(name or "").strip()
    if not name:
        return None
    parsed_type = _ngenie_parse_type(typ or "string")
    return {
        "name": name,
        "label": str(label or name).strip(),
        "type": parsed_type.get("raw") or str(typ or "string"),
        "kind": parsed_type.get("kind") or "string",
        "target": parsed_type.get("target") or "",
        "item_type": parsed_type.get("item_type"),
    }


def _ngenie_parse_data_structure(text: Any) -> Dict[str, Any]:
    # One canonical parser is shared with Solutions validators.  Keeping this
    # wrapper preserves all existing routes call-sites while preventing the
    # validator and Wizard/runtime grammars from drifting apart.
    from .data_structure import parse_data_structure
    return parse_data_structure(text)


def _ngenie_legacy_inline_table_bindings(parsed: Dict[str, Any], data: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Bind the old unnamed ``[A|a:..., B|b:...]`` table to a real data field.

    Older configurations declared one inline table without a field name while
    the UI stored it under a concrete Table id such as ``lines``.  The parser
    correctly exposes that schema in ``virtual_tables``, but nGenie previously
    ignored it during reference resolution, so Node fields were persisted as
    raw ``{"query": ...}`` dictionaries.  When there is exactly one unnamed
    table and one plausible list-valued data field, treat that field as the
    inline table.  This affects only nGenie data normalization, not layout
    rendering or other views.
    """
    if not isinstance(parsed, dict) or not isinstance(data, dict):
        return []
    virtual = [x for x in (parsed.get("virtual_tables") or []) if isinstance(x, dict)]
    if len(virtual) != 1:
        return []
    named = {
        str(x.get("name") or "").strip()
        for x in (parsed.get("tables") or [])
        if isinstance(x, dict) and str(x.get("name") or "").strip()
    }
    candidates = [
        str(k) for k, v in data.items()
        if isinstance(v, list) and str(k) not in named and not str(k).startswith("_")
    ]
    if len(candidates) != 1:
        common = [x for x in candidates if x.lower() in {"lines", "rows", "items", "positions"}]
        if len(common) != 1:
            return []
        candidates = common
    table = dict(virtual[0])
    table.update({
        "name": candidates[0],
        "label": table.get("label") or candidates[0],
        "kind": "inline_table",
        "relation": "inline",
        "inline": True,
    })
    return [table]


def _ngenie_has_data_field(cls_cfg: Dict[str, Any], field_name: str) -> bool:
    parsed = _ngenie_parse_data_structure(cls_cfg.get("data_structure") or "")
    for fld in parsed.get("fields") or []:
        if str(fld.get("name") or "") == str(field_name or ""):
            return True
    return False


def _ngenie_find_parent_link_field(child_cfg: Dict[str, Any], parent_class: str) -> str:
    """
    Ищет в классе строки явное поле-ссылку на родительский документ.

    Пример:
      ContainerLine:
        parent_doc|parent_doc: Node("Container")

    Тогда при добавлении строки в Container.positions nGenie должен заполнить:
      parent_doc = <uid Container>

    Это важнее, чем _parent: _parent — служебная иерархия, а parent_doc —
    бизнес-ссылка, заданная пользователем в DataStructure.
    """
    parent_class_l = str(parent_class or "").strip().lower()
    if not parent_class_l:
        return ""

    parsed = _ngenie_parse_data_structure(child_cfg.get("data_structure") or "")
    candidates: List[Tuple[int, str]] = []

    for fld in parsed.get("fields") or []:
        if str(fld.get("kind") or "").lower() != "node":
            continue

        target = str(fld.get("target") or "").strip().lower()
        if target != parent_class_l:
            continue

        name = str(fld.get("name") or "").strip()
        label = str(fld.get("label") or "").strip()
        if not name:
            continue

        probe = f"{name} {label}".lower()
        score = 10

        if name in {"parent_doc", "parent_document", "document", "doc", "owner_doc"}:
            score += 100
        if "parent" in probe:
            score += 60
        if "doc" in probe or "document" in probe or "док" in probe:
            score += 30
        if parent_class_l in probe:
            score += 20

        candidates.append((score, name))

    if not candidates:
        return ""

    candidates.sort(reverse=True)
    return candidates[0][1]

def _ngenie_all_repos(config_uid: str = "") -> List[models.Repo]:
    uid = _ngenie_effective_user_id()
    q = models.Repo.query
    if uid:
        q = q.filter_by(user_id=uid)
    elif not str(config_uid or "").strip():
        return []
    if str(config_uid or "").strip():
        q = q.filter_by(config_uid=str(config_uid).strip())
    return q.order_by(models.Repo.name.asc(), models.Repo.id.asc()).all()


def _ngenie_data_node_classes(parsed: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
    out: Dict[str, Dict[str, Any]] = {}
    for name, cls in (parsed.get("classes") or {}).items():
        if not isinstance(cls, dict):
            continue
        if _class_type_value(cls) != "data_node" or cls.get("hidden"):
            continue
        out[str(name)] = cls
    return out


def _ngenie_compact_sample_value(value: Any, depth: int = 0) -> Any:
    """Keep samples JSON-native while bounding their size.

    Older code converted every list/dict (including document table parts such as
    ``lines``) to a JSON *string*.  The LLM then quite reasonably generated code
    that iterated strings and called ``.get`` on their characters/items.  Keep the
    real shape instead: arrays remain arrays and rows remain dictionaries.
    """
    if depth >= 3:
        if isinstance(value, (dict, list, tuple)):
            return "…"
        return str(value)[:220] if isinstance(value, str) else value
    if isinstance(value, dict):
        out: Dict[str, Any] = {}
        for idx, (k, v) in enumerate(value.items()):
            if idx >= 16:
                out["__truncated__"] = True
                break
            out[str(k)] = _ngenie_compact_sample_value(v, depth + 1)
        return out
    if isinstance(value, (list, tuple)):
        result = [_ngenie_compact_sample_value(v, depth + 1) for v in list(value)[:20]]
        if len(value) > 20:
            result.append({"__truncated__": True, "remaining": len(value) - 20})
        return result
    if isinstance(value, str):
        return value[:1000]
    if value is None or isinstance(value, (bool, int, float)):
        return value
    try:
        json.dumps(value, ensure_ascii=False)
        return value
    except Exception:
        return str(value)[:220]


def _ngenie_short_node_samples_info(repo: models.Repo, class_name: str, limit: int = 5, user=None) -> Tuple[List[Dict[str, Any]], bool, int]:
    """Return compact samples plus whether the sample set is complete.

    Fetch one extra row so the model can distinguish a complete small dataset from
    a truncated sample. ``samples_complete`` is authoritative even when there are
    only a few rows: four returned rows with ``samples_complete=true`` means that
    the class really contains four accessible rows.
    """
    safe_limit = max(1, int(limit or 5))
    rows = _fetch_nodes_for_class(
        repo, config_uid=repo.config_uid, class_name=class_name, q="", limit=safe_limit + 1, user=user
    )
    complete = len(rows) <= safe_limit
    visible_rows = rows[:safe_limit]
    out: List[Dict[str, Any]] = []
    for r in visible_rows:
        data = r.get("_data") if isinstance(r, dict) else {}
        if not isinstance(data, dict):
            data = r if isinstance(r, dict) else {}
        small: Dict[str, Any] = {}
        for k, v in data.items():
            if str(k).startswith("_") and k not in {"_id", "_class"}:
                continue
            small[k] = _ngenie_compact_sample_value(v)
            if len(small) >= 16:
                break
        if small:
            out.append(small)
    # total_count is exact only for complete small sets. For truncated samples it
    # is deliberately reported as a lower bound through samples_complete=False.
    return out, complete, len(rows) if complete else safe_limit + 1


def _ngenie_short_node_sample(repo: models.Repo, class_name: str, limit: int = 5, user=None) -> List[Dict[str, Any]]:
    return _ngenie_short_node_samples_info(repo, class_name, limit, user=user)[0]


def _ngenie_collect_context(config_uid: str = "", include_samples: bool = True, sample_limit: int = 5) -> Tuple[List[Dict[str, Any]], Dict[Tuple[str, str], Tuple[models.Repo, Dict[str, Any]]]]:
    classes: List[Dict[str, Any]] = []
    lookup: Dict[Tuple[str, str], Tuple[models.Repo, Dict[str, Any]]] = {}
    actor = _ngenie_effective_user()
    for repo in _ngenie_all_repos(config_uid):
        parsed = get_parsed_config(repo, models.db, user=actor) or {}
        for cls_name, cls_cfg in _ngenie_data_node_classes(parsed).items():
            parsed_ds = _ngenie_parse_data_structure(cls_cfg.get("data_structure") or "")
            samples, samples_complete, sample_count = (
                _ngenie_short_node_samples_info(repo, cls_name, sample_limit, user=actor)
                if include_samples else ([], False, 0)
            )
            item = {
                "config_uid": repo.config_uid,
                "config_name": repo.display_name or repo.name or repo.config_uid,
                "class_name": cls_name,
                "display_name": cls_cfg.get("display_name") or cls_name,
                "record_view": cls_cfg.get("record_view") or "",
                "data_structure": cls_cfg.get("data_structure") or "",
                "ngenie_role": cls_cfg.get("ngenie_role") or cls_cfg.get("ngenieRole") or cls_cfg.get("nGenieRole") or "",
                "ngenie_description": cls_cfg.get("ngenie_description") or cls_cfg.get("ngenieDescription") or cls_cfg.get("nGenieDescription") or cls_cfg.get("description") or "",
                "ngenie_prompt": cls_cfg.get("ngenie_prompt") or cls_cfg.get("nGeniePrompt") or cls_cfg.get("prompt") or "",
                "fields": parsed_ds.get("fields") or [],
                "tables": parsed_ds.get("tables") or [],
                "virtual_tables": parsed_ds.get("virtual_tables") or [],
                "indexes": cls_cfg.get("indexes") or cls_cfg.get("indexes_json") or [],
                "file_fields": _ngenie_file_fields_from_class(cls_cfg),
                "samples": samples,
                "samples_complete": bool(samples_complete),
                "sample_count": int(sample_count),
                "sample_limit": int(sample_limit),
            }
            classes.append(item)
            lookup[(repo.config_uid, cls_name)] = (repo, cls_cfg)
    return classes, lookup


def _ngenie_collect_config_prompts(config_uid: str = "") -> List[Dict[str, str]]:
    prompts: List[Dict[str, str]] = []
    seen = set()
    actor = _ngenie_effective_user()
    for repo in _ngenie_all_repos(config_uid):
        try:
            parsed = get_parsed_config(repo, models.db, user=actor) or {}
            cfg = parsed.get("cfg") or {}
            prompt = str(cfg.get("ngenie_prompt") or cfg.get("nGeniePrompt") or "").strip()
            key = (repo.config_uid, prompt)
            if prompt and key not in seen:
                seen.add(key)
                prompts.append({
                    "config_uid": repo.config_uid,
                    "config_name": repo.display_name or repo.name or repo.config_uid,
                    "prompt": prompt,
                })
        except Exception:
            continue
    return prompts


def _ngenie_parse_plugins_json(value: Any) -> List[Dict[str, Any]]:
    raw = value
    if isinstance(raw, str):
        raw = raw.strip()
        if not raw:
            return []
        try:
            raw = json.loads(raw)
        except Exception:
            return []
    if isinstance(raw, dict):
        raw = [raw]
    if not isinstance(raw, list):
        return []
    out: List[Dict[str, Any]] = []
    for item in raw:
        if isinstance(item, dict):
            out.append(dict(item))
    return out


def _ngenie_file_fields_from_class(cls_cfg: Dict[str, Any]) -> List[Dict[str, str]]:
    """Return file/media gallery fields declared by web plugins.

    FileGallery/MediaGallery are not DataStructure fields, but they store their
    filenames in _data under the plugin id. Expose them to nGenie so an attached
    file can be added to the right list.
    """
    out: List[Dict[str, str]] = []
    seen = set()
    plugins: List[Dict[str, Any]] = []
    plugins.extend(_ngenie_parse_plugins_json((cls_cfg or {}).get("plug_in_web") or ""))
    plugins.extend(_ngenie_parse_plugins_json((cls_cfg or {}).get("plug_in") or ""))
    for plugin in plugins:
        typ = str(plugin.get("type") or "").strip()
        if typ not in {"FileGallery", "MediaGallery"}:
            continue
        fid = str(plugin.get("id") or ("pic_files" if typ == "MediaGallery" else "files")).strip()
        if not fid or fid in seen:
            continue
        seen.add(fid)
        out.append({"field": fid, "type": typ})
    return out


def _ngenie_normalize_attachments(raw: Any) -> List[Dict[str, Any]]:
    """Normalize attachment metadata and preserve server-side extraction data.

    Browsers first upload the file to UserFiles and then send this compact
    object to the chat endpoint.  The endpoint enriches it with extracted_text
    so the LLM receives the file contents, not only the filename.
    """
    if not isinstance(raw, list):
        return []
    out: List[Dict[str, Any]] = []
    for item in raw[:20]:
        if not isinstance(item, dict):
            continue
        filename = str(item.get("filename") or item.get("url") or "").strip()
        if not filename:
            continue
        normalized = {
            "filename": filename,
            "url": str(item.get("url") or filename).strip(),
            "original_name": str(item.get("original_name") or item.get("name") or "").strip(),
            "content_type": str(item.get("content_type") or item.get("type") or "").strip(),
            "size": item.get("size") if isinstance(item.get("size"), (int, float)) else None,
            "s3": bool(item.get("s3")),
        }
        extracted_text = item.get("extracted_text")
        if isinstance(extracted_text, str) and extracted_text:
            normalized["extracted_text"] = extracted_text
        extraction_format = str(item.get("extraction_format") or "").strip()
        if extraction_format:
            normalized["extraction_format"] = extraction_format
        extraction_error = str(item.get("extraction_error") or "").strip()
        if extraction_error:
            normalized["extraction_error"] = extraction_error
        if item.get("truncated"):
            normalized["truncated"] = True
        out.append(normalized)
    return out


_NGENIE_ATTACHMENT_MAX_BYTES = 25 * 1024 * 1024
_NGENIE_ATTACHMENT_MAX_TEXT_PER_FILE = 90000
_NGENIE_ATTACHMENT_MAX_TEXT_TOTAL = 180000


def _ngenie_attachment_display_name(item: Dict[str, Any]) -> str:
    return str(item.get("original_name") or item.get("filename") or item.get("url") or "attachment").strip()


def _ngenie_read_limited_stream(stream: Any, limit: int = _NGENIE_ATTACHMENT_MAX_BYTES) -> bytes:
    raw = stream.read(limit + 1)
    if len(raw) > limit:
        raise ValueError(f"file is larger than {limit // (1024 * 1024)} MiB")
    return raw


def _ngenie_attachment_bytes(item: Dict[str, Any], config_uid: str) -> Tuple[bytes, str]:
    """Read an attachment only from the selected configuration's UserFiles/S3.

    Attachment metadata is controlled by the browser, therefore arbitrary URLs
    and arbitrary local paths are deliberately rejected.
    """
    source = str(item.get("filename") or item.get("url") or "").strip()
    if not source:
        raise ValueError("attachment has no filename")

    parsed = urlparse(source)
    if parsed.scheme in {"http", "https"}:
        # Preferred path for files uploaded through the configured S3 backend.
        key_parser = getattr(main, "_s3_key_from_public_url", None)
        object_key = ""
        try:
            if callable(key_parser):
                object_key = str(key_parser(source) or "").strip()
        except Exception:
            object_key = ""
        if object_key:
            s3_client = getattr(main, "s3", None)
            bucket = getattr(main, "S3_BUCKET", None)
            if s3_client is None or not bucket:
                raise ValueError("S3 storage is not configured")
            obj = s3_client.get_object(Bucket=bucket, Key=object_key)
            return _ngenie_read_limited_stream(obj["Body"]), os.path.basename(unquote(parsed.path)) or _ngenie_attachment_display_name(item)

        # Some S3-compatible installations do not expose the helper above.
        # Allow a direct fetch only from the configured S3 endpoint host.
        endpoint = str(getattr(main, "S3_ENDPOINT", "") or "").strip()
        endpoint_host = urlparse(endpoint).netloc.lower() if endpoint else ""
        if endpoint_host and parsed.netloc.lower() == endpoint_host:
            response = requests.get(source, timeout=45, stream=True)
            response.raise_for_status()
            data = bytearray()
            for chunk in response.iter_content(chunk_size=65536):
                if not chunk:
                    continue
                data.extend(chunk)
                if len(data) > _NGENIE_ATTACHMENT_MAX_BYTES:
                    raise ValueError(f"file is larger than {_NGENIE_ATTACHMENT_MAX_BYTES // (1024 * 1024)} MiB")
            return bytes(data), os.path.basename(unquote(parsed.path)) or _ngenie_attachment_display_name(item)
        raise ValueError("remote attachment is not from configured UserFiles/S3 storage")

    config_uid = str(config_uid or "").strip()
    if not config_uid:
        raise ValueError("select a configuration before reading attachments")
    name = _safe_filename(unquote(os.path.basename(source)))
    if not name:
        raise ValueError("invalid attachment filename")
    base_dir = os.path.realpath(os.path.join(_userfiles_root(), config_uid))
    path = os.path.realpath(os.path.join(base_dir, name))
    if os.path.commonpath([base_dir, path]) != base_dir or not os.path.isfile(path):
        raise ValueError("uploaded file was not found in the selected configuration")
    with open(path, "rb") as fh:
        return _ngenie_read_limited_stream(fh), name


def _ngenie_decode_text_bytes(raw: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1251", "cp866", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _ngenie_trim_extracted_text(text: Any, max_chars: int) -> Tuple[str, bool]:
    value = str(text or "").replace("\x00", "").strip()
    if len(value) <= max_chars:
        return value, False
    suffix = "\n\n[Attachment content truncated by nGenie backend]"
    return value[:max(0, max_chars - len(suffix))].rstrip() + suffix, True


def _ngenie_cell_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.isoformat()
    if hasattr(value, "isoformat") and not isinstance(value, (str, bytes)):
        try:
            return value.isoformat()
        except Exception:
            pass
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    text = str(value)
    return text.replace("\r", " ").replace("\n", " ").replace("\t", " ").strip()


def _ngenie_extract_xlsx(raw: bytes, max_chars: int) -> Tuple[str, bool]:
    try:
        from openpyxl import load_workbook
    except Exception as exc:
        raise ValueError(f"openpyxl is not installed: {exc}")
    wb = load_workbook(filename=BytesIO(raw), read_only=True, data_only=True)
    parts: List[str] = []
    used = 0
    truncated = False
    try:
        for ws in wb.worksheets:
            header = f"Workbook sheet: {ws.title}"
            if parts:
                header = "\n" + header
            parts.append(header)
            used += len(header)
            nonempty_rows = 0
            for row_no, row in enumerate(ws.iter_rows(values_only=True), start=1):
                values = [_ngenie_cell_text(v) for v in row]
                while values and values[-1] == "":
                    values.pop()
                if not values or not any(values):
                    continue
                nonempty_rows += 1
                line = f"row {row_no}:\t" + "\t".join(values)
                if used + len(line) + 1 > max_chars:
                    truncated = True
                    break
                parts.append(line)
                used += len(line) + 1
            if nonempty_rows == 0:
                parts.append("[empty sheet]")
                used += 13
            if truncated:
                break
    finally:
        try:
            wb.close()
        except Exception:
            pass
    text = "\n".join(parts)
    if truncated:
        text += "\n[Workbook rows truncated by nGenie backend]"
    return text, truncated


def _ngenie_extract_pdf(raw: bytes, max_chars: int) -> Tuple[str, bool]:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise ValueError(f"pypdf is not installed: {exc}")
    reader = PdfReader(BytesIO(raw))
    parts: List[str] = []
    used = 0
    truncated = False
    for page_no, page in enumerate(reader.pages, start=1):
        text = str(page.extract_text() or "").strip()
        block = f"PDF page {page_no}:\n{text}" if text else f"PDF page {page_no}: [no extractable text]"
        if used + len(block) + 2 > max_chars:
            remain = max(0, max_chars - used - 2)
            if remain:
                parts.append(block[:remain])
            truncated = True
            break
        parts.append(block)
        used += len(block) + 2
    result = "\n\n".join(parts)
    if truncated:
        result += "\n\n[PDF text truncated by nGenie backend]"
    return result, truncated


def _ngenie_extract_docx(raw: bytes, max_chars: int) -> Tuple[str, bool]:
    try:
        from docx import Document
    except Exception as exc:
        raise ValueError(f"python-docx is not installed: {exc}")
    doc = Document(BytesIO(raw))
    parts: List[str] = []
    for paragraph in doc.paragraphs:
        text = str(paragraph.text or "").strip()
        if text:
            parts.append(text)
    for table_no, table in enumerate(doc.tables, start=1):
        parts.append(f"Table {table_no}:")
        for row_no, row in enumerate(table.rows, start=1):
            values = [_ngenie_cell_text(cell.text) for cell in row.cells]
            parts.append(f"row {row_no}:\t" + "\t".join(values))
    return _ngenie_trim_extracted_text("\n".join(parts), max_chars)


def _ngenie_extract_attachment(item: Dict[str, Any], config_uid: str, user_message: str, max_chars: int) -> Dict[str, Any]:
    result: Dict[str, Any] = {}
    raw, stored_name = _ngenie_attachment_bytes(item, config_uid)
    display_name = _ngenie_attachment_display_name(item) or stored_name
    ext = os.path.splitext(display_name)[1].lower() or os.path.splitext(stored_name)[1].lower()
    content_type = str(item.get("content_type") or mimetypes.guess_type(display_name)[0] or "").lower()

    if ext in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        text, truncated = _ngenie_extract_xlsx(raw, max_chars)
        result.update(extracted_text=text, extraction_format="excel", truncated=truncated)
        return result
    if ext == ".xls":
        raise ValueError("legacy .xls is not supported; save the workbook as .xlsx")
    if ext == ".pdf" or content_type == "application/pdf":
        text, truncated = _ngenie_extract_pdf(raw, max_chars)
        result.update(extracted_text=text, extraction_format="pdf", truncated=truncated)
        return result
    if ext == ".docx" or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        text, truncated = _ngenie_extract_docx(raw, max_chars)
        result.update(extracted_text=text, extraction_format="docx", truncated=truncated)
        return result
    if content_type.startswith("image/") or ext in {".jpg", ".jpeg", ".png", ".webp"}:
        mime = content_type if content_type in {"image/jpeg", "image/png", "image/webp"} else (mimetypes.guess_type(display_name)[0] or "image/jpeg")
        if mime == "image/jpg":
            mime = "image/jpeg"
        if len(raw) > 8 * 1024 * 1024:
            raise ValueError("image is larger than 8 MiB; compress it before attaching")
        data_url = _ngenie_validate_image_data_url("data:" + mime + ";base64," + base64.b64encode(raw).decode("ascii"))
        vision_prompt = (
            "Extract all visible text, tables, identifiers, quantities and other factual data needed for this user request. "
            "Return detailed plain text only; do not perform NodaLogic operations. User request: " + str(user_message or "")
        )
        text = _ngenie_generate_vision(vision_prompt, data_url, str(config_uid or ""))
        text, truncated = _ngenie_trim_extracted_text(text, max_chars)
        result.update(extracted_text=text, extraction_format="image_analysis", truncated=truncated)
        return result

    text_exts = {
        ".txt", ".csv", ".tsv", ".json", ".xml", ".html", ".htm", ".md", ".log",
        ".py", ".js", ".ts", ".java", ".kt", ".sql", ".yaml", ".yml", ".ini", ".cfg",
    }
    textual_mime = content_type.startswith("text/") or content_type in {
        "application/json", "application/xml", "application/javascript", "application/sql",
    }
    if ext in text_exts or textual_mime or (raw and b"\x00" not in raw[:4096]):
        text = _ngenie_decode_text_bytes(raw)
        text, truncated = _ngenie_trim_extracted_text(text, max_chars)
        fmt = "csv" if ext in {".csv", ".tsv"} else "text"
        result.update(extracted_text=text, extraction_format=fmt, truncated=truncated)
        return result
    raise ValueError(f"unsupported attachment format: {ext or content_type or 'binary'}")


def _ngenie_prepare_attachments_for_chat(raw: Any, config_uid: str, user_message: str) -> List[Dict[str, Any]]:
    attachments = _ngenie_normalize_attachments(raw)
    remaining = _NGENIE_ATTACHMENT_MAX_TEXT_TOTAL
    prepared: List[Dict[str, Any]] = []
    for item in attachments:
        enriched = dict(item)
        if enriched.get("extracted_text"):
            text, truncated = _ngenie_trim_extracted_text(enriched.get("extracted_text"), min(_NGENIE_ATTACHMENT_MAX_TEXT_PER_FILE, remaining))
            enriched["extracted_text"] = text
            if truncated:
                enriched["truncated"] = True
            remaining = max(0, remaining - len(text))
            prepared.append(enriched)
            continue
        if remaining <= 0:
            enriched["extraction_error"] = "total attachment text limit reached"
            prepared.append(enriched)
            continue
        try:
            details = _ngenie_extract_attachment(
                enriched,
                str(config_uid or "").strip(),
                str(user_message or ""),
                min(_NGENIE_ATTACHMENT_MAX_TEXT_PER_FILE, remaining),
            )
            enriched.update(details)
            remaining = max(0, remaining - len(str(enriched.get("extracted_text") or "")))
        except Exception as exc:
            enriched["extraction_error"] = str(exc)
        prepared.append(enriched)
    return prepared

def _ngenie_attachment_filenames(attachments: Any) -> List[str]:
    return [str(a.get("filename") or a.get("url") or "").strip() for a in _ngenie_normalize_attachments(attachments) if str(a.get("filename") or a.get("url") or "").strip()]


def _ngenie_first_file_field(cls_cfg: Dict[str, Any], prefer_media: bool = False) -> str:
    fields = _ngenie_file_fields_from_class(cls_cfg or {})
    if not fields:
        return ""
    if prefer_media:
        for f in fields:
            if f.get("type") == "MediaGallery":
                return str(f.get("field") or "")
    for f in fields:
        if f.get("type") == "FileGallery":
            return str(f.get("field") or "")
    return str((fields[0] or {}).get("field") or "")



def _ngenie_extract_json_object(text: Any) -> Dict[str, Any]:
    s = str(text or "").strip()
    if s.startswith("```"):
        s = re.sub(r"^```(?:json)?\s*", "", s, flags=re.I).strip()
        s = re.sub(r"\s*```$", "", s).strip()
    try:
        obj = json.loads(s)
        return obj if isinstance(obj, dict) else {}
    except Exception:
        pass
    start = s.find("{")
    end = s.rfind("}")
    if start >= 0 and end > start:
        try:
            obj = json.loads(s[start:end + 1])
            return obj if isinstance(obj, dict) else {}
        except Exception:
            return {}
    return {}


def _ngenie_call_deepseek(messages: List[Dict[str, str]]) -> Dict[str, Any]:
    """Call the shared provider from root credentials.json.

    Previously this path ignored credentials.json and always used the key from
    app.py plus api.deepseek.com.  That made timer/banner nGenie calls differ
    from nGenie Code and caused the reported TLS EOF failure.
    """
    data = _shared_chat_completion(
        messages,
        require_json=True,
        temperature=0.15,
        max_tokens=8000,
    )
    content = _shared_message_content(data)
    parsed = _ngenie_extract_json_object(content)
    if not parsed:
        parsed = {"reply": content or "LLM provider returned an empty answer", "operations": []}
    return parsed



def _ngenie_skill_context_summary(ctx: Dict[str, Any]) -> Dict[str, Any]:
    """Compact context for skill routing. Full class schema is sent only after
    selected skill prompts are known."""
    classes = []
    for c in (ctx.get("classes") or []):
        if not isinstance(c, dict):
            continue
        classes.append({
            "config_uid": c.get("config_uid") or "",
            "class_name": c.get("class_name") or c.get("name") or "",
            "display_name": c.get("display_name") or "",
            "role": c.get("ngenie_role") or "",
            "description": c.get("ngenie_description") or "",
            "file_fields": c.get("file_fields") or [],
        })
    return {
        "selected_config_uid": ctx.get("selected_config_uid") or "",
        "scope": ctx.get("scope") or "",
        "allow_catalog_create": bool(ctx.get("allow_catalog_create")),
        "current_node": {
            "present": bool(ctx.get("current_node")),
            "class_name": (ctx.get("current_node") or {}).get("class_name") if isinstance(ctx.get("current_node"), dict) else "",
        },
        "configuration_prompts": ctx.get("configuration_prompts") or [],
        "attachments": [
            {
                "name": _ngenie_attachment_display_name(a),
                "content_type": a.get("content_type") or "",
                "extraction_format": a.get("extraction_format") or "",
                "readable": bool(a.get("extracted_text")),
                "extraction_error": a.get("extraction_error") or "",
            }
            for a in _ngenie_normalize_attachments(ctx.get("attachments") or [])
        ],
        "classes": classes,
    }


def _ngenie_skill_selector_prompt() -> str:
    return """
Ты маршрутизатор навыков nGenie. Тебе дают пользовательский запрос, краткий контекст и каталог навыков.
Верни строго JSON:
{"skill_ids":["..."],"reason":"..."}

Правила:
- Выбери один или несколько skill_ids, достаточных для выполнения запроса.
- Не выбирай общие навыки программирования/nGenie code: это отдельный помощник, не этот чат. Исключение — специализированный навык wms_strategy_code, который формирует только модульную функцию planning_handler для текущего узла WMSStrategy и не меняет структуру конфигурации.
- Если пользователь просит создать/изменить данные и ещё подобрать ссылки, обычно нужен node_operations.
- Если приложен Excel/CSV/PDF/другой файл и пользователь просит импортировать его строки в узлы, создать документ или справочники по содержимому, выбирай node_operations.
- Если приложен файл и пользователь просит только анализ/сводку без изменения узлов, выбирай analysis_reports.
- Если пользователь просит только показать/найти список без изменения данных, обычно нужен search_display.
- Если пользователь просит аналитический отчёт/агрегацию/HTML-проекцию, обычно нужен analysis_reports.
- Если запрос относится к WMS-остаткам, QuantLedger, оборачиваемости, FEFO/FIFO, размещению или рекомендации перемещений, выбирай wms_ledger_decision. Для полноценного HTML-отчёта вместе/вместо него может понадобиться analysis_reports; для фактического изменения данных не позволяй runtime AI напрямую проводить QuantLedger.
- Если ни один навык не подходит, верни пустой массив skill_ids и краткую причину.
""".strip()


def _ngenie_select_skill_ids(user_message: str, ctx: Dict[str, Any], force_skill_ids: Any = None) -> Tuple[List[str], str]:
    forced = ngenie_skill_registry.normalize_skill_ids(force_skill_ids)
    if forced:
        return forced, "forced"
    catalog = ngenie_skill_registry.skill_catalog()
    if not catalog:
        return [], "skill catalog is empty"
    router_ctx = {
        "available_skills": catalog,
        "context": _ngenie_skill_context_summary(ctx),
        "user_message": str(user_message or ""),
    }
    try:
        answer = _ngenie_call_deepseek([
            {"role": "system", "content": _ngenie_skill_selector_prompt()},
            {"role": "user", "content": json.dumps(router_ctx, ensure_ascii=False, default=str)},
        ])
        ids = answer.get("skill_ids") or answer.get("skills") or answer.get("skillIds") or []
        if isinstance(ids, str):
            ids = [ids]
        selected = ngenie_skill_registry.normalize_skill_ids(ids)
        reason = str(answer.get("reason") or answer.get("reply") or "").strip()
        return selected, reason
    except Exception as e:
        traceback.print_exc()
        return [], f"skill selection failed: {e}"


def _ngenie_skill_blocks_for_messages(selected_skill_ids: Any, ctx: Dict[str, Any]) -> List[Dict[str, Any]]:
    return ngenie_skill_registry.selected_skill_payloads(selected_skill_ids, ctx)


def _ngenie_system_prompt() -> str:
    return """
Ты nGenie, AI-помощник NodaLogic для работы с узлами и анализа данных.
Отвечай строго JSON-объектом, без markdown вне JSON.

Теперь nGenie работает через файловые навыки. В контексте есть selected_skills:
- На первом внутреннем шаге backend отправляет маршрутизатору только DESCRIPTION всех навыков.
- В этот рабочий запрос backend добавляет только PROMPT выбранных навыков и их подготовленный контекст.
- Используй только выбранные навыки. Если selected_skills пустой, верни operations:[] и reply, что подходящего навыка для запроса нет.
- Не используй общие навыки nGenie code / программирования конфигураций: это отдельный помощник. Исключение — выбранный wms_strategy_code; он может вернуть в reply только готовую модульную функцию planning_handler для текущего WMSStrategy, без произвольного изменения классов и конфигурации.

Всегда учитывай:
- selected_config_uid задаёт единственную текущую конфигурацию; не смешивай её данные с другими конфигурациями;
- configuration_prompts, ngenie_description и ngenie_prompt класса как проектные правила;
- ngenie_role класса;
- scope: если scope=node_form, пользователь работает с текущим открытым узлом;
- scope=direct означает вызов функции ngenie() из выполняемого Python/NodaScript текущей конфигурации. Для явно запрошенной только сводки верни только непустой ключ summary, а остальные ключи оставь пустыми. Если кроме summary возвращаешь другие значимые данные, backend сохранит полный JSON.

Базовая JSON-схема ответа:
{
  "summary": "",
  "reply": "...",
  "operations": [],
  "data_requests": [],
  "resolve_requests": [],
  "clarification_requests": [],
  "display_requests": [],
  "analysis_html": "",
  "projection_title": "",
  "projection_method_code": "",
  "projection_parameters_layout": null,
  "projection_parameters_data": {},
  "projection_parameters_data_structure": "",
  "candidate_handler_code": "",
  "operation_handler_code": ""
}

Запрещено придумывать классы/поля/индексы. Используй только контекст NodaLogic и инструкции выбранных навыков.
""".strip()


def _ngenie_trim_text(value: Any, limit: int) -> str:
    text = str(value or "")
    if len(text) <= limit:
        return text
    return text[:limit] + "\n...[truncated]"


def _ngenie_normalize_projection_parameters_layout(value: Any) -> Any:
    if isinstance(value, str):
        text = value.strip()
        if not text:
            return None
        try:
            value = json.loads(text)
        except Exception:
            return None
    if isinstance(value, list):
        return value if value else None
    if isinstance(value, dict):
        return value if value else None
    return None


def _ngenie_normalize_conversation_artifact(raw: Any) -> Dict[str, Any]:
    if not isinstance(raw, dict):
        return {}
    layout = _ngenie_normalize_projection_parameters_layout(
        raw.get("projection_parameters_layout")
        or raw.get("parameters_layout")
        or raw.get("projection_params_layout")
    )
    params_data = (
        raw.get("projection_parameters_data")
        or raw.get("parameters_data")
        or raw.get("projection_params_data")
        or {}
    )
    if not isinstance(params_data, dict):
        params_data = {}
    out: Dict[str, Any] = {
        "projection_title": _ngenie_trim_text(raw.get("projection_title") or raw.get("title") or "", 1000),
        "projection_method_code": _ngenie_trim_text(raw.get("projection_method_code") or raw.get("method_code") or "", 70000),
        "analysis_html": _ngenie_trim_text(raw.get("analysis_html") or raw.get("projection_html") or raw.get("html") or "", 30000),
        "projection_parameters_layout": layout,
        "projection_parameters_data": dict(params_data),
        "projection_parameters_data_structure": _ngenie_trim_text(
            raw.get("projection_parameters_data_structure")
            or raw.get("parameters_data_structure")
            or raw.get("projection_params_data_structure")
            or "",
            10000,
        ),
    }
    return {k: v for k, v in out.items() if v not in (None, "", [], {})}


def _ngenie_latest_artifact_from_history(raw: Any) -> Dict[str, str]:
    """Recover the newest report artifact stored on a chat message.

    The browser keeps projection code on assistant messages.  A follow-up must
    still work if the top-level conversation_artifact was lost during a page
    refresh or by an older frontend build.
    """
    if not isinstance(raw, list):
        return {}
    for item in reversed(raw[-30:]):
        if isinstance(item, dict) and bool(item.get("projection_closed")):
            return {}
        artifact = _ngenie_normalize_conversation_artifact(item)
        if artifact:
            return artifact
    return {}


def _ngenie_history_mentions_report(raw: Any) -> bool:
    if not isinstance(raw, list):
        return False
    report_words = ("отчет", "отчёт", "report", "projection", "проекц")
    for item in reversed(raw[-12:]):
        if not isinstance(item, dict):
            continue
        if bool(item.get("projection_closed")):
            return False
        if _ngenie_normalize_conversation_artifact(item):
            return True
        text = str(item.get("content") or item.get("text") or item.get("reply") or "").lower()
        if any(word in text for word in report_words):
            return True
    return False


def _ngenie_message_requests_report_skill(message: Any, history: Any = None, artifact: Any = None) -> bool:
    """Deterministically route explicit report work to analysis_reports.

    Skill routing is not a semantic task here: phrases such as ``parameters of
    the report`` must never be sent to node_operations.  Keep the LLM router for
    genuinely ambiguous requests only.
    """
    text = str(message or "").strip().lower().replace("ё", "е")
    if not text:
        return False
    report_words = ("отчет", "report", "projection", "проекц")
    if any(word in text for word in report_words):
        return True
    continuation_words = (
        "добавь параметр", "добавить параметр", "добавь параметры", "добавить параметры",
        "добавь выбор", "добавить выбор", "выбор периода", "параметр периода",
        "фильтр периода", "добавь фильтр", "добавить фильтр", "измени параметры",
        "в параметры", "в него", "этот отчет", "этот отчёт",
    )
    has_prior_report = bool(_ngenie_normalize_conversation_artifact(artifact)) or _ngenie_history_mentions_report(history)
    return has_prior_report and any(word in text for word in continuation_words)


def _ngenie_message_is_report_edit(message: Any) -> bool:
    text = str(message or "").strip().lower().replace("ё", "е")
    if not text:
        return False
    edit_words = (
        "добав", "измен", "передел", "допол", "убер", "замен", "параметр",
        "фильтр", "выбор", "период", "колонк", "группиров", "сортиров",
    )
    report_words = ("отчет", "report", "projection", "проекц")
    direct_refs = ("в него", "к нему", "его параметр", "этот отчет")
    return (any(word in text for word in report_words) and any(word in text for word in edit_words)) or any(ref in text for ref in direct_refs)


_NGENIE_CHAT_ARTIFACTS: Dict[Tuple[str, str, str, str], Dict[str, Any]] = {}
_NGENIE_CHAT_ARTIFACTS_LOCK = threading.RLock()


def _ngenie_chat_artifact_key(config_uid: Any, scope: Any = "", node_context: Any = None) -> Tuple[str, str, str, str]:
    try:
        user_id = str(getattr(current_user, "id", "") or "")
    except Exception:
        user_id = ""
    node_id = ""
    if isinstance(node_context, dict):
        node_id = str(node_context.get("node_id") or node_context.get("id") or "").strip()
    return (user_id, str(config_uid or "").strip(), str(scope or "").strip(), node_id)


def _ngenie_cached_chat_artifact(config_uid: Any, scope: Any = "", node_context: Any = None) -> Dict[str, Any]:
    key = _ngenie_chat_artifact_key(config_uid, scope, node_context)
    with _NGENIE_CHAT_ARTIFACTS_LOCK:
        return dict(_NGENIE_CHAT_ARTIFACTS.get(key) or {})


def _ngenie_clear_chat_artifact(config_uid: Any, scope: Any = "", node_context: Any = None) -> None:
    key = _ngenie_chat_artifact_key(config_uid, scope, node_context)
    with _NGENIE_CHAT_ARTIFACTS_LOCK:
        _NGENIE_CHAT_ARTIFACTS.pop(key, None)


def _ngenie_remember_chat_artifact(config_uid: Any, scope: Any, node_context: Any, artifact: Any) -> Dict[str, Any]:
    normalized = _ngenie_normalize_conversation_artifact(artifact)
    if not normalized:
        return {}
    key = _ngenie_chat_artifact_key(config_uid, scope, node_context)
    with _NGENIE_CHAT_ARTIFACTS_LOCK:
        previous = dict(_NGENIE_CHAT_ARTIFACTS.get(key) or {})
        previous.update({k: v for k, v in normalized.items() if str(v or "").strip()})
        _NGENIE_CHAT_ARTIFACTS[key] = previous
        # Keep the process-local cache bounded for long-running cloud workers.
        if len(_NGENIE_CHAT_ARTIFACTS) > 256:
            for stale_key in list(_NGENIE_CHAT_ARTIFACTS.keys())[:64]:
                _NGENIE_CHAT_ARTIFACTS.pop(stale_key, None)
        return dict(previous)


def _ngenie_is_report_followup(message: Any, artifact: Any, history: Any = None) -> bool:
    """Return True when the user is modifying the report already open in chat.

    This is deliberately lexical and conservative.  The frontend artifact is
    the source of truth; without it, ordinary phrases such as "add parameters"
    must not be reinterpreted as a report edit.
    """
    current = _ngenie_normalize_conversation_artifact(artifact)
    if not current:
        return False
    text = str(message or "").strip().lower().replace("ё", "е")
    if not text:
        return False
    report_words = ("отчет", "отчета", "отчете", "отчету", "report", "projection", "проекц")
    direct_refs = (
        "этот отчет", "в этот отчет", "в отчете", "в него", "к нему",
        "добавь в него", "добавить в него", "измени его", "переделай его",
        "дополни его", "добавь выбор", "добавить выбор", "добавь параметр",
        "добавить параметр", "добавь фильтр", "добавить фильтр",
        "выбор периода", "параметр периода", "фильтр периода",
    )
    if any(x in text for x in direct_refs):
        return True
    action_words = ("добав", "измен", "передел", "допол", "убер", "замен", "сделай", "покаж", "выбор", "фильтр", "параметр")
    return any(r in text for r in report_words) and any(a in text for a in action_words)


def _ngenie_report_followup_instruction(message: Any, artifact: Dict[str, Any]) -> str:
    title = str((artifact or {}).get("projection_title") or "текущий отчёт").strip()
    return (
        "КРИТИЧЕСКИЙ КОНТЕКСТ ПРОДОЛЖЕНИЯ: это не новый запрос и не новый отчёт. "
        f"Пользователь продолжает изменять уже открытый отчёт «{title}». "
        "Нельзя спрашивать, какой отчёт имеется в виду. Возьми переданный ниже "
        "projection_method_code/analysis_html как текущую версию и верни обновлённый "
        "projection_method_code. Сохрани существующую структуру отчёта и внеси только "
        "запрошенное изменение. Если пользователь просит выбор периода, параметры "
        "начала/окончания НЕ должны попадать в _projection_html: это отдельный layout "
        "правой панели HTML-projection. Метод отчёта должен только прочитать эти значения "
        "из input_data/full_data или self._data и применить фильтрацию по подходящему полю "
        "даты из текущего кода и контекста класса. Если действительно существует несколько "
        "равноправных полей даты и выбор нельзя вывести из текущего кода, можно спросить "
        "только какое поле даты использовать, но не какой отчёт.\n\n"
        "Текущая просьба пользователя: " + str(message or "").strip()
    )


def _ngenie_message_requests_projection_parameters(message: Any, artifact: Any = None) -> bool:
    """Detect parameters both for a new report and for an existing report edit.

    The old implementation required an existing projection_method_code, so a
    one-shot request such as "сделай отчёт с периодом в параметрах" could never
    create the right-panel layout.
    """
    text = str(message or "").strip().lower().replace("ё", "е")
    if not text:
        return False
    parameter_words = (
        "параметр", "фильтр", "отбор", "выбор периода", "период для отбора",
        "date_from", "date_to", "начало периода", "конец периода",
        "report parameter", "report filter", "параметрами периода",
        "период в параметрах", "период в параметрах отчета",
    )
    action_words = ("добав", "измен", "сдел", "убер", "замен", "настрой")
    return any(word in text for word in parameter_words) and (
        any(word in text for word in action_words)
        or "выбор периода" in text
        or "параметры отчета" in text
        or "параметрах отчета" in text
        or "период в параметрах" in text
        or "с параметрами периода" in text
        or "параметры отчёта" in str(message or "").lower()
    )


def _ngenie_projection_parameter_fields_from_layout(layout: Any) -> List[Dict[str, str]]:
    fields: List[Dict[str, str]] = []
    seen = set()

    def walk(value: Any) -> None:
        if isinstance(value, list):
            for item in value:
                walk(item)
            return
        if not isinstance(value, dict):
            return
        typ = str(value.get("type") or "").strip()
        field_id = str(value.get("id") or "").strip()
        if field_id and typ in {"Input", "TextInput", "CheckBox", "Spinner", "NodeInput", "DatasetInput", "DatasetField"} and field_id not in seen:
            seen.add(field_id)
            input_type = str(value.get("input_type") or "").strip().upper()
            if typ == "CheckBox":
                data_type = "boolean"
            elif typ == "NodeInput":
                target = str(value.get("dataset") or "Node").strip() or "Node"
                data_type = f'Node("{target}")'
            elif typ in {"DatasetInput", "DatasetField"}:
                data_type = "string"
            elif input_type == "DATE":
                data_type = "date"
            elif input_type == "NUMBER":
                data_type = "number"
            else:
                data_type = "string"
            fields.append({
                "id": field_id,
                "caption": str(value.get("caption") or field_id).strip() or field_id,
                "type": data_type,
            })
        for key in ("layout", "tabs", "rows", "cols", "items", "children"):
            if key in value:
                walk(value.get(key))

    walk(layout)
    return fields


def _ngenie_projection_data_structure_from_layout(layout: Any) -> str:
    return "\n".join(
        f"{f['caption']}|{f['id']}: {f['type']}"
        for f in _ngenie_projection_parameter_fields_from_layout(layout)
    )


def _ngenie_projection_parameters_layout_valid(layout: Any, message: Any = "") -> bool:
    fields = _ngenie_projection_parameter_fields_from_layout(layout)
    if not fields:
        return False
    text = str(message or "").lower().replace("ё", "е")
    if "период" in text or "date_from" in text or "date_to" in text:
        ids = {str(f.get("id") or "").strip() for f in fields}
        # A period is a range.  Accept conventional or semantically equivalent
        # paired names, but reject a decorative/empty layout.
        conventional = {"date_from", "date_to"}.issubset(ids)
        starts = any(x in ids for x in ("date_from", "period_from", "period_start", "start_date", "date_start"))
        ends = any(x in ids for x in ("date_to", "period_to", "period_end", "end_date", "date_end"))
        return conventional or (starts and ends)
    return True


def _ngenie_projection_preserves_existing_report(old_code: Any, new_code: Any) -> bool:
    """Reject parameter edits that replace the working report with a new report.

    Parameter generation is delegated to nGenie Code, but the existing report is
    the authoritative artifact.  A valid edit keeps most of its meaningful HTML
    labels/CSS literals and has a recognizable code skeleton.
    """
    import difflib

    old = _ngenie_strip_method_code(old_code)
    new = _ngenie_strip_method_code(new_code)
    if not old or not new or "_projection_html" not in new:
        return False

    def meaningful_strings(body: str) -> List[str]:
        wrapped = "def __report(self, input_data):\n" + "\n".join(
            "    " + line if line.strip() else "" for line in body.splitlines()
        ) + "\n"
        try:
            tree = ast.parse(wrapped)
        except Exception:
            return []
        out: List[str] = []
        for node in ast.walk(tree):
            if not isinstance(node, ast.Constant) or not isinstance(node.value, str):
                continue
            value = re.sub(r"\s+", " ", node.value).strip()
            if len(value) < 8 or value.startswith("_projection_"):
                continue
            out.append(value)
        return out

    literals = meaningful_strings(old)
    if literals:
        retained = sum(1 for value in literals if value in new) / max(1, len(literals))
        if retained < 0.65:
            return False
    old_norm = re.sub(r"\s+", " ", old).strip()
    new_norm = re.sub(r"\s+", " ", new).strip()
    return difflib.SequenceMatcher(None, old_norm, new_norm).ratio() >= 0.32


def _ngenie_default_projection_parameters_for_request(message: Any) -> Dict[str, Any]:
    """Safe separate Projection layout when the optional Code bridge is absent."""
    text = str(message or "").lower().replace("ё", "е")
    if not any(x in text for x in ("период", "date_from", "date_to", "дата с", "дата по")):
        return {}
    layout = [
        [{"type": "Input", "id": "date_from", "caption": "Период с", "input_type": "DATE", "value": "@date_from"}],
        [{"type": "Input", "id": "date_to", "caption": "Период по", "input_type": "DATE", "value": "@date_to"}],
    ]
    return {
        "projection_parameters_layout": layout,
        "projection_parameters_data": {},
        "projection_parameters_data_structure": "Период с|date_from: date\nПериод по|date_to: date",
    }


def _ngenie_projection_parameters_bridge(
    message: Any,
    config_uid: str,
    artifact: Any,
) -> Dict[str, Any]:
    """Use optional nGenie Code knowledge to add real Projection parameters.

    The normal report skill owns data analysis.  nGenie Code is invoked only for
    the UI/code bridge because it contains the authoritative NodaLogic layout and
    Projection instructions.  A missing folder/credential simply falls back to
    the normal analysis_reports path.
    """
    current = _ngenie_normalize_conversation_artifact(artifact)
    current_method = str(current.get("projection_method_code") or "").strip()
    if not current_method:
        return {}
    try:
        import ngenie_code
        if not ngenie_code.available():
            return {}
    except Exception:
        return {}

    try:
        classes, _lookup = _ngenie_collect_context(config_uid, include_samples=False)
        compact_classes = []
        for cls in classes or []:
            if not isinstance(cls, dict):
                continue
            compact_classes.append({
                "class_name": cls.get("class_name") or "",
                "display_name": cls.get("display_name") or "",
                "ngenie_role": cls.get("ngenie_role") or "",
                "ngenie_description": cls.get("ngenie_description") or "",
                "ngenie_prompt": cls.get("ngenie_prompt") or "",
                "fields": [
                    {
                        "name": f.get("name") or "",
                        "label": f.get("label") or "",
                        "kind": f.get("kind") or "",
                        "target": f.get("target") or "",
                    }
                    for f in (cls.get("fields") or []) if isinstance(f, dict)
                ],
                "tables": [
                    {
                        "name": t.get("name") or "",
                        "label": t.get("label") or "",
                        "row_class": t.get("row_class") or t.get("target") or "",
                    }
                    for t in (cls.get("tables") or []) if isinstance(t, dict)
                ],
            })
        bridge_payload = {
            "mode": "projection_parameters_bridge",
            "selected_config_uid": str(config_uid or ""),
            "user_request": str(message or ""),
            "projection_title": current.get("projection_title") or "",
            "projection_method_code": current_method,
            "existing_projection_parameters_layout": current.get("projection_parameters_layout") or [],
            "existing_projection_parameters_data": current.get("projection_parameters_data") or {},
            "existing_projection_parameters_data_structure": current.get("projection_parameters_data_structure") or "",
            "configuration_classes": compact_classes,
        }
        bridge_system_prompt = ngenie_code.build_projection_parameters_bridge_system_prompt()
        raw = ngenie_code.call_llm(
            bridge_system_prompt,
            json.dumps(bridge_payload, ensure_ascii=False, default=str),
            max_tokens=12000,
            debug_stage="projection_parameters_bridge",
            debug_meta={"config_uid": str(config_uid or "")},
        )
        parsed = _ngenie_extract_json_object(raw)
        method_code = _ngenie_strip_method_code((parsed or {}).get("projection_method_code") or "") if isinstance(parsed, dict) else ""
        layout = _ngenie_normalize_projection_parameters_layout((parsed or {}).get("projection_parameters_layout")) if isinstance(parsed, dict) else None
        if (not method_code or not _ngenie_projection_parameters_layout_valid(layout, message) or not _ngenie_projection_preserves_existing_report(current_method, method_code)):
            repair_payload = dict(bridge_payload)
            repair_payload["validation_error"] = (
                "Предыдущий вариант был отклонён: он переписал существующий отчёт или потерял его HTML/группировки. "
                "Верни минимальное изменение исходного projection_method_code: сохрани все существующие строки HTML, "
                "CSS, группировки, колонки и итоги; добавь только чтение параметров и фильтрацию."
            )
            repair_payload["rejected_candidate"] = parsed if isinstance(parsed, dict) else {}
            raw = ngenie_code.call_llm(
                bridge_system_prompt,
                json.dumps(repair_payload, ensure_ascii=False, default=str),
                max_tokens=12000,
                debug_stage="projection_parameters_bridge_repair",
                debug_meta={"config_uid": str(config_uid or "")},
            )
            parsed = _ngenie_extract_json_object(raw)
            if not isinstance(parsed, dict):
                return {}
            method_code = _ngenie_strip_method_code(parsed.get("projection_method_code") or "")
            layout = _ngenie_normalize_projection_parameters_layout(parsed.get("projection_parameters_layout"))
        if not method_code or not _ngenie_projection_parameters_layout_valid(layout, message) or not _ngenie_projection_preserves_existing_report(current_method, method_code):
            return {}
        # Syntax only. Actual data access is verified by the normal preview pass.
        body_src = "def __projection_parameters_bridge(self, input_data):\n" + "\n".join(
            "    " + line if line.strip() else "" for line in method_code.splitlines()
        ) + "\n"
        ast.parse(body_src)
        defaults = dict(current.get("projection_parameters_data") or {})
        if isinstance(parsed.get("projection_parameters_data"), dict):
            defaults.update(parsed.get("projection_parameters_data") or {})
        data_structure = str(parsed.get("projection_parameters_data_structure") or "").strip()
        if not data_structure:
            data_structure = _ngenie_projection_data_structure_from_layout(layout)
        return {
            "summary": "",
            "reply": str(parsed.get("reply") or "Параметры отчёта добавлены.").strip(),
            "operations": [],
            "data_requests": [],
            "resolve_requests": [],
            "clarification_requests": [],
            "display_requests": [],
            "analysis_html": "",
            "projection_title": str(parsed.get("projection_title") or current.get("projection_title") or "nGenie analysis"),
            "projection_method_code": method_code,
            "projection_parameters_layout": layout,
            "projection_parameters_data": defaults,
            "projection_parameters_data_structure": data_structure,
            "candidate_handler_code": "",
            "operation_handler_code": "",
            "_ngenie_projection_parameters_bridge": True,
        }
    except Exception:
        traceback.print_exc()
        return {}


def _ngenie_render_projection_parameters_layout(config_uid: str, layout: Any, data: Any = None) -> str:
    layout_obj = _ngenie_normalize_projection_parameters_layout(layout)
    if layout_obj is None:
        return ""
    values = dict(data or {}) if isinstance(data, dict) else {}
    repo = None
    if config_uid:
        repo = models.Repo.query.filter_by(
            user_id=_ngenie_effective_user_id(), config_uid=str(config_uid or "")
        ).first()
    parsed = get_parsed_config(repo, models.db) if repo else None
    try:
        layout_obj = resolve_common_layout(parsed, layout_obj)
        if repo and parsed:
            _fill_nodeinput_views(repo, parsed, layout_obj, values)
        return str(render_nodalayout_html(
            layout_obj,
            values,
            assets_base_dir=_userfiles_dir_for_repo(repo) if repo else None,
            context=_nl_context(repo, class_name="", node_id="") if repo else {},
        ) or "")
    except Exception:
        traceback.print_exc()
        return ""


def _ngenie_repair_report_misroute(
    messages: List[Dict[str, str]],
    answer: Dict[str, Any],
    message: Any,
    history: Any = None,
    artifact: Any = None,
) -> Dict[str, Any]:
    """Reject the known wrong-skill/refused-report response and retry once."""
    if not _ngenie_message_requests_report_skill(message, history, artifact):
        return answer
    method_code = str((answer or {}).get("projection_method_code") or "").strip()
    if method_code:
        return answer
    reply = str((answer or {}).get("reply") or "").strip().lower().replace("ё", "е")
    wrong_markers = (
        "node_operations", "подходящего навыка", "нет подходящего навыка",
        "какой отчет", "какой именно отчет", "какой отчёт", "какой именно отчёт",
        "уточнить, какой отчет", "уточните, какой отчет",
    )
    if not any(marker in reply for marker in wrong_markers):
        return answer
    current = _ngenie_normalize_conversation_artifact(artifact)
    if not current:
        current = _ngenie_latest_artifact_from_history(history)
    correction = (
        "ОШИБКА МАРШРУТИЗАЦИИ ПРЕДЫДУЩЕГО ОТВЕТА. Запрос пользователя относится к "
        "навыку analysis_reports, а не node_operations. Нельзя отвечать, что подходящего "
        "навыка нет, и нельзя спрашивать, какой отчёт имеется в виду. "
    )
    if current:
        correction += (
            "В контексте уже передан текущий артефакт отчёта. Возьми его полный "
            "projection_method_code, внеси запрошенное изменение и верни ПОЛНЫЙ обновлённый "
            "projection_method_code, сохранив существующие группировки, оформление и итоги. "
        )
    else:
        correction += (
            "Сформируй отчёт по текущему запросу и верни projection_method_code. "
        )
    correction += "Верни строго полный JSON-ответ nGenie для текущей просьбы пользователя: " + str(message or "")
    retry_messages = list(messages) + [
        {"role": "assistant", "content": json.dumps(answer or {}, ensure_ascii=False, default=str)},
        {"role": "user", "content": correction},
    ]
    try:
        repaired = _ngenie_call_deepseek(retry_messages)
        return repaired if isinstance(repaired, dict) and repaired else answer
    except Exception:
        traceback.print_exc()
        return answer


def _ngenie_normalize_conversation_history(raw: Any, current_message: str = "") -> List[Dict[str, str]]:
    if not isinstance(raw, list):
        return []
    normalized: List[Dict[str, str]] = []
    for item in raw[-24:]:
        if not isinstance(item, dict):
            continue
        role = str(item.get("role") or "").strip().lower()
        if role not in {"user", "assistant"}:
            continue
        text = _ngenie_trim_text(item.get("content") or item.get("text") or item.get("reply") or "", 10000).strip()
        artifact = _ngenie_normalize_conversation_artifact(item)
        if artifact:
            artifact_note = (
                "\n\n[Артефакт этого ответа nGenie; используй его при ссылках "
                "«этот отчёт», «в него», «добавь параметры»:]\n"
                + json.dumps(artifact, ensure_ascii=False, default=str)
            )
            text = (text + artifact_note).strip()
        if text:
            normalized.append({"role": role, "content": text})

    # Frontend adds the current user bubble before sending. Do not duplicate it
    # in the actual prompt when it is present as the last history item.
    current = str(current_message or "").strip()
    if normalized and current and normalized[-1].get("role") == "user":
        last = str(normalized[-1].get("content") or "").strip()
        if last == current or last.startswith(current + "\n"):
            normalized.pop()
    return normalized


def _ngenie_build_messages(user_message: str, config_uid: str, node_context: Optional[Dict[str, Any]] = None, allow_catalog_create: bool = False, clarification_response: Optional[Dict[str, Any]] = None, scope: str = "", selected_skill_ids: Any = None, attachments: Any = None, conversation_history: Any = None, conversation_artifact: Any = None) -> List[Dict[str, str]]:
    direct_scope = str(scope or "").strip().lower() == "direct"
    classes, _lookup = _ngenie_collect_context(
        config_uid,
        include_samples=True,
        sample_limit=50 if direct_scope else 5,
    )
    ctx = {
        "selected_config_uid": config_uid or "",
        "allow_catalog_create": bool(allow_catalog_create),
        "scope": str(scope or "").strip(),
        "configuration_prompts": _ngenie_collect_config_prompts(config_uid),
        "classes": classes,
        "current_node": node_context or {},
        "attachments": _ngenie_normalize_attachments(attachments),
        "clarification_response": clarification_response or {},
        "hints": {
            "node_uid_format": "<config_uid>$<class_name>$<id>",
            "prefer_indexes": True,
            "search_order_before_create": ["class indexes", "text search", "regex/fuzzy local scan"],
            "helper_methods": ["Node.get_all(config_uid) returns dict id->Node", "ngenie_nodes", "ngenie_data", "ngenie_rows", "findByIndex", "getByIndex", "find(class_name, query) fallback when no suitable index exists", "findByGlobalIndex", "get_balance/getBalance/_get_balance", "class methods from nodes.py/handlers"],
            "resolve_candidate_limit": _ngenie_resolve_candidate_limit(),
        },
    }
    final_user_message = str(user_message or "")
    if clarification_response:
        final_user_message = (
            "Пользователь выбрал варианты в диалоге уточнения. Продолжи исходную задачу. "
            "Используй выбранные UID как готовые ссылки и не ищи их заново без необходимости.\n"
            + json.dumps(clarification_response, ensure_ascii=False, default=str)
        )

    normalized_history = _ngenie_normalize_conversation_history(conversation_history, current_message=user_message)
    history_artifact = _ngenie_latest_artifact_from_history(conversation_history)
    artifact = dict(history_artifact)
    artifact.update(_ngenie_normalize_conversation_artifact(conversation_artifact))
    report_skill_required = _ngenie_message_requests_report_skill(
        final_user_message, conversation_history, artifact
    )
    report_followup = _ngenie_is_report_followup(final_user_message, artifact, normalized_history)
    if not report_followup and _ngenie_message_is_report_edit(final_user_message):
        report_followup = bool(artifact) or _ngenie_history_mentions_report(conversation_history)
    selection_message = final_user_message
    if normalized_history:
        selection_message += "\n\nНедавний контекст диалога:\n" + "\n".join(
            str(item.get("content") or "")[:2000] for item in normalized_history[-4:]
        )
    if artifact:
        selection_message += "\n\nВ чате открыт существующий HTML-отчёт/projection, который пользователь может просить изменить."
    if report_followup:
        selection_message = _ngenie_report_followup_instruction(final_user_message, artifact) + "\n\n" + selection_message

    forced_skill_ids = selected_skill_ids
    # Explicit report creation/editing is deterministic routing.  Do not ask the
    # skill-selector model again: it occasionally chose node_operations merely
    # because the request contained the verb "add".
    if report_skill_required or report_followup:
        forced_skill_ids = ["analysis_reports"]
    selected_ids, skill_reason = _ngenie_select_skill_ids(selection_message, ctx, force_skill_ids=forced_skill_ids)
    try:
        g.ngenie_selected_skill_ids = selected_ids
    except Exception:
        pass
    skill_blocks = _ngenie_skill_blocks_for_messages(selected_ids, ctx)
    ctx["selected_skill_ids"] = selected_ids
    ctx["selected_skill_reason"] = skill_reason
    ctx["selected_skills"] = skill_blocks

    result: List[Dict[str, str]] = [
        {"role": "system", "content": _ngenie_system_prompt()},
        {"role": "user", "content": "Контекст NodaLogic и выбранные навыки:\n" + json.dumps(ctx, ensure_ascii=False, default=str)},
    ]
    if artifact:
        result.append({
            "role": "user",
            "content": (
                "Актуальный артефакт текущего веб-чата. Когда пользователь говорит "
                "«этот отчёт», «его», «в него», он имеет в виду этот артефакт. "
                "Изменяй существующий projection_method_code, а не спрашивай заново, какой отчёт:\n"
                + json.dumps(artifact, ensure_ascii=False, default=str)
            ),
        })
    result.extend(normalized_history)
    if report_followup:
        # Put the continuation instruction immediately before the current user
        # turn as well.  Long configuration context and history must not dilute
        # the reference to the already-open report.
        result.append({
            "role": "user",
            "content": _ngenie_report_followup_instruction(final_user_message, artifact),
        })
    else:
        result.append({"role": "user", "content": final_user_message})
    return result


def _ngenie_find_class(config_uid: str, class_name: str, lookup: Dict[Tuple[str, str], Tuple[models.Repo, Dict[str, Any]]]) -> Tuple[Optional[models.Repo], Optional[Dict[str, Any]], str]:
    class_name = str(class_name or "").strip()
    if config_uid and (config_uid, class_name) in lookup:
        repo, cfg = lookup[(config_uid, class_name)]
        return repo, cfg, class_name
    low = class_name.lower()
    for (cu, cn), (repo, cfg) in lookup.items():
        if config_uid and cu != config_uid:
            continue
        if cn.lower() == low or str(cfg.get("display_name") or "").strip().lower() == low:
            return repo, cfg, cn
    return None, None, class_name


def _ngenie_node_title(data: Dict[str, Any]) -> str:
    if not isinstance(data, dict):
        return ""
    for k in ("name", "title", "caption", "display_name", "article", "sku", "number"):
        v = data.get(k)
        if v is not None and str(v).strip():
            return str(v).strip()
    return ""


def _ngenie_get_internal_id_from_row(row: Dict[str, Any]) -> str:
    if not isinstance(row, dict):
        return ""
    for k in ("id", "node_id", "_id"):
        if row.get(k):
            try:
                return _nodes_mod.extract_internal_id(row.get(k))
            except Exception:
                return str(row.get(k))
    data = row.get("_data") if isinstance(row.get("_data"), dict) else row
    if isinstance(data, dict) and data.get("_id"):
        try:
            return _nodes_mod.extract_internal_id(data.get("_id"))
        except Exception:
            return str(data.get("_id"))
    return ""


def _ngenie_search_text_norm(value: Any) -> str:
    s = "" if value is None else str(value).strip().lower()
    return re.sub(r"\s+", " ", s)


def _ngenie_search_trigrams(value: Any) -> set:
    s = _ngenie_search_text_norm(value)
    if not s:
        return set()
    if len(s) < 3:
        return {s}
    padded = "  " + s + "  "
    return {padded[i:i + 3] for i in range(max(0, len(padded) - 2))}


def _ngenie_search_similarity(a: Any, b: Any) -> float:
    an = _ngenie_search_text_norm(a)
    bn = _ngenie_search_text_norm(b)
    if not an or not bn:
        return 0.0
    if an == bn:
        return 1.0
    if an in bn or bn in an:
        return 0.92
    ta = _ngenie_search_trigrams(an)
    tb = _ngenie_search_trigrams(bn)
    if not ta or not tb:
        return 0.0
    return len(ta & tb) / float(max(len(ta), len(tb), 1))


def _ngenie_node_search_strings(data: Dict[str, Any], cls_cfg: Optional[Dict[str, Any]] = None) -> List[str]:
    out: List[str] = []
    seen = set()

    def add(v: Any):
        if v is None or isinstance(v, (dict, list, tuple, set)):
            return
        s = str(v).strip()
        if not s or s in seen:
            return
        seen.add(s)
        out.append(s)

    # Prefer declared business fields, then fall back to all scalar fields.
    field_names: List[str] = []
    if isinstance(cls_cfg, dict):
        try:
            parsed = _ngenie_parse_data_structure(cls_cfg.get("data_structure") or "")
            field_names = [str(f.get("name") or "") for f in (parsed.get("fields") or []) if f.get("name")]
        except Exception:
            field_names = []
    if isinstance(data, dict):
        for key in field_names:
            add(data.get(key))
        for key in ("name", "title", "caption", "display_name", "article", "sku", "barcode", "code", "number"):
            add(data.get(key))
        for k, v in data.items():
            if str(k).startswith("_"):
                continue
            add(v)
    return out


def _ngenie_row_data(row: Any) -> Dict[str, Any]:
    if not isinstance(row, dict):
        return {}
    data = row.get("_data") if isinstance(row.get("_data"), dict) else row
    return data if isinstance(data, dict) else {}


def _ngenie_find_existing_node(repo: models.Repo, cls_cfg: Dict[str, Any], class_name: str, query_value: Any) -> Optional[str]:
    q = str(query_value or "").strip()
    if not q:
        return None
    indexes = cls_cfg.get("indexes") or cls_cfg.get("indexes_json") or []
    if not isinstance(indexes, list):
        indexes = []

    # 1) Exact/hash indexes are the fastest exact path. Text/trigram indexes
    # are candidate generators and must not be treated as a definitive match.
    for idx in indexes:
        if not isinstance(idx, dict):
            continue
        idx_name = str(idx.get("name") or "").strip()
        idx_kind = str(idx.get("kind") or idx.get("type") or "hash_index").strip().lower()
        if not idx_name or idx_kind not in {"hash", "hash_index", "exact"}:
            continue
        try:
            rows = _fetch_nodes_for_class(repo, config_uid=repo.config_uid, class_name=class_name, q="", limit=5, index_name=idx_name, index_value=q)
            if rows:
                internal_id = _ngenie_get_internal_id_from_row(rows[0])
                if internal_id:
                    return _nodes_mod.normalize_own_uid(repo.config_uid, class_name, internal_id)
        except Exception:
            continue

    # 2) Обычный поиск storage/search route (LIKE/text_index/text_index_full/remote q).
    try:
        rows = _fetch_nodes_for_class(repo, config_uid=repo.config_uid, class_name=class_name, q=q, limit=20)
        if rows:
            best_row = None
            best_score = -1.0
            for row in rows:
                data = _ngenie_row_data(row)
                strings = _ngenie_node_search_strings(data, cls_cfg)
                score = max([_ngenie_search_similarity(q, v) for v in strings] or [0.0])
                if score > best_score:
                    best_row, best_score = row, score
            if best_row:
                internal_id = _ngenie_get_internal_id_from_row(best_row)
                if internal_id:
                    return _nodes_mod.normalize_own_uid(repo.config_uid, class_name, internal_id)
    except Exception:
        pass

    # 3) Последний шанс перед созданием: локальный scan с trigram/fuzzy-оценкой.
    candidates: List[Tuple[float, str]] = []
    try:
        if _ngenie_local_repo(repo):
            parsed = get_parsed_config(repo, models.db) or {}
            _nodes_mod.set_runtime_context(repo.config_uid, parsed, system_user=_client_runtime_system_user_payload())
            node_cls = _load_server_node_class(repo.config_uid, class_name)
            nodes = node_cls.get_all(repo.config_uid) if node_cls and hasattr(node_cls, "get_all") else []
            node_values = list((nodes or {}).values()) if isinstance(nodes, dict) else list(nodes or [])
            for node in node_values[:2000]:
                try:
                    nid = _nodes_mod.extract_internal_id(getattr(node, "_id", "") or getattr(node, "id", ""))
                    data = node.get_data() if hasattr(node, "get_data") else getattr(node, "_data", {})
                    if not nid or not isinstance(data, dict):
                        continue
                    strings = _ngenie_node_search_strings(data, cls_cfg)
                    score = max([_ngenie_search_similarity(q, v) for v in strings] or [0.0])
                    if score >= 0.34:
                        candidates.append((score, nid))
                except Exception:
                    continue
        else:
            rows = _fetch_nodes_for_class(repo, config_uid=repo.config_uid, class_name=class_name, q="", limit=1000)
            for row in rows:
                internal_id = _ngenie_get_internal_id_from_row(row)
                data = _ngenie_row_data(row)
                if not internal_id or not data:
                    continue
                strings = _ngenie_node_search_strings(data, cls_cfg)
                score = max([_ngenie_search_similarity(q, v) for v in strings] or [0.0])
                if score >= 0.34:
                    candidates.append((score, internal_id))
    except Exception:
        candidates = []
    if candidates:
        candidates.sort(key=lambda x: x[0], reverse=True)
        return _nodes_mod.normalize_own_uid(repo.config_uid, class_name, candidates[0][1])
    return None




def _ngenie_safe_format(template: Any, values: Dict[str, Any]) -> str:
    template = str(template or "").strip()
    if not template:
        return ""

    class _SafeDict(dict):
        def __missing__(self, key):
            return ""

    try:
        return template.format_map(_SafeDict({k: "" if v is None else v for k, v in (values or {}).items()})).strip()
    except Exception:
        return template


def _ngenie_call_node_note_method(repo: models.Repo, class_name: str, internal_id: str, method_name: Any) -> Any:
    method_name = str(method_name or "").strip()
    if not method_name or "__" in method_name:
        return None
    # Notes are best-effort UI hints. Avoid remote method execution from here.
    try:
        if not _ngenie_local_repo(repo):
            return None
        parsed = get_parsed_config(repo, models.db) or {}
        _nodes_mod.set_runtime_context(repo.config_uid, parsed, system_user=_client_runtime_system_user_payload())
        node_cls = _load_server_node_class(repo.config_uid, class_name)
        node = node_cls.get(internal_id, repo.config_uid) if node_cls and hasattr(node_cls, "get") else None
        if node is None:
            return None
        fn = getattr(node, method_name, None)
        if not callable(fn):
            fn = getattr(node_cls, method_name, None)
        if not callable(fn):
            return None
        try:
            sig = inspect.signature(fn)
            required = [p for p in sig.parameters.values() if p.default is inspect._empty and p.kind in (p.POSITIONAL_ONLY, p.POSITIONAL_OR_KEYWORD, p.KEYWORD_ONLY)]
            if not required:
                return fn()
        except Exception:
            pass
        try:
            return fn({})
        except TypeError:
            return fn()
    except Exception:
        return None


def _ngenie_note_value_to_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list, tuple, set)):
        try:
            return json.dumps(value, ensure_ascii=False, default=str)
        except Exception:
            return str(value)
    return str(value)


def _ngenie_candidate_note(repo: models.Repo, class_name: str, internal_id: str, data: Dict[str, Any], request_obj: Dict[str, Any], explicit_note: Any = None) -> str:
    if explicit_note is not None and str(explicit_note).strip():
        return str(explicit_note).strip()
    data = data if isinstance(data, dict) else {}
    req = request_obj if isinstance(request_obj, dict) else {}

    method_name = req.get("note_method") or req.get("noteMethod") or req.get("candidate_note_method") or req.get("candidateNoteMethod")
    method_value = _ngenie_call_node_note_method(repo, class_name, internal_id, method_name) if method_name else None
    values = dict(data)
    if method_name:
        values[str(method_name)] = method_value
        values["method"] = method_value
        values["value"] = method_value

    template = req.get("note_template") or req.get("noteTemplate") or req.get("candidate_note_template") or req.get("candidateNoteTemplate")
    rendered = _ngenie_safe_format(template, values)
    if rendered:
        return rendered

    fields = req.get("note_fields") or req.get("noteFields") or req.get("candidate_note_fields") or req.get("candidateNoteFields") or []
    if isinstance(fields, str):
        fields = [x.strip() for x in re.split(r"[,;]", fields) if x.strip()]
    if isinstance(fields, list) and fields:
        parts = []
        label = str(req.get("note_label") or req.get("noteLabel") or "").strip()
        for f in fields:
            key = str(f or "").strip()
            if not key:
                continue
            val = data.get(key)
            if val is None or str(val).strip() == "":
                continue
            if len(fields) == 1 and label:
                parts.append(f"{label}: {_ngenie_note_value_to_text(val)}")
            else:
                parts.append(f"{key}: {_ngenie_note_value_to_text(val)}")
        if parts:
            return "; ".join(parts)

    if method_value is not None and str(method_value).strip():
        label = str(req.get("note_label") or req.get("noteLabel") or method_name or "").strip()
        val = _ngenie_note_value_to_text(method_value)
        return f"{label}: {val}" if label else val

    for key, label in (
        ("stock", "Остаток"),
        ("balance", "Остаток"),
        ("available", "Доступно"),
        ("qty", "Количество"),
        ("quantity", "Количество"),
        ("price", "Цена"),
        ("article", "Артикул"),
        ("barcode", "Штрихкод"),
        ("code", "Код"),
    ):
        val = data.get(key)
        if val is not None and str(val).strip():
            return f"{label}: {_ngenie_note_value_to_text(val)}"
    return ""



def _ngenie_resolve_candidate_limit() -> int:
    """Maximum candidate rows sent back to the LLM for one resolve step.

    Keep this technical limit generic. Subject-specific matching rules, such as
    how to interpret product dimensions, belong to class ngenie_prompt.
    """
    raw = getattr(main, "NGENIE_RESOLVE_CANDIDATE_LIMIT", None)
    if raw is None:
        raw = os.environ.get("NGENIE_RESOLVE_CANDIDATE_LIMIT", "10")
    try:
        val = int(raw)
    except Exception:
        val = 10
    return max(1, min(val, 50))


def _ngenie_resolve_request_cap(raw_requests: Any) -> List[Dict[str, Any]]:
    limit = _ngenie_resolve_candidate_limit()
    out: List[Dict[str, Any]] = []
    for req in (raw_requests if isinstance(raw_requests, list) else []):
        if not isinstance(req, dict):
            continue
        r = dict(req)
        try:
            requested = int(r.get("max_candidates") or r.get("maxCandidates") or limit)
        except Exception:
            requested = limit
        r["max_candidates"] = max(1, min(requested, limit))
        out.append(r)
    return out


def _ngenie_candidate_scalar_data(data: Any, max_fields: int = 30) -> Dict[str, Any]:
    if not isinstance(data, dict):
        return {}
    preferred = [
        "name", "title", "caption", "display_name", "article", "barcode", "code", "sku", "number",
        "brand", "series", "model", "unit", "qty", "stock", "balance", "available", "price",
    ]
    out: Dict[str, Any] = {}

    def add(k: str, v: Any):
        if len(out) >= max_fields:
            return
        if k in out or str(k).startswith("_"):
            return
        if v is None or isinstance(v, (dict, list, tuple, set)):
            return
        sv = str(v).strip()
        if not sv:
            return
        out[k] = v

    for k in preferred:
        add(k, data.get(k))
    for k, v in data.items():
        add(str(k), v)
    return out


def _ngenie_compact_candidate_for_llm(cand: Dict[str, Any]) -> Dict[str, Any]:
    data = cand.get("data") if isinstance(cand.get("data"), dict) else {}
    return {
        "uid": cand.get("uid") or "",
        "id": cand.get("id") or "",
        "class_name": cand.get("class") or cand.get("class_name") or "",
        "title": ((cand.get("view") or {}).get("title") if isinstance(cand.get("view"), dict) else "") or _ngenie_node_title(data),
        "score": cand.get("score"),
        "note": cand.get("note") or "",
        "data": _ngenie_candidate_scalar_data(data),
    }


def _ngenie_prepare_resolve_candidate_groups(raw_requests: Any, config_uid: str, lookup: Dict[Tuple[str, str], Tuple[models.Repo, Dict[str, Any]]]) -> List[Dict[str, Any]]:
    capped = _ngenie_resolve_request_cap(raw_requests)
    prepared = _ngenie_prepare_clarifications(capped, config_uid, lookup)
    groups: List[Dict[str, Any]] = []
    for g in prepared:
        if not isinstance(g, dict):
            continue
        cu = str(g.get("config_uid") or config_uid or "").strip()
        cls = str(g.get("class_name") or "").strip()
        repo, cls_cfg, cls_name = _ngenie_find_class(cu, cls, lookup) if cls else (None, None, cls)
        candidates = [_ngenie_compact_candidate_for_llm(c) for c in (g.get("candidates") or [])[:_ngenie_resolve_candidate_limit()] if isinstance(c, dict)]
        groups.append({
            "id": g.get("id") or "",
            "question": g.get("question") or "",
            "reason": g.get("reason") or "",
            "purpose": next((str(r.get("purpose") or "") for r in capped if isinstance(r, dict) and str(r.get("id") or r.get("key") or "") == str(g.get("id") or "")), ""),
            "config_uid": cu,
            "class_name": cls_name or cls,
            "query": g.get("query") or "",
            "context": g.get("context") if isinstance(g.get("context"), dict) else {},
            "ngenie_role": _ngenie_class_role(cls_cfg),
            "ngenie_prompt": (cls_cfg or {}).get("ngenie_prompt") or "",
            "candidate_limit": _ngenie_resolve_candidate_limit(),
            "candidates": candidates,
        })
    return groups


def _ngenie_build_resolve_messages(
    user_message: str,
    config_uid: str,
    node_context: Optional[Dict[str, Any]],
    allow_catalog_create: bool,
    scope: str,
    initial_answer: Dict[str, Any],
    resolve_groups: List[Dict[str, Any]],
    mode: str = "resolve_requests",
    attachments: Any = None,
) -> List[Dict[str, str]]:
    classes, _lookup = _ngenie_collect_context(config_uid, include_samples=False)
    ctx = {
        "selected_config_uid": config_uid or "",
        "allow_catalog_create": bool(allow_catalog_create),
        "scope": str(scope or "").strip(),
        "configuration_prompts": _ngenie_collect_config_prompts(config_uid),
        "classes": classes,
        "current_node": node_context or {},
        "attachments": _ngenie_normalize_attachments(attachments),
        "resolve_mode": mode,
        "resolve_candidate_limit": _ngenie_resolve_candidate_limit(),
        "initial_answer": initial_answer or {},
        "resolve_candidates": resolve_groups or [],
    }
    selected_ids, skill_reason = _ngenie_select_skill_ids(user_message, ctx, force_skill_ids=(initial_answer or {}).get("_ngenie_selected_skill_ids"))
    try:
        g.ngenie_selected_skill_ids = selected_ids
    except Exception:
        pass
    ctx["selected_skill_ids"] = selected_ids
    ctx["selected_skill_reason"] = skill_reason
    ctx["selected_skills"] = _ngenie_skill_blocks_for_messages(selected_ids, ctx)
    instruction = """
Это второй шаг nGenie: backend уже нашёл технических кандидатов для ссылочных узлов.
Твоя задача — применить ngenie_prompt соответствующего класса и исходную просьбу пользователя.
Не используй предметные правила из backend: вся предметная семантика находится в ngenie_prompt класса.

Правила ответа:
1. Если среди кандидатов есть один подходящий вариант, верни обычные operations и вставь выбранный UID в нужное Node-поле.
2. Если нужно добавить строку документа, верни append_table_rows с выбранным UID в поле строки и количеством/другими данными из исходного запроса.
3. Если выбрать нельзя или вариантов несколько по смыслу, верни clarification_requests с candidates, содержащими UID этих вариантов.
4. Не возвращай resolve_requests повторно на этом шаге.
5. Если кандидатов нет и allow_catalog_create=false, не создавай справочник, а сообщи что объект не найден.
6. Если scope=node_form, не создавай HTML-отчёты/проекции.
Отвечай строго JSON-объектом по основной схеме nGenie.
""".strip()
    return [
        {"role": "system", "content": _ngenie_system_prompt()},
        {"role": "user", "content": "Контекст NodaLogic и кандидаты resolve:\n" + json.dumps(ctx, ensure_ascii=False, default=str)},
        {"role": "user", "content": instruction + "\n\nИсходная просьба пользователя:\n" + str(user_message or "")},
    ]


def _ngenie_resolve_candidate_requests_with_llm(
    user_message: str,
    config_uid: str,
    node_context: Optional[Dict[str, Any]],
    allow_catalog_create: bool,
    scope: str,
    initial_answer: Dict[str, Any],
    raw_requests: Any,
    lookup: Dict[Tuple[str, str], Tuple[models.Repo, Dict[str, Any]]],
    mode: str = "resolve_requests",
    attachments: Any = None,
) -> Optional[Dict[str, Any]]:
    groups = _ngenie_prepare_resolve_candidate_groups(raw_requests, config_uid, lookup)
    if not groups:
        return None
    try:
        answer = _ngenie_call_deepseek(_ngenie_build_resolve_messages(
            user_message,
            config_uid,
            node_context,
            allow_catalog_create,
            scope,
            initial_answer,
            groups,
            mode=mode,
            attachments=attachments,
        ))
        if isinstance(answer, dict) and answer:
            # Keep the candidate list available in debug/raw responses without
            # forcing it into user-visible cards unless the model asks for them.
            answer.setdefault("_resolve_candidates", groups)
            return answer
    except Exception:
        traceback.print_exc()
    return None



def _ngenie_is_resolved_node_value(value: Any) -> bool:
    if not isinstance(value, str):
        return False
    raw = value.strip()
    if not raw or "$" not in raw:
        return False
    try:
        _cu, cls, iid = _nodes_mod.parse_uid_any(raw)
        return bool(cls and iid)
    except Exception:
        return False


def _ngenie_unresolved_node_ref_requests_from_data(
    repo: models.Repo,
    cls_cfg: Dict[str, Any],
    class_name: str,
    data: Dict[str, Any],
    lookup: Dict[Tuple[str, str], Tuple[models.Repo, Dict[str, Any]]],
    base_context: Dict[str, Any],
    prefix: str,
    field_defs: Optional[List[Dict[str, Any]]] = None,
) -> List[Dict[str, Any]]:
    if not repo or not isinstance(cls_cfg, dict) or not isinstance(data, dict):
        return []
    parsed = _ngenie_parse_data_structure(cls_cfg.get("data_structure") or "")
    fields = field_defs if isinstance(field_defs, list) else (parsed.get("fields") or [])
    out: List[Dict[str, Any]] = []
    for fld in fields:
        if not isinstance(fld, dict) or fld.get("kind") != "node":
            continue
        fname = str(fld.get("name") or "").strip()
        target = str(fld.get("target") or "").strip()
        if not fname or not target or fname not in data:
            continue
        val = data.get(fname)
        if val is None or isinstance(val, (list, tuple, set)):
            continue
        if isinstance(val, dict):
            uid, query, _create_data = _ngenie_node_ref_spec(val)
            if uid and _ngenie_is_resolved_node_value(uid):
                continue
            text = str(query or "").strip()
        else:
            text = str(val).strip()
        if not text or _ngenie_is_resolved_node_value(text):
            continue
        target_repo, target_cfg, target_name = _ngenie_find_class(repo.config_uid, target, lookup)
        if not target_repo or not target_cfg:
            continue
        ctx = dict(base_context or {})
        ctx.update({"source_class": class_name, "target_field": fname, "target_class": target_name})
        out.append({
            "id": f"{prefix}_{fname}_{len(out) + 1}",
            "question": f"Подобрать {target_name} для поля {fname}",
            "class_name": target_name,
            "config_uid": target_repo.config_uid,
            "query": text,
            "purpose": f"значение поля {fname} класса {class_name}",
            "context": ctx,
        })
    return out


def _ngenie_extract_resolve_requests_from_operations(
    raw_ops: Any,
    config_uid: str,
    node_context: Optional[Dict[str, Any]],
    lookup: Dict[Tuple[str, str], Tuple[models.Repo, Dict[str, Any]]],
) -> List[Dict[str, Any]]:
    """Find unresolved textual Node references in planned operations.

    This is a guard for cases where the LLM returned operations directly instead
    of resolve_requests. We still keep all domain decisions in the second LLM
    resolve step; this function only identifies which textual Node fields need
    candidate lookup.
    """
    ops = raw_ops if isinstance(raw_ops, list) else []
    requests_out: List[Dict[str, Any]] = []
    current_repo: Optional[models.Repo] = None
    current_class = ""
    if isinstance(node_context, dict) and node_context:
        current_class = str(node_context.get("class_name") or "").strip()
        cu = str(node_context.get("config_uid") or config_uid or "").strip()
        current_repo = models.Repo.query.filter_by(user_id=current_user.id, config_uid=cu).first()

    for op_idx, op in enumerate(ops):
        if not isinstance(op, dict):
            continue
        tool = str(op.get("tool") or op.get("action") or "").strip()
        if tool in {"", "none"}:
            continue

        if tool == "create_node":
            cu = str(op.get("config_uid") or config_uid or "").strip()
            cls = str(op.get("class_name") or op.get("class") or "").strip()
            repo, cls_cfg, cls_name = _ngenie_find_class(cu, cls, lookup)
            data = dict(op.get("data") or {}) if isinstance(op.get("data"), dict) else {}
            if repo and cls_cfg:
                base = {"tool": tool, "op_index": op_idx, "class_name": cls_name}
                requests_out.extend(_ngenie_unresolved_node_ref_requests_from_data(repo, cls_cfg, cls_name, data, lookup, base, f"op{op_idx}"))
                parsed = _ngenie_parse_data_structure(cls_cfg.get("data_structure") or "")
                table_defs = list(parsed.get("tables") or []) + _ngenie_legacy_inline_table_bindings(parsed, data)
                for t in table_defs:
                    field = str(t.get("name") or "").strip()
                    rows = data.get(field)
                    if not field or not isinstance(rows, list):
                        continue
                    is_inline = bool(t.get("inline")) or str(t.get("kind") or "").strip().lower() == "inline_table"
                    if is_inline:
                        for row_idx, row in enumerate(rows):
                            row_data = dict(row or {}) if isinstance(row, dict) else {}
                            base = {"tool": tool, "op_index": op_idx, "class_name": cls_name, "field": field, "row_index": row_idx, "row_class": ""}
                            requests_out.extend(_ngenie_unresolved_node_ref_requests_from_data(
                                repo, cls_cfg, cls_name, row_data, lookup, base,
                                f"op{op_idx}_row{row_idx}", field_defs=t.get("fields") or []
                            ))
                        continue
                    row_class = str(t.get("row_class") or t.get("target") or "").strip()
                    if not row_class:
                        continue
                    row_repo, row_cfg, row_class_name = _ngenie_find_class(repo.config_uid, row_class, lookup)
                    if not row_repo or not row_cfg:
                        continue
                    for row_idx, row in enumerate(rows):
                        row_data = dict(row or {}) if isinstance(row, dict) else {}
                        base = {"tool": tool, "op_index": op_idx, "class_name": cls_name, "field": field, "row_index": row_idx, "row_class": row_class_name}
                        requests_out.extend(_ngenie_unresolved_node_ref_requests_from_data(row_repo, row_cfg, row_class_name, row_data, lookup, base, f"op{op_idx}_row{row_idx}"))

        elif tool == "update_current_node":
            if not current_repo or not current_class:
                continue
            cls_cfg = (((get_parsed_config(current_repo, models.db) or {}).get("classes") or {}).get(current_class) or {})
            patch = dict(op.get("data") or {}) if isinstance(op.get("data"), dict) else {}
            base = {"tool": tool, "op_index": op_idx, "class_name": current_class}
            requests_out.extend(_ngenie_unresolved_node_ref_requests_from_data(current_repo, cls_cfg, current_class, patch, lookup, base, f"op{op_idx}"))

        elif tool == "append_table_rows":
            if not current_repo or not current_class:
                continue
            field = str(op.get("field") or op.get("name") or "").strip()
            rows = op.get("rows") if isinstance(op.get("rows"), list) else []
            parent_cfg = (((get_parsed_config(current_repo, models.db) or {}).get("classes") or {}).get(current_class) or {})
            parsed_parent = _ngenie_parse_data_structure(parent_cfg.get("data_structure") or "")
            table_defs = list(parsed_parent.get("tables") or [])
            table_def = next((t for t in table_defs if str(t.get("name")) == str(field)), {})
            if not table_def and len(parsed_parent.get("virtual_tables") or []) == 1 and field:
                table_def = dict((parsed_parent.get("virtual_tables") or [])[0])
                table_def.update({"name": field, "kind": "inline_table", "inline": True, "relation": "inline"})
            is_inline = bool(table_def.get("inline")) or str(table_def.get("kind") or "").strip().lower() == "inline_table"
            if is_inline:
                for row_idx, row in enumerate(rows):
                    row_data = dict(row or {}) if isinstance(row, dict) else {}
                    base = {"tool": tool, "op_index": op_idx, "field": field, "row_index": row_idx, "row_class": ""}
                    requests_out.extend(_ngenie_unresolved_node_ref_requests_from_data(
                        current_repo, parent_cfg, current_class, row_data, lookup, base,
                        f"op{op_idx}_row{row_idx}", field_defs=table_def.get("fields") or []
                    ))
                continue
            row_class = str(table_def.get("row_class") or table_def.get("target") or "").strip()
            row_repo, row_cfg, row_class_name = _ngenie_find_class(current_repo.config_uid, row_class, lookup)
            if not row_repo or not row_cfg:
                continue
            for row_idx, row in enumerate(rows):
                row_data = dict(row or {}) if isinstance(row, dict) else {}
                base = {"tool": tool, "op_index": op_idx, "field": field, "row_index": row_idx, "row_class": row_class_name}
                requests_out.extend(_ngenie_unresolved_node_ref_requests_from_data(row_repo, row_cfg, row_class_name, row_data, lookup, base, f"op{op_idx}_row{row_idx}"))

    # Deduplicate by target class + query + field context to avoid repeated LLM work.
    seen = set()
    unique: List[Dict[str, Any]] = []
    for req in requests_out:
        ctx = req.get("context") if isinstance(req.get("context"), dict) else {}
        key = (req.get("config_uid"), req.get("class_name"), req.get("query"), ctx.get("tool"), ctx.get("op_index"), ctx.get("row_index"), ctx.get("target_field"))
        if key in seen:
            continue
        seen.add(key)
        unique.append(req)
    return unique


def _ngenie_candidate_payload(repo: models.Repo, class_name: str, internal_id: str, score: float = 0.0, request_obj: Optional[Dict[str, Any]] = None, explicit_note: Any = None) -> Optional[Dict[str, Any]]:
    internal_id = str(internal_id or "").strip()
    class_name = str(class_name or "").strip()
    if not internal_id or not class_name:
        return None
    uid = _nodes_mod.normalize_own_uid(repo.config_uid, class_name, _nodes_mod.extract_internal_id(internal_id))
    try:
        obj = _projection_object_payload(repo, "__ngenie_clarify__", uid)
    except Exception:
        obj = None
    if not obj:
        data = _fetch_node_data_for_repo(repo, class_name, _nodes_mod.extract_internal_id(internal_id)) or {}
        obj = {
            "uid": uid,
            "repo_id": repo.id,
            "repo_uid": repo.config_uid,
            "class": class_name,
            "id": _nodes_mod.extract_internal_id(internal_id),
            "data": data if isinstance(data, dict) else {},
            "view": {"title": _ngenie_node_title(data if isinstance(data, dict) else {}) or _nodes_mod.extract_internal_id(internal_id)},
            "cover_html": "",
            "open_url": url_for("client.node_form", config_uid=repo.config_uid, class_name=class_name, node_id=_nodes_mod.extract_internal_id(internal_id)),
        }
    data = obj.get("data") if isinstance(obj.get("data"), dict) else {}
    obj["score"] = round(float(score or 0.0), 4)
    obj["note"] = _ngenie_candidate_note(repo, class_name, obj.get("id") or internal_id, data, request_obj or {}, explicit_note=explicit_note)
    return obj


def _ngenie_find_candidate_nodes(repo: models.Repo, cls_cfg: Dict[str, Any], class_name: str, query_value: Any, max_candidates: int = 8, min_score: float = 0.0, literal_query: Any = "") -> List[Dict[str, Any]]:
    q = str(query_value or "").strip()
    literal = str(literal_query or "").strip()
    search_queries: List[str] = []
    for candidate_query in (literal, q):
        if candidate_query and candidate_query not in search_queries:
            search_queries.append(candidate_query)
    if not search_queries and q:
        search_queries.append(q)
    max_candidates = max(1, min(int(max_candidates or 8), 30))
    min_score = float(min_score or 0.0)
    if q and min_score <= 0.0:
        min_score = 0.25
    indexes = cls_cfg.get("indexes") or cls_cfg.get("indexes_json") or []
    if not isinstance(indexes, list):
        indexes = []
    by_uid: Dict[str, Tuple[float, Dict[str, Any]]] = {}

    def add_row(row: Any, base_score: float = 0.0):
        if not isinstance(row, dict):
            return
        internal_id = _ngenie_get_internal_id_from_row(row)
        data = _ngenie_row_data(row)
        if not internal_id:
            return
        strings = _ngenie_node_search_strings(data, cls_cfg)
        sim = max([_ngenie_search_similarity(search_q, v) for search_q in (search_queries or [q]) for v in strings] or [0.0]) if (q or search_queries) else 0.0
        score = max(float(base_score or 0.0), sim)
        if q and score < min_score:
            return
        uid = _nodes_mod.normalize_own_uid(repo.config_uid, class_name, internal_id)
        old = by_uid.get(uid)
        if not old or score > old[0]:
            by_uid[uid] = (score, {"internal_id": internal_id})

    if q:
        for idx in indexes:
            if not isinstance(idx, dict):
                continue
            idx_name = str(idx.get("name") or "").strip()
            if not idx_name:
                continue
            try:
                idx_kind = str(idx.get("kind") or idx.get("type") or "hash_index").strip().lower()
                index_queries = [q] if idx_kind in {"hash", "hash_index", "exact"} else (search_queries or [q])
                # Hash indexes are exact by nature. Text/trigram/semantic indexes
                # are queried both with the literal user text and the planner's
                # query. This keeps nGenie search identical to Section search and
                # prevents a synonym rewrite (e.g. болгарка -> УШМ) from losing a
                # valid semantic-index result.
                base_score = 1.0 if idx_kind in {"hash", "hash_index", "exact"} else (0.7 if idx_kind in {"semantic", "semantic_index", "semanic_index"} else 0.0)
                for index_query in index_queries:
                    if not str(index_query or "").strip():
                        continue
                    rows = _fetch_nodes_for_class(repo, config_uid=repo.config_uid, class_name=class_name, q="", limit=max_candidates * 2, index_name=idx_name, index_value=str(index_query))
                    for row in rows or []:
                        add_row(row, base_score)
            except Exception:
                continue

        for normal_query in (search_queries or [q]):
            try:
                rows = _fetch_nodes_for_class(repo, config_uid=repo.config_uid, class_name=class_name, q=str(normal_query or ""), limit=max(max_candidates * 4, 20))
                for row in rows or []:
                    add_row(row, 0.0)
            except Exception:
                pass

    try:
        scan_limit = max(1000, max_candidates * 100)
        if _ngenie_local_repo(repo):
            parsed = get_parsed_config(repo, models.db) or {}
            _nodes_mod.set_runtime_context(repo.config_uid, parsed, system_user=_client_runtime_system_user_payload())
            node_cls = _load_server_node_class(repo.config_uid, class_name)
            nodes = node_cls.get_all(repo.config_uid) if node_cls and hasattr(node_cls, "get_all") else []
            for node in list(nodes or [])[:scan_limit]:
                try:
                    internal_id = _nodes_mod.extract_internal_id(getattr(node, "_id", "") or getattr(node, "id", ""))
                    data = node.get_data() if hasattr(node, "get_data") else getattr(node, "_data", {})
                    if not internal_id:
                        continue
                    add_row({"id": internal_id, "_data": data if isinstance(data, dict) else {}}, 0.0)
                except Exception:
                    continue
        else:
            rows = _fetch_nodes_for_class(repo, config_uid=repo.config_uid, class_name=class_name, q="", limit=scan_limit)
            for row in rows or []:
                add_row(row, 0.0)
    except Exception:
        pass

    items = sorted(by_uid.items(), key=lambda kv: kv[1][0], reverse=True)[:max_candidates]
    out: List[Dict[str, Any]] = []
    for uid, (score, meta) in items:
        obj = _ngenie_candidate_payload(repo, class_name, meta.get("internal_id") or uid, score=score)
        if obj:
            out.append(obj)
    return out




def _ngenie_strip_candidate_handler_code(code: Any) -> str:
    """Return Python source for generated nGenie candidate resolver handler."""
    text = str(code or "").strip()
    if not text:
        return ""
    if text.startswith("```"):
        text = re.sub(r"^```(?:python|py)?\s*", "", text, flags=re.I).strip()
        text = re.sub(r"\s*```$", "", text).strip()
    if "def resolve_candidates" not in text:
        body = text
        lines = body.splitlines() or ["return {'items': []}"]
        text = "def resolve_candidates(ctx):\n" + "\n".join("    " + line if line.strip() else "" for line in lines)
    return text.strip()


def _ngenie_validate_candidate_handler_code(src: str) -> None:
    """Small safety gate for LLM-generated candidate search handlers.

    This is not a generic sandbox; it only allows a tiny resolver function that
    can call backend-provided search helpers. Subject-specific matching logic is
    expected here, but file/network/process/import access is not.
    """
    tree = ast.parse(src or "")
    forbidden_nodes = (
        ast.Import, ast.ImportFrom, ast.Global, ast.Nonlocal, ast.Lambda,
        ast.ClassDef, ast.AsyncFunctionDef, ast.With, ast.AsyncWith,
        ast.Raise, ast.Delete, ast.While, ast.Await, ast.Yield,
        ast.YieldFrom,
    )
    forbidden_names = {
        "open", "exec", "eval", "compile", "__import__", "input",
        "globals", "locals", "vars", "dir", "getattr", "setattr", "delattr",
        "os", "sys", "subprocess", "socket", "requests", "pickle", "shutil",
        "pathlib", "Path", "builtins",
        "ConfigClass", "Configuration", "indexes_json",
    }
    allowed_func_names = {
        # Legacy helper names are kept only for backward compatibility.
        # New generated handlers should use the same names as nodes.py:
        # findByIndex/getByIndex.
        "find_by_index", "find", "candidate_limit", "node_title", "scalar",
        # Android/native handler names. On backend these are mapped to the
        # same safe resolver functions, so web nGenie keeps working while
        # mobile code can look like normal nodes.py handlers.
        "findByIndex", "getByIndex", "ngenie_node_payload", "ngenie_node_payloads", "ngenie_node_id",
        "len", "str", "int", "float", "bool", "dict", "list", "set", "tuple",
        "enumerate", "range", "min", "max", "sum", "any", "all", "sorted",
        "lower", "upper",
    }
    allowed_attr_calls = {
        "append", "extend", "insert", "get", "setdefault", "pop",
        "lower", "upper", "strip", "lstrip", "rstrip", "replace", "split",
        "join", "startswith", "endswith", "isdigit", "isalnum", "isalpha",
        "format", "items", "keys", "values", "sort",
    }
    funcs = [n for n in tree.body if isinstance(n, ast.FunctionDef)]
    if len(funcs) != 1 or funcs[0].name != "resolve_candidates":
        raise ValueError("candidate_handler_code must define exactly def resolve_candidates(ctx)")
    for node in ast.walk(tree):
        if isinstance(node, forbidden_nodes):
            raise ValueError(f"Forbidden syntax in candidate handler: {type(node).__name__}")
        if isinstance(node, ast.Name) and node.id in forbidden_names:
            raise ValueError(f"Forbidden name in candidate handler: {node.id}")
        if isinstance(node, ast.Attribute):
            attr = str(node.attr or "")
            if attr.startswith("__"):
                raise ValueError("Dunder attributes are forbidden in candidate handler")
            if attr in {"_indexes", "indexes", "indexes_json", "_get_defined_indexes", "ConfigClass", "Configuration"}:
                raise ValueError(f"Forbidden config/index attribute in candidate handler: {attr}")
        if isinstance(node, ast.Call):
            fn = node.func
            if isinstance(fn, ast.Name):
                if fn.id in forbidden_names:
                    raise ValueError(f"Forbidden call in candidate handler: {fn.id}")
                # Allow helper calls, safe builtins and re.* via Attribute below.
                if fn.id not in allowed_func_names:
                    # User code can call local functions only if we later decide to allow them.
                    # For now keep handlers simple and predictable.
                    raise ValueError(f"Unsupported function call in candidate handler: {fn.id}")
            elif isinstance(fn, ast.Attribute):
                if str(fn.attr or "").startswith("__"):
                    raise ValueError("Dunder calls are forbidden in candidate handler")
                if isinstance(fn.value, ast.Name) and fn.value.id == "re":
                    continue
                if fn.attr not in allowed_attr_calls:
                    raise ValueError(f"Unsupported method call in candidate handler: {fn.attr}")
            else:
                raise ValueError("Unsupported call expression in candidate handler")



def _ngenie_repair_invalid_candidate_handler(
    messages: Any,
    answer: Any,
    user_message: str = "",
    max_attempts: int = 2,
) -> Dict[str, Any]:
    """Repair an unsafe/invalid LLM candidate resolver before execution.

    Candidate handlers are deliberately tiny and cannot import modules.  A model
    sometimes puts random document creation (``import random``/``datetime``)
    into this first search stage.  Previously the AST validation error leaked to
    the user.  Feed the concrete validation error back to the model and require
    either a pure lookup resolver or ``data_requests`` for arbitrary/full sets.
    """
    current = dict(answer) if isinstance(answer, dict) else {}
    attempts = max(1, min(int(max_attempts or 1), 3))
    for _attempt in range(attempts):
        code = (
            current.get("candidate_handler_code")
            or current.get("candidateHandlerCode")
            or current.get("resolve_handler_code")
            or current.get("resolveHandlerCode")
            or ""
        )
        if not str(code or "").strip():
            return current
        try:
            _ngenie_validate_candidate_handler_code(_ngenie_strip_candidate_handler_code(code))
            return current
        except Exception as exc:
            validation_error = str(exc or "Invalid candidate handler")

        repair_messages = list(messages or []) + [
            {"role": "assistant", "content": json.dumps(current, ensure_ascii=False, default=str)},
            {
                "role": "user",
                "content": (
                    "Исправь полный JSON-план для исходного запроса. Backend отклонил "
                    "candidate_handler_code: " + validation_error + ". "
                    "candidate_handler_code предназначен ТОЛЬКО для безопасного поиска существующих "
                    "ссылочных узлов через findByIndex/getByIndex/find и не может содержать import/from, "
                    "random, datetime, uuid, создание или изменение узлов. Не переноси создание заказов "
                    "в candidate handler. Если нужны произвольные существующие клиенты/товары или полный "
                    "набор класса, верни data_requests; backend прочитает доступные rows и вызовет второй "
                    "шаг, где нужно сразу создать/изменить данные. Не спрашивай разрешение на get_all/findAll. "
                    "Сохрани все требования исходного запроса: количество документов, число строк, даты, "
                    "цены и суммы. Верни только полный исправленный JSON."
                ),
            },
        ]
        repaired = _ngenie_call_deepseek(repair_messages)
        if not isinstance(repaired, dict) or not repaired:
            return current
        # Preserve diagnostics/skill routing used by the caller.
        if current.get("_ngenie_selected_skill_ids") and not repaired.get("_ngenie_selected_skill_ids"):
            repaired["_ngenie_selected_skill_ids"] = current.get("_ngenie_selected_skill_ids")
        current = repaired
    return current


def _ngenie_candidate_rows_to_payloads(repo: models.Repo, class_name: str, rows: Any, request_obj: Optional[Dict[str, Any]] = None, score: float = 0.0, limit: Optional[int] = None) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    lim = max(1, min(int(limit or _ngenie_resolve_candidate_limit()), 50))
    for row in rows or []:
        if not isinstance(row, dict):
            continue
        internal_id = _ngenie_get_internal_id_from_row(row)
        if not internal_id:
            continue
        obj = _ngenie_candidate_payload(repo, class_name, internal_id, score=score, request_obj=request_obj or {})
        if obj:
            out.append(obj)
        if len(out) >= lim:
            break
    return out


def _ngenie_execute_candidate_handler(
    code: Any,
    user_message: str,
    config_uid: str,
    node_context: Optional[Dict[str, Any]],
    allow_catalog_create: bool,
    scope: str,
    initial_answer: Optional[Dict[str, Any]],
    lookup: Dict[Tuple[str, str], Tuple[models.Repo, Dict[str, Any]]],
    attachments: Any = None,
) -> Dict[str, Any]:
    src = _ngenie_strip_candidate_handler_code(code)
    if not src:
        return {"items": []}
    _ngenie_validate_candidate_handler_code(src)

    limit_default = _ngenie_resolve_candidate_limit()

    def _limit(value: Any = None) -> int:
        try:
            n = int(value if value is not None else limit_default)
        except Exception:
            n = limit_default
        return max(1, min(n, limit_default))

    def _resolve_class(class_name: Any, config_uid_arg: Any = "") -> Tuple[models.Repo, Dict[str, Any], str]:
        cls = str(class_name or "").strip()
        cu = str(config_uid_arg or config_uid or "").strip()
        repo, cls_cfg, cls_name = _ngenie_find_class(cu, cls, lookup)
        if not repo or not cls_cfg:
            raise ValueError(f"Class not found for nGenie handler: {cls}")
        return repo, cls_cfg, cls_name

    def _find_index(cls_cfg: Dict[str, Any], index_name: str) -> Optional[Dict[str, Any]]:
        """Find a declared index for generated nGenie handlers.

        LLM prompts sometimes mention old virtual helper names such as
        llm_name_trigram while the real saved index is llm_name. A missing
        index must not crash the whole assistant request: the handler can then
        continue to the next fallback search branch. If the requested name is a
        suffix alias of a real index, use the real index name.
        """
        idx_name = str(index_name or "").strip()
        indexes = cls_cfg.get("indexes") or cls_cfg.get("indexes_json") or []
        if not isinstance(indexes, list):
            indexes = []
        by_name: Dict[str, Dict[str, Any]] = {}
        for idx in indexes:
            if not isinstance(idx, dict):
                continue
            name = str(idx.get("name") or "").strip()
            if name:
                by_name[name] = idx
        if idx_name in by_name:
            return by_name[idx_name]

        # Soft compatibility with prompts that add implementation suffixes to a
        # business index name. Example: prompt says llm_name_trigram, config has
        # one semantic_index named llm_name.
        for suffix in ("_trigram", "_text", "_fulltext", "_semantic"):
            if idx_name.endswith(suffix):
                base = idx_name[:-len(suffix)]
                if base in by_name:
                    return by_name[base]

        # Exact aliases are intentionally stricter: use only if a true exact/hash
        # index with the base name exists, otherwise return None and let the
        # generated handler fall through to the next search method.
        for suffix in ("_exact", "_hash"):
            if idx_name.endswith(suffix):
                base = idx_name[:-len(suffix)]
                idx = by_name.get(base)
                kind = str((idx or {}).get("kind") or (idx or {}).get("type") or "").strip().lower()
                if idx and kind in {"hash", "hash_index", "exact"}:
                    return idx
                return None
        return None

    def find_by_index(class_name: Any, index_name: Any, value: Any, config_uid: Any = "", limit: Any = None):  # noqa: A002 - helper API
        repo, cls_cfg, cls_name = _resolve_class(class_name, config_uid)
        idx = _find_index(cls_cfg, str(index_name or ""))
        if not idx:
            return []
        lim = _limit(limit)
        rows = _fetch_nodes_for_class(
            repo,
            config_uid=repo.config_uid,
            class_name=cls_name,
            q="",
            limit=lim,
            index_name=str(idx.get("name") or index_name or "").strip(),
            index_value=str(value or ""),
        )
        kind = str(idx.get("kind") or idx.get("type") or "hash_index").strip().lower()
        score = 1.0 if kind in {"hash", "hash_index", "exact"} else (0.7 if kind in {"semantic", "semantic_index", "semanic_index"} else 0.0)
        return _ngenie_candidate_rows_to_payloads(repo, cls_name, rows, score=score, limit=lim)

    def find(class_name: Any, query: Any, config_uid: Any = "", limit: Any = None):  # noqa: A001 - helper API
        repo, _cls_cfg, cls_name = _resolve_class(class_name, config_uid)
        lim = _limit(limit)
        rows = _fetch_nodes_for_class(
            repo,
            config_uid=repo.config_uid,
            class_name=cls_name,
            q=str(query or ""),
            limit=lim,
        )
        payloads = _ngenie_candidate_rows_to_payloads(repo, cls_name, rows, score=0.0, limit=lim)
        # Keep the same native contract as findByIndex: generated handlers work
        # with Node-like objects and serialize them only through ngenie_node_payload(s).
        return [_CandidateNode(p) for p in payloads]

    def candidate_limit():
        return limit_default

    def node_title(candidate: Any) -> str:
        if not isinstance(candidate, dict):
            return ""
        view = candidate.get("view") if isinstance(candidate.get("view"), dict) else {}
        return str(view.get("title") or _ngenie_node_title(candidate.get("data") if isinstance(candidate.get("data"), dict) else {}) or "")

    def scalar(candidate: Any, field: Any, default: Any = "") -> Any:
        if isinstance(candidate, _CandidateNode):
            return candidate._data.get(str(field or ""), default)
        if not isinstance(candidate, dict):
            return default
        data = candidate.get("data") if isinstance(candidate.get("data"), dict) else {}
        return data.get(str(field or ""), default)

    class _CandidateNode:
        def __init__(self, payload: Dict[str, Any]):
            self._payload = payload if isinstance(payload, dict) else {}
            self._data = self._payload.get("data") if isinstance(self._payload.get("data"), dict) else {}
            self.uid = self._payload.get("uid") or self._payload.get("id") or self._data.get("_id") or ""
            if self.uid and not self._data.get("_id"):
                self._data["_id"] = self.uid

    def _payload_from_candidate_node(node: Any) -> Dict[str, Any]:
        if isinstance(node, _CandidateNode):
            return dict(node._payload)
        if isinstance(node, dict):
            return dict(node)
        return {}

    def findByIndex(class_name: Any, index_name: Any, value: Any):
        return [_CandidateNode(p) for p in find_by_index(class_name, index_name, value)]

    def getByIndex(class_name: Any, index_name: Any, value: Any):
        rows = findByIndex(class_name, index_name, value)
        return rows[0] if rows else None

    def ngenie_node_id(node: Any) -> str:
        if isinstance(node, _CandidateNode):
            return str(node._data.get("_id") or node.uid or "")
        if isinstance(node, dict):
            data = node.get("data") if isinstance(node.get("data"), dict) else node
            return str(data.get("_id") or node.get("uid") or node.get("id") or "")
        return ""

    def ngenie_node_payload(node: Any, score: float = 0.0, note: Any = None) -> Dict[str, Any]:
        payload = _payload_from_candidate_node(node)
        if not payload:
            return {}
        payload.setdefault("id", payload.get("uid") or payload.get("node_id") or payload.get("_id") or payload.get("id"))
        payload.setdefault("uid", payload.get("id"))
        if note is not None:
            payload["note"] = str(note)
        return payload

    def ngenie_node_payloads(nodes: Any, limit: Any = None, score: float = 0.0) -> List[Dict[str, Any]]:
        try:
            lim = _limit(limit)
        except Exception:
            lim = limit_default
        out = []
        for node in nodes or []:
            p = ngenie_node_payload(node, score=score)
            if p:
                out.append(p)
            if len(out) >= lim:
                break
        return out

    safe_builtins = {
        "len": len, "str": str, "int": int, "float": float, "bool": bool,
        "dict": dict, "list": list, "set": set, "tuple": tuple,
        "Exception": Exception,
        "enumerate": enumerate, "range": range, "min": min, "max": max,
        "sum": sum, "any": any, "all": all, "sorted": sorted,
    }
    ns: Dict[str, Any] = {
        "__builtins__": safe_builtins,
        "re": re,
        "find_by_index": find_by_index,
        "find": find,
        "findByIndex": findByIndex,
        "getByIndex": getByIndex,
        "ngenie_node_id": ngenie_node_id,
        "ngenie_node_payload": ngenie_node_payload,
        "ngenie_node_payloads": ngenie_node_payloads,
        "candidate_limit": candidate_limit,
        "node_title": node_title,
        "scalar": scalar,
    }
    compiled = compile(src, "<ngenie_candidate_handler>", "exec")
    exec(compiled, ns, ns)
    fn = ns.get("resolve_candidates")
    if not callable(fn):
        raise ValueError("resolve_candidates is not callable")
    ctx = {
        "message": str(user_message or ""),
        "selected_config_uid": config_uid or "",
        "allow_catalog_create": bool(allow_catalog_create),
        "scope": str(scope or ""),
        "current_node": node_context or {},
        "attachments": _ngenie_normalize_attachments(attachments),
        "initial_answer": initial_answer or {},
        "candidate_limit": limit_default,
    }
    result = fn(ctx)
    if result is None:
        result = {"items": []}
    if isinstance(result, list):
        result = {"items": result}
    if not isinstance(result, dict):
        raise ValueError("candidate handler must return dict or list")
    result.setdefault("items", [])
    result["handler_code"] = src if bool((initial_answer or {}).get("debug_handler")) else ""
    return result


def _ngenie_handler_item_status(item: Dict[str, Any]) -> str:
    return str(item.get("status") or item.get("state") or "").strip().lower()


def _ngenie_handler_blocking_messages(result: Dict[str, Any]) -> List[str]:
    messages: List[str] = []
    for item in result.get("items") or []:
        if not isinstance(item, dict):
            continue
        status = _ngenie_handler_item_status(item)
        if status in {"not_found", "notfound", "fail", "failed", "error"}:
            msg = str(item.get("message") or item.get("reply") or item.get("reason") or "Объект не найден").strip()
            if msg:
                messages.append(msg)
    return messages


def _ngenie_handler_ambiguities_to_clarifications(result: Dict[str, Any], config_uid: str, lookup: Dict[Tuple[str, str], Tuple[models.Repo, Dict[str, Any]]]) -> List[Dict[str, Any]]:
    raw: List[Dict[str, Any]] = []
    for idx, item in enumerate(result.get("items") or [], start=1):
        if not isinstance(item, dict):
            continue
        status = _ngenie_handler_item_status(item)
        if status not in {"ambiguous", "clarify", "clarification", "select"}:
            continue
        cls = str(item.get("class_name") or item.get("class") or "").strip()
        cu = str(item.get("config_uid") or config_uid or "").strip()
        candidates = item.get("candidates") if isinstance(item.get("candidates"), list) else []
        raw.append({
            "id": str(item.get("id") or f"handler_clarify_{idx}"),
            "question": str(item.get("question") or item.get("title") or "Уточните вариант"),
            "reason": str(item.get("reason") or item.get("message") or ""),
            "class_name": cls,
            "config_uid": cu,
            "query": str(item.get("query") or ""),
            "candidates": candidates,
            "context": item.get("context") if isinstance(item.get("context"), dict) else {},
            "note_fields": item.get("note_fields") or item.get("noteFields") or [],
            "note_template": item.get("note_template") or item.get("noteTemplate") or "",
            "note_method": item.get("note_method") or item.get("noteMethod") or "",
            "note_label": item.get("note_label") or item.get("noteLabel") or "",
        })
    return _ngenie_prepare_clarifications(raw, config_uid, lookup) if raw else []


def _ngenie_compact_handler_result_for_llm(result: Dict[str, Any]) -> Dict[str, Any]:
    out_items: List[Dict[str, Any]] = []
    for item in result.get("items") or []:
        if not isinstance(item, dict):
            continue
        x = dict(item)
        cands = x.get("candidates") if isinstance(x.get("candidates"), list) else []
        if cands:
            x["candidates"] = [_ngenie_compact_candidate_for_llm(c) for c in cands if isinstance(c, dict)]
        out_items.append(x)
    out = dict(result)
    out["items"] = out_items
    out.pop("handler_code", None)
    return out


def _ngenie_build_candidate_handler_result_messages(
    user_message: str,
    config_uid: str,
    node_context: Optional[Dict[str, Any]],
    allow_catalog_create: bool,
    scope: str,
    initial_answer: Dict[str, Any],
    handler_result: Dict[str, Any],
    attachments: Any = None,
) -> List[Dict[str, str]]:
    classes, _lookup = _ngenie_collect_context(config_uid, include_samples=False)
    ctx = {
        "selected_config_uid": config_uid or "",
        "allow_catalog_create": bool(allow_catalog_create),
        "scope": str(scope or "").strip(),
        "configuration_prompts": _ngenie_collect_config_prompts(config_uid),
        "classes": classes,
        "current_node": node_context or {},
        "attachments": _ngenie_normalize_attachments(attachments),
        "candidate_limit": _ngenie_resolve_candidate_limit(),
        "initial_answer": initial_answer or {},
        "candidate_handler_result": _ngenie_compact_handler_result_for_llm(handler_result or {}),
    }
    selected_ids, skill_reason = _ngenie_select_skill_ids(user_message, ctx, force_skill_ids=(initial_answer or {}).get("_ngenie_selected_skill_ids"))
    try:
        g.ngenie_selected_skill_ids = selected_ids
    except Exception:
        pass
    ctx["selected_skill_ids"] = selected_ids
    ctx["selected_skill_reason"] = skill_reason
    ctx["selected_skills"] = _ngenie_skill_blocks_for_messages(selected_ids, ctx)
    instruction = """
Это второй шаг nGenie после выполнения candidate_handler_code.
Backend уже выполнил обработчик поиска кандидатов, который ты написал на первом шаге.

Правила ответа:
1. Используй items со status=resolved как уже выбранные UID.
2. Если status=review или status=ambiguous и есть candidates, сначала применяй ngenie_prompt класса, исходный запрос и поля кандидатов: отфильтруй нерелевантные варианты, выбери UID при уверенном совпадении либо верни clarification_requests только с релевантными кандидатами.
3. Не выполняй новый поиск и не возвращай candidate_handler_code/resolve_requests повторно.
4. Верни итоговое действие для исходной задачи. Если это mobile-финализация, предпочтителен один operation_handler_code; для backend/web допустимы operations. Для append_table_rows подставь выбранные UID в Node-поля строк.
5. Если scope=node_form, не создавай HTML-отчёты/проекции.
Отвечай строго JSON-объектом по основной схеме nGenie.
""".strip()
    return [
        {"role": "system", "content": _ngenie_system_prompt()},
        {"role": "user", "content": "Контекст NodaLogic и результат candidate handler:\n" + json.dumps(ctx, ensure_ascii=False, default=str)},
        {"role": "user", "content": instruction + "\n\nИсходная просьба пользователя:\n" + str(user_message or "")},
    ]


def _ngenie_finalize_candidate_handler_with_llm(
    user_message: str,
    config_uid: str,
    node_context: Optional[Dict[str, Any]],
    allow_catalog_create: bool,
    scope: str,
    initial_answer: Dict[str, Any],
    handler_result: Dict[str, Any],
    attachments: Any = None,
) -> Optional[Dict[str, Any]]:
    try:
        answer = _ngenie_call_deepseek(_ngenie_build_candidate_handler_result_messages(
            user_message,
            config_uid,
            node_context,
            allow_catalog_create,
            scope,
            initial_answer,
            handler_result,
            attachments=attachments,
        ))
        if isinstance(answer, dict) and answer:
            answer.setdefault("_candidate_handler_result", _ngenie_compact_handler_result_for_llm(handler_result or {}))
            return answer
    except Exception:
        traceback.print_exc()
    return None

def _ngenie_prepare_clarifications(raw_requests: Any, config_uid: str, lookup: Dict[Tuple[str, str], Tuple[models.Repo, Dict[str, Any]]], literal_query: Any = "") -> List[Dict[str, Any]]:
    requests_in = raw_requests if isinstance(raw_requests, list) else []
    prepared: List[Dict[str, Any]] = []
    for idx, req in enumerate(requests_in, start=1):
        if not isinstance(req, dict):
            continue
        cls = str(req.get("class_name") or req.get("class") or req.get("target_class") or req.get("targetClass") or "").strip()
        cu = str(req.get("config_uid") or req.get("configUid") or config_uid or "").strip()
        repo, cls_cfg, cls_name = _ngenie_find_class(cu, cls, lookup) if cls else (None, None, cls)
        query = str(req.get("query") or req.get("search") or req.get("value") or "").strip()
        max_candidates = req.get("max_candidates") or req.get("maxCandidates") or 8
        min_score = req.get("min_score") or req.get("minScore") or 0.0

        explicit_notes: Dict[str, Any] = {}
        candidates: List[Dict[str, Any]] = []
        raw_candidates = req.get("candidates") if isinstance(req.get("candidates"), list) else []
        if repo and cls_cfg and raw_candidates:
            for c in raw_candidates:
                if not isinstance(c, dict):
                    continue
                uid = str(c.get("uid") or c.get("node_uid") or c.get("id") or "").strip()
                note = c.get("note") or c.get("explanation") or c.get("reason")
                if note:
                    explicit_notes[uid] = note
                if not uid:
                    continue
                try:
                    ccu, ccls, internal_id = _nodes_mod.parse_uid_any(uid)
                    cand_repo = _repo_for_config_uid(repo, ccu or repo.config_uid) if ccu else repo
                    cand_cls = ccls or cls_name
                except Exception:
                    cand_repo = repo
                    cand_cls = cls_name
                    internal_id = uid
                obj = _ngenie_candidate_payload(cand_repo, cand_cls, internal_id, score=float(c.get("score") or 0.0), request_obj=req, explicit_note=note)
                if obj:
                    candidates.append(obj)
        if repo and cls_cfg and not candidates:
            try:
                candidates = _ngenie_find_candidate_nodes(repo, cls_cfg, cls_name, query, max_candidates=max_candidates, min_score=float(min_score or 0.0), literal_query=literal_query)
            except Exception:
                candidates = []
        # Re-apply request-specific notes/templates after candidate search.
        if repo and cls_cfg:
            enriched = []
            for cand in candidates:
                uid = str(cand.get("uid") or "")
                explicit_note = explicit_notes.get(uid) or explicit_notes.get(str(cand.get("id") or ""))
                cand = dict(cand)
                cand["note"] = _ngenie_candidate_note(
                    _repo_for_config_uid(repo, cand.get("repo_uid") or repo.config_uid),
                    cand.get("class") or cls_name,
                    cand.get("id") or "",
                    cand.get("data") if isinstance(cand.get("data"), dict) else {},
                    req,
                    explicit_note=explicit_note,
                )
                enriched.append(cand)
            candidates = enriched
        req_id = str(req.get("id") or req.get("key") or f"clarify_{idx}").strip()
        prepared.append({
            "id": req_id,
            "question": str(req.get("question") or req.get("title") or "Уточните вариант").strip(),
            "reason": str(req.get("reason") or "").strip(),
            "config_uid": (repo.config_uid if repo else cu),
            "class_name": cls_name or cls,
            "query": query,
            "context": req.get("context") if isinstance(req.get("context"), dict) else {},
            "required": bool(req.get("required", True)),
            "display_only": bool(req.get("display_only") or req.get("displayOnly") or req.get("mode") == "display"),
            "candidates": candidates,
        })
    return prepared

def _ngenie_local_repo(repo: models.Repo) -> bool:
    base_url = (repo.base_url or "").strip().rstrip("/")
    current = (request.host_url or "").rstrip("/")
    return not base_url or base_url == current


def _ngenie_pick_name_field(cls_cfg: Dict[str, Any]) -> str:
    parsed = _ngenie_parse_data_structure(cls_cfg.get("data_structure") or "")
    names = [f.get("name") for f in (parsed.get("fields") or [])]
    for candidate in ("name", "title", "caption", "display_name", "article"):
        if candidate in names:
            return candidate
    return str(names[0] or "name") if names else "name"


def _ngenie_node_ref_view(repo: models.Repo, node_uid: Any, default: str = "") -> str:
    """Return human-readable view for a Node reference UID.

    nGenie often writes linked fields programmatically. The normal node_form UI
    fills <field>_view lazily, but table rows need this value immediately after
    creation; otherwise the parent table shows an empty cell until the child row
    is opened and saved by hand.
    """
    raw = str(node_uid or "").strip()
    if not raw:
        return str(default or "")
    try:
        uid_cfg, cls_name, internal_id = _nodes_mod.parse_uid_any(raw)
        cls_name = str(cls_name or "").strip()
        internal_id = str(internal_id or "").strip()
        if not cls_name or not internal_id:
            return str(default or raw)
        obj_repo = _repo_for_config_uid(repo, uid_cfg or repo.config_uid) or repo
        data = _fetch_node_data_for_repo(obj_repo, cls_name, internal_id) or {}
        parsed = get_parsed_config(obj_repo, models.db) or {}
        view = _render_class_record_view(parsed, cls_name, internal_id, data if isinstance(data, dict) else {})
        return str(view or default or raw)
    except Exception:
        return str(default or raw)


def _ngenie_create_node(repo: models.Repo, class_name: str, data: Dict[str, Any]) -> str:
    payload = dict(data or {})
    payload.setdefault("_class", class_name)
    if _ngenie_local_repo(repo):
        parsed = get_parsed_config(repo, models.db) or {}
        _nodes_mod.set_runtime_context(repo.config_uid, parsed, system_user=_client_runtime_system_user_payload())
        internal_id = _node_local_create(repo.config_uid, class_name, payload)
        return _nodes_mod.normalize_own_uid(repo.config_uid, class_name, internal_id)
    result = _api_post_remote(repo, f"/api/config/{repo.config_uid}/node/{class_name}", json_data=payload)
    if isinstance(result, dict):
        raw_id = result.get("_id") or result.get("id") or result.get("node_id")
        if raw_id:
            try:
                return _nodes_mod.normalize_own_uid(repo.config_uid, class_name, _nodes_mod.extract_internal_id(raw_id))
            except Exception:
                return str(raw_id)
    raw_id = payload.get("_id") or str(uuid.uuid4())
    return _nodes_mod.normalize_own_uid(repo.config_uid, class_name, _nodes_mod.extract_internal_id(raw_id))


def _ngenie_update_node(repo: models.Repo, class_name: str, node_id: str, patch: Dict[str, Any]) -> Dict[str, Any]:
    internal_id = _nodes_mod.extract_internal_id(node_id)
    old = _fetch_node_data_for_repo(repo, class_name, internal_id) or {}
    merged = dict(old if isinstance(old, dict) else {})
    merged.update(dict(patch or {}))
    merged.setdefault("_class", class_name)
    if _ngenie_local_repo(repo):
        parsed = get_parsed_config(repo, models.db) or {}
        _nodes_mod.set_runtime_context(repo.config_uid, parsed, system_user=_client_runtime_system_user_payload())
        node = _node_local_update_data(repo.config_uid, class_name, internal_id, merged)
        try:
            fresh = node.get_data() or {}
            if isinstance(fresh, dict):
                return fresh
        except Exception:
            pass
    else:
        _api_post_remote(repo, f"/api/config/{repo.config_uid}/node/{class_name}/{internal_id}", json_data=merged)
        try:
            fresh = _fetch_node_data_for_repo(repo, class_name, internal_id) or {}
            if isinstance(fresh, dict) and fresh:
                return fresh
        except Exception:
            pass
    return merged


def _ngenie_normalize_role(value: Any) -> str:
    role = str(value or "").strip().lower().replace("-", "_").replace(" ", "_")
    aliases = {
        "ref": "catalog",
        "refs": "catalog",
        "reference": "catalog",
        "directory": "catalog",
        "directories": "catalog",
        "catalogue": "catalog",
        "dict": "catalog",
        "dictionary": "catalog",
        "справочник": "catalog",
        "справочники": "catalog",
        "doc": "document",
        "документ": "document",
        "line": "document_line",
        "row": "document_line",
        "position": "document_line",
        "documentrow": "document_line",
        "document_row": "document_line",
        "documentline": "document_line",
        "строка": "document_line",
        "строка_документа": "document_line",
        "register_record": "register",
        "movement": "register",
        "движение": "register",
        "регистр": "register",
        "projection": "report",
        "html_projection": "report",
        "отчет": "report",
        "отчёт": "report",
        "technical": "service",
        "system": "service",
        "технический": "service",
    }
    return aliases.get(role, role)


def _ngenie_class_role(cls_cfg: Optional[Dict[str, Any]]) -> str:
    if not isinstance(cls_cfg, dict):
        return ""
    return _ngenie_normalize_role(
        cls_cfg.get("ngenie_role")
        or cls_cfg.get("ngenieRole")
        or cls_cfg.get("nGenieRole")
        or cls_cfg.get("nGenie_role")
        or ""
    )


def _ngenie_role_label(role: str) -> str:
    role = _ngenie_normalize_role(role)
    return {
        "catalog": "справочник",
        "document": "документ",
        "document_line": "строка документа",
        "register": "регистр/движение",
        "report": "отчет/проекция",
        "service": "технический класс",
    }.get(role, "узел")


def _ngenie_all_row_classes(lookup: Dict[Tuple[str, str], Tuple[models.Repo, Dict[str, Any]]]) -> set:
    rows = set()
    for (_cu, _cn), (_repo, cfg) in (lookup or {}).items():
        try:
            parsed = _ngenie_parse_data_structure(cfg.get("data_structure") or "")
            for t in parsed.get("tables") or []:
                rc = str(t.get("row_class") or t.get("target") or "").strip()
                if rc:
                    rows.add(rc.lower())
        except Exception:
            continue
    return rows


def _ngenie_is_row_class(class_name: str, cls_cfg: Optional[Dict[str, Any]], lookup: Dict[Tuple[str, str], Tuple[models.Repo, Dict[str, Any]]]) -> bool:
    role = _ngenie_class_role(cls_cfg)
    if role == "document_line":
        return True
    if role:
        return False
    low = str(class_name or "").strip().lower()
    if low and low in _ngenie_all_row_classes(lookup):
        return True
    display = str((cls_cfg or {}).get("display_name") or "").lower()
    text = low + " " + display
    return any(x in text for x in ("line", "row", "position", "строк", "позици"))


def _ngenie_is_catalog_like_class(class_name: str, cls_cfg: Optional[Dict[str, Any]]) -> bool:
    role = _ngenie_class_role(cls_cfg)
    if role == "catalog":
        return True
    if role:
        return False
    low = str(class_name or "").strip().lower()
    display = str((cls_cfg or {}).get("display_name") or "").lower()
    text = low + " " + display
    keywords = (
        "goods", "good", "product", "item", "sku", "catalog", "directory", "nomenclature",
        "counterparty", "contractor", "customer", "supplier", "warehouse", "stock", "partner",
        "товар", "номенклатур", "материал", "справочник", "контрагент", "клиент", "поставщик", "склад",
    )
    return any(k in text for k in keywords)


def _ngenie_may_create_class(class_name: str, cls_cfg: Optional[Dict[str, Any]], lookup: Dict[Tuple[str, str], Tuple[models.Repo, Dict[str, Any]]], allow_catalog_create: bool) -> bool:
    role = _ngenie_class_role(cls_cfg)
    if role == "catalog":
        return bool(allow_catalog_create)
    if role in {"service", "system", "report", "projection"}:
        return False
    if role in {"document", "document_line", "register"}:
        return True
    if bool(allow_catalog_create):
        return True
    if _ngenie_is_row_class(class_name, cls_cfg, lookup):
        return True
    if _ngenie_is_catalog_like_class(class_name, cls_cfg):
        return False
    return True


def _ngenie_node_ref_spec(value: Any) -> Tuple[str, Any, Dict[str, Any]]:
    """Normalize a Node field value produced by nGenie.

    Besides an already resolved UID or a plain search string, the web runtime
    accepts a structured value:
      {"query":"ADI.1015.T.25", "create_data":{"name":"...", "article":"...", "barcode":"..."}}

    This is important for imports: the backend can find an existing catalog item
    or create it with all supplied attributes and immediately put its UID into the
    document row in one operation.
    """
    if not isinstance(value, dict):
        return "", value, {}

    uid = str(
        value.get("uid")
        or value.get("node_uid")
        or value.get("nodeUid")
        or value.get("_id")
        or ""
    ).strip()
    create_data_raw = value.get("create_data")
    if not isinstance(create_data_raw, dict):
        create_data_raw = value.get("createData")
    if not isinstance(create_data_raw, dict):
        create_data_raw = value.get("data")
    create_data = dict(create_data_raw or {}) if isinstance(create_data_raw, dict) else {}

    query = value.get("query")
    if query is None or str(query).strip() == "":
        query = value.get("value")
    if query is None or str(query).strip() == "":
        for key in ("article", "code", "barcode", "name", "title"):
            if key in create_data and str(create_data.get(key) or "").strip():
                query = create_data.get(key)
                break
    if query is None or str(query).strip() == "":
        query = value.get("name") or value.get("title") or ""
    return uid, query, create_data


def _ngenie_resolve_node_reference(
    repo: models.Repo,
    fld: Dict[str, Any],
    value: Any,
    lookup: Dict[Tuple[str, str], Tuple[models.Repo, Dict[str, Any]]],
    allow_catalog_create: bool = False,
) -> Tuple[bool, Any, str, List[str]]:
    """Resolve one Node(...) field.

    Returns (handled, resolved_value, resolved_view, notes).  resolved_value=None
    means that the field must be removed because no target was found and catalog
    creation is disabled.
    """
    target = str((fld or {}).get("target") or "").strip()
    target_repo, target_cfg, target_name = _ngenie_find_class(repo.config_uid, target, lookup)
    if not target_repo or not target_cfg:
        return False, value, "", []

    uid, query, create_data = _ngenie_node_ref_spec(value)
    if uid and _ngenie_is_resolved_node_value(uid):
        return True, uid, _ngenie_node_ref_view(target_repo, uid, default=query or uid), []
    if isinstance(value, str) and _ngenie_is_resolved_node_value(value):
        return True, value, _ngenie_node_ref_view(target_repo, value, default=value), []

    query_text = str(query or "").strip()
    if not query_text:
        return True, value, "", []

    found = _ngenie_find_existing_node(target_repo, target_cfg, target_name, query_text)
    if found:
        return True, found, _ngenie_node_ref_view(target_repo, found, default=query_text), []

    if _ngenie_may_create_class(target_name, target_cfg, lookup, allow_catalog_create):
        data_to_create = {
            str(k): v for k, v in dict(create_data or {}).items()
            if str(k) not in {"_id", "_class", "_created_date", "_last_change_date"}
        }
        name_field = _ngenie_pick_name_field(target_cfg)
        if not str(data_to_create.get(name_field) or "").strip():
            data_to_create[name_field] = query_text
        created_uid = _ngenie_create_node(target_repo, target_name, data_to_create)
        view = _ngenie_node_ref_view(target_repo, created_uid, default=query_text)
        return True, created_uid, view, [f"Создан {_ngenie_role_label(_ngenie_class_role(target_cfg))} {target_name}: {view or query_text}"]

    return True, None, "", [f"Не найден справочник {target_name}: {query_text}. Создание справочников выключено."]


def _ngenie_resolve_inline_row_refs(
    repo: models.Repo,
    table_def: Dict[str, Any],
    row_data: Dict[str, Any],
    lookup: Dict[Tuple[str, str], Tuple[models.Repo, Dict[str, Any]]],
    allow_catalog_create: bool = False,
) -> Tuple[Dict[str, Any], List[str]]:
    """Resolve Node fields declared inside a named inline table such as
    lines:[Product|product: Node("Goods"), Quantity|qty:number].
    """
    out = dict(row_data or {})
    notes: List[str] = []
    for fld in (table_def or {}).get("fields") or []:
        if not isinstance(fld, dict) or fld.get("kind") != "node":
            continue
        fname = str(fld.get("name") or "").strip()
        if not fname or fname not in out:
            continue
        val = out.get(fname)
        if val is None or (not isinstance(val, dict) and str(val).strip() == ""):
            continue
        handled, resolved, view, ref_notes = _ngenie_resolve_node_reference(
            repo, fld, val, lookup, allow_catalog_create=allow_catalog_create
        )
        notes.extend(ref_notes)
        if not handled:
            continue
        view_key = f"{fname}_view"
        if resolved is None:
            out.pop(fname, None)
            out.pop(view_key, None)
        else:
            out[fname] = resolved
            if view and not str(out.get(view_key) or "").strip():
                out[view_key] = view
    return out, notes


def _ngenie_resolve_refs(repo: models.Repo, cls_cfg: Dict[str, Any], data: Dict[str, Any], lookup: Dict[Tuple[str, str], Tuple[models.Repo, Dict[str, Any]]], allow_catalog_create: bool = False) -> Tuple[Dict[str, Any], List[Dict[str, Any]], List[str]]:
    parsed = _ngenie_parse_data_structure(cls_cfg.get("data_structure") or "")
    out = dict(data or {})
    table_rows: List[Dict[str, Any]] = []
    notes: List[str] = []
    table_defs = list(parsed.get("tables") or []) + _ngenie_legacy_inline_table_bindings(parsed, out)
    for t in table_defs:
        name = t.get("name")
        if name in out and isinstance(out.get(name), list):
            table_rows.append({
                "field": name,
                "row_class": t.get("row_class") or t.get("target") or "",
                "rows": out.pop(name),
                "table_def": t,
            })
    for fld in parsed.get("fields") or []:
        fname = str(fld.get("name") or "").strip()
        if fld.get("kind") != "node" or not fname or fname not in out:
            continue
        val = out.get(fname)
        if val is None or (not isinstance(val, dict) and str(val).strip() == ""):
            continue
        handled, resolved, view, ref_notes = _ngenie_resolve_node_reference(
            repo, fld, val, lookup, allow_catalog_create=allow_catalog_create
        )
        notes.extend(ref_notes)
        if not handled:
            continue
        view_key = f"{fname}_view"
        if resolved is None:
            out.pop(fname, None)
            out.pop(view_key, None)
        else:
            out[fname] = resolved
            if view and not str(out.get(view_key) or "").strip():
                out[view_key] = view
    return out, table_rows, notes



def _ngenie_append_table_rows(repo: models.Repo, parent_class: str, parent_id: str, parent_data: Dict[str, Any], field_name: str, rows: List[Any], lookup: Dict[Tuple[str, str], Tuple[models.Repo, Dict[str, Any]]], allow_catalog_create: bool = False) -> Tuple[Dict[str, Any], List[str], List[str]]:
    parent_cfg = (((get_parsed_config(repo, models.db) or {}).get("classes") or {}).get(parent_class) or {})
    parsed_parent = _ngenie_parse_data_structure(parent_cfg.get("data_structure") or "")
    table_defs = parsed_parent.get("tables") or []
    table_def = next((t for t in table_defs if str(t.get("name")) == str(field_name)), {})
    if not table_def and len(parsed_parent.get("virtual_tables") or []) == 1 and str(field_name or "").strip():
        table_def = dict((parsed_parent.get("virtual_tables") or [])[0])
        table_def.update({
            "name": str(field_name),
            "label": table_def.get("label") or str(field_name),
            "kind": "inline_table",
            "relation": "inline",
            "inline": True,
        })

    row_class = str(table_def.get("row_class") or table_def.get("target") or "").strip()
    item_type = table_def.get("item_type") or {}
    row_relation = str(table_def.get("relation") or item_type.get("kind") or "").strip().lower()

    row_repo, row_cfg, row_class_name = _ngenie_find_class(repo.config_uid, row_class, lookup)

    parent_uid = _nodes_mod.normalize_own_uid(
        repo.config_uid,
        parent_class,
        _nodes_mod.extract_internal_id(parent_id),
    )

    created: List[str] = []
    notes: List[str] = []
    current = dict(parent_data or {})

    existing = current.get(field_name)
    if isinstance(existing, list):
        existing = list(existing)
    else:
        existing = []

    for raw in rows or []:
        row_data = dict(raw or {}) if isinstance(raw, dict) else {"name": str(raw)}

        is_inline_table = bool(table_def.get("inline")) or str(table_def.get("kind") or "").strip().lower() == "inline_table"
        if is_inline_table:
            row_data, ref_notes = _ngenie_resolve_inline_row_refs(
                repo, table_def, row_data, lookup, allow_catalog_create=allow_catalog_create
            )
            notes.extend(ref_notes)
            existing.append(row_data)
            continue

        if row_repo and row_cfg and row_class_name:
            parent_link_field = _ngenie_find_parent_link_field(row_cfg, parent_class)

            if parent_link_field:
                row_data.setdefault(parent_link_field, parent_uid)

            if row_relation == "childnode":
                row_data.setdefault("_parent", parent_uid)

            if not parent_link_field and _ngenie_has_data_field(row_cfg, "parent"):
                row_data.setdefault("parent", parent_uid)

            row_data, nested, ref_notes = _ngenie_resolve_refs(row_repo, row_cfg, row_data, lookup, allow_catalog_create=allow_catalog_create)
            notes.extend(ref_notes)
            uid = _ngenie_create_node(row_repo, row_class_name, row_data)

            created.append(uid)
            existing.append(uid)

            for nr in nested:
                _, nested_created, nested_notes = _ngenie_append_table_rows(
                    row_repo,
                    row_class_name,
                    uid,
                    row_data,
                    nr.get("field"),
                    nr.get("rows") or [],
                    lookup,
                    allow_catalog_create=allow_catalog_create,
                )
                created.extend(nested_created)
                notes.extend(nested_notes)
        else:
            existing.append(row_data)

    current[field_name] = existing
    current = _ngenie_update_node(repo, parent_class, parent_id, current)

    return current, created, notes



def _ngenie_extra_table_rows_from_operation(op: Dict[str, Any], cls_cfg: Dict[str, Any]) -> List[Dict[str, Any]]:
    parsed = _ngenie_parse_data_structure((cls_cfg or {}).get("data_structure") or "")
    table_defs = parsed.get("tables") or []
    if not table_defs or not isinstance(op, dict):
        return []
    table_names = {str(t.get("name") or "").strip(): t for t in table_defs if str(t.get("name") or "").strip()}
    out: List[Dict[str, Any]] = []
    seen = set()

    def add(field: Any, rows: Any):
        fname = str(field or "").strip()
        if not fname and len(table_defs) == 1:
            fname = str(table_defs[0].get("name") or "").strip()
        if not fname or fname not in table_names or not isinstance(rows, list):
            return
        key = (fname, id(rows))
        if key in seen:
            return
        seen.add(key)
        t = table_names.get(fname) or {}
        out.append({"field": fname, "row_class": t.get("row_class") or t.get("target") or "", "rows": rows})

    for spec in (op.get("table_rows") or op.get("tableRows") or []):
        if isinstance(spec, dict):
            add(spec.get("field") or spec.get("name"), spec.get("rows") or spec.get("items") or [])

    children = op.get("children") or op.get("child_rows") or op.get("childRows")
    if isinstance(children, dict):
        for field, rows in children.items():
            add(field, rows)
    elif isinstance(children, list):
        add(op.get("field") or op.get("table") or "", children)

    # Common LLM shortcut: create_node with top-level rows/lines/items for the
    # only table in the parent class. This is especially useful outside a node
    # form where append_table_rows cannot be used yet because the parent node
    # does not exist before create_node.
    for key in ("rows", "lines", "items", "positions"):
        if key in op and isinstance(op.get(key), list):
            add(key if key in table_names else "", op.get(key))

    return out


def _ngenie_apply_attachments_to_data(cls_cfg: Dict[str, Any], data: Dict[str, Any], attachments: Any, field: Any = "", append: bool = True) -> Tuple[Dict[str, Any], str, int]:
    files = _ngenie_attachment_filenames(attachments)
    out = dict(data or {})
    if not files:
        return out, "", 0
    fname = str(field or "").strip() or _ngenie_first_file_field(cls_cfg, prefer_media=any(str(f).lower().rsplit('.', 1)[-1] in {"png", "jpg", "jpeg", "gif", "webp", "bmp", "svg"} for f in files))
    if not fname:
        return out, "", 0
    cur = out.get(fname)
    if append and isinstance(cur, list):
        next_list = list(cur)
    else:
        next_list = []
    for fn in files:
        if fn and fn not in next_list:
            next_list.append(fn)
    out[fname] = next_list
    return out, fname, len(files)


def _ngenie_execute_operations(raw_ops: Any, config_uid: str = "", node_context: Optional[Dict[str, Any]] = None, allow_catalog_create: bool = False, attachments: Any = None) -> Dict[str, Any]:
    ops = raw_ops if isinstance(raw_ops, list) else []
    _classes, lookup = _ngenie_collect_context(config_uid, include_samples=False)
    notes: List[str] = []
    attachments_list = _ngenie_normalize_attachments(attachments)
    attachment_files = _ngenie_attachment_filenames(attachments_list)
    object_uids: List[str] = []
    changed_current: Optional[Dict[str, Any]] = None
    current_repo: Optional[models.Repo] = None
    current_class = ""
    current_id = ""
    current_data: Dict[str, Any] = {}
    last_created_repo: Optional[models.Repo] = None
    last_created_class = ""
    last_created_id = ""
    last_created_data: Dict[str, Any] = {}
    if isinstance(node_context, dict) and node_context:
        current_class = str(node_context.get("class_name") or "").strip()
        current_id = str(node_context.get("node_id") or node_context.get("id") or "").strip()
        cu = str(node_context.get("config_uid") or config_uid or "").strip()
        current_repo = models.Repo.query.filter_by(user_id=current_user.id, config_uid=cu).first()
        current_data = dict(node_context.get("data") or {}) if isinstance(node_context.get("data"), dict) else {}
    for op in ops:
        if not isinstance(op, dict):
            continue
        tool = str(op.get("tool") or op.get("action") or "").strip()
        if tool in {"", "none"}:
            continue
        if tool == "create_node":
            cu = str(op.get("config_uid") or config_uid or "").strip()
            cls = str(op.get("class_name") or op.get("class") or "").strip()
            repo, cls_cfg, cls_name = _ngenie_find_class(cu, cls, lookup)
            if not repo or not cls_cfg:
                notes.append(f"Не найден класс {cls}")
                continue
            if not _ngenie_may_create_class(cls_name, cls_cfg, lookup, allow_catalog_create):
                role_label = _ngenie_role_label(_ngenie_class_role(cls_cfg))
                notes.append(f"Нельзя создать {role_label} {cls_name}: создание справочников/служебных классов выключено или запрещено ролью класса.")
                continue
            data = dict(op.get("data") or {}) if isinstance(op.get("data"), dict) else {}
            op_attachments = op.get("attachments") if isinstance(op.get("attachments"), list) else (attachments_list if op.get("attach_files") or op.get("attachFiles") else [])
            attach_field = op.get("file_field") or op.get("fileField") or op.get("gallery_field") or op.get("galleryField") or ""
            if op_attachments:
                data, used_attach_field, attach_count = _ngenie_apply_attachments_to_data(cls_cfg, data, op_attachments, field=attach_field, append=True)
                if attach_count and used_attach_field:
                    notes.append(f"Файлы добавлены в {cls_name}.{used_attach_field}: {attach_count}")
            data, delayed, ref_notes = _ngenie_resolve_refs(repo, cls_cfg, data, lookup, allow_catalog_create=allow_catalog_create)
            delayed.extend(_ngenie_extra_table_rows_from_operation(op, cls_cfg))
            notes.extend(ref_notes)
            uid = _ngenie_create_node(repo, cls_name, data)
            object_uids.append(uid)
            title = _ngenie_node_title(data) or uid
            notes.append(f"Создан {cls_name}: {title}")
            for d in delayed:
                parent_data, child_uids, child_notes = _ngenie_append_table_rows(repo, cls_name, uid, data, d.get("field"), d.get("rows") or [], lookup, allow_catalog_create=allow_catalog_create)
                object_uids.extend(child_uids)
                notes.extend(child_notes)
                data = parent_data
            last_created_repo = repo
            last_created_class = cls_name
            last_created_id = uid
            last_created_data = dict(data or {})
        elif tool in {"bulk_update_nodes", "update_nodes", "bulkUpdateNodes"}:
            cu = str(op.get("config_uid") or config_uid or "").strip()
            cls = str(op.get("class_name") or op.get("class") or "").strip()
            repo, cls_cfg, cls_name = _ngenie_find_class(cu, cls, lookup)
            if not repo or not cls_cfg:
                notes.append(f"Не найден класс {cls}")
                continue
            updates = op.get("updates") if isinstance(op.get("updates"), list) else op.get("items")
            if not isinstance(updates, list):
                notes.append(f"Для массового изменения {cls_name} не передан список updates")
                continue
            changed_count = 0
            skipped_count = 0
            for item in updates[:1000]:
                if not isinstance(item, dict):
                    skipped_count += 1
                    continue
                uid_raw = str(
                    item.get("uid") or item.get("node_uid") or item.get("nodeUid")
                    or item.get("_id") or item.get("id") or ""
                ).strip()
                patch = item.get("data") if isinstance(item.get("data"), dict) else item.get("patch")
                patch = dict(patch or {}) if isinstance(patch, dict) else {}
                if not uid_raw or not patch:
                    skipped_count += 1
                    continue
                target_internal_id = ""
                try:
                    if "$" in uid_raw:
                        uid_cu, uid_cls, uid_iid = _nodes_mod.parse_uid_any(uid_raw)
                        if uid_cu and str(uid_cu) != str(repo.config_uid):
                            skipped_count += 1
                            continue
                        if uid_cls and str(uid_cls).lower() != str(cls_name).lower():
                            skipped_count += 1
                            continue
                        target_internal_id = str(uid_iid or "").strip()
                    else:
                        target_internal_id = _nodes_mod.extract_internal_id(uid_raw)
                except Exception:
                    target_internal_id = _nodes_mod.extract_internal_id(uid_raw)
                if not target_internal_id or not _fetch_node_data_for_repo(repo, cls_name, target_internal_id):
                    skipped_count += 1
                    continue
                patch = {
                    str(k): v for k, v in patch.items()
                    if str(k) not in {"_id", "_class", "_created_date", "_created_user"}
                }
                patch, delayed, ref_notes = _ngenie_resolve_refs(
                    repo, cls_cfg, patch, lookup, allow_catalog_create=allow_catalog_create
                )
                notes.extend(ref_notes)
                changed = _ngenie_update_node(repo, cls_name, target_internal_id, patch)
                full_uid = _nodes_mod.normalize_own_uid(repo.config_uid, cls_name, target_internal_id)
                object_uids.append(full_uid)
                changed_count += 1
                for d in delayed:
                    changed, child_uids, child_notes = _ngenie_append_table_rows(
                        repo, cls_name, target_internal_id, changed, d.get("field"),
                        d.get("rows") or [], lookup, allow_catalog_create=allow_catalog_create
                    )
                    object_uids.extend(child_uids)
                    notes.extend(child_notes)
            notes.append(f"Изменено узлов {cls_name}: {changed_count}")
            if skipped_count:
                notes.append(f"Пропущено узлов {cls_name}: {skipped_count}")
        elif tool == "update_current_node":
            if not current_repo or not current_class or not current_id:
                notes.append("Нет текущего узла для изменения")
                continue
            patch = dict(op.get("data") or {}) if isinstance(op.get("data"), dict) else {}
            cls_cfg = (((get_parsed_config(current_repo, models.db) or {}).get("classes") or {}).get(current_class) or {})
            op_attachments = op.get("attachments") if isinstance(op.get("attachments"), list) else (attachments_list if op.get("attach_files") or op.get("attachFiles") else [])
            attach_field = op.get("file_field") or op.get("fileField") or op.get("gallery_field") or op.get("galleryField") or ""
            if op_attachments:
                patch_base = dict(current_data or {})
                patch_base.update(patch)
                patch, used_attach_field, attach_count = _ngenie_apply_attachments_to_data(cls_cfg, patch_base, op_attachments, field=attach_field, append=True)
                if attach_count and used_attach_field:
                    notes.append(f"Файлы добавлены в {current_class}.{used_attach_field}: {attach_count}")
            patch, delayed, ref_notes = _ngenie_resolve_refs(current_repo, cls_cfg, patch, lookup, allow_catalog_create=allow_catalog_create)
            notes.extend(ref_notes)
            changed_current = _ngenie_update_node(current_repo, current_class, current_id, patch)
            current_data = changed_current
            object_uids.append(_nodes_mod.normalize_own_uid(current_repo.config_uid, current_class, _nodes_mod.extract_internal_id(current_id)))
            notes.append("Текущий узел изменён")
            for d in delayed:
                current_data, child_uids, child_notes = _ngenie_append_table_rows(current_repo, current_class, current_id, current_data, d.get("field"), d.get("rows") or [], lookup, allow_catalog_create=allow_catalog_create)
                object_uids.extend(child_uids)
                notes.extend(child_notes)
                changed_current = current_data
        elif tool == "append_table_rows":
            field = str(op.get("field") or op.get("name") or "").strip()
            rows = op.get("rows") if isinstance(op.get("rows"), list) else []
            target_repo = current_repo
            target_class = current_class
            target_id = current_id
            target_data = current_data
            target_is_current = bool(current_repo and current_class and current_id)
            if not target_repo or not target_class or not target_id:
                # In the general nGenie window there is no opened current node.
                # If the model produced a common two-step plan
                #   create_node(parent) -> append_table_rows(...),
                # attach those rows to the node created immediately before this
                # operation instead of dropping them.
                if last_created_repo and last_created_class and last_created_id:
                    target_repo = last_created_repo
                    target_class = last_created_class
                    target_id = last_created_id
                    target_data = last_created_data
                    target_is_current = False
                else:
                    notes.append("Нет текущего узла для добавления строк")
                    continue
            if not field:
                cls_cfg_for_rows = (((get_parsed_config(target_repo, models.db) or {}).get("classes") or {}).get(target_class) or {})
                table_defs_for_rows = _ngenie_parse_data_structure(cls_cfg_for_rows.get("data_structure") or "").get("tables") or []
                if len(table_defs_for_rows) == 1:
                    field = str(table_defs_for_rows[0].get("name") or "").strip()
            new_parent_data, child_uids, child_notes = _ngenie_append_table_rows(target_repo, target_class, target_id, target_data, field, rows, lookup, allow_catalog_create=allow_catalog_create)
            if target_is_current:
                current_data = new_parent_data
                changed_current = current_data
            if last_created_repo == target_repo and last_created_class == target_class and last_created_id == target_id:
                last_created_data = new_parent_data
            object_uids.extend(child_uids)
            notes.extend(child_notes)
            notes.append(f"Добавлены строки в {field}: {len(rows)}")
        elif tool in {"attach_files", "attachFiles", "add_files", "addFiles"}:
            target_repo = current_repo
            target_class = current_class
            target_id = current_id
            target_data = current_data
            uid_raw = str(op.get("uid") or op.get("node_uid") or op.get("nodeUid") or "").strip()
            if uid_raw:
                try:
                    cu2, cls2, iid2 = _nodes_mod.parse_uid_any(uid_raw)
                    target_repo = _repo_for_config_uid(current_repo, cu2 or config_uid) if current_repo else models.Repo.query.filter_by(user_id=_ngenie_effective_user_id(), config_uid=(cu2 or config_uid)).first()
                    target_class = cls2 or target_class
                    target_id = iid2 or target_id
                    if target_repo and target_class and target_id:
                        target_data = _fetch_node_data_for_repo(target_repo, target_class, target_id) or {}
                except Exception:
                    pass
            if not target_repo or not target_class or not target_id:
                notes.append("Нет узла для добавления файлов")
                continue
            cls_cfg = (((get_parsed_config(target_repo, models.db) or {}).get("classes") or {}).get(target_class) or {})
            op_attachments = op.get("attachments") if isinstance(op.get("attachments"), list) else attachments_list
            field = op.get("field") or op.get("file_field") or op.get("fileField") or op.get("gallery_field") or op.get("galleryField") or ""
            patch, used_attach_field, attach_count = _ngenie_apply_attachments_to_data(cls_cfg, target_data or {}, op_attachments, field=field, append=True)
            if not attach_count or not used_attach_field:
                notes.append("Файлы не добавлены: нет вложений или FileGallery/MediaGallery поля")
                continue
            changed = _ngenie_update_node(target_repo, target_class, target_id, patch)
            if target_repo == current_repo and target_class == current_class and _nodes_mod.extract_internal_id(target_id) == _nodes_mod.extract_internal_id(current_id):
                current_data = changed
                changed_current = current_data
            object_uids.append(_nodes_mod.normalize_own_uid(target_repo.config_uid, target_class, _nodes_mod.extract_internal_id(target_id)))
            notes.append(f"Файлы добавлены в {target_class}.{used_attach_field}: {attach_count}")
    base_repo = None
    if object_uids:
        first_uid = object_uids[0]
        try:
            cu, _cn, _iid = _nodes_mod.parse_uid_any(first_uid)
        except Exception:
            cu = config_uid
        base_repo = models.Repo.query.filter_by(user_id=current_user.id, config_uid=cu or config_uid).first()
    if not base_repo:
        base_repo = current_repo or (models.Repo.query.filter_by(user_id=_ngenie_effective_user_id()).first())
    objects: List[Dict[str, Any]] = []
    if base_repo:
        for uid in object_uids:
            try:
                obj = _projection_object_payload(base_repo, "__ngenie__", uid)
                if obj:
                    objects.append(obj)
            except Exception:
                pass
    return {"notes": notes, "objects": objects, "node_data": changed_current}



def _ngenie_render_node_layout(repo: Optional[models.Repo], class_name: str, node_id: str, node_data: Optional[Dict[str, Any]]) -> str:
    """Render the current node form layout after nGenie changed data.

    Important: node_form can build its screen in onShowWeb, not only from
    init_screen_layout. The earlier nGenie refresh used only init_screen_layout;
    for such classes it returned an empty string and the frontend replaced the
    whole form with nothing. This mirrors the node_form route enough for refresh.
    """
    if not repo or not class_name or not node_id:
        return ""
    try:
        parsed = get_parsed_config(repo, models.db) or {}
        cls = ((parsed.get("classes") or {}).get(class_name) or {})
        layout = None
        data_for_layout = dict(node_data or {}) if isinstance(node_data, dict) else {}

        # Local nodes may define their screen dynamically in onShowWeb.
        try:
            if _ngenie_local_repo(repo):
                node_class = _load_server_node_class(repo.config_uid, class_name)
                internal_id = _nodes_mod.extract_internal_id(node_id)
                node = node_class.get(internal_id, repo.config_uid)
                if node:
                    _nodes_mod.CURRENT_NODE = node
                    try:
                        stored = node.get_data() or {}
                        if isinstance(stored, dict) and stored:
                            data_for_layout = dict(stored)
                    except Exception:
                        pass
                    try:
                        node._data_cache = dict(data_for_layout)
                    except Exception:
                        pass
                    for ev in (cls.get("events") or []):
                        if (ev.get("event") or "") != "onShowWeb":
                            continue
                        for action in (ev.get("actions") or []):
                            m = str((action or {}).get("method") or "").strip()
                            try:
                                if m.lower() == "nodascript":
                                    code = action.get("methodText") or action.get("method_text") or action.get("text") or action.get("code") or ""
                                    _run_web_nodascript_action(code, node._data_cache, node, repo.config_uid)
                                elif m and hasattr(node, m):
                                    getattr(node, m)({})
                                post_m = str(action.get("postExecuteMethod") or action.get("post_execute_method") or "").strip().lower()
                                if post_m == "nodascript":
                                    post_code = action.get("postExecuteMethodText") or action.get("post_execute_method_text") or action.get("postText") or action.get("post_text") or ""
                                    _run_web_nodascript_action(post_code, node._data_cache, node, repo.config_uid)
                            except Exception:
                                traceback.print_exc()
                    try:
                        if not _layout_is_empty(getattr(node, "_ui_layout", None)):
                            layout = node._ui_layout
                    except Exception:
                        pass
                    try:
                        if isinstance(getattr(node, "_data_cache", None), dict):
                            data_for_layout = dict(node._data_cache)
                    except Exception:
                        pass
        except Exception:
            traceback.print_exc()

        if _layout_is_empty(layout):
            data_layout = data_for_layout.get("_layout") if isinstance(data_for_layout, dict) else None
            layout = _first_usable_layout(
                data_layout,
                cls.get("init_screen_layout_web"),
                cls.get("init_screen_layout"),
            )
        layout = resolve_common_layout(parsed, layout)
        if layout is None:
            return ""
        if isinstance(data_for_layout, dict):
            _fill_nodeinput_views(repo, parsed, layout, data_for_layout)
        return render_nodalayout_html(
            layout,
            data_for_layout if isinstance(data_for_layout, dict) else {},
            assets_base_dir=_userfiles_dir_for_repo(repo),
            context=_nl_context(repo, class_name=class_name, node_id=node_id),
        )
    except Exception:
        traceback.print_exc()
        return ""

def _ngenie_projection_ref_view(value: Any, default: str = "") -> str:
    """Compatibility wrapper around the platform-level ``nodes.node_view``."""
    try:
        return str(_nodes_mod.node_view(value, default=default))
    except Exception:
        pass
    if value is None:
        return str(default or "")
    if isinstance(value, dict):
        explicit = str(value.get("_view") or value.get("view") or "").strip()
        if explicit:
            return explicit
        value = value.get("_id") or value.get("id") or value.get("uid") or ""
    raw = str(value or "").strip()
    if not raw:
        return str(default or "")
    try:
        config_uid, class_name, internal_id = _nodes_mod.parse_uid_any(raw)
    except Exception:
        return str(default or raw)
    if not class_name or not internal_id:
        return str(default or raw)
    try:
        q = models.Repo.query.filter_by(config_uid=str(config_uid or ""))
        try:
            uid = _ngenie_effective_user_id()
            if uid:
                q = q.filter_by(user_id=uid)
        except Exception:
            pass
        repo = q.order_by(models.Repo.id.asc()).first()
        if not repo:
            return str(default or raw)
        normalized = _nodes_mod.normalize_own_uid(repo.config_uid, class_name, internal_id)
        return _ngenie_node_ref_view(repo, normalized, default=str(default or raw))
    except Exception:
        return str(default or raw)


def _ngenie_projection_hydrate_link_views(data: Dict[str, Any]) -> Dict[str, Any]:
    """Return a shallow copy with missing ``<field>_view`` values resolved."""
    if not isinstance(data, dict):
        return {}
    out = dict(data)
    for key, value in list(data.items()):
        skey = str(key or "")
        if not skey or skey.startswith("_") or skey.endswith("_view"):
            continue
        if not isinstance(value, (str, dict)):
            continue
        raw = value
        if isinstance(value, dict):
            raw = value.get("_id") or value.get("id") or value.get("uid") or ""
        if not isinstance(raw, str) or "$" not in raw:
            continue
        try:
            _cu, cls_name, internal_id = _nodes_mod.parse_uid_any(raw)
        except Exception:
            continue
        if not cls_name or not internal_id:
            continue
        view_key = f"{skey}_view"
        if str(out.get(view_key) or "").strip():
            continue
        out[view_key] = _ngenie_projection_ref_view(raw, default=raw)
    return out


def _ngenie_projection_nodes(value: Any) -> List[Any]:
    """Normalize Node.get_all()/lists/single nodes for generated reports."""
    if value is None:
        return []
    if isinstance(value, str):
        raw = value.strip()
        if not raw:
            return []
        try:
            return _ngenie_projection_nodes(json.loads(raw))
        except Exception:
            return []
    if isinstance(value, dict):
        # A hydrated storage row or plain data dict is one object; get_all() is a
        # mapping id -> Node and should be iterated through its values.
        if "_data" in value or "_id" in value or "_class" in value:
            return [value]
        return list(value.values())
    if isinstance(value, (list, tuple, set)):
        return list(value)
    return [value]


def _ngenie_projection_data(value: Any) -> Dict[str, Any]:
    """Return a data dictionary from a Node, storage payload or JSON string."""
    if value is None:
        return {}
    if isinstance(value, str):
        raw = value.strip()
        if not raw:
            return {}
        try:
            parsed = json.loads(raw)
        except Exception:
            return {}
        return _ngenie_projection_data(parsed)
    if isinstance(value, dict):
        nested = value.get("_data")
        data = nested if isinstance(nested, dict) else value
        return _ngenie_projection_hydrate_link_views(data)
    try:
        data = value.get_data()
        if isinstance(data, dict):
            return _ngenie_projection_hydrate_link_views(data)
    except Exception:
        pass
    try:
        data = getattr(value, "_data", None)
        if isinstance(data, dict):
            return _ngenie_projection_hydrate_link_views(data)
    except Exception:
        pass
    return {}


def _ngenie_projection_rows(value: Any) -> List[Dict[str, Any]]:
    """Normalize an inline table/list, including legacy JSON-string values."""
    if value is None or value == "":
        return []
    if isinstance(value, str):
        raw = value.strip()
        if not raw:
            return []
        try:
            value = json.loads(raw)
        except Exception:
            return []
    if isinstance(value, dict):
        if "_data" in value or "_id" in value or "_class" in value:
            values = [value]
        elif value and all(isinstance(v, (dict, str)) or hasattr(v, "get_data") for v in value.values()):
            values = list(value.values())
        else:
            values = [value]
    elif isinstance(value, (list, tuple, set)):
        values = list(value)
    else:
        values = [value]
    rows: List[Dict[str, Any]] = []
    for item in values:
        data = _ngenie_projection_data(item)
        if data:
            rows.append(data)
    return rows


def _ngenie_repair_projection_after_preview_error(
    messages: List[Dict[str, str]], answer: Dict[str, Any], error: Exception
) -> Dict[str, Any]:
    """Ask the planner once to repair only a failing projection method.

    Report operations have already been executed by this point, so the repair pass
    is explicitly forbidden from returning data-changing operations.
    """
    if not isinstance(answer, dict) or not str(answer.get("projection_method_code") or "").strip():
        return answer
    repair_prompt = (
        "Предварительный запуск projection_method_code завершился ошибкой:\n"
        + str(error)
        + "\nИсправь только код HTML-отчёта и верни обычный JSON nGenie. "
          "operations, resolve_requests, clarification_requests и display_requests должны быть пустыми. "
          "Не повторяй и не меняй данные. Помни: Class.get_all(config_uid) возвращает dict id->Node; "
          "используй .values() либо ngenie_nodes(...), для данных ngenie_data(...), "
          "для табличной части ngenie_rows(...), для NodeLink node_view(uid). "
          "Текущая конфигурация внутри метода доступна как self._config_uid; если код использует "
          "selected_config_uid, сначала явно задай selected_config_uid = self._config_uid. "
          "Никогда не вызывай .get у строки/id и не выводи UID ссылки вместо представления."
    )
    retry_messages = list(messages) + [
        {"role": "assistant", "content": json.dumps(answer, ensure_ascii=False, default=str)},
        {"role": "user", "content": repair_prompt},
    ]
    fixed = _ngenie_call_deepseek(retry_messages)
    if not isinstance(fixed, dict) or not str(fixed.get("projection_method_code") or "").strip():
        return answer
    merged = dict(answer)
    for key in ("projection_method_code", "projection_title", "analysis_html", "reply", "summary"):
        if key in fixed:
            merged[key] = fixed.get(key)
    return merged


def _ngenie_run_projection_method_preview(config_uid: str, title: str, method_code: Any, input_data: Any = None) -> str:
    body = _ngenie_strip_method_code(method_code)
    if not body:
        return ""
    # Generated report code historically used selected_config_uid as though it
    # were a global supplied by the planner context.  A Projection method is an
    # ordinary Node method, so make the current configuration explicit inside
    # the method.  This also keeps old generated reports compatible.
    runtime_alias = (
        "selected_config_uid = str(getattr(self, '_config_uid', '') "
        "or (getattr(self, '_data', {}) or {}).get('_config_uid') or '')"
    )
    # Projection/node_form sends parameters inside payload.full_data.  Flatten
    # them into input_data as well, while retaining full_data for handlers that
    # explicitly use it.  Thus both input_data.get('date_from') and
    # input_data.get('full_data', {}).get('date_from') are valid.
    runtime_input = (
        "input_data = dict(input_data or {}) if isinstance(input_data, dict) else {}\n"
        "_projection_full_data = input_data.get('full_data')\n"
        "if isinstance(_projection_full_data, dict):\n"
        "    _projection_input = dict(_projection_full_data)\n"
        "    _projection_input.update(input_data)\n"
        "    input_data = _projection_input"
    )
    body = runtime_alias + "\n" + runtime_input + "\n" + body
    src = "def __ngenie_preview_onRunProjection(self, input_data):\n" + "\n".join("    " + line if line.strip() else "" for line in body.splitlines()) + "\n"
    ast.parse(src)
    repo = None
    if config_uid:
        uid = _ngenie_effective_user_id()
        q = models.Repo.query.filter_by(config_uid=str(config_uid or ""))
        if uid:
            q = q.filter_by(user_id=uid)
        repo = q.order_by(models.Repo.id.asc()).first()
    parsed = get_parsed_config(repo, models.db) if repo else None
    ns: Dict[str, Any] = {}
    context_tokens = None
    try:
        if config_uid:
            context_tokens = _nodes_mod.set_runtime_context(
                config_uid,
                parsed or {},
                system_user={"id": int(getattr(repo, "user_id", 0) or 0)} if repo else None,
            )
        ns.update(_load_server_handlers_ns(config_uid, parsed or {}) if config_uid else {})
        ns.setdefault("Node", _nodes_mod.Node)
        ns.setdefault("json", json)
        ns.setdefault("re", re)
        ns.setdefault("math", math)
        ns.setdefault("datetime", datetime)
        ns.setdefault("timedelta", timedelta)
        ns.setdefault("timezone", timezone)
        ns.setdefault("html_escape", escape)
        ns.setdefault("ngenie_nodes", getattr(_nodes_mod, "ngenie_nodes", _ngenie_projection_nodes))
        ns.setdefault("ngenie_data", getattr(_nodes_mod, "ngenie_data", _ngenie_projection_data))
        ns.setdefault("ngenie_rows", getattr(_nodes_mod, "ngenie_rows", _ngenie_projection_rows))
        ns.setdefault("node_view", getattr(_nodes_mod, "node_view", _ngenie_projection_ref_view))
        ns.setdefault("get_node_view", getattr(_nodes_mod, "get_node_view", _ngenie_projection_ref_view))
        ns.setdefault("ngenie_ref_view", getattr(_nodes_mod, "ngenie_ref_view", _ngenie_projection_ref_view))
        exec(src, ns, ns)

        class _PreviewNode:
            def __init__(self):
                self._id = "__ngenie_preview__"
                self._class = "__ngenie_preview__"
                self._config_uid = str(config_uid or "")
                self._data = {
                    "_config_uid": str(config_uid or ""),
                    "_projection_type": "html_projection",
                    "_projection_header": str(title or "nGenie analysis"),
                    "_projection_html": "",
                }

            def get_data(self):
                return self._data

            def save(self):
                return self

        preview_values = dict(input_data or {}) if isinstance(input_data, dict) else {}
        node = _PreviewNode()
        node._data.update(preview_values)
        event_payload = dict(preview_values)
        event_payload["full_data"] = dict(preview_values)
        result = ns["__ngenie_preview_onRunProjection"](node, event_payload)
        data = result if isinstance(result, dict) else getattr(node, "_data", {})
        return str((data or {}).get("_projection_html") or (data or {}).get("analysis_html") or "")
    finally:
        if context_tokens is not None:
            try:
                _nodes_mod.reset_runtime_context(context_tokens)
            except Exception:
                pass


def _ngenie_strip_method_code(code: Any) -> str:
    text = str(code or "").strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:python|py)?\s*", "", text, flags=re.I).strip()
        text = re.sub(r"\s*```$", "", text).strip()
    m = re.search(r"def\s+onRunProjection\s*\([^)]*\):\s*\n(?P<body>.*)$", text, flags=re.S)
    if m:
        body = m.group("body")
        lines = body.splitlines()
        min_indent = None
        for line in lines:
            if line.strip():
                indent = len(line) - len(line.lstrip())
                min_indent = indent if min_indent is None else min(min_indent, indent)
        if min_indent:
            lines = [line[min_indent:] if len(line) >= min_indent else line for line in lines]
        text = "\n".join(lines).strip()
    return text


def _ngenie_make_static_projection_code(title: str, html: str) -> str:
    safe_title = json.dumps(str(title or "nGenie HTML"), ensure_ascii=False)
    safe_html = json.dumps(str(html or ""), ensure_ascii=False)
    return (
        "self._data['_projection_type'] = 'html_projection'\n"
        f"self._data['_projection_header'] = {safe_title}\n"
        f"self._data['_projection_html'] = {safe_html}\n"
        "return self._data"
    )


def _ngenie_slug_class_name(title: str) -> str:
    raw = str(title or "NGenieHTML").strip()
    translit = {
        "а":"a","б":"b","в":"v","г":"g","д":"d","е":"e","ё":"e","ж":"zh","з":"z","и":"i","й":"y","к":"k","л":"l","м":"m","н":"n","о":"o","п":"p","р":"r","с":"s","т":"t","у":"u","ф":"f","х":"h","ц":"c","ч":"ch","ш":"sh","щ":"sch","ъ":"","ы":"y","ь":"","э":"e","ю":"yu","я":"ya"
    }
    s = "".join(translit.get(ch.lower(), ch) for ch in raw)
    s = re.sub(r"[^0-9a-zA-Z_]+", "_", s).strip("_") or "HTMLProjection"
    if s[0].isdigit():
        s = "P_" + s
    if not s.lower().startswith("ngenie"):
        s = "NGenie_" + s
    return s[:80]


def _ngenie_ensure_handlers_method(config: Any, class_name: str, method_body: str) -> str:
    body = _ngenie_strip_method_code(method_body)
    if not body:
        body = _ngenie_make_static_projection_code(class_name, "")
    # Persist the same compatibility alias used by preview.  Saved HTML
    # Projections are regular Node handlers and must not depend on a request-
    # local planner variable existing in module globals.
    runtime_alias = (
        "selected_config_uid = str(getattr(self, '_config_uid', '') "
        "or (getattr(self, '_data', {}) or {}).get('_config_uid') or '')"
    )
    # Projection/node_form sends parameters inside payload.full_data.  Flatten
    # them into input_data as well, while retaining full_data for handlers that
    # explicitly use it.  Thus both input_data.get('date_from') and
    # input_data.get('full_data', {}).get('date_from') are valid.
    runtime_input = (
        "input_data = dict(input_data or {}) if isinstance(input_data, dict) else {}\n"
        "_projection_full_data = input_data.get('full_data')\n"
        "if isinstance(_projection_full_data, dict):\n"
        "    _projection_input = dict(_projection_full_data)\n"
        "    _projection_input.update(input_data)\n"
        "    input_data = _projection_input"
    )
    body = runtime_alias + "\n" + runtime_input + "\n" + body
    body_lines = body.splitlines() or ["return self._data"]
    indented = "\n".join("        " + (line if line.strip() else "") for line in body_lines)
    method_text = f"\nclass {class_name}(Node):\n    def onRunProjection(self, input_data):\n{indented}\n"
    existing = ""
    raw = getattr(config, "nodes_server_handlers", None)
    if raw:
        try:
            existing = base64.b64decode(raw).decode("utf-8")
        except Exception:
            existing = str(raw or "")
    if "from nodes import" not in existing:
        existing = "from nodes import *\n" + existing
    pattern = rf"\nclass\s+{re.escape(class_name)}\s*\(Node\):\n(?:    .*\n|\n)*?(?=\nclass\s+\w+\s*\(Node\):|\Z)"
    if re.search(pattern, existing, flags=re.S):
        new_code = re.sub(pattern, method_text, existing, flags=re.S)
    else:
        new_code = existing.rstrip() + "\n" + method_text
    ast.parse(new_code)
    config.nodes_server_handlers = base64.b64encode(new_code.encode("utf-8")).decode("ascii")
    return new_code


def _ngenie_write_handlers_file(config_uid: str, code: str) -> None:
    try:
        hdir = Path("Handlers") / str(config_uid)
        hdir.mkdir(parents=True, exist_ok=True)
        (hdir / "handlers.py").write_text(code, encoding="utf-8")
    except Exception:
        pass


def _ngenie_missing_requested_table_operation(
    answer: Dict[str, Any],
    user_message: str,
    config_uid: str,
    node_context: Optional[Dict[str, Any]],
    scope: str,
) -> List[str]:
    """Return current table names when a node-form answer clearly stopped after
    creating reference/catalog nodes and forgot the requested document rows.

    This is a narrow validation guard.  It does not try to infer business data;
    it only asks the LLM for one corrected answer before any operations execute.
    """
    if str(scope or "").strip().lower() != "node_form" or not isinstance(node_context, dict):
        return []
    text = str(user_message or "").lower()
    row_intent = bool(re.search(
        r"(?:добав\w*|заполн\w*|загруз\w*|импорт\w*|перенес\w*)[^\n]{0,80}(?:строк|позици|товар)|"
        r"(?:строк|позици|товар)[^\n]{0,80}(?:в заказ|в документ)|"
        r"(?:заказ|документ)[^\n]{0,80}(?:строк|позици|товар)|"
        r"(?:add|append|import|load)[^\n]{0,80}(?:row|line|item)",
        text,
        flags=re.I,
    ))
    if not row_intent:
        return []

    cls_name = str(node_context.get("class_name") or "").strip()
    cu = str(node_context.get("config_uid") or config_uid or "").strip()
    repo = models.Repo.query.filter_by(user_id=_ngenie_effective_user_id(), config_uid=cu).first()
    if not repo or not cls_name:
        return []
    cls_cfg = (((get_parsed_config(repo, models.db) or {}).get("classes") or {}).get(cls_name) or {})
    table_names = [
        str(t.get("name") or "").strip()
        for t in (_ngenie_parse_data_structure(cls_cfg.get("data_structure") or "").get("tables") or [])
        if str(t.get("name") or "").strip()
    ]
    if not table_names:
        return []

    # Do not interfere with the normal candidate/clarification pipeline.
    if any(answer.get(k) for k in (
        "candidate_handler_code", "candidateHandlerCode", "resolve_requests", "resolveRequests",
        "clarification_requests", "clarifications",
    )):
        return []

    ops = answer.get("operations") if isinstance(answer.get("operations"), list) else []
    if not ops:
        return table_names
    has_current_table_change = False
    has_catalog_create = False
    for op in ops:
        if not isinstance(op, dict):
            continue
        tool = str(op.get("tool") or op.get("action") or "").strip()
        if tool == "append_table_rows":
            field = str(op.get("field") or op.get("name") or "").strip()
            if (not field and len(table_names) == 1) or field in table_names:
                has_current_table_change = True
        elif tool == "update_current_node":
            patch = op.get("data") if isinstance(op.get("data"), dict) else {}
            if any(name in patch and isinstance(patch.get(name), list) for name in table_names):
                has_current_table_change = True
        elif tool == "create_node":
            has_catalog_create = True
    if has_current_table_change:
        return []
    return table_names if has_catalog_create or ops else []


def _ngenie_repair_missing_table_operation(
    messages: List[Dict[str, str]],
    answer: Dict[str, Any],
    user_message: str,
    config_uid: str,
    node_context: Optional[Dict[str, Any]],
    scope: str,
) -> Dict[str, Any]:
    table_names = _ngenie_missing_requested_table_operation(answer, user_message, config_uid, node_context, scope)
    if not table_names:
        return answer
    repair = (
        "Проверка выполнения обнаружила ошибку: пользователь просил добавить строки в текущий открытый узел, "
        "но твой JSON не изменяет ни одну его табличную часть. Табличные части текущего класса: "
        + ", ".join(table_names)
        + ". Верни ПОЛНЫЙ исправленный JSON-ответ для исходной задачи. "
          "Для именованной inline-таблицы используй append_table_rows. Если из файла надо создать отсутствующие "
          "справочные узлы и тут же вставить ссылки, не останавливайся на create_node: используй в Node-поле строки "
          "структуру {\"query\":\"...\",\"create_data\":{...}}, чтобы backend нашёл/создал справочник и добавил строку одной операцией."
    )
    retry_messages = list(messages) + [
        {"role": "assistant", "content": json.dumps(answer, ensure_ascii=False, default=str)},
        {"role": "user", "content": repair},
    ]
    try:
        return _ngenie_call_deepseek(retry_messages)
    except Exception:
        traceback.print_exc()
        return answer


@client_bp.route("/ngenie")
@login_required
def ngenie_page():
    if not _client_ngenie_enabled():
        flash("nGenie is disabled in web client settings", "warning")
        return redirect(url_for("client.client_settings"))
    repos = models.Repo.query.filter_by(user_id=_ngenie_effective_user_id()).order_by(models.Repo.name.asc(), models.Repo.id.asc()).all()
    return render_template(
        "client/ngenie.html",
        title=f"{APP_TITLE} — nGenie",
        repos=repos,
        purple=NGENIE_PURPLE,
        api_chat=url_for("client.api_ngenie_chat"),
        api_save_projection=url_for("client.api_ngenie_save_html_projection"),
        api_preview_projection=url_for("client.api_ngenie_preview_html_projection"),
    )


@client_bp.route("/api/ngenie/chat", methods=["POST"])
@login_required
def api_ngenie_chat():
    if not _client_ngenie_enabled():
        return jsonify({"ok": False, "error": "nGenie is disabled in web client settings"}), 403
    j = request.get_json(force=True) or {}
    message = str(j.get("message") or "").strip()
    config_uid = str(j.get("config_uid") or "").strip()
    node_context = j.get("node_context") if isinstance(j.get("node_context"), dict) else None
    allow_catalog_create = bool(j.get("allow_catalog_create"))
    raw_attachments = j.get("attachments") or []
    conversation_history = j.get("chat_history") or j.get("conversation_history") or []
    conversation_artifact = j.get("conversation_artifact") if isinstance(j.get("conversation_artifact"), dict) else {}
    scope = str(j.get("scope") or "").strip()
    cached_artifact = _ngenie_cached_chat_artifact(config_uid, scope, node_context)
    has_prior_user_turn = any(
        isinstance(item, dict) and str(item.get("role") or "").strip().lower() == "user"
        for item in (conversation_history if isinstance(conversation_history, list) else [])
    )
    if cached_artifact and (conversation_artifact or has_prior_user_turn):
        merged_artifact = dict(cached_artifact)
        merged_artifact.update({k: v for k, v in conversation_artifact.items() if str(v or "").strip()})
        conversation_artifact = merged_artifact
    is_node_form_scope = scope == "node_form"
    if not message:
        return jsonify({"ok": False, "error": "empty message"}), 400
    if config_uid and not models.Repo.query.filter_by(user_id=current_user.id, config_uid=config_uid).first():
        return jsonify({"ok": False, "error": "configuration not found"}), 404
    try:
        attachments = _ngenie_prepare_attachments_for_chat(raw_attachments, config_uid, message)
        clarification_response = j.get("clarification_response") if isinstance(j.get("clarification_response"), dict) else None
        chat_messages = _ngenie_build_messages(
            message,
            config_uid,
            node_context,
            allow_catalog_create=allow_catalog_create,
            clarification_response=clarification_response,
            scope=scope,
            attachments=attachments,
            conversation_history=conversation_history,
            conversation_artifact=conversation_artifact,
        )
        requests_projection_parameters = _ngenie_message_requests_projection_parameters(
            message, conversation_artifact
        )
        prior_artifact = _ngenie_normalize_conversation_artifact(conversation_artifact)
        bridge_answer = {}
        # For an existing report the Code bridge can work immediately.  For a
        # one-shot request (create a report already with parameters) the report
        # method does not exist yet, so first let analysis_reports create it and
        # invoke the same bridge below on the newly produced artifact.
        if requests_projection_parameters and str(prior_artifact.get("projection_method_code") or "").strip():
            bridge_answer = _ngenie_projection_parameters_bridge(
                message, config_uid, prior_artifact
            )
        answer = bridge_answer or _ngenie_call_deepseek(chat_messages)
        answer = _ngenie_repair_report_misroute(
            chat_messages, answer, message, conversation_history, conversation_artifact
        )
        if not isinstance(answer, dict):
            answer = {}

        # New report with parameters: now that projection_method_code exists, use
        # the optional nGenie Code knowledge to build the standard Projection
        # right-panel layout and minimally adapt the handler.
        if (
            requests_projection_parameters
            and not bridge_answer
            and str(answer.get("projection_method_code") or "").strip()
        ):
            bridge_artifact = dict(prior_artifact)
            bridge_artifact.update({
                "projection_title": answer.get("projection_title") or prior_artifact.get("projection_title") or "",
                "projection_method_code": answer.get("projection_method_code") or "",
                "analysis_html": answer.get("analysis_html") or prior_artifact.get("analysis_html") or "",
                "projection_parameters_layout": answer.get("projection_parameters_layout") or prior_artifact.get("projection_parameters_layout"),
                "projection_parameters_data": answer.get("projection_parameters_data") or prior_artifact.get("projection_parameters_data") or {},
                "projection_parameters_data_structure": answer.get("projection_parameters_data_structure") or prior_artifact.get("projection_parameters_data_structure") or "",
            })
            generated_bridge = _ngenie_projection_parameters_bridge(
                message, config_uid, bridge_artifact
            )
            if generated_bridge:
                answer = generated_bridge

        # Parameters belong to the Projection shell, not to report HTML. Preserve
        # an existing layout across unrelated report edits unless the answer
        # explicitly replaces it.
        for artifact_key in (
            "projection_parameters_layout",
            "projection_parameters_data",
            "projection_parameters_data_structure",
        ):
            if artifact_key not in answer and prior_artifact.get(artifact_key) not in (None, "", [], {}):
                answer[artifact_key] = prior_artifact.get(artifact_key)
        if requests_projection_parameters and not _ngenie_projection_parameters_layout_valid(answer.get("projection_parameters_layout"), message):
            defaults = _ngenie_default_projection_parameters_for_request(message)
            for key, value in defaults.items():
                answer[key] = value
        if not clarification_response:
            answer = _ngenie_repair_missing_table_operation(chat_messages, answer, message, config_uid, node_context, scope)
            answer = _ngenie_repair_bulk_data_confirmation(chat_messages, answer, message, config_uid)
            data_final = _ngenie_finalize_data_requests_with_llm(
                message,
                config_uid,
                node_context,
                allow_catalog_create,
                scope,
                answer,
                attachments=attachments,
                mobile_mode=False,
            )
            if data_final:
                answer = data_final
            answer = _ngenie_repair_invalid_candidate_handler(
                chat_messages, answer, user_message=message
            )
            # A repaired plan may correctly switch from an invalid candidate
            # handler to data_requests. Execute that normal two-step path now.
            if answer.get('data_requests') or answer.get('dataRequests'):
                repaired_data_final = _ngenie_finalize_data_requests_with_llm(
                    message,
                    config_uid,
                    node_context,
                    allow_catalog_create,
                    scope,
                    answer,
                    attachments=attachments,
                    mobile_mode=False,
                )
                if repaired_data_final:
                    answer = repaired_data_final
        _ngenie_remember_chat_artifact(config_uid, scope, node_context, answer)
        selected_skill_ids: List[str] = []
        try:
            selected_skill_ids = list(getattr(g, "ngenie_selected_skill_ids", []) or [])
            answer["_ngenie_selected_skill_ids"] = selected_skill_ids
        except Exception:
            selected_skill_ids = []
        _classes, lookup = _ngenie_collect_context(config_uid, include_samples=False)

        # Candidate resolving can be controlled by the LLM itself. The preferred
        # path is candidate_handler_code: the first LLM call sees class prompts,
        # DataStructure and indexes, writes a tiny resolver handler, and backend
        # executes it with safe search helpers. This lets every catalog define its
        # own strategy in ngenie_prompt (article/code exact, ngenie_name, etc.)
        # without hard-coding product semantics in Python.
        if not clarification_response:
            handler_code = (
                answer.get("candidate_handler_code")
                or answer.get("candidateHandlerCode")
                or answer.get("resolve_handler_code")
                or answer.get("resolveHandlerCode")
            )
            if handler_code:
                handler_result = _ngenie_execute_candidate_handler(
                    handler_code,
                    message,
                    config_uid,
                    node_context,
                    allow_catalog_create,
                    scope,
                    answer,
                    lookup,
                    attachments=attachments,
                )
                blocking_messages = _ngenie_handler_blocking_messages(handler_result)
                handler_clarifications = _ngenie_handler_ambiguities_to_clarifications(handler_result, config_uid, lookup)
                if blocking_messages:
                    return jsonify({
                        "ok": True,
                        "reply": "\n".join(blocking_messages),
                        "clarifications": [],
                        "created_objects": [],
                        "node_data": None,
                        "layout_html": "",
                        "analysis_html": "",
                        "projection_title": answer.get("projection_title") or "",
                        "projection_method_code": answer.get("projection_method_code") or "",
                        "raw": {**answer, "_candidate_handler_result": _ngenie_compact_handler_result_for_llm(handler_result)} if bool(j.get("debug")) else None,
                    })
                if handler_clarifications:
                    return jsonify({
                        "ok": True,
                        "reply": str(answer.get("reply") or "Нужно уточнение выбора."),
                        "clarifications": handler_clarifications,
                        "created_objects": [],
                        "node_data": None,
                        "layout_html": "",
                        "analysis_html": "",
                        "projection_title": answer.get("projection_title") or "",
                        "projection_method_code": answer.get("projection_method_code") or "",
                        "raw": {**answer, "_candidate_handler_result": _ngenie_compact_handler_result_for_llm(handler_result)} if bool(j.get("debug")) else None,
                    })
                finalized = _ngenie_finalize_candidate_handler_with_llm(
                    message,
                    config_uid,
                    node_context,
                    allow_catalog_create,
                    scope,
                    answer,
                    handler_result,
                    attachments=attachments,
                )
                if finalized:
                    answer = finalized
            else:
                # Backward-compatible fallback: if no generated resolver handler
                # was produced, keep the older resolve_requests flow.
                resolve_raw = answer.get("resolve_requests") or answer.get("resolveRequests") or []
                if resolve_raw:
                    resolved_answer = _ngenie_resolve_candidate_requests_with_llm(
                        message,
                        config_uid,
                        node_context,
                        allow_catalog_create,
                        scope,
                        answer,
                        resolve_raw,
                        lookup,
                        mode="resolve_requests",
                        attachments=attachments,
                    )
                    if resolved_answer:
                        answer = resolved_answer
                else:
                    initial_clar = answer.get("clarification_requests") or answer.get("clarifications") or []
                    if initial_clar:
                        resolved_answer = _ngenie_resolve_candidate_requests_with_llm(
                            message,
                            config_uid,
                            node_context,
                            allow_catalog_create,
                            scope,
                            answer,
                            initial_clar,
                            lookup,
                            mode="clarification_review",
                            attachments=attachments,
                        )
                        if resolved_answer:
                            answer = resolved_answer

                op_resolve_raw = _ngenie_extract_resolve_requests_from_operations(answer.get("operations") or [], config_uid, node_context, lookup)
                if op_resolve_raw:
                    resolved_answer = _ngenie_resolve_candidate_requests_with_llm(
                        message,
                        config_uid,
                        node_context,
                        allow_catalog_create,
                        scope,
                        answer,
                        op_resolve_raw,
                        lookup,
                        mode="operation_ref_resolve",
                        attachments=attachments,
                    )
                    if resolved_answer:
                        answer = resolved_answer

        _ngenie_remember_chat_artifact(config_uid, scope, node_context, answer)
        clarifications = _ngenie_prepare_clarifications(answer.get("clarification_requests") or answer.get("clarifications") or [], config_uid, lookup)
        display_raw = answer.get("display_requests") or answer.get("displayRequests") or answer.get("candidate_requests") or answer.get("candidateRequests") or []
        if isinstance(display_raw, list):
            display_raw = [dict(x, display_only=True) if isinstance(x, dict) else x for x in display_raw]
        display_groups = _ngenie_prepare_clarifications(display_raw, config_uid, lookup, literal_query=message)
        clarifications = clarifications + display_groups
        reply = str(answer.get("reply") or "").strip()
        if clarifications:
            has_selectable_clarification = any(not bool(g.get("display_only")) for g in clarifications if isinstance(g, dict))
            return jsonify({
                "ok": True,
                "reply": reply or ("Нужно уточнение выбора." if has_selectable_clarification else "Показал найденные варианты."),
                "clarifications": clarifications,
                "created_objects": [],
                "node_data": None,
                "layout_html": "",
                "analysis_html": "",
                "projection_title": answer.get("projection_title") or "",
                "projection_method_code": answer.get("projection_method_code") or "",
                "raw": answer if bool(j.get("debug")) else None,
            })

        exec_result = _ngenie_execute_operations(answer.get("operations") or [], config_uid, node_context, allow_catalog_create=allow_catalog_create, attachments=attachments)
        notes = exec_result.get("notes") or []

        layout_html = ""
        if isinstance(node_context, dict) and isinstance(exec_result.get("node_data"), dict):
            cu = str(node_context.get("config_uid") or config_uid or "").strip()
            repo = models.Repo.query.filter_by(user_id=current_user.id, config_uid=cu).first()
            layout_html = _ngenie_render_node_layout(
                repo,
                str(node_context.get("class_name") or ""),
                str(node_context.get("node_id") or node_context.get("id") or ""),
                exec_result.get("node_data"),
            )

        method_code = answer.get("projection_method_code") or ""
        analysis_html = answer.get("analysis_html") or ""
        if is_node_form_scope and (method_code or analysis_html):
            # In node_form nGenie works with the currently opened document.
            # Do not create/show configuration-level reports here; list/search results
            # should be returned through display_requests instead.
            if not reply:
                reply = "В форме узла отчёт/проекцию не создаю. Для списка узлов попроси показать кандидатов в чате."
            method_code = ""
            analysis_html = ""
        if method_code:
            try:
                preview_html = _ngenie_run_projection_method_preview(
                    config_uid,
                    answer.get("projection_title") or "nGenie analysis",
                    method_code,
                    input_data=answer.get("projection_parameters_data") or {},
                )
                if preview_html:
                    analysis_html = preview_html
            except Exception as preview_exc:
                traceback.print_exc()
                repaired_ok = False
                try:
                    repaired = _ngenie_repair_projection_after_preview_error(chat_messages, answer, preview_exc)
                    repaired_code = (repaired.get("projection_method_code") or "") if isinstance(repaired, dict) else ""
                    if repaired_code and str(repaired_code).strip() != str(method_code).strip():
                        repaired_html = _ngenie_run_projection_method_preview(
                            config_uid,
                            repaired.get("projection_title") or answer.get("projection_title") or "nGenie analysis",
                            repaired_code,
                            input_data=(repaired.get("projection_parameters_data") or answer.get("projection_parameters_data") or {}),
                        )
                        answer = repaired
                        method_code = repaired_code
                        if repaired_html:
                            analysis_html = repaired_html
                        elif repaired.get("analysis_html"):
                            analysis_html = repaired.get("analysis_html") or ""
                        if not reply and repaired.get("reply"):
                            reply = str(repaired.get("reply") or "").strip()
                        repaired_ok = bool(analysis_html)
                except Exception:
                    traceback.print_exc()
                if not repaired_ok:
                    notes.append(f"HTML preview не выполнен: {preview_exc}")

        if notes:
            reply = (reply + "\n" if reply else "") + "\n".join(notes)

        parameters_layout = _ngenie_normalize_projection_parameters_layout(
            answer.get("projection_parameters_layout")
        )
        if parameters_layout is not None and not _ngenie_projection_parameters_layout_valid(parameters_layout, message):
            parameters_layout = None
        parameters_data = answer.get("projection_parameters_data") if isinstance(answer.get("projection_parameters_data"), dict) else {}
        parameters_data_structure = str(answer.get("projection_parameters_data_structure") or "").strip()
        parameters_layout_html = _ngenie_render_projection_parameters_layout(
            config_uid, parameters_layout, parameters_data
        ) if parameters_layout is not None else ""
        close_projection = bool(
            prior_artifact.get("projection_method_code")
            and not method_code
            and "analysis_reports" not in selected_skill_ids
        )
        if close_projection:
            _ngenie_clear_chat_artifact(config_uid, scope, node_context)
        else:
            _ngenie_remember_chat_artifact(config_uid, scope, node_context, {
                "projection_title": answer.get("projection_title") or "",
                "projection_method_code": method_code,
                "analysis_html": analysis_html or "",
                "projection_parameters_layout": parameters_layout,
                "projection_parameters_data": parameters_data,
                "projection_parameters_data_structure": parameters_data_structure,
            })
        return jsonify({
            "ok": True,
            "reply": reply or "Готово.",
            "created_objects": exec_result.get("objects") or [],
            "node_data": exec_result.get("node_data"),
            "layout_html": layout_html,
            "analysis_html": analysis_html or "",
            "projection_title": answer.get("projection_title") or "",
            "projection_method_code": method_code,
            "projection_parameters_layout": parameters_layout,
            "projection_parameters_layout_html": parameters_layout_html,
            "projection_parameters_data": parameters_data,
            "projection_parameters_data_structure": parameters_data_structure,
            "close_projection": close_projection,
            "raw": answer if bool(j.get("debug")) else None,
        })
    except Exception as exc:
        traceback.print_exc()
        return jsonify({"ok": False, "error": str(exc)}), 500


@client_bp.route("/api/ngenie/preview-html-projection", methods=["POST"])
@login_required
def api_ngenie_preview_html_projection():
    if not _client_ngenie_enabled():
        return jsonify({"ok": False, "error": "nGenie is disabled in web client settings"}), 403
    j = request.get_json(force=True) or {}
    config_uid = str(j.get("config_uid") or "").strip()
    title = str(j.get("title") or j.get("projection_title") or "nGenie analysis").strip()
    method_code = _ngenie_strip_method_code(j.get("method_code") or j.get("projection_method_code") or "")
    parameters_data = j.get("projection_parameters_data") or j.get("parameters_data") or {}
    if not isinstance(parameters_data, dict):
        parameters_data = {}
    if not config_uid:
        return jsonify({"ok": False, "error": "select a configuration first"}), 400
    if not models.Repo.query.filter_by(user_id=_ngenie_effective_user_id(), config_uid=config_uid).first():
        return jsonify({"ok": False, "error": "configuration not found"}), 404
    if not method_code:
        return jsonify({"ok": False, "error": "projection method is empty"}), 400
    try:
        html = _ngenie_run_projection_method_preview(
            config_uid, title, method_code, input_data=parameters_data
        )
        return jsonify({
            "ok": True,
            "analysis_html": html or "",
            "projection_title": title,
            "projection_parameters_data": parameters_data,
        })
    except Exception as exc:
        traceback.print_exc()
        return jsonify({"ok": False, "error": str(exc)}), 500


def _ngenie_html_to_plain_text(value: Any) -> str:
    text = str(value or "")
    text = re.sub(r"<\s*br\s*/?>", "\n", text, flags=re.I)
    text = re.sub(r"</\s*(p|div|li|tr|h[1-6])\s*>", "\n", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    return re.sub(r"\n{3,}", "\n\n", unescape(text)).strip()


def _ngenie_value_is_meaningful(value: Any) -> bool:
    return value not in (None, "", [], {})


def _ngenie_prompt_requests_summary_only(prompt: str) -> bool:
    """True when the direct call asks for a summary, not a report/JSON artifact."""
    text = str(prompt or "").strip().lower()
    if not text:
        return False
    summary_requested = bool(re.search(r"\bsummary\b|сводк|резюм|кратк(?:о|ую|ий)|итог(?:и|овую)?", text, flags=re.I))
    if not summary_requested:
        return False
    # These words indicate that summary is only one part of a compound result.
    other_artifact = bool(re.search(
        r"\bhtml\b|хтмл|\bjson\b|таблиц|график|диаграм|проекц|отч[её]т|список\s+узл|"
        r"создай\s+узл|измени\s+узл|удали\s+узл|покажи\s+узл",
        text,
        flags=re.I,
    ))
    return not other_artifact


def _ngenie_summary_from_computed_result(prompt: str, computed_html: str) -> str:
    plain = _ngenie_html_to_plain_text(computed_html)
    if not plain:
        return ""
    # A second, small model pass turns the already computed report into the exact
    # string requested by ngenie().  It receives no database tools and therefore
    # cannot replace actual values with invented ones.
    try:
        answer = _ngenie_call_deepseek([
            {
                "role": "system",
                "content": (
                    "Ты получаешь уже вычисленный результат анализа данных NodaLogic. "
                    "Верни строго JSON-объект вида {\"summary\":\"...\"}. "
                    "В summary дай только содержательную сводку по фактическим данным ниже. "
                    "Не упоминай код, проекцию, необходимость дополнительного анализа или отсутствие доступа к базе. "
                    "Не добавляй другие ключи."
                ),
            },
            {
                "role": "user",
                "content": (
                    "Исходный запрос пользователя:\n" + str(prompt or "") +
                    "\n\nВычисленный результат:\n" + plain[:50000]
                ),
            },
        ])
        value = answer.get("summary") if isinstance(answer, dict) else None
        if _ngenie_value_is_meaningful(value):
            return value if isinstance(value, str) else json.dumps(value, ensure_ascii=False)
    except Exception:
        traceback.print_exc()
    return plain


def _ngenie_direct_result_string(answer: Any, prompt: str, config_uid: str) -> str:
    if not isinstance(answer, dict):
        return str(answer or "")

    summary_only = _ngenie_prompt_requests_summary_only(prompt)
    method_code = answer.get("projection_method_code") or ""

    # A direct summary may need all rows.  Execute the generated projection first,
    # then summarize its real output.  Previously the presence of summary caused an
    # early JSON return and leaked projection code into banner().
    if method_code and config_uid:
        try:
            preview = _ngenie_run_projection_method_preview(
                config_uid,
                answer.get("projection_title") or "nGenie analysis",
                method_code,
            )
            if preview:
                if summary_only:
                    return _ngenie_summary_from_computed_result(prompt, preview)
                wants_html = bool(re.search(r"\bhtml\b|хтмл|html[- ]?отч", str(prompt or ""), flags=re.I))
                return preview if wants_html else _ngenie_html_to_plain_text(preview)
        except Exception as exc:
            traceback.print_exc()
            if summary_only and not _ngenie_value_is_meaningful(answer.get("summary")):
                return f"nGenie не смог выполнить анализ: {exc}"

    # For an explicit summary-only direct call, summary is the public result even
    # if the model unnecessarily filled technical JSON fields.  Compound requests
    # still return the complete JSON object.
    if summary_only and _ngenie_value_is_meaningful(answer.get("summary")):
        value = answer.get("summary")
        return value if isinstance(value, str) else json.dumps(value, ensure_ascii=False)

    # Generic summary contract for prompts that are not recognised as summary-only.
    if "summary" in answer:
        meaningful = {k: v for k, v in answer.items() if _ngenie_value_is_meaningful(v)}
        if set(meaningful) == {"summary"}:
            value = meaningful.get("summary")
            return value if isinstance(value, str) else json.dumps(value, ensure_ascii=False)
        return json.dumps(answer, ensure_ascii=False)

    analysis_html = answer.get("analysis_html")
    if analysis_html:
        wants_html = bool(re.search(r"\bhtml\b|хтмл|html[- ]?отч", str(prompt or ""), flags=re.I))
        return str(analysis_html) if wants_html else _ngenie_html_to_plain_text(analysis_html)

    for key in ("result", "reply", "text", "html", "content"):
        if answer.get(key) is not None:
            value = answer.get(key)
            return value if isinstance(value, str) else json.dumps(value, ensure_ascii=False)
    return json.dumps(answer, ensure_ascii=False)



def _ngenie_data_value(data: Dict[str, Any], field: str) -> Any:
    cur: Any = data
    for part in str(field or "").split("."):
        if not part:
            continue
        if not isinstance(cur, dict) or part not in cur:
            return None
        cur = cur.get(part)
    return cur


def _ngenie_bool_value(value: Any) -> Optional[bool]:
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)) and value in (0, 1):
        return bool(value)
    if isinstance(value, str):
        low = value.strip().lower()
        if low in {"1", "true", "yes", "on", "да", "истина"}:
            return True
        if low in {"0", "false", "no", "off", "нет", "ложь", ""}:
            return False
    return None


def _ngenie_values_equal(left: Any, right: Any) -> bool:
    lb = _ngenie_bool_value(left)
    rb = _ngenie_bool_value(right)
    if lb is not None and rb is not None:
        return lb == rb
    if isinstance(left, (int, float)) and isinstance(right, (int, float)):
        return float(left) == float(right)
    return left == right or str(left) == str(right)


def _ngenie_row_matches_filters(data: Dict[str, Any], filters: Any) -> bool:
    for flt in filters or []:
        if not isinstance(flt, dict):
            continue
        field = str(flt.get("field") or "").strip()
        op = str(flt.get("op") or "eq").strip().lower()
        expected = flt.get("value")
        actual = _ngenie_data_value(data, field)
        if op in {"eq", "=", "=="}:
            ok = _ngenie_values_equal(actual, expected)
        elif op in {"ne", "!=", "<>"}:
            ok = not _ngenie_values_equal(actual, expected)
        elif op == "in":
            vals = expected if isinstance(expected, list) else [expected]
            ok = any(_ngenie_values_equal(actual, x) for x in vals)
        elif op in {"not_in", "notin"}:
            vals = expected if isinstance(expected, list) else [expected]
            ok = not any(_ngenie_values_equal(actual, x) for x in vals)
        elif op in {"contains", "has"}:
            if isinstance(actual, (list, tuple, set)):
                ok = any(_ngenie_values_equal(x, expected) for x in actual)
            else:
                ok = str(expected or "").lower() in str(actual or "").lower()
        elif op == "truthy":
            b = _ngenie_bool_value(actual)
            ok = b is True if b is not None else bool(actual)
        elif op == "falsy":
            b = _ngenie_bool_value(actual)
            ok = b is False if b is not None else not bool(actual)
        else:
            ok = False
        if not ok:
            return False
    return True


def _ngenie_normalize_storage_row(row: Any) -> Dict[str, Any]:
    if not isinstance(row, dict):
        return {}
    data = row.get("_data") if isinstance(row.get("_data"), dict) else row
    out = dict(data or {})
    for key in ("_id", "_class", "id", "node_id"):
        if key in row and key not in out:
            out[key] = row.get(key)
    return out


def _ngenie_fetch_real_rows(repo: models.Repo, class_name: str, max_rows: int = 10000) -> Tuple[List[Dict[str, Any]], bool]:
    """Read real nodes for a direct nGenie analysis, with ACL in request context."""
    max_rows = max(1, min(int(max_rows or 10000), 20000))
    page_size = 500
    rows: List[Dict[str, Any]] = []
    truncated = False
    base_url = str(repo.base_url or "").strip().rstrip("/")
    current = ""
    if has_request_context():
        try:
            current = str(request.host_url or "").rstrip("/")
        except Exception:
            current = ""

    storage_path = os.path.join("node_storage", f"{class_name}_{repo.config_uid}.sqlite")
    # Presence of the class storage is authoritative even in timer execution,
    # where there is no Flask request.host_url to compare with repo.base_url.
    is_local = os.path.exists(storage_path) or (not base_url) or (bool(current) and base_url == current)
    if is_local:
        actor = _ngenie_effective_user()
        offset = 0
        while len(rows) < max_rows:
            limit = min(page_size, max_rows - len(rows))
            chunk = _nodes_storage_page(repo.config_uid, class_name, offset=offset, limit=limit, q="")
            if actor is not None:
                try:
                    if not _client_user_can_access_class(repo.config_uid, class_name, user=actor):
                        return [], False
                    chunk = _client_filter_nodes_for_acl(repo.config_uid, class_name, chunk, user=actor)
                except Exception:
                    chunk = []
            if not chunk:
                break
            rows.extend(_ngenie_normalize_storage_row(x) for x in chunk)
            offset += len(chunk)
            if len(chunk) < limit:
                break
        if len(rows) >= max_rows:
            extra = _nodes_storage_page(repo.config_uid, class_name, offset=max_rows, limit=1, q="")
            truncated = bool(extra)
        return [x for x in rows if x], truncated

    # Remote repository. Page through its normal API without relying on request.host.
    offset = 0
    while len(rows) < max_rows:
        limit = min(page_size, max_rows - len(rows))
        try:
            payload = _api_get_remote(
                repo,
                f"/api/config/{repo.config_uid}/node/{class_name}/page",
                params={"offset": offset, "limit": limit},
            )
            chunk = payload.get("items", []) if isinstance(payload, dict) else []
        except Exception:
            chunk = []
        if not isinstance(chunk, list) or not chunk:
            break
        rows.extend(_ngenie_normalize_storage_row(x) for x in chunk)
        offset += len(chunk)
        if len(chunk) < limit:
            break
    truncated = len(rows) >= max_rows
    return [x for x in rows if x], truncated


def _ngenie_class_catalog_for_planner(classes: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    out = []
    for c in classes:
        out.append({
            "class_name": c.get("class_name"),
            "display_name": c.get("display_name"),
            "ngenie_role": c.get("ngenie_role"),
            "ngenie_description": c.get("ngenie_description"),
            "ngenie_prompt": c.get("ngenie_prompt"),
            "fields": [
                {
                    "name": f.get("name"),
                    "label": f.get("label"),
                    "kind": f.get("kind"),
                    "target": f.get("target"),
                }
                for f in (c.get("fields") or [])
            ],
            "tables": [
                {"name": f.get("name"), "label": f.get("label"), "kind": f.get("kind")}
                for f in (c.get("tables") or [])
            ],
        })
    return out


def _ngenie_guess_data_request(prompt: str, classes: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Deterministic safety net when the LLM planner omitted data_requests."""
    text = str(prompt or "").lower()
    words = [w for w in re.findall(r"[a-zа-яё0-9_]+", text, flags=re.I) if len(w) >= 4]
    scored: List[Tuple[int, Dict[str, Any]]] = []
    for cls in classes:
        hay = " ".join(str(cls.get(k) or "") for k in (
            "class_name", "display_name", "ngenie_description", "ngenie_prompt", "ngenie_role"
        )).lower()
        score = 0
        for word in words:
            stem = word[:5]
            if word in hay:
                score += 8
            elif len(stem) >= 4 and stem in hay:
                score += 4
        if str(cls.get("ngenie_role") or "").lower() == "document" and any(x in text for x in ("заказ", "документ")):
            score += 3
        if score:
            scored.append((score, cls))
    if not scored:
        return []
    scored.sort(key=lambda x: x[0], reverse=True)
    cls = scored[0][1]
    field_names = [str(f.get("name") or "") for f in (cls.get("fields") or []) if str(f.get("name") or "")]
    filters: List[Dict[str, Any]] = []
    if "active" in field_names and re.search(r"активн|active", text, flags=re.I):
        filters.append({"field": "active", "op": "eq", "value": True})
    return [{
        "class_name": cls.get("class_name"),
        "filters": filters,
        "fields": field_names,
    }]


def _ngenie_plan_data_summary(prompt: str, config_uid: str, classes: List[Dict[str, Any]]) -> Dict[str, Any]:
    base_ctx = {
        "selected_config_uid": config_uid,
        "scope": "direct",
        "classes": classes,
        "configuration_prompts": _ngenie_collect_config_prompts(config_uid),
    }
    skill_payloads = ngenie_skill_registry.selected_skill_payloads(["data_summary"], base_ctx)
    return _ngenie_call_deepseek([
        {
            "role": "system",
            "content": (
                "Ты планировщик запроса к фактическим узлам NodaLogic. "
                "Верни строго JSON без markdown: {\"data_requests\":[...]}. "
                "На этом шаге нельзя писать summary, projection_method_code, operations или придумывать данные. "
                "class_name и поля должны существовать в переданном каталоге."
            ),
        },
        {
            "role": "user",
            "content": json.dumps({
                "selected_config_uid": config_uid,
                "configuration_prompts": _ngenie_collect_config_prompts(config_uid),
                "classes": _ngenie_class_catalog_for_planner(classes),
                "selected_skills": skill_payloads,
                "user_request": str(prompt or ""),
            }, ensure_ascii=False, default=str),
        },
    ])


def _ngenie_validate_data_requests(
    requests_raw: Any,
    config_uid: str,
    classes: List[Dict[str, Any]],
    lookup: Dict[Tuple[str, str], Tuple[models.Repo, Dict[str, Any]]],
) -> List[Dict[str, Any]]:
    class_meta = {str(c.get("class_name") or ""): c for c in classes}
    out: List[Dict[str, Any]] = []
    if not isinstance(requests_raw, list):
        return out
    for req in requests_raw[:8]:
        if not isinstance(req, dict):
            continue
        repo, _cfg, resolved = _ngenie_find_class(config_uid, req.get("class_name") or "", lookup)
        if not repo or not resolved or resolved not in class_meta:
            continue
        meta = class_meta[resolved]
        declared = {str(f.get("name") or "") for f in (meta.get("fields") or []) if str(f.get("name") or "")}
        declared.update(str(f.get("name") or "") for f in (meta.get("tables") or []) if str(f.get("name") or ""))
        fields = [str(x) for x in (req.get("fields") or []) if str(x) in declared]
        # Class prompts are authoritative field semantics. If they explicitly
        # mention a declared field (for example title), include it even when the
        # planner omitted it. This prevents a generic aggregate from replacing
        # the requested business meaning.
        semantic_text = " ".join([
            str(meta.get("ngenie_description") or ""),
            str(meta.get("ngenie_prompt") or ""),
        ]).lower()
        for field_name in sorted(declared):
            if re.search(rf"(?<![a-zA-Z0-9_]){re.escape(field_name.lower())}(?![a-zA-Z0-9_])", semantic_text) and field_name not in fields:
                fields.append(field_name)
        if not fields:
            fields = sorted(declared)
        filters = []
        for flt in req.get("filters") or []:
            if not isinstance(flt, dict):
                continue
            field = str(flt.get("field") or "").strip()
            op = str(flt.get("op") or "eq").strip().lower()
            if field not in declared or op not in {"eq", "=", "==", "ne", "!=", "<>", "in", "not_in", "notin", "contains", "has", "truthy", "falsy"}:
                continue
            filters.append({"field": field, "op": op, "value": flt.get("value")})
            if field not in fields:
                fields.append(field)
        out.append({
            "repo": repo,
            "class_name": resolved,
            "meta": meta,
            "fields": fields,
            "filters": filters,
        })
    return out


def _ngenie_execute_data_summary_requests(requests_plan: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    results: List[Dict[str, Any]] = []
    for req in requests_plan:
        repo = req["repo"]
        class_name = req["class_name"]
        all_rows, source_truncated = _ngenie_fetch_real_rows(repo, class_name, max_rows=10000)
        matched = [row for row in all_rows if _ngenie_row_matches_filters(row, req.get("filters"))]
        fields = list(req.get("fields") or [])
        # Preserve exact counts while limiting the LLM payload. The final prompt is
        # explicit when the row list was truncated.
        payload_rows = []
        payload_limit = 1000
        for row in matched[:payload_limit]:
            item = {field: _ngenie_data_value(row, field) for field in fields}
            if row.get("_id") is not None:
                item["_id"] = row.get("_id")
            payload_rows.append(item)
        meta = req.get("meta") or {}
        results.append({
            "config_uid": repo.config_uid,
            "class_name": class_name,
            "display_name": meta.get("display_name") or class_name,
            "ngenie_role": meta.get("ngenie_role") or "",
            "ngenie_description": meta.get("ngenie_description") or "",
            "ngenie_prompt": meta.get("ngenie_prompt") or "",
            "filters": req.get("filters") or [],
            "fields": fields,
            "scanned_count": len(all_rows),
            "matched_count": len(matched),
            "source_truncated": bool(source_truncated),
            "rows_truncated": len(matched) > payload_limit,
            "rows": payload_rows,
        })
    return results


def _ngenie_answer_data_requests(answer: Any) -> List[Dict[str, Any]]:
    if not isinstance(answer, dict):
        return []
    raw = answer.get("data_requests") or answer.get("dataRequests") or []
    return [dict(x) for x in raw if isinstance(x, dict)] if isinstance(raw, list) else []


def _ngenie_answer_has_data_action(answer: Any) -> bool:
    if not isinstance(answer, dict):
        return False
    if answer.get("candidate_handler_code") or answer.get("candidateHandlerCode"):
        return True
    if answer.get("operation_handler_code") or answer.get("operationHandlerCode"):
        return True
    if isinstance(answer.get("operations"), list) and any(
        isinstance(x, dict)
        and str(x.get("tool") or x.get("action") or "").strip().lower()
        not in {"", "none", "findall", "find_all", "getall", "get_all"}
        for x in answer.get("operations") or []
    ):
        return True
    return bool(_ngenie_answer_data_requests(answer))


def _ngenie_repair_bulk_data_confirmation(
    messages: List[Dict[str, str]],
    answer: Dict[str, Any],
    user_message: Any,
    config_uid: str = "",
) -> Dict[str, Any]:
    """Repair the loop where the model asks permission to call findAll/get_all.

    Reading accessible rows is an internal planning step, not a user-facing
    permission dialog.  A real destructive action still has to be explicit in
    the user's request; this guard only activates for clear all/every-node bulk
    wording and a reply that asks to confirm a technical read operation.
    """
    if not isinstance(answer, dict) or _ngenie_answer_has_data_action(answer):
        return answer
    request_text = str(user_message or "").strip().lower().replace("ё", "е")
    reply = str(answer.get("reply") or "").strip().lower().replace("ё", "е")
    bulk_intent = bool(re.search(
        r"\bвсе(?:х|м|ми)?\b|\bкажд\w*\b|полный\s+список|весь\s+справочник|"
        r"all\s+(?:nodes|clients|customers|records)|every\s+(?:node|client|record)",
        request_text,
        flags=re.I,
    ))
    pseudo_full_read = any(
        isinstance(x, dict)
        and str(x.get("tool") or x.get("action") or "").strip().lower()
        in {"findall", "find_all", "getall", "get_all"}
        for x in (answer.get("operations") or [])
    ) if isinstance(answer.get("operations"), list) else False
    technical_confirmation = pseudo_full_read or bool(re.search(
        r"подтверд\w*|разреш\w*|могу\s+выполнить|сначала\s+получить|"
        r"findall|get_all|полный\s+список",
        reply,
        flags=re.I,
    ))
    if not (bulk_intent and technical_confirmation):
        return answer
    correction = (
        "Ты зациклился на техническом подтверждении чтения данных. Пользователь уже явно поставил "
        "массовую задачу. Не спрашивай разрешение на findAll/get_all и не используй вымышленный findAll. "
        "Верни полный исправленный JSON для исходной просьбы. Если нужны все существующие узлы класса, "
        "сейчас верни data_requests с class_name, fields и filters. Backend сам прочитает доступные узлы "
        "текущей конфигурации и вызовет второй шаг. После data_results выполни изменение сразу."
    )
    repaired: Dict[str, Any] = answer
    try:
        candidate = _ngenie_call_deepseek(list(messages) + [
            {"role": "assistant", "content": json.dumps(answer, ensure_ascii=False, default=str)},
            {"role": "user", "content": correction},
        ])
        if isinstance(candidate, dict) and candidate:
            repaired = candidate
    except Exception:
        traceback.print_exc()

    if _ngenie_answer_has_data_action(repaired):
        return repaired

    # Do not depend on the model escaping the same repeated sentence.  For an
    # explicit all/every-node request, deterministically select the matching
    # class from the real configuration and start the server-side read step.
    # The second LLM call still decides the actual business values/patches.
    if str(config_uid or "").strip():
        try:
            classes, _lookup = _ngenie_collect_context(config_uid, include_samples=False)
            guessed = _ngenie_guess_data_request(request_text, classes)
            if guessed:
                forced = dict(repaired or answer)
                forced.update({
                    "summary": "",
                    "reply": "",
                    "operations": [],
                    "data_requests": guessed,
                    "resolve_requests": [],
                    "clarification_requests": [],
                    "display_requests": [],
                    "candidate_handler_code": "",
                    "operation_handler_code": "",
                })
                return forced
        except Exception:
            traceback.print_exc()
    return repaired


def _ngenie_build_data_operation_messages(
    user_message: str,
    config_uid: str,
    node_context: Optional[Dict[str, Any]],
    allow_catalog_create: bool,
    scope: str,
    initial_answer: Dict[str, Any],
    data_results: List[Dict[str, Any]],
    attachments: Any = None,
    mobile_mode: bool = False,
) -> List[Dict[str, str]]:
    classes, _lookup = _ngenie_collect_context(config_uid, include_samples=False)
    selected_ids = ngenie_skill_registry.normalize_skill_ids(
        (initial_answer or {}).get("_ngenie_selected_skill_ids")
        or list(getattr(g, "ngenie_selected_skill_ids", []) or [])
        or ["node_operations"]
    )
    ctx = {
        "selected_config_uid": config_uid or "",
        "allow_catalog_create": bool(allow_catalog_create),
        "scope": str(scope or "").strip(),
        "configuration_prompts": _ngenie_collect_config_prompts(config_uid),
        "classes": classes,
        "current_node": node_context or {},
        "attachments": _ngenie_normalize_attachments(attachments),
        "selected_skill_ids": selected_ids,
        "selected_skills": _ngenie_skill_blocks_for_messages(selected_ids, {
            "selected_config_uid": config_uid or "",
            "allow_catalog_create": bool(allow_catalog_create),
            "scope": str(scope or "").strip(),
            "classes": classes,
            "current_node": node_context or {},
            "attachments": _ngenie_normalize_attachments(attachments),
        }),
        "initial_answer": initial_answer or {},
        "data_results": data_results,
        "execution_target": "android" if mobile_mode else "web_backend",
    }
    if mobile_mode:
        final_rules = (
            "Это второй шаг после безопасного чтения фактических узлов на сервере. "
            "data_results уже содержит полный доступный набор и точные _id. Не возвращай data_requests повторно, "
            "не спрашивай подтверждение и не вызывай get_all/findAll на устройстве. Выполни исходную задачу сейчас. "
            "Для изменения данных Android верни один operation_handler_code с def apply_operations(ctx). "
            "Для каждого существующего узла используй точный UID из data_results: node = GetNode(uid), затем измени "
            "node._data и вызови node._save(). Не создавай новые узлы вместо обновления существующих. "
            "Если есть Node-поля, записывай только разрешённые UID, никогда не сохраняй {'query': ...}."
        )
    else:
        final_rules = (
            "Это второй шаг после безопасного чтения фактических узлов. data_results уже содержит реальные строки "
            "и точные _id. Не возвращай data_requests повторно, не спрашивай подтверждение и не используй findAll. "
            "Выполни исходную задачу сейчас. Для изменения нескольких существующих узлов верни одну операцию "
            "bulk_update_nodes: {'tool':'bulk_update_nodes','class_name':'Client','updates':["
            "{'uid':'<точный _id из data_results>','data':{'name':'...'}}]}. "
            "Не меняй _id/_class и не создавай новые узлы вместо обновления существующих. Используй только data_results."
        )
    return [
        {"role": "system", "content": _ngenie_system_prompt()},
        {"role": "user", "content": "Контекст NodaLogic и фактические данные для операции:\n" + json.dumps(ctx, ensure_ascii=False, default=str)},
        {"role": "user", "content": final_rules + "\n\nИсходная просьба пользователя:\n" + str(user_message or "")},
    ]


def _ngenie_finalize_data_requests_with_llm(
    user_message: str,
    config_uid: str,
    node_context: Optional[Dict[str, Any]],
    allow_catalog_create: bool,
    scope: str,
    initial_answer: Dict[str, Any],
    attachments: Any = None,
    mobile_mode: bool = False,
) -> Optional[Dict[str, Any]]:
    raw_requests = _ngenie_answer_data_requests(initial_answer)
    if not raw_requests:
        return None
    classes, lookup = _ngenie_collect_context(config_uid, include_samples=False)
    plan = _ngenie_validate_data_requests(raw_requests, config_uid, classes, lookup)
    if not plan:
        return None
    data_results = _ngenie_execute_data_summary_requests(plan)
    if any(bool(x.get("source_truncated") or x.get("rows_truncated")) for x in data_results):
        return {
            "summary": "",
            "reply": "Массовая операция не выполнена: полный набор данных превышает безопасный лимит nGenie. Уточните отбор.",
            "operations": [],
            "data_requests": [],
            "resolve_requests": [],
            "clarification_requests": [],
            "display_requests": [],
            "candidate_handler_code": "",
            "operation_handler_code": "",
        }
    try:
        messages = _ngenie_build_data_operation_messages(
            user_message, config_uid, node_context, allow_catalog_create, scope,
            initial_answer, data_results, attachments=attachments, mobile_mode=mobile_mode,
        )
        answer = _ngenie_call_deepseek(messages)
        if not isinstance(answer, dict) or not answer:
            return None
        if _ngenie_answer_data_requests(answer) and not _ngenie_answer_has_data_action({
            **answer, "data_requests": []
        }):
            answer = _ngenie_call_deepseek(messages + [
                {"role": "assistant", "content": json.dumps(answer, ensure_ascii=False, default=str)},
                {"role": "user", "content": "data_results уже переданы. Не запрашивай их снова. Верни конечное изменение данных для исходной задачи."},
            ])
        if isinstance(answer, dict) and answer:
            answer["_ngenie_selected_skill_ids"] = ngenie_skill_registry.normalize_skill_ids(
                (initial_answer or {}).get("_ngenie_selected_skill_ids")
                or list(getattr(g, "ngenie_selected_skill_ids", []) or [])
                or ["node_operations"]
            )
            answer["_data_request_result_meta"] = [
                {
                    "class_name": x.get("class_name"),
                    "matched_count": x.get("matched_count"),
                    "fields": x.get("fields") or [],
                }
                for x in data_results
            ]
            return answer
    except Exception:
        traceback.print_exc()
    return None


def _ngenie_generate_data_summary(prompt: str, config_uid: str) -> str:
    classes, lookup = _ngenie_collect_context(config_uid, include_samples=False)
    if not classes:
        try:
            print("[nGenie direct summary] no data classes", json.dumps({
                "config_uid": config_uid,
                "effective_user_id": _ngenie_effective_user_id(),
                "repos": [int(r.id) for r in _ngenie_all_repos(config_uid)],
            }, ensure_ascii=False))
        except Exception:
            pass
        return "В текущей конфигурации нет доступных классов данных."

    planner_answer = _ngenie_plan_data_summary(prompt, config_uid, classes)
    requests_raw = planner_answer.get("data_requests") if isinstance(planner_answer, dict) else None
    if not requests_raw:
        requests_raw = _ngenie_guess_data_request(prompt, classes)
    plan = _ngenie_validate_data_requests(requests_raw, config_uid, classes, lookup)
    if not plan:
        return "nGenie не смог сопоставить запрос с классом и полями текущей конфигурации. Проверьте ngenie_description/ngenie_prompt класса."

    data_results = _ngenie_execute_data_summary_requests(plan)
    try:
        print("[nGenie direct summary]", json.dumps({
            "config_uid": config_uid,
            "requests": [
                {
                    "class_name": x.get("class_name"),
                    "filters": x.get("filters") or [],
                    "scanned_count": x.get("scanned_count"),
                    "matched_count": x.get("matched_count"),
                    "source_truncated": x.get("source_truncated"),
                    "rows_truncated": x.get("rows_truncated"),
                }
                for x in data_results
            ],
        }, ensure_ascii=False, default=str))
    except Exception:
        pass
    answer = _ngenie_call_deepseek([
        {
            "role": "system",
            "content": (
                "Ты формируешь итог по уже прочитанным фактическим узлам NodaLogic. "
                "Верни строго JSON вида {\"summary\":\"...\"} и никаких других ключей. "
                "Используй только data_results. Не придумывай поля, статусы, суммы, средние, максимумы или другие факты. "
                "Смысл полей бери из ngenie_description/ngenie_prompt класса. "
                "Если rows_truncated или source_truncated=true, честно укажи, что сводка по переданной части данных; "
                "не выдавай неполную выборку за полную. Если matched_count=0, сообщи, что подходящих узлов не найдено."
            ),
        },
        {
            "role": "user",
            "content": json.dumps({
                "user_request": str(prompt or ""),
                "selected_config_uid": config_uid,
                "data_results": data_results,
            }, ensure_ascii=False, default=str),
        },
    ])
    value = answer.get("summary") if isinstance(answer, dict) else None
    if _ngenie_value_is_meaningful(value):
        return value if isinstance(value, str) else json.dumps(value, ensure_ascii=False)
    return "nGenie прочитал данные, но не смог сформировать сводку."



def _ngenie_validate_image_data_url(value: Any) -> str:
    raw = str(value or "").strip()
    match = re.match(r"^data:(image/(?:jpeg|jpg|png|webp));base64,([A-Za-z0-9+/=\s]+)$", raw, flags=re.I)
    if not match:
        raise ValueError("nGenie image must be a JPEG, PNG, or WebP data URL")
    try:
        decoded = base64.b64decode(re.sub(r"\s+", "", match.group(2)), validate=True)
    except Exception as exc:
        raise ValueError("invalid nGenie image base64") from exc
    if not decoded:
        raise ValueError("empty nGenie image")
    if len(decoded) > 8 * 1024 * 1024:
        raise ValueError("nGenie image is larger than 8 MiB after compression")
    mime = match.group(1).lower().replace("image/jpg", "image/jpeg")
    return "data:" + mime + ";base64," + base64.b64encode(decoded).decode("ascii")


def _ngenie_local_image_data_url(file_path: Any) -> str:
    path = str(file_path or "").strip()
    if not path or not os.path.isfile(path):
        return ""
    mime = str(mimetypes.guess_type(path)[0] or "").lower()
    if mime not in {"image/jpeg", "image/png", "image/webp"}:
        return ""
    with open(path, "rb") as fh:
        raw = fh.read(8 * 1024 * 1024 + 1)
    if len(raw) > 8 * 1024 * 1024:
        raise ValueError("nGenie image is larger than 8 MiB; compress it before calling ngenie()")
    return _ngenie_validate_image_data_url(
        "data:" + mime + ";base64," + base64.b64encode(raw).decode("ascii")
    )


def _ngenie_vision_credentials() -> Tuple[str, str, str]:
    key = str(getattr(main, "NGENIE_VISION_API_KEY", "") or os.environ.get("NGENIE_VISION_API_KEY", "")).strip()
    url = str(getattr(main, "NGENIE_VISION_API_URL", "") or os.environ.get("NGENIE_VISION_API_URL", "")).strip()
    model = str(getattr(main, "NGENIE_VISION_MODEL", "") or os.environ.get("NGENIE_VISION_MODEL", "")).strip()

    # Reuse the existing nGenie Code provider configuration when a dedicated
    # runtime vision configuration is not present.  The file may optionally
    # define vision_api_key / vision_url / vision_model.
    if not (key and url and model):
        try:
            credentials_path = Path(__file__).resolve().parent.parent / "credentials.json"
            raw = json.loads(credentials_path.read_text(encoding="utf-8")) if credentials_path.is_file() else {}
        except Exception:
            raw = {}
        key = key or str(raw.get("vision_api_key") or raw.get("deepseek_api_key") or "").strip()
        url = url or str(raw.get("vision_url") or raw.get("deepseek_url") or "").strip()
        model = model or str(raw.get("vision_model") or "").strip()

    if not key or not url or not model:
        raise RuntimeError(
            "nGenie image analysis is not configured: set NGENIE_VISION_API_KEY, "
            "NGENIE_VISION_API_URL and NGENIE_VISION_MODEL (or vision_* in credentials.json in the project root)"
        )
    return key, url, model


def _ngenie_vision_result_string(content: Any) -> str:
    text = str(content or "").strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json|text)?\s*", "", text, flags=re.I).strip()
        text = re.sub(r"\s*```$", "", text).strip()
    try:
        parsed = json.loads(text)
    except Exception:
        return text
    if isinstance(parsed, dict):
        meaningful = [(k, v) for k, v in parsed.items() if _ngenie_value_is_meaningful(v)]
        if len(meaningful) == 1 and meaningful[0][0] in {"result", "answer", "reply", "summary", "count", "value"}:
            value = meaningful[0][1]
            return value if isinstance(value, str) else json.dumps(value, ensure_ascii=False)
    return json.dumps(parsed, ensure_ascii=False)


def _ngenie_generate_vision(prompt: str, image_data_url: str, config_uid: str = "") -> str:
    key, url, model = _ngenie_vision_credentials()
    system_text = (
        "Analyze the attached image and return only the result requested by the user. "
        "Do not describe your reasoning and do not add Markdown fences."
    )
    config_prompts = _ngenie_collect_config_prompts(str(config_uid or "").strip()) if config_uid else []
    if config_prompts:
        system_text += " Configuration instructions: " + " | ".join(
            str(item.get("prompt") or "").strip() for item in config_prompts if str(item.get("prompt") or "").strip()
        )
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_text},
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": str(prompt or "")},
                    {"type": "image_url", "image_url": {"url": image_data_url}},
                ],
            },
        ],
        "temperature": 0.1,
        "max_tokens": 512,
        "stream": False,
    }
    resp = requests.post(
        url,
        headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
        json=payload,
        timeout=120,
    )
    resp.raise_for_status()
    data = resp.json()
    content = (((data.get("choices") or [{}])[0].get("message") or {}).get("content") or "")
    if not content:
        raise RuntimeError("nGenie vision provider returned an empty answer")
    return _ngenie_vision_result_string(content)


def _ngenie_generate_text(prompt: str, file_path=None, config_uid: str = "", image_data_url: Any = None):
    message = str(prompt or "").strip()
    image_payload = ""
    if image_data_url:
        image_payload = _ngenie_validate_image_data_url(image_data_url)
    elif file_path:
        image_payload = _ngenie_local_image_data_url(file_path)

    if image_payload:
        return _ngenie_generate_vision(message, image_payload, str(config_uid or "").strip())

    if file_path:
        try:
            with open(str(file_path), "rb") as fh:
                raw = fh.read(2 * 1024 * 1024)
            message += "\n\nAttached file content:\n" + raw.decode("utf-8", errors="replace")
        except Exception as exc:
            message += f"\n\nAttachment could not be read: {exc}"

    config_uid = str(config_uid or "").strip()
    if config_uid and _ngenie_prompt_requests_summary_only(message):
        # A summary must be based on real rows, not on examples or LLM-generated
        # projection code. Plan a validated class/filter request, read the current
        # configuration's storage, then ask the model only to summarize those rows.
        return _ngenie_generate_data_summary(message, config_uid)
    if config_uid:
        # Non-summary direct calls keep the general skill-router flow.
        messages = _ngenie_build_messages(message, config_uid, scope="direct")
        result = _ngenie_call_deepseek(messages)
    else:
        messages = [
            {"role": "system", "content": "Return only the requested result. It may be plain text, HTML, or JSON. Do not wrap it in Markdown fences unless explicitly requested."},
            {"role": "user", "content": message},
        ]
        result = _ngenie_call_deepseek(messages)

    return _ngenie_direct_result_string(result, message, config_uid)

try:
    _nodes_mod.set_ngenie_runner(_ngenie_generate_text)
except Exception:
    pass

@client_bp.route("/api/ngenie/generate", methods=["POST"])
@_ngenie_api_auth_required
def api_ngenie_generate():
    j = request.get_json(force=True) or {}
    prompt = str(j.get("prompt") or j.get("message") or "").strip()
    if not prompt:
        return jsonify({"ok": False, "error": "empty prompt"}), 400
    try:
        return jsonify({"ok": True, "result": _ngenie_generate_text(
            prompt,
            None,
            str(j.get("config_uid") or ""),
            image_data_url=j.get("image_data_url") or j.get("image"),
        )})
    except Exception as exc:
        traceback.print_exc()
        return jsonify({"ok": False, "error": str(exc)}), 500

@client_bp.route("/api/ngenie/mobile/plan", methods=["POST"])
@_ngenie_api_auth_required
def api_ngenie_mobile_plan():
    j = request.get_json(force=True) or {}
    message = str(j.get("message") or "").strip()
    config_uid = str(j.get("config_uid") or "").strip()
    node_context = j.get("node_context") if isinstance(j.get("node_context"), dict) else None
    allow_catalog_create = bool(j.get("allow_catalog_create"))
    clarification_response = j.get("clarification_response") if isinstance(j.get("clarification_response"), dict) else None
    scope = str(j.get("scope") or "").strip()
    client_capabilities = j.get("client_capabilities") if isinstance(j.get("client_capabilities"), dict) else {}
    if not message:
        return jsonify({"ok": False, "error": "empty message"}), 400
    if not config_uid:
        return jsonify({"ok": False, "error": "empty config_uid"}), 400
    try:
        res = ngenie_core.build_mobile_plan(
            rt=sys.modules[__name__],
            message=message,
            config_uid=config_uid,
            node_context=node_context,
            allow_catalog_create=allow_catalog_create,
            clarification_response=clarification_response,
            scope=scope,
            debug=bool(j.get("debug")),
            client_capabilities=client_capabilities,
        )
        return jsonify(res)
    except Exception as exc:
        traceback.print_exc()
        return jsonify({"ok": False, "error": str(exc)}), 500


@client_bp.route("/api/ngenie/mobile/finalize", methods=["POST"])
@_ngenie_api_auth_required
def api_ngenie_mobile_finalize():
    j = request.get_json(force=True) or {}
    state = j.get("state") if isinstance(j.get("state"), dict) else {}
    handler_result = (
        j.get("candidate_handler_result")
        or j.get("handler_result")
        or j.get("result")
        or {}
    )
    if not isinstance(handler_result, dict):
        handler_result = {}
    try:
        res = ngenie_core.finalize_mobile_plan(
            rt=sys.modules[__name__],
            state=state,
            handler_result=handler_result,
            debug=bool(j.get("debug")),
        )
        return jsonify(res)
    except Exception as exc:
        traceback.print_exc()
        return jsonify({"ok": False, "error": str(exc)}), 500


@client_bp.route("/api/ngenie/save-html-projection", methods=["POST"])
@login_required
def api_ngenie_save_html_projection():
    j = request.get_json(force=True) or {}
    config_uid = str(j.get("config_uid") or "").strip()
    title = str(j.get("title") or j.get("projection_title") or "nGenie HTML").strip()
    method_code = _ngenie_strip_method_code(j.get("method_code") or j.get("projection_method_code") or "")
    html = str(j.get("html") or "")
    parameters_layout = _ngenie_normalize_projection_parameters_layout(
        j.get("projection_parameters_layout")
        if "projection_parameters_layout" in j
        else j.get("parameters_layout")
    )
    parameters_data = j.get("projection_parameters_data") or j.get("parameters_data") or {}
    if not isinstance(parameters_data, dict):
        parameters_data = {}
    parameters_data_structure = str(
        j.get("projection_parameters_data_structure")
        or j.get("parameters_data_structure")
        or ""
    ).strip()
    if parameters_layout is not None and not parameters_data_structure:
        parameters_data_structure = _ngenie_projection_data_structure_from_layout(parameters_layout)
    if not config_uid:
        return jsonify({"ok": False, "error": "select a configuration first"}), 400
    repo = models.Repo.query.filter_by(user_id=current_user.id, config_uid=config_uid).first()
    if not repo:
        return jsonify({"ok": False, "error": "configuration not found"}), 404
    root = _ngenie_root_models()
    config = root.Configuration.query.filter_by(uid=config_uid, user_id=current_user.id).first()
    if not config:
        return jsonify({"ok": False, "error": "configuration is not local or is not editable by this user"}), 404
    if not method_code:
        method_code = _ngenie_make_static_projection_code(title, html)
    base_name = _ngenie_slug_class_name(title)
    existing = {c.name for c in (config.classes or [])}
    class_name = base_name
    i = 2
    while class_name in existing:
        class_name = f"{base_name}_{i}"
        i += 1
    cls = root.ConfigClass(
        name=class_name,
        config=config,
        class_type=PROJECTION_CLASS_TYPE,
        projection_type=PROJECTION_HTML_TYPE,
        has_storage=False,
        display_name=title,
        record_view="@_projection_header",
        hidden=False,
        use_standard_commands=True,
    )
    if parameters_layout is not None:
        layout_json = json.dumps(parameters_layout, ensure_ascii=False)
        # This is the ordinary Projection settings panel used by node_form:
        # output on the left, NodaLogic layout and Generate button on the right.
        cls.init_screen_layout = layout_json
        cls.init_screen_layout_web = layout_json
        cls.data_structure = parameters_data_structure or ""
    cls.methods.append(root.ClassMethod(name="onRunProjection", source="internal", engine="server_python", code="onRunProjection"))
    ev = root.ClassEvent(event="onInputWeb", listener="onRunProjection")
    ev.actions.append(root.EventAction(action="run", source="internal", method="onRunProjection", order=1))
    cls.event_objs.append(ev)
    config.classes.append(cls)
    code = _ngenie_ensure_handlers_method(config, class_name, method_code)
    config.last_modified = datetime.now(timezone.utc)
    root.db.session.add(config)
    root.db.session.commit()
    _ngenie_write_handlers_file(config_uid, code)
    try:
        cfg_json = fetch_config_from_local_db(config_uid)
        rc = models.RepoConfig.query.filter_by(repo_id=repo.id).first()
        if rc:
            rc.config_json = json.dumps(cfg_json, ensure_ascii=False)
            rc.updated_at = datetime.now(timezone.utc)
            models.db.session.commit()
        _invalidate_repo_config_mem(repo.id)
    except Exception:
        traceback.print_exc()
    return jsonify({
        "ok": True,
        "class_name": class_name,
        "display_name": title,
        "parameters_saved": parameters_layout is not None,
        "projection_parameters_data": parameters_data,
        "open_url": url_for("client.node_form", config_uid=config_uid, class_name=class_name, node_id=f"{config_uid}${class_name}$singleton"),
    })

# ---------- UI routes ----------

@client_bp.app_context_processor
def _inject_globals():
    # Scanner WS settings are used by base.html to auto-connect.
    # Stored per-user in client.sqlite.
    scanner_ws_url = _get_setting("scanner_ws_url", "").strip()
    scanner_ws_enabled = (_get_setting("scanner_ws_enabled", "1").strip() or "1")
    scanner_ws_enabled = (scanner_ws_enabled not in ("0", "false", "False", "no", "off"))
    return {
        "nl_css": DEFAULT_NL_CSS,
        "client_app_title": APP_TITLE,
        "scanner_ws_url": scanner_ws_url,
        "scanner_ws_enabled": scanner_ws_enabled,
        "ngenie_enabled": _client_ngenie_enabled(),
    }


@client_bp.route("/")
@login_required
def home():
    repos = models.Repo.query.filter_by(user_id=_ngenie_effective_user_id()).all()
    sections = _with_received_nodes_section(build_global_sections(repos, models.db), repos)
    section_code = request.args.get("section", None)
    if section_code is None:
        section_code = _default_section_code(sections)
    scode_url = section_code if section_code != "" else "__empty__"
    return redirect(url_for("client.section_view", section_code=scode_url))


@client_bp.route("/sections")
@login_required
def sections_home():
    repos = models.Repo.query.filter_by(user_id=_ngenie_effective_user_id()).all()
    sections = _with_received_nodes_section(build_global_sections(repos, models.db), repos) if repos else []
    if not sections:
        return render_template(
            "client/section.html",
            title=f"{APP_TITLE} — Sections",
            repos=repos,
            sections=[],
            section_code="",
            section_name="",
            auto_refresh=AUTO_REFRESH_SECONDS,
            no_repos=(len(repos) == 0),
        )
    first = _default_section_code(sections)
    scode_url = first if first != "" else "__empty__"
    return redirect(url_for("client.section_view", section_code=scode_url))


def _get_repo_or_404(repo_id: int) -> models.Repo:
    repo = models.Repo.query.filter_by(id=repo_id, user_id=current_user.id).first()
    if not repo:
        abort(404)
    return repo


def _get_class_cfg(repo: models.Repo, class_name: str) -> Optional[Dict[str, Any]]:
    parsed = get_parsed_config(repo, models.db)
    if not parsed:
        return None
    return (parsed.get("classes") or {}).get(class_name)




def _node_local_get_data(config_uid: str, class_name: str, node_id: str) -> Dict[str, Any]:
    node_class = _load_server_node_class(config_uid, class_name)
    node = node_class.get(node_id, config_uid)
    if not node:
        return {}
    return node.get_data() or {}


def _node_local_update_data(config_uid: str, class_name: str, node_id: str, data: Dict[str, Any], user_modification: Optional[Dict[str, Any]] = None):
    node_class = _load_server_node_class(config_uid, class_name)
    node = node_class.get(node_id, config_uid)
    if not node:
        raise ValueError("node not found")

    try:
        node._schema_class_name = class_name
    except Exception:
        pass

    merged = dict(data or {})
    merged.setdefault("_class", class_name)

    # Full-form saves should keep replacement semantics, but validation/events
    # must still see the real previous saved_state.  Prime _data_cache instead
    # of saving before update_data(); Node.update_data() will run onAcceptServer,
    # persist, and then run onAfterAcceptServer.
    try:
        node._data_cache = dict(merged)
    except Exception:
        pass

    input_data = dict(merged)
    if isinstance(user_modification, dict) and user_modification:
        input_data["_user_modification"] = user_modification

    node.update_data(input_data)
    return node

def _node_local_delete(config_uid: str, class_name: str, node_id: str) -> None:
    node_class = _load_server_node_class(config_uid, class_name)
    node = node_class.get(node_id, config_uid)
    if not node:
        return

    node.delete()


def _node_local_create(config_uid: str, class_name: str, initial_data: Optional[Dict[str, Any]] = None) -> str:

    node_class = _load_server_node_class(config_uid, class_name)

    data = initial_data or {}
    
    user_data = dict(data or {})

    #node_id = (data.get("_id") or str(uuid.uuid4()))
    #node = node_class(node_id, config_uid)
    raw_id = data.get("_id")
    node_id = _nodes_mod.extract_internal_id(raw_id) if raw_id else str(uuid.uuid4())
    node = node_class(node_id, config_uid)
    if user_data:
        # update_data() already persists the node and runs onAcceptServer/onAfterAcceptServer.
        # Calling _save() again here would duplicate hooks and, for schedule cell creation,
        # could lose the original _user_modification payload.
        node.update_data(user_data)
    return node_id


def _register_nodes_to_room_local(config_uid: str, class_name: str, room_uid: str, node_ids: List[str]) -> Tuple[int, Dict[str, Any]]:
    """Register selected nodes in a room locally (without HTTP self-calls).

    Mirrors: /api/config/<config_uid>/node/<class_name>/register/<room_uid>
    but uses direct python calls to avoid deadlocks.
    """
    room_uid = (room_uid or "").strip()
    if not room_uid:
        return 0, {"ok": False, "error": "room uid is empty"}

    node_class = _load_server_node_class(config_uid, class_name)
    nodes_data: List[Dict[str, Any]] = []

    for nid in (node_ids or []):
        nid = str(nid or "").strip()
        if not nid:
            continue
        try:
            node = node_class.get(nid, config_uid)
        except Exception:
            node = None
        if not node:
            continue
        try:
            d = node.to_dict() if hasattr(node, "to_dict") else {}
        except Exception:
            d = {}
        if not isinstance(d, dict):
            d = {}
        d.setdefault("_id", nid)
        nodes_data.append(d)

    if not nodes_data:
        return 0, {"ok": False, "error": "nodes not found"}

    # Delegate to server helper that queues objects in the room
    try:
        rv = main.handle_room_objects(config_uid, class_name, room_uid, nodes_data)
        response_obj = rv[0] if isinstance(rv, tuple) and rv else rv
        result: Dict[str, Any] = {}
        if hasattr(response_obj, "get_json"):
            result = response_obj.get_json(silent=True) or {}
        elif isinstance(response_obj, dict):
            result = dict(response_obj)
        return len(nodes_data), result
    except Exception as e:
        # if it fails, still return count=0 to the caller
        return 0, {"ok": False, "error": str(e)}


def _resolve_class_default_room_uid(parsed: Optional[Dict[str, Any]], cls_cfg: Dict[str, Any]) -> str:
    """Resolve default room uid for Migration registration.

    New style: class stores migration_default_room_alias, mapping stored in cfg['rooms'].
    Backward compatible: falls back to migration_default_room_uid.
    """
    try:
        alias = str(cls_cfg.get("migration_default_room_alias") or "").strip()
        if alias:
            rooms_map = (parsed or {}).get("rooms") if isinstance(parsed, dict) else None
            if isinstance(rooms_map, dict):
                ru = str(rooms_map.get(alias) or "").strip()
                if ru:
                    return ru
        # fallback
        return str(cls_cfg.get("migration_default_room_uid") or "").strip()
    except Exception:
        return str(cls_cfg.get("migration_default_room_uid") or "").strip()


def _normalize_custom_process_uid(config_uid: str, class_name: str, node_id: str) -> str:
    """Ensure custom_process uid is always 3-part: cfg$Class$singleton.

    Backward compatible with older 2-part form: cfg$Class.
    """
    config_uid = str(config_uid or "").strip()
    class_name = str(class_name or "").strip()
    raw = str(node_id or "").strip()
    if not config_uid or not class_name:
        return raw

    parts = raw.split("$") if raw else []

    # Already normalized: cfg$Class$something
    if len(parts) >= 3 and parts[0] == config_uid and parts[1] == class_name:
        return raw

    # Old form: cfg$Class
    if len(parts) == 2 and parts[0] == config_uid and parts[1] == class_name:
        try:
            return _nodes_mod.normalize_own_uid(config_uid, class_name, "singleton")
        except Exception:
            return f"{config_uid}${class_name}$singleton"

    # Fallback: force singleton
    try:
        return _nodes_mod.normalize_own_uid(config_uid, class_name, "singleton")
    except Exception:
        return f"{config_uid}${class_name}$singleton"
    
def _node_local_upsert_custom_process(config_uid: str, class_name: str, node_id: str, data: Dict[str, Any]) -> str:
    """Create (if missing) and save a custom_process singleton node locally."""
    node_uid = _normalize_custom_process_uid(config_uid, class_name, node_id)
    node_class = _load_server_node_class(config_uid, class_name)

    node = None
    try:
        node = node_class.get(node_uid, config_uid)
    except Exception:
        node = None

    if not node:
        try:
            internal_id = _nodes_mod.extract_internal_id(node_uid) or "singleton"
        except Exception:
            internal_id = "singleton"
        node = node_class(internal_id, config_uid)

    merged = dict(data or {})
    merged.setdefault("_class", class_name)
    merged.setdefault("_id", node_uid)

    try:
        node._data = merged
    except Exception:
        node.update_data(merged)

    try:
        node._schema_class_name = class_name
    except Exception:
        pass

    if hasattr(node, "_save") and callable(getattr(node, "_save")):
        node._save()
    return node_uid    


def _parse_projection_kanban_columns(cls_cfg: Dict[str, Any]) -> List[Dict[str, str]]:
    raw = (cls_cfg or {}).get("projection_kanban_columns") or (cls_cfg or {}).get("_kanban_columns") or []
    obj = []
    if isinstance(raw, list):
        obj = raw
    elif isinstance(raw, str):
        try:
            parsed = json.loads(raw) if raw.strip() else []
            obj = parsed if isinstance(parsed, list) else []
        except Exception:
            obj = []
    out: List[Dict[str, str]] = []
    seen = set()
    for col in obj:
        if not isinstance(col, dict):
            continue
        cid = str(col.get("id") or col.get("key") or "").strip()
        if not cid or cid in seen:
            continue
        seen.add(cid)
        out.append({"id": cid, "caption": str(col.get("caption") or col.get("title") or cid)})
    return out


def _apply_projection_defaults_to_data(cls_cfg: Dict[str, Any], data: Dict[str, Any], config_uid: str, class_name: str, node_id: str) -> None:
    if not isinstance(data, dict) or not _is_projection_class_type(cls_cfg):
        return
    projection_type = str((cls_cfg or {}).get("projection_type") or PROJECTION_KANBAN_TYPE).strip() or PROJECTION_KANBAN_TYPE
    data.setdefault("_projection_type", projection_type)
    data.setdefault("_projection_uid", _normalize_custom_process_uid(config_uid, class_name, node_id))
    if projection_type == PROJECTION_KANBAN_TYPE:
        # Columns are a class-level projection setting.  Older projection nodes may
        # already have _kanban_columns stored in their own _data; if we keep that
        # value, editing the class from 3 columns to 5 columns still renders the
        # old 3-column snapshot.  When the class config contains a columns field,
        # always re-sync it from the class; only preserve node/client-supplied
        # columns for legacy/raw projections where the class has no such field.
        class_has_columns_field = ("projection_kanban_columns" in (cls_cfg or {})) or ("_kanban_columns" in (cls_cfg or {}))
        if class_has_columns_field or "_kanban_columns" not in data:
            data["_kanban_columns"] = _parse_projection_kanban_columns(cls_cfg)



def _normalize_print_targets(raw: Any) -> List[str]:
    if raw is None:
        return []
    if isinstance(raw, list):
        return [str(x).strip() for x in raw if str(x or '').strip()]
    s = str(raw or '').strip()
    if not s:
        return []
    try:
        parsed = json.loads(s)
        if isinstance(parsed, list):
            return [str(x).strip() for x in parsed if str(x or '').strip()]
    except Exception:
        pass
    return [x.strip() for x in re.split(r"[,;\n]+", s) if x.strip()]


def _print_forms_for_class(parsed: Optional[Dict[str, Any]], class_name: str) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    classes = (parsed or {}).get("classes") or {}
    for pf in classes.values():
        if not isinstance(pf, dict) or not _is_print_form_class_type(pf):
            continue
        targets = _normalize_print_targets(pf.get("print_target_classes") or pf.get("printTargetClasses"))
        if class_name in targets:
            out.append(pf)
    out.sort(key=lambda x: str(x.get("display_name") or x.get("name") or "").lower())
    return out


def _print_qr_data_url(value: Any) -> str:
    if qrcode is None:
        return ""
    try:
        img = qrcode.make(str(value or ""))
        buf = BytesIO()
        img.save(buf, format="PNG")
        return "data:image/png;base64," + base64.b64encode(buf.getvalue()).decode("ascii")
    except Exception:
        return ""


def _print_barcode_data_url(value: Any, bar_height: Any = 44, bar_width: Any = 1.15, human_readable: Any = False) -> str:
    """Return a Code128 barcode as an SVG data URL for PrintForm templates.

    SVG is used first because it does not depend on ReportLab's optional
    renderPM/Cairo PNG backend. Unicode values are encoded as UTF-8 bytes via
    Code128 FNC4; plain ASCII still uses ReportLab's standard compact encoder.
    """
    raw = str(value or "").strip()
    if not raw:
        return ""
    if createBarcodeDrawing is None:
        print(
            "PrintForm Code128 is unavailable: "
            f"{type(_REPORTLAB_IMPORT_ERROR).__name__}: {_REPORTLAB_IMPORT_ERROR}",
            file=sys.stderr,
        )
        return ""

    height = max(8.0, min(float(bar_height or 44), 300.0))
    width = max(0.2, min(float(bar_width or 1.15), 4.0))
    drawing = None
    try:
        if all(ord(char) <= 127 for char in raw):
            drawing = createBarcodeDrawing(
                "Code128",
                value=raw,
                barHeight=height,
                barWidth=width,
                humanReadable=bool(human_readable),
            )
        elif _Utf8BarcodeCode128 is not None and _ReportLabDrawing is not None:
            barcode_widget = _Utf8BarcodeCode128(
                value=raw,
                barHeight=height,
                barWidth=width,
                humanReadable=bool(human_readable),
            )
            barcode_widget.validate()
            x1, y1, x2, y2 = barcode_widget.getBounds()
            drawing = _ReportLabDrawing(
                width=float(x2 - x1),
                height=float(y2 - y1),
                transform=[1, 0, 0, 1, -float(x1), -float(y1)],
            )
            drawing.add(barcode_widget, "_bc")
        else:
            return ""

        svg = drawing.asString("svg")
        if isinstance(svg, str):
            svg = svg.encode("utf-8")
        return "data:image/svg+xml;base64," + base64.b64encode(svg).decode("ascii")
    except Exception as svg_error:
        # Keep a PNG fallback for installations where SVG export is unavailable.
        if drawing is not None:
            try:
                png = drawing.asString("png")
                return "data:image/png;base64," + base64.b64encode(png).decode("ascii")
            except Exception as png_error:
                print(
                    "PrintForm Code128 render failed: "
                    f"value={raw[:120]!r}; SVG={type(svg_error).__name__}: {svg_error}; "
                    f"PNG={type(png_error).__name__}: {png_error}",
                    file=sys.stderr,
                )
                return ""
        print(
            f"PrintForm Code128 build failed: value={raw[:120]!r}; "
            f"{type(svg_error).__name__}: {svg_error}",
            file=sys.stderr,
        )
        return ""


def _print_image_src(repo: models.Repo, value: Any) -> str:
    raw = str(value or "").strip()
    if not raw:
        return ""
    if raw.startswith("data:") or raw.startswith("blob:") or raw.startswith("http://") or raw.startswith("https://"):
        return raw
    return url_for("client.api_userfiles_raw", repo_id=repo.id, filename=raw)


def _print_table_rows(value: Any) -> List[Any]:
    if isinstance(value, list):
        return value
    if isinstance(value, dict):
        return list(value.values())
    return []


def _build_print_jinja_context(repo: models.Repo, data: Dict[str, Any]) -> Dict[str, Any]:
    data = data if isinstance(data, dict) else {}
    wrapped_data = _print_attr_tree(data)
    ctx: Dict[str, Any] = {
        "_data": wrapped_data,
        "data": wrapped_data,
        "qr": _print_qr_data_url,
        "barcode": _print_barcode_data_url,
        "image_src": lambda value: _print_image_src(repo, value),
        "table_rows": _print_table_rows,
    }
    for k, v in data.items():
        key = str(k or "")
        if not key:
            continue
        wrapped_value = _print_attr_tree(v)
        ctx[key] = wrapped_value
        ctx["_" + key.lstrip("_")] = wrapped_value
    return ctx


def _render_print_html(repo: models.Repo, print_cls: Dict[str, Any], data: Dict[str, Any]) -> str:
    template_type = str((print_cls or {}).get("print_template_type") or PRINT_FORM_TEMPLATE_HTML_JINJA).strip()
    html_template = _decode_print_html_template((print_cls or {}).get("print_html_template") or "")
    if template_type != PRINT_FORM_TEMPLATE_HTML_JINJA:
        return f"<div class='alert alert-warning'>Unsupported PrintForm template type: {escape(template_type)}</div>"
    try:
        env = _PrintSandboxedEnvironment(autoescape=select_autoescape(["html", "xml"]))
        env.globals.update(
            qr=_print_qr_data_url,
            barcode=_print_barcode_data_url,
            image_src=lambda value: _print_image_src(repo, value),
            table_rows=_print_table_rows,
        )
        return env.from_string(html_template).render(**_build_print_jinja_context(repo, data))
    except Exception as e:
        return f"<div class='alert alert-danger'>PrintForm render error: {escape(str(e))}</div>"


def _print_form_node_id(config_uid: str, print_class_name: str, base_class_name: str, base_node_id: str) -> str:
    digest = hashlib.sha1(f"{base_class_name}:{base_node_id}".encode("utf-8", errors="ignore")).hexdigest()[:16]
    return _nodes_mod.normalize_own_uid(config_uid, print_class_name, f"print_{digest}")


def _execute_print_form_start_handler(repo: models.Repo, parsed: Dict[str, Any], print_cls: Dict[str, Any], print_node_id: str, base_class_name: str, base_node_id: str, base_data: Dict[str, Any]) -> Dict[str, Any]:
    """Create an ephemeral PrintForm node, inject _basement_data, run onInputWeb/onStartForm, return _data."""
    print_class_name = str((print_cls or {}).get("name") or "").strip()
    node_data: Dict[str, Any] = {
        "_id": print_node_id,
        "_class": print_class_name,
        "_basement_class": base_class_name,
        "_basement_id": base_node_id,
        "_basement_data": dict(base_data or {}),
    }

    actions: List[Dict[str, Any]] = []
    for ev in (print_cls.get("events") or []):
        if (ev.get("event") or "") not in ("onInputWeb", "onInput"):
            continue
        listener = str(ev.get("listener") or "").strip()
        if listener and listener != "onStartForm":
            continue
        # Prefer the explicit web event, but keep onInput as a fallback for older configs.
        if (ev.get("event") or "") == "onInputWeb" or not actions:
            actions.extend(ev.get("actions") or [])

    if not print_class_name:
        return node_data

    base_url = (repo.base_url or "").strip().rstrip("/")
    current = (request.host_url or "").rstrip("/")

    try:
        if base_url and base_url != current:
            # Remote PrintForm rendering can still use the template in cached configuration.
            # If handlers are remote, call selected methods and merge returned data when the API supports it.
            payload = {"listener": "onStartForm", "_basement_data": base_data, "base_class_name": base_class_name, "base_node_id": base_node_id}
            for a in actions:
                m = str((a or {}).get("method") or "").strip()
                if not m:
                    continue
                try:
                    r = _api_post_remote(repo, f"/api/config/{repo.config_uid}/node/{print_class_name}/{print_node_id}/{m}", json_data=payload)
                    if isinstance(r, dict) and isinstance(r.get("data"), dict):
                        data = r.get("data") or {}
                        if isinstance(data.get("_data"), dict):
                            node_data.update(data.get("_data") or {})
                        else:
                            node_data.update(data)
                except Exception:
                    continue
            node_data.setdefault("_basement_data", dict(base_data or {}))
            return node_data

        node_class = _load_server_node_class(repo.config_uid, print_class_name)
        node = node_class(_nodes_mod.extract_internal_id(print_node_id) or print_node_id, repo.config_uid)
        try:
            node._schema_class_name = print_class_name
        except Exception:
            pass
        try:
            node._data_cache = dict(node_data)
            node._data = dict(node_data)
        except Exception:
            pass

        prev_current = getattr(_nodes_mod, "CURRENT_NODE", None)
        setattr(_nodes_mod, "CURRENT_NODE", node)
        try:
            payload = {"listener": "onStartForm", "_basement_data": base_data, "base_class_name": base_class_name, "base_node_id": base_node_id}
            for a in actions:
                m = str((a or {}).get("method") or "").strip()
                if m and hasattr(node, m):
                    getattr(node, m)(payload)
        finally:
            setattr(_nodes_mod, "CURRENT_NODE", prev_current)

        try:
            if isinstance(getattr(node, "_data_cache", None), dict):
                node_data.update(node._data_cache or {})
            elif isinstance(getattr(node, "_data", None), dict):
                node_data.update(node._data or {})
        except Exception:
            pass

        # PrintForm is deliberately ephemeral: remove the runtime node row that
        # Node.__init__ may have created so the form does not persist data.
        try:
            if getattr(node, "_storage", None) is not None and getattr(node, "_id", None) in node._storage:
                del node._storage[node._id]
        except Exception:
            pass

        node_data.setdefault("_basement_data", dict(base_data or {}))
        node_data.setdefault("_id", print_node_id)
        node_data.setdefault("_class", print_class_name)
        return node_data
    except Exception as e:
        node_data["_print_error"] = str(e)
        return node_data


def _build_print_form_runtime(repo: models.Repo, parsed: Dict[str, Any], print_class_name: str, base_class_name: str, base_node_id: str) -> Tuple[Dict[str, Any], str, str, Dict[str, Any]]:
    classes = (parsed or {}).get("classes") or {}
    print_cls = classes.get(print_class_name) or {}
    if not print_cls or not _is_print_form_class_type(print_cls):
        abort(404)
    targets = _normalize_print_targets(print_cls.get("print_target_classes"))
    if base_class_name not in targets:
        abort(403)

    base_data = _fetch_node_data_for_repo(repo, base_class_name, base_node_id) or {}
    if isinstance(base_data, dict):
        base_data.setdefault("_id", base_node_id)
        base_data.setdefault("_class", base_class_name)
    print_node_id = _print_form_node_id(repo.config_uid, print_class_name, base_class_name, base_node_id)
    print_data = _execute_print_form_start_handler(repo, parsed, print_cls, print_node_id, base_class_name, base_node_id, base_data)
    html = _render_print_html(repo, print_cls, print_data)
    return print_cls, print_node_id, html, print_data

def _strip_projection_runtime_fields_for_save(cls_cfg: Dict[str, Any], data: Dict[str, Any]) -> Dict[str, Any]:
    """Return projection data without generated/report-only object lists.

    The projection node stores parameters only.  The current set of objects can be
    returned by onRunProjection and kept in the browser for the current render,
    but a normal Save of the projection must not persist those UIDs.
    """
    if not isinstance(data, dict):
        return {}
    if not _is_projection_class_type(cls_cfg):
        return data
    out = dict(data)
    for key in PROJECTION_TRANSIENT_SAVE_FIELDS:
        out.pop(key, None)
    return out


def _collect_runtime_messages_payload(node: Any = None) -> Dict[str, Any]:
    """Collect messages produced while saving a projection object."""
    messages: List[Dict[str, str]] = []
    seen = set()

    def add(msg: Any) -> None:
        if not msg:
            return
        if isinstance(msg, list):
            for one in msg:
                add(one)
            return
        if isinstance(msg, dict):
            text = str(msg.get("text") or msg.get("message") or "").strip()
            level = str(msg.get("level") or "info").strip() or "info"
        else:
            text = str(msg).strip()
            level = "info"
        if not text:
            return
        if level == "error":
            level = "danger"
        key = (text, level)
        if key in seen:
            return
        seen.add(key)
        messages.append({"text": text, "level": level})

    try:
        runtime_messages = getattr(_nodes_mod, "RUNTIME_MESSAGES", None)
        if runtime_messages is not None:
            add(runtime_messages.get())
    except Exception:
        pass
    try:
        add(getattr(node, "_ui_message", None))
    except Exception:
        pass

    if not messages:
        return {}
    return {"messages": messages, "message": messages[-1]}


def _projection_accept_error_payload(e: Exception) -> Dict[str, Any]:
    payload = getattr(e, "payload", None) or {}
    if not isinstance(payload, dict):
        payload = {"error": str(e)}
    msg = payload.get("message")
    if not isinstance(msg, dict):
        msg = {"text": str(payload.get("error") or str(e) or "Save rejected"), "level": "danger"}
    if msg.get("level") == "error":
        msg["level"] = "danger"
    return {"ok": False, "error": payload.get("error") or str(e) or "rejected", "message": msg}


def _projection_move_success_payload(save_meta: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    out: Dict[str, Any] = {"ok": True}
    if isinstance(save_meta, dict):
        out.update(save_meta)
    return out


def _normalize_projection_object_ids(value: Any) -> List[str]:
    if isinstance(value, str):
        s = value.strip()
        if not s:
            return []
        try:
            parsed = json.loads(s)
            value = parsed
        except Exception:
            value = [x.strip() for x in re.split(r"[\n,;]+", s) if x.strip()]

    if isinstance(value, dict):
        # Accept both a single record {uid/_id/id: ...} and a mapping
        # {id: uid_or_record}. This mirrors nodes.to_uid(get_all()).
        if any(k in value for k in ("uid", "_uid", "_id", "id")):
            value = [value]
        else:
            value = list(value.values())
    elif isinstance(value, (tuple, set)):
        value = list(value)

    if not isinstance(value, list):
        return []
    out: List[str] = []
    seen = set()
    for item in value:
        uid = ""
        if isinstance(item, str):
            uid = item.strip()
        elif isinstance(item, dict):
            uid = str(item.get("uid") or item.get("_uid") or item.get("_id") or item.get("id") or "").strip()
        if uid and uid not in seen:
            seen.add(uid)
            out.append(uid)
    return out


def _repo_for_config_uid(fallback: models.Repo, config_uid: str) -> models.Repo:
    if not config_uid or str(config_uid) == str(fallback.config_uid):
        return fallback
    repo = models.Repo.query.filter_by(config_uid=str(config_uid), user_id=current_user.id).first()
    return repo or fallback


def _get_projection_node_data(repo: models.Repo, cls_cfg: Dict[str, Any], class_name: str, node_id: str) -> Dict[str, Any]:
    cfg_uid = repo.config_uid
    node_uid = _normalize_custom_process_uid(cfg_uid, class_name, node_id)
    defaults = (cls_cfg.get("_data") or {}) if isinstance(cls_cfg, dict) else {}
    if not isinstance(defaults, dict):
        defaults = {}
    data = dict(defaults)
    try:
        node_class = _load_server_node_class(cfg_uid, class_name)
        node = node_class.get(node_uid, cfg_uid)
        if node:
            stored = node.get_data() or {}
            if isinstance(stored, dict):
                data.update(stored)
    except Exception:
        pass
    if _is_projection_class_type(cls_cfg):
        for key in PROJECTION_TRANSIENT_SAVE_FIELDS:
            data.pop(key, None)
    _apply_projection_defaults_to_data(cls_cfg, data, cfg_uid, class_name, node_uid)
    data.setdefault("_id", node_uid)
    data.setdefault("_class", class_name)
    return data


def _projection_key_aliases(projection_uid: str) -> List[str]:
    raw = str(projection_uid or "").strip()
    out: List[str] = []
    if raw:
        out.append(raw)
        parts = raw.split("$")
        # Backward compatibility with older singleton ids: cfg$Class
        if len(parts) >= 3 and parts[0] and parts[1]:
            legacy = f"{parts[0]}${parts[1]}"
            if legacy not in out:
                out.append(legacy)
    return out



def _projection_value_for_data(data: Dict[str, Any], projection_uid: str) -> Any:
    if not isinstance(data, dict):
        return None
    vals = data.get("_projection_values")
    if isinstance(vals, dict):
        for key in _projection_key_aliases(projection_uid):
            if key in vals:
                return vals.get(key)

    # Backward/shortcut compatibility: some handlers use singular
    # _projection_value either as a map by projection uid or as the direct value.
    single = data.get("_projection_value")
    if isinstance(single, dict):
        for key in _projection_key_aliases(projection_uid):
            if key in single:
                return single.get(key)
        direct_markers = {
            "id", "column_id", "resource_id", "doctor_id", "task_id", "parent",
            "start", "end", "period_start", "period_end", "x1", "y1", "x2", "y2",
        }
        if any(k in single for k in direct_markers):
            return single
    elif single is not None:
        return single
    return None


def _normalize_projection_timer(value: Any) -> int:
    try:
        n = int(float(value))
    except Exception:
        return 0
    return n if n > 0 else 0


def _boolish_projection_value(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return value != 0
    if isinstance(value, str):
        return value.strip().lower() in {"1", "true", "yes", "y", "on", "да", "истина"}
    return False


def _normalize_diagram_projection_value(value: Any, index: int = 0) -> Dict[str, Any]:
    if isinstance(value, str):
        try:
            parsed = json.loads(value) if value.strip() else {}
            value = parsed if isinstance(parsed, dict) else {}
        except Exception:
            value = {}
    if not isinstance(value, dict):
        value = {}

    def num(name: str, default: float) -> float:
        try:
            return float(value.get(name, default))
        except Exception:
            return float(default)

    x1 = num("x1", num("x", 24 + (index % 10) * 110))
    y1 = num("y1", num("y", 24 + (index // 10) * 74))
    w = num("width", 90)
    h = num("height", 52)
    x2 = num("x2", x1 + w)
    y2 = num("y2", y1 + h)
    if x2 <= x1:
        x2 = x1 + max(20, w)
    if y2 <= y1:
        y2 = y1 + max(20, h)

    figure = str(value.get("figure") or "rectangle").strip().lower()
    if figure not in {"rectangle", "circle", "svg"}:
        figure = "rectangle"
    bg = str(value.get("background") or "#ffffff").strip()
    if not re.match(r"^#[0-9a-fA-F]{3}([0-9a-fA-F]{3})?$", bg):
        bg = "#ffffff"

    text_color = str(value.get("text_color") or value.get("color") or "#212529").strip()
    if not re.match(r"^#[0-9a-fA-F]{3}([0-9a-fA-F]{3})?$", text_color):
        text_color = "#212529"
    border_color = str(value.get("border_color") or "rgba(0,0,0,.25)").strip()
    tooltip = str(value.get("tooltip") or value.get("title") or "")

    return {
        "x1": x1,
        "y1": y1,
        "x2": x2,
        "y2": y2,
        "figure": figure,
        "svg": str(value.get("svg") or ""),
        "background": bg,
        "text": str(value.get("text") or ""),
        "text_color": text_color,
        "border_color": border_color,
        "tooltip": tooltip,
    }

def _projection_object_payload(repo: models.Repo, projection_uid: str, object_uid: str) -> Optional[Dict[str, Any]]:
    try:
        cfg_uid, cls_name, internal_id = _nodes_mod.parse_uid_any(object_uid)
    except Exception:
        cfg_uid, cls_name, internal_id = None, None, None
    cls_name = str(cls_name or "").strip()
    internal_id = str(internal_id or "").strip()
    if not cls_name or not internal_id:
        return None
    obj_repo = _repo_for_config_uid(repo, cfg_uid or repo.config_uid)
    data = _fetch_node_data_for_repo(obj_repo, cls_name, internal_id) or {}
    if not isinstance(data, dict):
        data = {}
    if data.get("_hidden"):
        return None
    normalized_uid = _nodes_mod.normalize_own_uid(obj_repo.config_uid, cls_name, internal_id)
    projection_value = _projection_value_for_data(data, projection_uid)
    col = "__empty__"
    if projection_value is not None and not isinstance(projection_value, dict):
        col = str(projection_value or "__empty__")
    try:
        cover_html = _node_cover_html(obj_repo, cls_name, internal_id)
    except Exception:
        cover_html = ""
    parsed = get_parsed_config(obj_repo, models.db) or {}
    view = _render_class_record_view(parsed, cls_name, internal_id, data)
    return {
        "uid": normalized_uid,
        "repo_id": obj_repo.id,
        "repo_uid": obj_repo.config_uid,
        "class": cls_name,
        "id": internal_id,
        "column_id": col or "__empty__",
        "projection_value": projection_value,
        "data": data,
        "view": view,
        "cover_html": cover_html,
        "open_url": url_for("client.node_form", config_uid=obj_repo.config_uid, class_name=cls_name, node_id=internal_id),
    }


def _projection_object_diagram_payload(repo: models.Repo, projection_uid: str, object_uid: str) -> Optional[Dict[str, Any]]:
    """Return a lightweight payload for diagram projections.

    Diagram projections render their own shapes from _projection_values and do not
    need card covers or full class record views. Avoiding _node_cover_html() and
    _render_class_record_view() keeps the Loading... stage fast for large
    diagrams such as warehouse maps.
    """
    try:
        cfg_uid, cls_name, internal_id = _nodes_mod.parse_uid_any(object_uid)
    except Exception:
        cfg_uid, cls_name, internal_id = None, None, None
    cls_name = str(cls_name or "").strip()
    internal_id = str(internal_id or "").strip()
    if not cls_name or not internal_id:
        return None

    obj_repo = _repo_for_config_uid(repo, cfg_uid or repo.config_uid)
    data = _fetch_node_data_for_repo(obj_repo, cls_name, internal_id) or {}
    if not isinstance(data, dict):
        data = {}
    if data.get("_hidden"):
        return None

    projection_value = _projection_value_for_data(data, projection_uid)
    if projection_value is None:
        return None

    title = str(
        data.get("caption")
        or data.get("title")
        or data.get("name")
        or internal_id
        or ""
    )
    normalized_uid = _nodes_mod.normalize_own_uid(obj_repo.config_uid, cls_name, internal_id)
    return {
        "uid": normalized_uid,
        "repo_id": obj_repo.id,
        "repo_uid": obj_repo.config_uid,
        "class": cls_name,
        "id": internal_id,
        "projection_value": projection_value,
        "view": {"title": title},
        "open_url": url_for("client.node_form", config_uid=obj_repo.config_uid, class_name=cls_name, node_id=internal_id),
    }



def _parse_projection_datetime(value: Any, fallback: Optional[datetime] = None) -> Optional[datetime]:
    if isinstance(value, datetime):
        return value.replace(tzinfo=None)
    if isinstance(value, (int, float)):
        try:
            # Accept Unix seconds or milliseconds.
            n = float(value)
            if n > 10_000_000_000:
                n = n / 1000.0
            return datetime.fromtimestamp(n).replace(tzinfo=None)
        except Exception:
            return fallback
    raw = str(value or "").strip()
    if not raw:
        return fallback
    if raw.endswith("Z"):
        raw = raw[:-1]
    raw = raw.replace(" ", "T")
    for fmt in ("%Y%m%dT%H%M%S", "%Y%m%d%H%M%S", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%dT%H:%M", "%Y-%m-%d"):
        try:
            dt = datetime.strptime(raw[:len(datetime.now().strftime(fmt))] if "%" in fmt else raw, fmt)
            return dt.replace(tzinfo=None)
        except Exception:
            pass
    try:
        return datetime.fromisoformat(raw).replace(tzinfo=None)
    except Exception:
        return fallback


def _projection_datetime_iso(value: Optional[datetime]) -> str:
    if not value:
        return ""
    return value.replace(microsecond=0).isoformat()


def _projection_day_key(value: Optional[datetime]) -> str:
    return value.strftime("%Y-%m-%d") if value else ""


def _projection_read_jsonish(value: Any, fallback: Any) -> Any:
    if isinstance(value, str):
        raw = value.strip()
        if not raw:
            return fallback
        try:
            return json.loads(raw)
        except Exception:
            return raw
    return value if value is not None else fallback


def _normalize_projection_orientation(value: Any) -> str:
    raw = str(value or "vertical").strip().lower()
    if raw in {"horizontal", "h", "row", "rows", "time-horizontal"}:
        return "horizontal"
    return "vertical"


def _projection_float(value: Any, fallback: float) -> float:
    try:
        n = float(value)
        if math.isfinite(n):
            return n
    except Exception:
        pass
    return fallback


def _projection_schedule_create_class(data: Dict[str, Any]) -> str:
    data = data if isinstance(data, dict) else {}
    return str(data.get("_projection_create_class") or "").strip()


def _projection_schedule_default_interval_hours(data: Dict[str, Any]) -> float:
    data = data if isinstance(data, dict) else {}
    n = _projection_float(data.get("_projection_default_interval_hours"), 0.25)
    return min(24.0, max(1.0 / 60.0, n))


def _normalize_schedule_columns(value: Any, period_start: datetime, period_end: datetime, selected_date: str = "") -> Tuple[str, List[Dict[str, Any]], str]:
    raw = _projection_read_jsonish(value, [])
    mode = "resources"
    projection_id = ""
    if isinstance(raw, str) and raw.strip().lower() == "days":
        mode = "days"
    elif isinstance(raw, dict):
        if str(raw.get("mode") or raw.get("type") or "").strip().lower() == "days":
            mode = "days"
        projection_id = str(raw.get("id") or raw.get("projection_id") or "").strip()
        raw = raw.get("columns") or raw.get("items") or []
    elif isinstance(raw, list) and len(raw) == 1 and str(raw[0]).strip().lower() == "days":
        mode = "days"

    if mode == "days":
        base_date = _parse_projection_datetime(selected_date) or period_start or datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
        start_day = period_start.date() if period_start else base_date.date()
        end_day = period_end.date() if period_end and period_end.date() > start_day else (start_day + timedelta(days=6))
        rows = []
        d = start_day
        while d <= end_day and len(rows) < 120:
            key = d.isoformat()
            rows.append({"id": key, "caption": key, "date": key, "areas": []})
            d += timedelta(days=1)
        return mode, rows, projection_id

    cols: List[Dict[str, Any]] = []
    if isinstance(raw, list):
        seen = set()
        for col in raw:
            if not isinstance(col, dict):
                continue
            cid = str(col.get("id") or col.get("key") or col.get("uid") or "").strip()
            if not cid or cid in seen:
                continue
            seen.add(cid)
            areas = col.get("areas") or col.get("availability") or []
            if isinstance(areas, dict):
                areas = [areas]
            norm_areas = []
            if isinstance(areas, list):
                for a in areas:
                    if not isinstance(a, dict):
                        continue
                    a_start = _parse_projection_datetime(a.get("period_start") or a.get("perior_start") or a.get("start"), None)
                    a_end = _parse_projection_datetime(a.get("period_end") or a.get("perior_end") or a.get("end"), None)
                    norm_areas.append({
                        "start": _projection_datetime_iso(a_start) if a_start else str(a.get("period_start") or a.get("perior_start") or a.get("start") or ""),
                        "end": _projection_datetime_iso(a_end) if a_end else str(a.get("period_end") or a.get("perior_end") or a.get("end") or ""),
                        "color": str(a.get("color") or a.get("background") or "#e9ecef"),
                        "available": bool(a.get("available", a.get("availible", True))),
                    })
            cols.append({
                "id": cid,
                "caption": str(col.get("caption") or col.get("title") or cid),
                "areas": norm_areas,
            })
    return mode, cols, projection_id


def _normalize_schedule_projection_value(value: Any, projection_uid: str, mode: str = "resources", index: int = 0) -> Optional[Dict[str, Any]]:
    value = _projection_read_jsonish(value, value)
    if not isinstance(value, dict):
        return None
    row_id = str(value.get("id") or value.get("column_id") or value.get("resource_id") or value.get("doctor_id") or value.get("row_id") or "").strip()
    start = _parse_projection_datetime(value.get("start") or value.get("period_start") or value.get("begin") or value.get("from"), None)
    end = _parse_projection_datetime(value.get("end") or value.get("period_end") or value.get("finish") or value.get("to"), None)
    if not start:
        return None
    if not end or end <= start:
        end = start + timedelta(minutes=30)
    if mode == "days":
        row_id = _projection_day_key(start)
    return {
        "id": row_id,
        "start": _projection_datetime_iso(start),
        "end": _projection_datetime_iso(end),
        "color": str(value.get("color") or value.get("background") or ""),
        "caption": str(value.get("caption") or value.get("title") or ""),
    }


def _projection_object_schedule_payload(repo: models.Repo, projection_uid: str, object_uid: str, mode: str = "resources", index: int = 0) -> Optional[Dict[str, Any]]:
    obj = _projection_object_payload(repo, projection_uid, object_uid)
    if not obj:
        return None
    val = _normalize_schedule_projection_value(obj.get("projection_value"), projection_uid, mode, index)
    if not val:
        return None
    obj["schedule"] = val
    return obj


def _normalize_gantt_tasks(value: Any) -> List[Dict[str, Any]]:
    raw = _projection_read_jsonish(value, [])
    if isinstance(raw, dict):
        raw = raw.get("tasks") or raw.get("items") or []
    out: List[Dict[str, Any]] = []
    seen = set()
    if isinstance(raw, list):
        for item in raw:
            if not isinstance(item, dict):
                continue
            tid = str(item.get("id") or item.get("task_id") or item.get("key") or "").strip()
            if not tid or tid in seen:
                continue
            seen.add(tid)
            out.append({
                "id": tid,
                "caption": str(item.get("caption") or item.get("title") or item.get("name") or tid),
                "parent": str(item.get("parent") or item.get("parent_id") or "").strip(),
                "color": str(item.get("color") or item.get("background") or ""),
            })
    return out


def _normalize_gantt_projection_value(value: Any, index: int = 0) -> Optional[Dict[str, Any]]:
    value = _projection_read_jsonish(value, value)
    if not isinstance(value, dict):
        return None
    task_id = str(value.get("id") or value.get("task_id") or value.get("row_id") or "").strip()
    start = _parse_projection_datetime(value.get("start") or value.get("period_start") or value.get("begin") or value.get("from"), None)
    end = _parse_projection_datetime(value.get("end") or value.get("period_end") or value.get("finish") or value.get("to"), None)
    if not start:
        return None
    if not end or end <= start:
        end = start + timedelta(days=1)
    return {
        "id": task_id,
        "start": _projection_datetime_iso(start),
        "end": _projection_datetime_iso(end),
        "title": str(value.get("title") or value.get("caption") or ""),
        "parent": str(value.get("parent") or value.get("parent_id") or "").strip(),
        "color": str(value.get("color") or value.get("background") or ""),
    }


def _projection_object_gantt_payload(repo: models.Repo, projection_uid: str, object_uid: str, index: int = 0) -> Optional[Dict[str, Any]]:
    obj = _projection_object_diagram_payload(repo, projection_uid, object_uid)
    if not obj:
        return None
    val = _normalize_gantt_projection_value(obj.get("projection_value"), index)
    if not val:
        return None
    if not val.get("title"):
        val["title"] = str((obj.get("view") or {}).get("title") or obj.get("id") or "")
    obj["gantt"] = val
    return obj


def _save_projection_object_data(obj_repo: models.Repo, cls_name: str, internal_id: str, data: Dict[str, Any], user_modification: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    payload = dict(data or {})
    if isinstance(user_modification, dict) and user_modification:
        payload["_user_modification"] = user_modification
    base_url = (obj_repo.base_url or "").strip().rstrip("/")
    current = (request.host_url or "").rstrip("/")
    if not base_url or base_url == current:
        parsed_ctx = get_parsed_config(obj_repo, models.db) or {}
        tokens = _nodes_mod.set_runtime_context(obj_repo.config_uid, parsed_ctx, system_user=_client_runtime_system_user_payload())
        try:
            node = _node_local_update_data(obj_repo.config_uid, cls_name, internal_id, data, user_modification=user_modification)
            return _collect_runtime_messages_payload(node)
        finally:
            _nodes_mod.reset_runtime_context(tokens)

    url = obj_repo.base_url.rstrip("/") + f"/api/config/{obj_repo.config_uid}/node/{cls_name}/{internal_id}"
    resp = requests.put(url, json=payload, auth=_auth_tuple(obj_repo), headers=_client_remote_system_user_headers(), timeout=20)
    resp.raise_for_status()
    try:
        remote_payload = resp.json()
    except Exception:
        remote_payload = None
    if isinstance(remote_payload, dict) and remote_payload.get("status") is False:
        err_payload = remote_payload.get("data") if isinstance(remote_payload.get("data"), dict) else remote_payload
        raise _nodes_mod.AcceptRejected(err_payload)
    if isinstance(remote_payload, dict):
        out: Dict[str, Any] = {}
        if isinstance(remote_payload.get("message"), dict):
            out["message"] = remote_payload.get("message")
        if isinstance(remote_payload.get("messages"), list):
            out["messages"] = remote_payload.get("messages")
            if "message" not in out and out["messages"]:
                out["message"] = out["messages"][-1]
        return out
    return {}

def _resolve_projection_object(repo: models.Repo, object_uid: str) -> Tuple[models.Repo, str, str]:
    try:
        cfg_uid, cls_name, internal_id = _nodes_mod.parse_uid_any(object_uid)
    except Exception:
        cfg_uid, cls_name, internal_id = None, None, None
    cls_name = str(cls_name or "").strip()
    internal_id = str(internal_id or "").strip()
    if not cls_name or not internal_id:
        raise ValueError("bad object uid")
    return _repo_for_config_uid(repo, cfg_uid or repo.config_uid), cls_name, internal_id

@client_bp.route("/api/node_views/batch", methods=["POST"])
@login_required
def api_node_views_batch():
    """Resolve NodeLink labels for a whole virtual table in one request.

    Existing ``<field>_view`` values never reach this endpoint.  The browser
    sends only missing references, grouped here by configuration and class so
    each class store is opened once rather than once per table cell.
    """
    j = request.get_json(force=True) or {}
    repo_id = int(j.get("repo_id") or 0)
    raw_items = j.get("uids") or []
    if not repo_id or not isinstance(raw_items, list):
        return jsonify({"ok": False, "error": "bad args"}), 400

    repo = _get_repo_or_404(repo_id)
    uids: List[str] = []
    seen = set()
    for value in raw_items:
        uid = str(value or "").strip()
        if uid and uid not in seen:
            seen.add(uid)
            uids.append(uid)
        if len(uids) >= 1000:
            break

    groups: Dict[Tuple[str, str], List[Tuple[str, str]]] = {}
    views: Dict[str, str] = {}
    for uid in uids:
        try:
            uid_cfg, cls_name, internal_id = _nodes_mod.parse_uid_any(uid)
        except Exception:
            uid_cfg, cls_name, internal_id = None, None, None
        eff_cfg = str(uid_cfg or repo.config_uid or "").strip()
        cls_name = str(cls_name or "").strip()
        internal_id = str(internal_id or "").strip()
        if not eff_cfg or not cls_name or not internal_id:
            continue
        if cls_name == "_User":
            user_view = _client_system_user_view(uid)
            if user_view:
                views[uid] = user_view
            continue
        if not _client_user_can_access_config(eff_cfg, current_user):
            continue
        if not _client_user_can_access_class(eff_cfg, cls_name, current_user):
            continue
        groups.setdefault((eff_cfg, cls_name), []).append((uid, internal_id))

    parsed_cache: Dict[str, Dict[str, Any]] = {}
    repo_cache: Dict[str, Optional[models.Repo]] = {str(repo.config_uid or ""): repo}

    def target_repo_for(config_uid: str) -> Optional[models.Repo]:
        if config_uid in repo_cache:
            return repo_cache[config_uid]
        target = models.Repo.query.filter_by(
            user_id=current_user.id, config_uid=config_uid
        ).first()
        repo_cache[config_uid] = target
        return target

    def parsed_for(config_uid: str) -> Dict[str, Any]:
        if config_uid in parsed_cache:
            return parsed_cache[config_uid]
        parsed: Dict[str, Any] = {}
        target = target_repo_for(config_uid)
        try:
            if target is not None:
                parsed = get_parsed_config(target, models.db) or {}
            else:
                cfg = fetch_config_from_local_db(config_uid, user=current_user)
                parsed = build_parsed_config(cfg or {}) if cfg else {}
        except Exception:
            parsed = {}
        parsed_cache[config_uid] = parsed
        return parsed

    for (config_uid, cls_name), entries in groups.items():
        raw_uids = [uid for uid, _ in entries]
        data_by_uid = _node_storage_data_batch_direct(config_uid, cls_name, raw_uids)
        parsed = parsed_for(config_uid)
        target_repo = target_repo_for(config_uid)
        for uid, internal_id in entries:
            data = data_by_uid.get(uid)
            if not isinstance(data, dict) and target_repo is not None:
                try:
                    data = _fetch_node_data_for_repo(target_repo, cls_name, internal_id) or {}
                except Exception:
                    data = {}
            if not isinstance(data, dict) or not data:
                continue
            try:
                view = _render_class_record_view(parsed, cls_name, internal_id, data)
            except Exception:
                view = str(data.get("_view") or data.get("name") or data.get("title") or internal_id)
            views[uid] = str(view or internal_id)

    return jsonify({"ok": True, "views": views})


@client_bp.route("/api/nodalayout/render", methods=["POST"])
@login_required
def api_nodalayout_render():
    j = request.get_json(force=True) or {}
    repo_id = int(j.get("repo_id") or 0)
    class_name = str(j.get("class_name") or "")
    node_id = str(j.get("node_id") or "")
    layout = j.get("layout")
    data = j.get("data")

    if not repo_id or not class_name or not node_id:
        return jsonify({"ok": False, "error": "bad args"}), 400
    if layout is None:
        return jsonify({"ok": False, "error": "layout required"}), 400
    if data is None or not isinstance(data, dict):
        data = {}

    repo = _get_repo_or_404(repo_id)
    try:
        html = render_nodalayout_html(
            layout,
            data,
            assets_base_dir=_userfiles_dir_for_repo(repo),
            context=_nl_context(repo, class_name=class_name, node_id=node_id),
        )
        return jsonify({"ok": True, "layout_html": html or ""})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@client_bp.route("/api/node/save", methods=["POST"])
@login_required
def api_node_save():
    j = request.get_json(force=True) or {}
    repo_id = int(j.get("repo_id") or 0)
    class_name = str(j.get("class_name") or "")
    node_id = str(j.get("node_id") or "")
    data = j.get("data") or {}

    if not repo_id or not class_name or not node_id or not isinstance(data, dict):
        return jsonify({"ok": False, "error": "bad args"}), 400

    repo = _get_repo_or_404(repo_id)
    cfg_uid = repo.config_uid
    parsed_ctx = get_parsed_config(repo, models.db) or {}
    cls_cfg = (parsed_ctx.get("classes") or {}).get(class_name) or {}
    if not cls_cfg:
        return jsonify({"ok": False, "error": "Forbidden"}), 403
    # For existing nodes RLS also controls updates. New data is checked below by
    # root RLS after save; here we prevent editing rows outside the user's RLS set.
    if _client_repo_is_local(repo) and node_id and not _client_user_can_access_node(cfg_uid, class_name, node_id, data):
        return jsonify({"ok": False, "error": "Forbidden"}), 403
    data = _strip_projection_runtime_fields_for_save(cls_cfg, data)
    is_custom_process = _is_singleton_class_type(cls_cfg)
    _ctx_tokens = _nodes_mod.set_runtime_context(cfg_uid, parsed_ctx, system_user=_client_runtime_system_user_payload())

    @after_this_request
    def _reset_ctx(resp):
        _nodes_mod.reset_runtime_context(_ctx_tokens)
        return resp


    base_url = (repo.base_url or "").strip().rstrip("/")
    current = (request.host_url or "").rstrip("/")

    try:
        if not base_url or base_url == current:
            if is_custom_process:
                # custom_process nodes are singletons; create on first save
                node_id = _node_local_upsert_custom_process(cfg_uid, class_name, node_id, data)
            else:
                _node_local_update_data(cfg_uid, class_name, node_id, data)
            # Optional: register in default room after save (Migration tab)
            reg_count = 0
            room_uid = ""
            try:
                parsed = get_parsed_config(repo, models.db) or {}
                cls_cfg = (parsed.get("classes") or {}).get(class_name) or {}
                if bool(cls_cfg.get("migration_register_on_save")):
                    room_uid = _resolve_class_default_room_uid(parsed, cls_cfg)
                    if room_uid:
                        reg_count, reg_result = _register_nodes_to_room_local(cfg_uid, class_name, room_uid, [node_id])
            except Exception:
                reg_count = 0
                reg_result = {}

            out = {"ok": True}
            if reg_count:
                out["registered"] = reg_count
                out["room_uid"] = room_uid
                push = reg_result.get("push") if isinstance(reg_result, dict) else None
                transport = str((reg_result or {}).get("transport") or "").strip().lower()
                if transport == "fcm" and isinstance(push, dict) and not push.get("ok"):
                    err = str(push.get("error") or reg_result.get("delivery_error") or "FCM push failed")
                    out["delivery_ok"] = False
                    out["delivery"] = reg_result
                    out["message"] = {"text": f"Saved and queued, but FCM delivery failed: {err}", "level": "danger"}
                else:
                    out["delivery_ok"] = True
                    out["message"] = {"text": f"Registered in Room: {room_uid}", "level": "success"}
            # attach runtime messages (from nodes.push_message)
            try:
                msgs = getattr(_nodes_mod, "RUNTIME_MESSAGES", None)
                msgs = msgs.get() if msgs else []
                if isinstance(msgs, list) and msgs:
                    out.setdefault("messages", msgs)
                    out.setdefault("message", msgs[-1])
            except Exception:
                pass
            return jsonify(out)
        # remote
        if is_custom_process:
            node_id = _normalize_custom_process_uid(cfg_uid, class_name, node_id)
        _api_post_remote(
            repo,
            f"/api/config/{cfg_uid}/node/{class_name}/{node_id}/save",
            {"_data": data},
        )
        msgs = getattr(_nodes_mod, "RUNTIME_MESSAGES", None)
        msgs = msgs.get() if msgs else []
        if isinstance(msgs, list) and msgs:
            last = msgs[-1]
            if isinstance(last, dict) and last.get("level") == "error":
                last["level"] = "danger"
            return jsonify({"ok": True, "message": last if isinstance(last, dict) else {"text": str(last), "level": "info"}, "messages": msgs})
        return jsonify({"ok": True})
    except _nodes_mod.AcceptRejected as e:

        payload = getattr(e, 'payload', None) or {}

        msg = payload.get('message')

        # fallback to runtime messages if handler used nodes.push_message
        if not isinstance(msg, dict):
            try:
                msgs = getattr(_nodes_mod, "RUNTIME_MESSAGES", None)
                msgs = msgs.get() if msgs else []
                if isinstance(msgs, list) and msgs:
                    msg = msgs[-1]
            except Exception:
                msg = None

        if not isinstance(msg, dict):
            msg = {'text': payload.get('error') or 'Save rejected', 'level': 'danger'}
        # normalize level
        if msg.get('level') == 'error':
            msg['level'] = 'danger'

        return jsonify({'ok': False, 'error': payload.get('error') or 'rejected', 'message': msg}), 200

    except Exception as e:
        return jsonify({"ok": False, 
                        "error": str(e),
                        "message": {"text": f"Handler error: {e}", "level": "error"},
                        }), 500


@client_bp.route("/api/projection/kanban/data", methods=["POST"])
@login_required
def api_projection_kanban_data():
    j = request.get_json(force=True) or {}
    repo_id = int(j.get("repo_id") or 0)
    class_name = str(j.get("class_name") or "").strip()
    node_id = str(j.get("node_id") or "").strip()

    if not repo_id or not class_name or not node_id:
        return jsonify({"ok": False, "error": "bad args"}), 400

    repo = _get_repo_or_404(repo_id)
    parsed = get_parsed_config(repo, models.db) or {}
    cls_cfg = (parsed.get("classes") or {}).get(class_name) or {}
    if not _is_projection_class_type(cls_cfg):
        return jsonify({"ok": False, "error": "class is not a projection"}), 400

    projection_type = str(cls_cfg.get("projection_type") or PROJECTION_KANBAN_TYPE).strip() or PROJECTION_KANBAN_TYPE
    if projection_type != PROJECTION_KANBAN_TYPE:
        return jsonify({"ok": False, "error": "only kanban_projection is implemented"}), 400

    projection_uid = _normalize_custom_process_uid(repo.config_uid, class_name, node_id)
    data = _get_projection_node_data(repo, cls_cfg, class_name, projection_uid)

    # Raw-node projection handlers return updated projection _data to the browser,
    # but the normal projection data endpoint reads from repo storage. Accept the
    # just-returned projection contract fields from the client so Generate works
    # for raw embedded classes as well.
    client_data = j.get("node_data")
    if isinstance(client_data, dict):
        for k in ("_projection_objects", "_projection_timer", "_projection_uid", "_projection_type", "_kanban_columns"):
            if k in client_data:
                data[k] = client_data.get(k)
        _apply_projection_defaults_to_data(cls_cfg, data, repo.config_uid, class_name, projection_uid)

    columns = data.get("_kanban_columns") if isinstance(data.get("_kanban_columns"), list) else _parse_projection_kanban_columns(cls_cfg)
    object_ids = _normalize_projection_object_ids(data.get("_projection_objects"))

    objects = []
    for uid in object_ids:
        obj = _projection_object_payload(repo, projection_uid, uid)
        if obj:
            objects.append(obj)

    return jsonify({
        "ok": True,
        "projection_uid": projection_uid,
        "projection_type": projection_type,
        "columns": columns,
        "objects": objects,
        "node_data": data,
        "timer": _normalize_projection_timer(data.get("_projection_timer")),
        "empty_column": {"id": "__empty__", "caption": "No column"},
    })



@client_bp.route("/api/projection/diagram/data", methods=["POST"])
@login_required
def api_projection_diagram_data():
    j = request.get_json(force=True) or {}
    repo_id = int(j.get("repo_id") or 0)
    class_name = str(j.get("class_name") or "").strip()
    node_id = str(j.get("node_id") or "").strip()

    if not repo_id or not class_name or not node_id:
        return jsonify({"ok": False, "error": "bad args"}), 400

    repo = _get_repo_or_404(repo_id)
    parsed = get_parsed_config(repo, models.db) or {}
    cls_cfg = (parsed.get("classes") or {}).get(class_name) or {}
    if not _is_projection_class_type(cls_cfg):
        return jsonify({"ok": False, "error": "class is not a projection"}), 400

    projection_type = str(cls_cfg.get("projection_type") or PROJECTION_KANBAN_TYPE).strip() or PROJECTION_KANBAN_TYPE
    if projection_type != PROJECTION_DIAGRAM_TYPE:
        return jsonify({"ok": False, "error": "class is not a diagram_projection"}), 400

    projection_uid = _normalize_custom_process_uid(repo.config_uid, class_name, node_id)
    data = _get_projection_node_data(repo, cls_cfg, class_name, projection_uid)

    # Same raw-node bridge as kanban: use contract fields returned by the event
    # handler before reading objects for visualization.
    client_data = j.get("node_data")
    if isinstance(client_data, dict):
        for k in ("_projection_objects", "_projection_timer", "_projection_uid", "_projection_type", "_projection_header", "_projection_editor", "_projection_legend", "_projection_note"):
            if k in client_data:
                data[k] = client_data.get(k)
        _apply_projection_defaults_to_data(cls_cfg, data, repo.config_uid, class_name, projection_uid)

    object_ids = _normalize_projection_object_ids(data.get("_projection_objects"))

    objects = []
    for idx, uid in enumerate(object_ids):
        obj = _projection_object_diagram_payload(repo, projection_uid, uid)
        if not obj:
            continue
        # Diagram objects are linked only when this projection has a dict-like
        # value in the object's _projection_values. Clearing that value is the
        # non-destructive "unlink" operation.
        obj["diagram"] = _normalize_diagram_projection_value(obj.get("projection_value"), idx)
        objects.append(obj)

    return jsonify({
        "ok": True,
        "projection_uid": projection_uid,
        "projection_type": projection_type,
        "objects": objects,
        "node_data": data,
        "header": str(data.get("_projection_header") or ""),
        "note": str(data.get("_projection_note") or ""),
        "legend": data.get("_projection_legend") if isinstance(data.get("_projection_legend"), list) else [],
        "editor": _boolish_projection_value(data.get("_projection_editor")),
        "timer": _normalize_projection_timer(data.get("_projection_timer")),
    })




@client_bp.route("/api/projection/schedule/data", methods=["POST"])
@login_required
def api_projection_schedule_data():
    j = request.get_json(force=True) or {}
    repo_id = int(j.get("repo_id") or 0)
    class_name = str(j.get("class_name") or "").strip()
    node_id = str(j.get("node_id") or "").strip()
    selected_date = str(j.get("selected_date") or "").strip()

    if not repo_id or not class_name or not node_id:
        return jsonify({"ok": False, "error": "bad args"}), 400

    repo = _get_repo_or_404(repo_id)
    parsed = get_parsed_config(repo, models.db) or {}
    cls_cfg = (parsed.get("classes") or {}).get(class_name) or {}
    if not _is_projection_class_type(cls_cfg):
        return jsonify({"ok": False, "error": "class is not a projection"}), 400

    projection_type = str(cls_cfg.get("projection_type") or PROJECTION_KANBAN_TYPE).strip() or PROJECTION_KANBAN_TYPE
    if projection_type != PROJECTION_SCHEDULE_TYPE:
        return jsonify({"ok": False, "error": "class is not a schedule_projection"}), 400

    projection_uid = _normalize_custom_process_uid(repo.config_uid, class_name, node_id)
    data = _get_projection_node_data(repo, cls_cfg, class_name, projection_uid)

    client_data = j.get("node_data")
    if isinstance(client_data, dict):
        for k in ("_projection_objects", "_projection_timer", "_projection_uid", "_projection_type", "_projection_columns", "_projection_period_start", "_projection_period_end", "_projection_id", "_projection_header", "_projection_orientation", "_projection_create_class", "_projection_default_interval_hours"):
            if k in client_data:
                data[k] = client_data.get(k)
        _apply_projection_defaults_to_data(cls_cfg, data, repo.config_uid, class_name, projection_uid)

    now = datetime.now().replace(second=0, microsecond=0)
    default_start = now.replace(hour=8, minute=0)
    default_end = now.replace(hour=18, minute=0)
    period_start = _parse_projection_datetime(data.get("_projection_period_start"), default_start) or default_start
    period_end = _parse_projection_datetime(data.get("_projection_period_end"), default_end) or default_end
    if period_end <= period_start:
        period_end = period_start + timedelta(hours=10)

    # For resource schedule, selected date switches the day but keeps configured hours.
    sel_dt = _parse_projection_datetime(selected_date, None)
    if sel_dt:
        day = sel_dt.date()
        duration = period_end - period_start
        period_start = datetime.combine(day, period_start.time())
        period_end = period_start + duration

    mode, columns, projection_id = _normalize_schedule_columns(data.get("_projection_columns"), period_start, period_end, selected_date)
    if not projection_id:
        projection_id = str(data.get("_projection_id") or "").strip()

    objects = []
    for idx, uid in enumerate(_normalize_projection_object_ids(data.get("_projection_objects"))):
        obj = _projection_object_schedule_payload(repo, projection_uid, uid, mode, idx)
        if obj:
            objects.append(obj)

    orientation = _normalize_projection_orientation(data.get("_projection_orientation"))
    default_interval_hours = _projection_schedule_default_interval_hours(data)
    create_class = _projection_schedule_create_class(data)

    return jsonify({
        "ok": True,
        "projection_uid": projection_uid,
        "projection_type": projection_type,
        "mode": mode,
        "projection_id": projection_id,
        "columns": columns,
        "objects": objects,
        "node_data": data,
        "header": str(data.get("_projection_header") or ""),
        "projection_orientation": orientation,
        "projection_create_class": create_class,
        "projection_default_interval_hours": default_interval_hours,
        "slot_minutes": max(1, int(round(default_interval_hours * 60))),
        "period_start": _projection_datetime_iso(period_start),
        "period_end": _projection_datetime_iso(period_end),
        "selected_date": _projection_day_key(period_start),
        "timer": _normalize_projection_timer(data.get("_projection_timer")),
    })


@client_bp.route("/api/projection/schedule/move", methods=["POST"])
@login_required
def api_projection_schedule_move():
    j = request.get_json(force=True) or {}
    repo_id = int(j.get("repo_id") or 0)
    projection_uid = str(j.get("projection_uid") or "").strip()
    object_uid = str(j.get("object_uid") or "").strip()
    fields = j.get("fields") or {}
    if not isinstance(fields, dict):
        fields = {}
    if not repo_id or not projection_uid or not object_uid:
        return jsonify({"ok": False, "error": "bad args"}), 400

    repo = _get_repo_or_404(repo_id)
    try:
        obj_repo, cls_name, internal_id = _resolve_projection_object(repo, object_uid)
        data = _fetch_node_data_for_repo(obj_repo, cls_name, internal_id) or {}
        if not isinstance(data, dict):
            data = {}
        vals = data.get("_projection_values")
        if not isinstance(vals, dict):
            vals = {}
        existing = None
        for key in _projection_key_aliases(projection_uid):
            if key in vals:
                existing = vals.get(key)
                break
        value = _projection_read_jsonish(existing, {})
        if not isinstance(value, dict):
            value = {}
        if "id" in fields or "row_id" in fields:
            value["id"] = str(fields.get("id") or fields.get("row_id") or "")
        if "start" in fields:
            value["start"] = _projection_datetime_iso(_parse_projection_datetime(fields.get("start"), None)) or str(fields.get("start") or "")
        if "end" in fields:
            value["end"] = _projection_datetime_iso(_parse_projection_datetime(fields.get("end"), None)) or str(fields.get("end") or "")
        vals[projection_uid] = value
        for key in _projection_key_aliases(projection_uid)[1:]:
            vals.pop(key, None)
        data["_projection_values"] = vals
        save_meta = _save_projection_object_data(obj_repo, cls_name, internal_id, data, {
            "source": "projection",
            "projection_type": PROJECTION_SCHEDULE_TYPE,
            "projection_uid": projection_uid,
            "action": "move",
            "object_uid": object_uid,
            "fields": fields,
        })
    except _nodes_mod.AcceptRejected as e:
        return jsonify(_projection_accept_error_payload(e)), 200
    except Exception as e:
        return jsonify({"ok": False, "error": str(e), "message": {"text": str(e), "level": "danger"}}), 200
    return jsonify(_projection_move_success_payload(save_meta))


@client_bp.route("/api/projection/gantt/data", methods=["POST"])
@login_required
def api_projection_gantt_data():
    j = request.get_json(force=True) or {}
    repo_id = int(j.get("repo_id") or 0)
    class_name = str(j.get("class_name") or "").strip()
    node_id = str(j.get("node_id") or "").strip()
    if not repo_id or not class_name or not node_id:
        return jsonify({"ok": False, "error": "bad args"}), 400

    repo = _get_repo_or_404(repo_id)
    parsed = get_parsed_config(repo, models.db) or {}
    cls_cfg = (parsed.get("classes") or {}).get(class_name) or {}
    if not _is_projection_class_type(cls_cfg):
        return jsonify({"ok": False, "error": "class is not a projection"}), 400
    projection_type = str(cls_cfg.get("projection_type") or PROJECTION_KANBAN_TYPE).strip() or PROJECTION_KANBAN_TYPE
    if projection_type != PROJECTION_GANTT_TYPE:
        return jsonify({"ok": False, "error": "class is not a gantt_projection"}), 400

    projection_uid = _normalize_custom_process_uid(repo.config_uid, class_name, node_id)
    data = _get_projection_node_data(repo, cls_cfg, class_name, projection_uid)
    client_data = j.get("node_data")
    if isinstance(client_data, dict):
        for k in ("_projection_objects", "_projection_timer", "_projection_uid", "_projection_type", "_projection_tasks", "_projection_period_start", "_projection_period_end", "_projection_header", "_projection_scale"):
            if k in client_data:
                data[k] = client_data.get(k)
        _apply_projection_defaults_to_data(cls_cfg, data, repo.config_uid, class_name, projection_uid)

    now = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
    period_start = _parse_projection_datetime(data.get("_projection_period_start"), now) or now
    period_end = _parse_projection_datetime(data.get("_projection_period_end"), period_start + timedelta(days=30)) or (period_start + timedelta(days=30))
    if period_end <= period_start:
        period_end = period_start + timedelta(days=30)

    tasks = _normalize_gantt_tasks(data.get("_projection_tasks") or data.get("tasks"))
    objects = []
    generated = {t["id"]: t for t in tasks}
    for idx, uid in enumerate(_normalize_projection_object_ids(data.get("_projection_objects"))):
        obj = _projection_object_gantt_payload(repo, projection_uid, uid, idx)
        if not obj:
            continue
        gv = obj.get("gantt") or {}
        tid = str(gv.get("id") or obj.get("id") or "").strip()
        if not tid:
            tid = str(obj.get("id") or "")
            gv["id"] = tid
        if tid and tid not in generated:
            generated[tid] = {
                "id": tid,
                "caption": str(gv.get("title") or tid),
                "parent": str(gv.get("parent") or ""),
                "color": str(gv.get("color") or ""),
            }
        objects.append(obj)
    tasks = list(generated.values())

    return jsonify({
        "ok": True,
        "projection_uid": projection_uid,
        "projection_type": projection_type,
        "tasks": tasks,
        "objects": objects,
        "node_data": data,
        "header": str(data.get("_projection_header") or ""),
        "period_start": _projection_datetime_iso(period_start),
        "period_end": _projection_datetime_iso(period_end),
        "scale": str(data.get("_projection_scale") or "day"),
        "timer": _normalize_projection_timer(data.get("_projection_timer")),
    })



@client_bp.route("/api/projection/node-list/data", methods=["POST"])
@login_required
def api_projection_node_list_data():
    j = request.get_json(force=True) or {}
    repo_id = int(j.get("repo_id") or 0)
    class_name = str(j.get("class_name") or "").strip()
    node_id = str(j.get("node_id") or "").strip()
    if not repo_id or not class_name or not node_id:
        return jsonify({"ok": False, "error": "bad args"}), 400
    repo = _get_repo_or_404(repo_id)
    parsed = get_parsed_config(repo, models.db) or {}
    cls_cfg = (parsed.get("classes") or {}).get(class_name) or {}
    if not _is_projection_class_type(cls_cfg):
        return jsonify({"ok": False, "error": "class is not a projection"}), 400
    projection_type = str(cls_cfg.get("projection_type") or PROJECTION_NODE_LIST_TYPE).strip() or PROJECTION_NODE_LIST_TYPE
    if not _is_node_list_projection_type(projection_type):
        return jsonify({"ok": False, "error": "class is not a node_list_projection"}), 400
    projection_uid = _normalize_custom_process_uid(repo.config_uid, class_name, node_id)
    data = _get_projection_node_data(repo, cls_cfg, class_name, projection_uid)
    client_data = j.get("node_data")
    if isinstance(client_data, dict):
        for k in ("_projection_objects", "_projection_timer", "_projection_uid", "_projection_type", "_projection_header"):
            if k in client_data:
                data[k] = client_data.get(k)
        _apply_projection_defaults_to_data(cls_cfg, data, repo.config_uid, class_name, projection_uid)
    objects = []
    for uid in _normalize_projection_object_ids(data.get("_projection_objects")):
        obj = _projection_object_payload(repo, projection_uid, uid)
        if obj:
            objects.append(obj)
    return jsonify({
        "ok": True,
        "projection_uid": projection_uid,
        "projection_type": projection_type,
        "objects": objects,
        "node_data": data,
        "header": str(data.get("_projection_header") or ""),
        "timer": _normalize_projection_timer(data.get("_projection_timer")),
    })


@client_bp.route("/api/projection/html/data", methods=["POST"])
@login_required
def api_projection_html_data():
    j = request.get_json(force=True) or {}
    repo_id = int(j.get("repo_id") or 0)
    class_name = str(j.get("class_name") or "").strip()
    node_id = str(j.get("node_id") or "").strip()
    if not repo_id or not class_name or not node_id:
        return jsonify({"ok": False, "error": "bad args"}), 400
    repo = _get_repo_or_404(repo_id)
    parsed = get_parsed_config(repo, models.db) or {}
    cls_cfg = (parsed.get("classes") or {}).get(class_name) or {}
    if not _is_projection_class_type(cls_cfg):
        return jsonify({"ok": False, "error": "class is not a projection"}), 400
    projection_type = str(cls_cfg.get("projection_type") or PROJECTION_HTML_TYPE).strip() or PROJECTION_HTML_TYPE
    if not _is_html_projection_type(projection_type):
        return jsonify({"ok": False, "error": "class is not an html_projection"}), 400
    projection_uid = _normalize_custom_process_uid(repo.config_uid, class_name, node_id)
    data = _get_projection_node_data(repo, cls_cfg, class_name, projection_uid)
    client_data = j.get("node_data")
    if isinstance(client_data, dict):
        for k in ("_projection_html", "_html", "html", "_projection_timer", "_projection_uid", "_projection_type", "_projection_header"):
            if k in client_data:
                data[k] = client_data.get(k)
        _apply_projection_defaults_to_data(cls_cfg, data, repo.config_uid, class_name, projection_uid)
    html = data.get("_projection_html")
    if html is None:
        html = data.get("_html")
    if html is None:
        html = data.get("html")
    return jsonify({
        "ok": True,
        "projection_uid": projection_uid,
        "projection_type": projection_type,
        "html": str(html or ""),
        "node_data": data,
        "header": str(data.get("_projection_header") or ""),
        "timer": _normalize_projection_timer(data.get("_projection_timer")),
    })


@client_bp.route("/api/projection/gantt/move", methods=["POST"])
@login_required
def api_projection_gantt_move():
    j = request.get_json(force=True) or {}
    repo_id = int(j.get("repo_id") or 0)
    projection_uid = str(j.get("projection_uid") or "").strip()
    object_uid = str(j.get("object_uid") or "").strip()
    fields = j.get("fields") or {}
    if not isinstance(fields, dict):
        fields = {}
    if not repo_id or not projection_uid or not object_uid:
        return jsonify({"ok": False, "error": "bad args"}), 400
    repo = _get_repo_or_404(repo_id)
    try:
        obj_repo, cls_name, internal_id = _resolve_projection_object(repo, object_uid)
        data = _fetch_node_data_for_repo(obj_repo, cls_name, internal_id) or {}
        if not isinstance(data, dict):
            data = {}
        vals = data.get("_projection_values")
        if not isinstance(vals, dict):
            vals = {}
        existing = None
        for key in _projection_key_aliases(projection_uid):
            if key in vals:
                existing = vals.get(key)
                break
        value = _projection_read_jsonish(existing, {})
        if not isinstance(value, dict):
            value = {}
        if "id" in fields or "task_id" in fields:
            value["id"] = str(fields.get("id") or fields.get("task_id") or "")
        if "parent" in fields:
            value["parent"] = str(fields.get("parent") or "")
        if "start" in fields:
            value["start"] = _projection_datetime_iso(_parse_projection_datetime(fields.get("start"), None)) or str(fields.get("start") or "")
        if "end" in fields:
            value["end"] = _projection_datetime_iso(_parse_projection_datetime(fields.get("end"), None)) or str(fields.get("end") or "")
        vals[projection_uid] = value
        for key in _projection_key_aliases(projection_uid)[1:]:
            vals.pop(key, None)
        data["_projection_values"] = vals
        save_meta = _save_projection_object_data(obj_repo, cls_name, internal_id, data, {
            "source": "projection",
            "projection_type": PROJECTION_GANTT_TYPE,
            "projection_uid": projection_uid,
            "action": "move",
            "object_uid": object_uid,
            "fields": fields,
        })
    except _nodes_mod.AcceptRejected as e:
        return jsonify(_projection_accept_error_payload(e)), 200
    except Exception as e:
        return jsonify({"ok": False, "error": str(e), "message": {"text": str(e), "level": "danger"}}), 200
    return jsonify(_projection_move_success_payload(save_meta))

@client_bp.route("/api/projection/kanban/move", methods=["POST"])
@login_required
def api_projection_kanban_move():
    j = request.get_json(force=True) or {}
    repo_id = int(j.get("repo_id") or 0)
    projection_uid = str(j.get("projection_uid") or "").strip()
    object_uid = str(j.get("object_uid") or "").strip()
    column_id = str(j.get("column_id") or "").strip() or "__empty__"

    if not repo_id or not projection_uid or not object_uid:
        return jsonify({"ok": False, "error": "bad args"}), 400

    repo = _get_repo_or_404(repo_id)
    try:
        cfg_uid, cls_name, internal_id = _nodes_mod.parse_uid_any(object_uid)
    except Exception:
        cfg_uid, cls_name, internal_id = None, None, None
    cls_name = str(cls_name or "").strip()
    internal_id = str(internal_id or "").strip()
    if not cls_name or not internal_id:
        return jsonify({"ok": False, "error": "bad object uid"}), 400

    obj_repo = _repo_for_config_uid(repo, cfg_uid or repo.config_uid)
    data = _fetch_node_data_for_repo(obj_repo, cls_name, internal_id) or {}
    if not isinstance(data, dict):
        data = {}
    vals = data.get("_projection_values")
    if not isinstance(vals, dict):
        vals = {}
    if column_id == "__empty__":
        for key in _projection_key_aliases(projection_uid):
            vals.pop(key, None)
    else:
        vals[projection_uid] = column_id
        for key in _projection_key_aliases(projection_uid)[1:]:
            vals.pop(key, None)
    data["_projection_values"] = vals

    try:
        save_meta = _save_projection_object_data(obj_repo, cls_name, internal_id, data, {
            "source": "projection",
            "projection_type": PROJECTION_KANBAN_TYPE,
            "projection_uid": projection_uid,
            "action": "move",
            "object_uid": object_uid,
            "column_id": column_id,
        })
    except _nodes_mod.AcceptRejected as e:
        return jsonify(_projection_accept_error_payload(e)), 200
    except Exception as e:
        return jsonify({"ok": False, "error": str(e), "message": {"text": str(e), "level": "danger"}}), 200

    return jsonify(_projection_move_success_payload(save_meta))


@client_bp.route("/api/projection/diagram/move", methods=["POST"])
@login_required
def api_projection_diagram_move():
    j = request.get_json(force=True) or {}
    repo_id = int(j.get("repo_id") or 0)
    projection_uid = str(j.get("projection_uid") or "").strip()
    object_uid = str(j.get("object_uid") or "").strip()
    action = str(j.get("action") or "move").strip().lower()
    coords = j.get("coords") or {}
    fields = j.get("fields") or {}
    if not isinstance(coords, dict):
        coords = {}
    if not isinstance(fields, dict):
        fields = {}

    if not repo_id or not projection_uid or not object_uid:
        return jsonify({"ok": False, "error": "bad args"}), 400

    repo = _get_repo_or_404(repo_id)
    try:
        cfg_uid, cls_name, internal_id = _nodes_mod.parse_uid_any(object_uid)
    except Exception:
        cfg_uid, cls_name, internal_id = None, None, None
    cls_name = str(cls_name or "").strip()
    internal_id = str(internal_id or "").strip()
    if not cls_name or not internal_id:
        return jsonify({"ok": False, "error": "bad object uid"}), 400

    obj_repo = _repo_for_config_uid(repo, cfg_uid or repo.config_uid)
    data = _fetch_node_data_for_repo(obj_repo, cls_name, internal_id) or {}
    if not isinstance(data, dict):
        data = {}

    vals = data.get("_projection_values")
    if not isinstance(vals, dict):
        vals = {}

    if action == "unlink":
        for key in _projection_key_aliases(projection_uid):
            vals.pop(key, None)
        data["_projection_values"] = vals
    else:
        existing = None
        for key in _projection_key_aliases(projection_uid):
            if key in vals:
                existing = vals.get(key)
                break
        value = _normalize_diagram_projection_value(existing, 0)

        source = fields if action == "update" else coords
        if action == "move" and not source and fields:
            source = fields
        if not isinstance(source, dict):
            source = {}

        for key in ("x1", "y1", "x2", "y2"):
            if key in source:
                try:
                    value[key] = float(source.get(key))
                except Exception:
                    pass

        if action == "update":
            if "figure" in source:
                figure = str(source.get("figure") or "rectangle").strip().lower()
                if figure not in {"rectangle", "circle", "svg"}:
                    figure = "rectangle"
                value["figure"] = figure
            if "background" in source:
                bg = str(source.get("background") or "#ffffff").strip()
                if not re.match(r"^#[0-9a-fA-F]{3}([0-9a-fA-F]{3})?$", bg):
                    bg = "#ffffff"
                value["background"] = bg
            if "text" in source:
                value["text"] = str(source.get("text") or "")
            if "svg" in source:
                value["svg"] = str(source.get("svg") or "")

        # Keep the shape valid after manual editing.
        try:
            if float(value.get("x2", 0)) <= float(value.get("x1", 0)):
                value["x2"] = float(value.get("x1", 0)) + 20
            if float(value.get("y2", 0)) <= float(value.get("y1", 0)):
                value["y2"] = float(value.get("y1", 0)) + 20
        except Exception:
            pass

        vals[projection_uid] = value
        for key in _projection_key_aliases(projection_uid)[1:]:
            vals.pop(key, None)
        data["_projection_values"] = vals

    try:
        save_meta = _save_projection_object_data(obj_repo, cls_name, internal_id, data, {
            "source": "projection",
            "projection_type": PROJECTION_DIAGRAM_TYPE,
            "projection_uid": projection_uid,
            "action": action,
            "object_uid": object_uid,
            "fields": fields,
            "coords": coords,
        })
    except _nodes_mod.AcceptRejected as e:
        return jsonify(_projection_accept_error_payload(e)), 200
    except Exception as e:
        return jsonify({"ok": False, "error": str(e), "message": {"text": str(e), "level": "danger"}}), 200

    return jsonify(_projection_move_success_payload(save_meta))


@client_bp.route("/api/node/register", methods=["POST"])
@login_required
def api_node_register():
    """Register one node in a room defined in class Migration tab."""
    j = request.get_json(force=True) or {}
    repo_id = int(j.get("repo_id") or 0)
    class_name = str(j.get("class_name") or "")
    node_id = str(j.get("node_id") or "")
    # Allow override, but default is from config
    room_uid_req = str(j.get("room_uid") or "").strip()
    room_alias_req = str(j.get("room_alias") or "").strip()

    if not repo_id or not class_name or not node_id:
        return jsonify({"ok": False, "error": "bad args"}), 400

    repo = _get_repo_or_404(repo_id)
    cfg_uid = repo.config_uid

    room_uid = room_uid_req
    if not room_uid and room_alias_req:
        try:
            parsed = get_parsed_config(repo, models.db) or {}
            rooms_map = (parsed.get("rooms") or {}) if isinstance(parsed, dict) else {}
            room_uid = str((rooms_map or {}).get(room_alias_req) or "").strip()
        except Exception:
            room_uid = ""

    if not room_uid:
        try:
            parsed = get_parsed_config(repo, models.db) or {}
            cls_cfg = (parsed.get("classes") or {}).get(class_name) or {}
            room_uid = _resolve_class_default_room_uid(parsed, cls_cfg)
        except Exception:
            room_uid = ""

    if not room_uid:
        return jsonify({"ok": False, "error": "Room not specified"}), 400

    base_url = (repo.base_url or "").strip().rstrip("/")
    current = (request.host_url or "").rstrip("/")

    try:
        # Local register (preferred: avoids HTTP self-calls)
        if not base_url or base_url == current:
            cnt, delivery = _register_nodes_to_room_local(cfg_uid, class_name, room_uid, [node_id])
            if not cnt:
                err = str((delivery or {}).get("error") or "node not found or registration failed")
                return jsonify({"ok": False, "error": err, "delivery": delivery}), 404

            push = delivery.get("push") if isinstance(delivery, dict) else None
            transport = str((delivery or {}).get("transport") or "").strip().lower()
            if transport == "fcm" and isinstance(push, dict) and not push.get("ok"):
                err = str(push.get("error") or delivery.get("delivery_error") or "FCM push failed")
                return jsonify({
                    "ok": False,
                    "queued": True,
                    "count": cnt,
                    "room_uid": room_uid,
                    "delivery": delivery,
                    "error": err,
                    "message": {"text": f"Node queued, but FCM delivery failed: {err}", "level": "danger"}
                }), 202
            return jsonify({
                "ok": True,
                "count": cnt,
                "room_uid": room_uid,
                "delivery": delivery,
                "message": {"text": f"Registered in the room: {room_uid}", "level": "success"}
            })

        # Remote repo: fall back to remote API (no local deadlocks)
        delivery = _api_post_remote(repo, f"/api/config/{cfg_uid}/node/{class_name}/register/{room_uid}", json_data=[node_id])
        push = delivery.get("push") if isinstance(delivery, dict) else None
        transport = str((delivery or {}).get("transport") or "").strip().lower() if isinstance(delivery, dict) else ""
        if transport == "fcm" and isinstance(push, dict) and not push.get("ok"):
            err = str(push.get("error") or delivery.get("delivery_error") or "FCM push failed")
            return jsonify({
                "ok": False,
                "queued": True,
                "count": 1,
                "room_uid": room_uid,
                "delivery": delivery,
                "error": err,
                "message": {"text": f"Node queued, but FCM delivery failed: {err}", "level": "danger"}
            }), 202
        return jsonify({
            "ok": True,
            "count": 1,
            "room_uid": room_uid,
            "delivery": delivery,
            "message": {"text": f"Registered in the room: {room_uid}", "level": "success"}
        })
    except Exception as e:
        return jsonify({"ok": False, "error": str(e), "message": {"text": str(e), "level": "danger"}}), 500


# @client_bp.route("/node/<int:repo_id>/<path:class_name>/<path:node_id>")
# @login_required
# def node_view(repo_id: int, class_name: str, node_id: str):
#     repo = _get_repo_or_404(repo_id)
#     class_cfg = _get_class_cfg(repo, class_name) or {}
#     use_std = bool(class_cfg.get("use_standard_commands"))

#     base_url = (repo.base_url or "").strip().rstrip("/")
#     current = (request.host_url or "").rstrip("/")

#     data = {}
#     try:
#         if not base_url or base_url == current:
#             data = _node_local_get_data(repo.config_uid, class_name, node_id)
#         else:
#             payload = _api_get_remote(repo, f"/api/config/{repo.config_uid}/node/{class_name}/{node_id}")
#             data = (payload or {}).get("_data") or {}
#     except Exception:
#         data = {}

#     data_json = "{}"
#     try:
#         data_json = json.dumps(data or {}, ensure_ascii=False, indent=2)
#     except Exception:
#         data_json = "{}"    

#     return render_template(
#         "client/node.html",
#         title=f"{APP_TITLE} — {class_name}/{node_id}",
#         repo=repo,
#         repo_id=repo.id,
#         class_name=class_name,
#         node_id=node_id,
#         data=data or {},
#         use_standard_commands=use_std,
#         data_json=data_json,
#     )


@client_bp.route("/section/<path:section_code>")
@login_required
def section_view(section_code: str):
    repos = models.Repo.query.filter_by(user_id=_ngenie_effective_user_id()).all()
    if not repos:
        return redirect(url_for("client.sections_home"))

    code = "" if section_code == "__empty__" else section_code
    sections = _with_received_nodes_section(build_global_sections(repos, models.db), repos)
    valid_codes = {str(s.get("code") or "") for s in sections}
    if sections and code not in valid_codes:
        first = _default_section_code(sections)
        return redirect(url_for("client.section_view", section_code=(first if first != "" else "__empty__")))
    sec_name = DASHBOARD_SECTION_NAME if code == DASHBOARD_SECTION_CODE else ("<...>" if code == "" else next((s["name"] for s in sections if s["code"] == code), code))

    return render_template(
        "client/section.html",
        title=f"{APP_TITLE} — {sec_name}",
        repos=repos,
        sections=sections,
        section_code=code,
        section_name=sec_name,
        auto_refresh=AUTO_REFRESH_SECONDS,
        no_repos=(len(repos) == 0),
    )


@client_bp.route("/api/available_configs")
@login_required
def api_available_configs():
 
    Configuration = main.Configuration
    UserConfigAccess = getattr(main, "UserConfigAccess", None)

    
    own_cfgs = models.db.session.execute(
        select(Configuration).where(Configuration.user_id == current_user.id)
    ).scalars().all()

    
    shared_cfgs = []
    if UserConfigAccess is not None:
        shared_cfgs = models.db.session.execute(
            select(Configuration)
            .join(UserConfigAccess, UserConfigAccess.config_id == Configuration.id)
            .where(UserConfigAccess.user_id == current_user.id)
        ).scalars().all()

    # merge unique by uid
    by_uid = {}
    for c in list(own_cfgs) + list(shared_cfgs):
        try:
            by_uid[c.uid] = c
        except Exception:
            pass

    cfgs = list(by_uid.values())

    out = []
    for c in cfgs:
        out.append({
            "uid": c.uid,
            "name": c.name or "",
            "vendor": c.vendor or "",
            "version": getattr(c, "version", "") or "",
            "server_name": getattr(c, "server_name", "") or "",
            "last_modified": c.last_modified.isoformat() if getattr(c, "last_modified", None) else "",
        })
    
    out.sort(key=lambda x: (x["name"].lower(), x["uid"]))
    return jsonify(out)

# ---------- Repos management ----------

@client_bp.route("/repos")
@login_required
def repos_manage():
    repos = models.Repo.query.filter_by(user_id=_ngenie_effective_user_id()).order_by(models.Repo.id.desc()).all()
    return render_template("client/repos.html", title=f"{APP_TITLE} — Репозитории", repos=repos)


@client_bp.route("/settings", methods=["GET", "POST"])
@login_required
def client_settings():
    """Client settings page.

    Currently:
      - scanner_ws_url: ws://127.0.0.1:8765
      - scanner_ws_enabled: on/off
    """
    if request.method == "POST":
        ws_url = (request.form.get("scanner_ws_url") or "").strip()
        enabled = "1" if (request.form.get("scanner_ws_enabled") in ("1", "on", "true", "True")) else "0"
        _set_setting("scanner_ws_url", ws_url)
        _set_setting("scanner_ws_enabled", enabled)
        ngenie_enabled = "1" if (request.form.get("ngenie_enabled") in ("1", "on", "true", "True")) else "0"
        _set_setting("ngenie_enabled", ngenie_enabled)
        if _client_is_admin():
            show_json = "1" if (request.form.get("show_node_json") in ("1", "on", "true", "True")) else "0"
            _set_setting("show_node_json", show_json)
        flash("Settings saved", "success")
        return redirect(url_for("client.client_settings"))

    return render_template(
        "client/settings.html",
        title=f"{APP_TITLE} — Settings",
        scanner_ws_url=_get_setting("scanner_ws_url", "ws://127.0.0.1:8765"),
        scanner_ws_enabled=(_get_setting("scanner_ws_enabled", "1") not in ("0", "false", "False", "no", "off")),
        ngenie_enabled=_client_ngenie_enabled(),
        is_admin=_client_is_admin(),
        show_node_json=_client_show_node_json(),
    )


@client_bp.route("/repos/add", methods=["POST"])
@login_required
def repos_add():
    config_url = (request.form.get("config_url") or "").strip()

    try:
        base_url, config_uid, normalized_url = parse_config_url(config_url)
    except Exception as e:
        flash(f"Invalid ref: {e}", "error")
        return redirect(url_for("client.repos_manage"))

    exists = models.Repo.query.filter_by(user_id=current_user.id, config_uid=config_uid).first()
    if exists:
        flash("This configuration has already been added to the repository.", "info")
        return redirect(url_for("client.repos_manage"))

    try:
        current = (request.host_url or "").rstrip("/")
        is_current_server = (base_url.rstrip("/") == current)

        if is_current_server:
            # Local config on this server: enforce access list.
            if hasattr(_client_root_app_module(), "user_can_access_config") and not _client_root_app_module().user_can_access_config(current_user, config_uid):
                flash("You do not have access to this configuration.", "error")
                return redirect(url_for("client.repos_manage"))
            cfg = fetch_config_from_local_db(config_uid)
        else:
            cfg = fetch_config(normalized_url)
    except Exception as e:
        flash(f"Failed to read configuration: {e}", "error")
        return redirect(url_for("client.repos_manage"))

    vendor = cfg.get("vendor") or cfg.get("provider") or ""
    version = cfg.get("version") or ""
    display_name = cfg.get("display_name") or cfg.get("name") or ""
    name = display_name or f"{base_url} · {config_uid[:8]}"

    # Empty base_url means "read from this Designer DB".  Preserve the host
    # for an external link; otherwise the repo is falsely classified as local.
    stored_base_url = "" if is_current_server else base_url.rstrip("/")

    r = models.Repo(
        user_id=current_user.id,
        name=name,
        base_url=stored_base_url,
        config_uid=config_uid,
        config_url=normalized_url,
        vendor=vendor,
        version=version,
        display_name=display_name,
        username="",
        password="",
        config_json=json.dumps(cfg, ensure_ascii=False),
        config_cached_at=datetime.now(timezone.utc),
    )

    models.db.session.add(r)
    models.db.session.commit()

    row = models.RepoConfig.query.filter_by(repo_id=r.id).first()
    if not row:
        row = models.RepoConfig(repo_id=r.id, config_json=json.dumps(cfg, ensure_ascii=False))
        models.db.session.add(row)
    else:
        row.config_json = json.dumps(cfg, ensure_ascii=False)
    row.updated_at = datetime.now(timezone.utc)
    models.db.session.commit()

    _invalidate_repo_config_mem(r.id)
    return redirect(url_for("client.repos_manage"))

@client_bp.route("/repos/add_local", methods=["POST"])
@login_required
def repos_add_local():
    config_uid = (request.form.get("config_uid") or "").strip()
    if not config_uid:
        flash("config_uid not selected", "error")
        return redirect(url_for("client.repos_manage"))

    # enforce access list for local configs
    if hasattr(_client_root_app_module(), "user_can_access_config") and not _client_root_app_module().user_can_access_config(current_user, config_uid):
        flash("You do not have access to this configuration.", "error")
        return redirect(url_for("client.repos_manage"))

    
    exists = models.Repo.query.filter_by(user_id=current_user.id, config_uid=config_uid).first()
    if exists:
        flash("This configuration has already been added to the repository.", "info")
        return redirect(url_for("client.repos_manage"))

    try:
        cfg = fetch_config_from_local_db(config_uid)
    except Exception as e:
        flash(f"Failed to read configuration from database: {e}", "error")
        return redirect(url_for("client.repos_manage"))

    normalized_url = (request.host_url or "").rstrip("/") + f"/api/config/{config_uid}"
    vendor = cfg.get("vendor") or cfg.get("provider") or ""
    version = cfg.get("version") or ""
    display_name = cfg.get("display_name") or cfg.get("name") or ""
    name = display_name or f"local · {config_uid[:8]}"

    r = models.Repo(
        user_id=current_user.id,
        name=name,
        base_url="",  
        config_uid=config_uid,
        config_url=normalized_url,
        vendor=vendor,
        version=version,
        display_name=display_name,
        username="",
        password="",
        config_json=json.dumps(cfg, ensure_ascii=False),
        config_cached_at=datetime.now(timezone.utc),
    )

    models.db.session.add(r)
    models.db.session.commit()

    # cache table
    row = models.RepoConfig.query.filter_by(repo_id=r.id).first()
    if not row:
        row = models.RepoConfig(repo_id=r.id, config_json=json.dumps(cfg, ensure_ascii=False))
        models.db.session.add(row)
    else:
        row.config_json = json.dumps(cfg, ensure_ascii=False)
    row.updated_at = datetime.now(timezone.utc)
    models.db.session.commit()

    _invalidate_repo_config_mem(r.id)
    flash("Configuration added from the current server", "success")
    return redirect(url_for("client.repos_manage"))

@client_bp.route("/repos/<int:repo_id>/remove", methods=["POST"])
@login_required
def repos_remove(repo_id: int):
    r = models.Repo.query.get_or_404(repo_id)
    if r.user_id != current_user.id:
        abort(403)
    models.RepoConfig.query.filter_by(repo_id=r.id).delete(synchronize_session=False)
    models.db.session.delete(r)
    models.db.session.commit()
    _invalidate_repo_config_mem(r.id)
    return redirect(url_for("client.repos_manage"))


@client_bp.route("/repos/<int:repo_id>/update_api", methods=["POST"])
@login_required
def repo_update_api(repo_id: int):
    repo = models.Repo.query.get_or_404(repo_id)
    if repo.user_id != current_user.id:
        abort(403)

    repo.base_url = (request.form.get("base_url") or "").strip().rstrip("/")
    repo.username = request.form.get("username") or ""
    repo.password = request.form.get("password") or ""

    models.db.session.commit()
    _invalidate_repo_config_mem(repo.id)
    flash("API parameters saved", "success")
    return redirect(url_for("client.repos_manage"))


@client_bp.route("/repos/refresh_all", methods=["POST"])
@login_required
def repos_refresh_all():
    repos = models.Repo.query.filter_by(user_id=_ngenie_effective_user_id()).all()
    ok, fail = 0, 0

    for r in repos:
        try:
             
            if _client_repo_is_local(r):
                cfg = fetch_config_from_local_db(r.config_uid)
            else:
                cfg = fetch_config(r.config_url)

            row = models.RepoConfig.query.filter_by(repo_id=r.id).first()
            if not row:
                row = models.RepoConfig(repo_id=r.id, config_json=json.dumps(cfg, ensure_ascii=False))
                models.db.session.add(row)
            else:
                row.config_json = json.dumps(cfg, ensure_ascii=False)
            row.updated_at = datetime.now(timezone.utc)

            r.config_json = json.dumps(cfg, ensure_ascii=False)
            r.config_cached_at = datetime.now(timezone.utc)
            # Keep repository labels synchronized with the configuration itself.
            display_name = str(cfg.get("display_name") or cfg.get("name") or "").strip()
            if display_name:
                r.name = display_name
                r.display_name = display_name
            r.vendor = str(cfg.get("vendor") or cfg.get("provider") or r.vendor or "")
            r.version = str(cfg.get("version") or r.version or "")

            models.db.session.commit()
            _invalidate_repo_config_mem(r.id)
            ok += 1
        except Exception:
            fail += 1

    flash(f"Configurations have been updated: ok={ok}, fail={fail}", "success" if fail == 0 else "info")
    #return redirect(url_for("client.repos_manage"))
    return jsonify({"ok": True, "message": "Configurations have been updated"}), 200


def _parse_display_image_table(spec: str, data: Dict[str, Any]) -> Tuple[List[str], Dict[str, str]]:

    spec = (spec or "").strip()
    if not spec:
        return [], {}

    parts = [p.strip() for p in spec.split(",") if p.strip()]
    headers: List[str] = []
    values: Dict[str, str] = {}


    for p in parts:
        if "|" in p:
            title, expr = p.split("|", 1)
        else:
            
            chunks = p.split(None, 1)
            title = chunks[0]
            expr = chunks[1] if len(chunks) > 1 else ""

        title = (title or "").strip()
        expr = (expr or "").strip()

        if not title:
            continue

        headers.append(title)

        if expr.startswith("@"):
            key = expr[1:]
            v = (data or {}).get(key)
            values[title] = "" if v is None else (v if isinstance(v, str) else json.dumps(v, ensure_ascii=False))
        else:
            values[title] = expr

    return headers, values

# ---------- API ----------

def _client_section_json_safe(value: Any) -> Any:
    """Normalize values that are valid Node/runtime values but not stdlib JSON values.

    The application uses a custom JSON provider backed by ``json.dumps``.  Some
    generated/runtime handlers legitimately return ``Decimal`` (quantities,
    capacities, ledger values).  Keep this fix local to the section payload so
    one Decimal cannot break the whole client section refresh.
    """
    try:
        from decimal import Decimal
        if isinstance(value, Decimal):
            # Match Flask's normal Decimal behavior: preserve exact decimal text
            # instead of silently losing precision via float().
            return str(value)
    except Exception:
        pass
    if isinstance(value, dict):
        return {key: _client_section_json_safe(item) for key, item in value.items()}
    if isinstance(value, list):
        return [_client_section_json_safe(item) for item in value]
    if isinstance(value, tuple):
        return [_client_section_json_safe(item) for item in value]
    return value


@client_bp.route("/api/dashboard_layout", methods=["GET", "POST"])
@login_required
def api_dashboard_layout():
    key = "dashboard_layout_order"
    if request.method == "POST":
        payload = request.get_json(force=True) or {}
        order = payload.get("order") or []
        if not isinstance(order, list):
            order = []
        clean = [str(x) for x in order if str(x or "").strip()]
        _set_setting(key, json.dumps(clean, ensure_ascii=False))
        return jsonify({"ok": True, "order": clean})
    raw = _get_setting(key, "[]")
    try:
        order = json.loads(raw)
        if not isinstance(order, list):
            order = []
    except Exception:
        order = []
    return jsonify({"ok": True, "order": [str(x) for x in order]})


@client_bp.route("/api/class_metadata")
@login_required
def api_class_metadata():
    """Return full class JSON only when an explicit export needs it.

    Normal section refreshes carry compact card data; generated handlers,
    events and layouts are no longer repeated in the main response.
    """
    try:
        repo_id = int(request.args.get("repo_id") or 0)
    except Exception:
        repo_id = 0
    class_name = str(request.args.get("class_name") or "").strip()
    if not repo_id or not class_name:
        return jsonify({"ok": False, "error": "repo_id and class_name are required"}), 400
    repo = _get_repo_or_404(repo_id)
    parsed = get_parsed_config(repo, models.db) or {}
    cls = ((parsed.get("classes") or {}).get(class_name) or {})
    if not cls:
        return jsonify({"ok": False, "error": "class not found"}), 404
    return jsonify({"ok": True, "class_obj": cls})


@client_bp.route("/api/section_data")
@login_required
def api_section_data():
    section_code = request.args.get("section_code", "")
    # Preserve the original query for embedding models.  q_lower is used only
    # by the local lexical checks for singleton/custom-process nodes.
    q = (request.args.get("q") or "").strip()
    q_lower = q.lower()
    index_name = (request.args.get("index_name") or "").strip()
    index_value = request.args.get("index_value")
    tag_filter = (request.args.get("tag") or "").strip()

    if section_code == RAW_NODES_SECTION_CODE:
        items, meta = _build_raw_node_items(q=q)
        return jsonify(_client_section_json_safe({
            "ok": True,
            "items": items,
            "count": len(items),
            "nl_css": DEFAULT_NL_CSS,
            "meta": meta,
        }))

    repos = models.Repo.query.filter_by(user_id=_ngenie_effective_user_id()).all()
    is_dashboard = section_code == DASHBOARD_SECTION_CODE
    # The browser may still be on a section that became unavailable after role
    # changes.  Redirect it to the current default section instead of returning
    # an empty payload that looks like data disappeared.
    try:
        available_sections = _with_received_nodes_section(build_global_sections(repos, models.db), repos) if repos else []
        valid_codes = {str(s.get("code") or "") for s in available_sections}
        if available_sections and section_code not in valid_codes and section_code != RAW_NODES_SECTION_CODE:
            first = _default_section_code(available_sections)
            return jsonify({
                "ok": True,
                "items": [],
                "count": 0,
                "nl_css": DEFAULT_NL_CSS,
                "redirect_url": url_for("client.section_view", section_code=(first if first != "" else "__empty__")),
                "meta": {"classes_ui": [], "table_headers": [], "start_menu_cmds_ui": [], "timers_ui": [], "filter_indexes": [], "tag_filter": [], "selected_tag": tag_filter},
            })
    except Exception as e:
        print("section validity check failed:", e)

    merged: List[Dict[str, Any]] = []
    any_desc = False

    classes_ui: List[Dict[str, Any]] = []
    std_map: Dict[Tuple[int, str], bool] = {}            # (repo_id, class)->use_standard_commands
    display_name_map: Dict[Tuple[int, str], str] = {}    # (repo_id, class)->display_name
    commands_map: Dict[Tuple[int, str], str] = {}        # (repo_id, class)->commands string

    table_headers: List[str] = []
    table_headers_set: set = set()
    filter_indexes_map: Dict[str, Dict[str, Any]] = {}
    tag_filter_map: Dict[str, Dict[str, str]] = {}

    def remember_tags(data: Dict[str, Any], enabled: bool = True) -> List[str]:
        # _tags are global UI metadata now: collect and render them regardless of
        # the legacy show_tag_cloud class flag.
        tags = _normalize_node_tags(data)
        for tag in tags:
            tid = str(tag.get("id") or "")
            if tid and tid not in tag_filter_map:
                tag_filter_map[tid] = tag
        return [str(tag.get("id") or "") for tag in tags if str(tag.get("id") or "")]

    def parse_table_spec(spec: str) -> List[Tuple[str, str, bool]]:
        """
        "Title|@field,Title2|value" -> [(Title, field, True), (Title2, 'value', False)]
        If no '|': treat whole token as Title, value=''
        """
        out = []
        for raw in (spec or "").split(","):
            raw = (raw or "").strip()
            if not raw:
                continue
            if "|" in raw:
                t, v = raw.split("|", 1)
                t = (t or "").strip()
                v = (v or "").strip()
            else:
                t, v = raw, ""
            is_field = v.startswith("@")
            field = v[1:].strip() if is_field else v
            out.append((t, field, is_field))
        return out

    def build_table_values(spec: str, data: dict) -> Tuple[List[str], Dict[str, str]]:
        specs = parse_table_spec(spec)
        headers = []
        values = {}
        for title, key_or_val, is_field in specs:
            title = title or ""
            headers.append(title)
            if is_field:
                # NodeInput/NodeLink fields keep the raw UID in <field> and the
                # human-readable caption in <field>_view.  A table cover should
                # show the caption whenever it is available.
                view_value = data.get(key_or_val + "_view")
                v = view_value if view_value not in (None, "") else data.get(key_or_val)
                if v is None:
                    s = ""
                else:
                    try:
                        s = v if isinstance(v, str) else json.dumps(v, ensure_ascii=False)
                    except Exception:
                        s = str(v)
            else:
                s = key_or_val or ""
            values[title] = s
        return headers, values

    def build_cover_html(layout_to_use: Any, data: Dict[str, Any], class_name: str, node_id: str, show_tags: bool = False, render_cache: Optional[Dict[str, Any]] = None) -> str:
        """Cover renderer with per-node override via _data['_cover'].

        Priority:
        1) _cover in data
        2) display_image_web
        3) cover_image
        """
        rc = render_cache if isinstance(render_cache, dict) else {}
        assets_base_dir = rc.get("assets_base_dir")
        if not assets_base_dir:
            assets_base_dir = _userfiles_dir_for_repo(repo)
            rc["assets_base_dir"] = assets_base_dir
        nl_context = _nl_context(repo, class_name=class_name, node_id=node_id, shared_cache=rc)

        # 1) _cover override
        try:
            cov = data.get("_cover") if isinstance(data, dict) else None
            if cov:
                if isinstance(cov, (dict, list)):
                    return _cover_with_tags(_wrap_client_tpl_html(str(render_nodalayout_html(cov, data, assets_base_dir=assets_base_dir, context=nl_context) or ""), data), data, show_tags)
                if isinstance(cov, str):
                    s = cov.strip()
                    if s.startswith("[") or s.startswith("{"):
                        return _cover_with_tags(_wrap_client_tpl_html(str(render_nodalayout_html(s, data, assets_base_dir=assets_base_dir, context=nl_context) or ""), data), data, show_tags)
                    pic_layout = [[{"type": "Picture", "value": s, "width": -1}]]
                    return _cover_with_tags(_wrap_client_tpl_html(str(render_nodalayout_html(pic_layout, data, assets_base_dir=assets_base_dir, context=nl_context) or ""), data), data, show_tags)
        except Exception:
            pass

        # 2) class cover layouts (existing)
        try:
            if layout_to_use is not None and isinstance(data, dict):
                _fill_nodeinput_views(repo, parsed, layout_to_use, data, shared_cache=rc)
            return _cover_with_tags(_wrap_client_tpl_html(str(render_nodalayout_html(layout_to_use, data, assets_base_dir=assets_base_dir, context=nl_context) or ""), data), data, show_tags)
        except Exception:
            return _render_tags_html(data)

    start_menu_cmds_ui: List[Dict[str, Any]] = []
    timers_ui: List[Dict[str, Any]] = []

    for repo in repos:
        parsed = get_parsed_config(repo, models.db)
        if not parsed:
            continue

        cfg_obj = (parsed.get("cfg") or {}) if isinstance(parsed, dict) else {}
        repo_render_cache: Dict[str, Any] = {"parsed_by_config": {str(repo.config_uid or ""): parsed}}
        # runtime=client means the Android runtime.  The browser must never run
        # those actions: Android handlers and worker helpers are not available
        # in the web/server handler namespace.

        classes_by_section = parsed["classes_by_section"]
        if is_dashboard:
            cls_in_section = [c for c in ((parsed.get("cfg") or {}).get("classes") or []) if _class_dashboard_enabled(c)]
        else:
            cls_in_section = classes_by_section.get(section_code, []) if section_code != "" else classes_by_section.get("", [])

        cls_in_section = [
            c for c in (cls_in_section or [])
            if not bool(c.get("hidden"))
            and not bool(c.get('hide_web_client') or c.get('hideWebClient'))
            and not _is_print_form_class_type(c)
            and _client_user_can_access_class(repo.config_uid, str(c.get("name") or ""))
        ]

        sec_cmds = ""
        for s in (parsed.get("sections") or []):
            if (s.get("code") or "") == section_code:
                sec_cmds = (s.get("commands") or "").strip()
                break

        if sec_cmds and not is_dashboard:
            for raw in sec_cmds.split(","):
                raw = (raw or "").strip()
                if not raw:
                    continue
                if "|" in raw:
                    t, k = raw.split("|", 1)
                else:
                    t, k = raw, raw
                title = (t or "").strip()
                key = (k or "").strip()
                if not title or not key:
                    continue
                start_menu_cmds_ui.append({
                    "repo_id": repo.id,
                    "repo": repo.name,
                    "title": title,
                    "key": key,
                })

        
        for c in cls_in_section:
            cn = (c.get("name") or "").strip()
            if not cn:
                continue
            use_std = bool(c.get("use_standard_commands"))
            disp = (c.get("display_name") or cn).strip()
            cmds = (c.get("commands") or "").strip()

            std_map[(repo.id, cn)] = use_std
            display_name_map[(repo.id, cn)] = disp
            commands_map[(repo.id, cn)] = cmds
            for idx in (c.get("indexes") or []):
                if not isinstance(idx, dict) or not idx.get("filter_enabled"):
                    continue
                iname = str(idx.get("name") or "").strip()
                if not iname:
                    continue
                cur = filter_indexes_map.get(iname)
                item = {
                    "name": iname,
                    "kind": str(idx.get("kind") or "hash_index"),
                    "keys": str(idx.get("keys") or ""),
                    "filter_type": str(idx.get("filter_type") or "string"),
                    "filter_label": str(idx.get("filter_label") or "").strip(),
                    "filter_list_enabled": bool(idx.get("filter_list_enabled")),
                    "classes": [],
                }
                if not cur:
                    filter_indexes_map[iname] = item
                    cur = item
                if cn not in cur["classes"]:
                    cur["classes"].append(cn)

            classes_ui.append({
                "repo": repo.name,
                "repo_id": repo.id,
                "class": cn,
                "display_name": disp,
                "use_standard_commands": use_std,
                "commands": cmds,
                "repo_uid": repo.config_uid,
                "class_type": _class_type_value(c),
                "projection_type": str(c.get("projection_type") or ""),
                "dashboard_enabled": _class_dashboard_enabled(c),
                "dashboard_width": str(c.get("dashboard_width") or c.get("dashboardWidth") or "100"),
                "dashboard_top": bool(c.get("dashboard_top") or c.get("dashboardTop")),
            })

        # items
        for c in cls_in_section:
            cn = c.get("name")
            if not cn:
                continue

            ctype = c.get("class_type") or "data_node"

            cover_layout = _prepare_nodalayout(c.get("cover_image"))
            cover_web_layout = _prepare_nodalayout(c.get("display_image_web") or "")
            cover_render_layout = cover_web_layout if (not isinstance(cover_web_layout, str) or cover_web_layout.strip()) else cover_layout
            cover_table_layout = c.get("display_image_table") or ""

            # custom_process
            if _is_singleton_class_type(ctype):
                #node_id = f"{repo.config_uid}${cn}"
                node_id = _nodes_mod.normalize_own_uid(repo.config_uid, cn, "singleton")
                data = (c.get("_data") or {}).copy()
                try:
                    node_class = _load_server_node_class(repo.config_uid, cn)
                    node = node_class.get(node_id, repo.config_uid)
                    if node:
                        saved = node.get_data() or {}
                        if isinstance(saved, dict):
                            data.update(saved)   # сохранённое поверх дефолта
                except Exception:
                    pass

                
                _apply_projection_defaults_to_data(c, data, repo.config_uid, cn, node_id)
                data.setdefault("_id", node_id)

                if isinstance(data, dict) and data.get("_hidden"):
                    continue

                if not _client_user_can_access_node(repo.config_uid, cn, node_id, data if isinstance(data, dict) else {}):
                    continue

                if q:
                    sidx = data.get("_search_index")
                    if isinstance(sidx, str):
                        if q_lower not in sidx.lower():
                            continue
                    else:
                        try:
                            if q_lower not in json.dumps(data, ensure_ascii=False).lower():
                                continue
                        except Exception:
                            continue

                item_tag_ids = remember_tags(data, bool(c.get("show_tag_cloud")))
                if tag_filter and tag_filter not in item_tag_ids:
                    continue

                # cover html for web (priority: display_image_web else cover_image)
                display_image_html = ""
                display_image_html = build_cover_html(cover_render_layout, data, cn, node_id, bool(c.get("show_tag_cloud")), repo_render_cache)
                # try:
                #     if (cover_web_layout or "").strip():
                #         display_image_html = str(render_nodalayout_html(cover_web_layout, data))
                #     else:
                #         display_image_html = str(render_nodalayout_html(cover_layout, data))
                # except Exception:
                #     display_image_html = ""

                # table values (display_image_table spec)
                tv_headers, tv = build_table_values(cover_table_layout, data)
                for h in tv_headers:
                    if h not in table_headers_set:
                        table_headers_set.add(h)
                        table_headers.append(h)

                merged.append({
                    "repo": repo.name,
                    "repo_id": repo.id,
                    "class": cn,
                    "id": node_id,
                    "data": data,
                    "class_key": f"{int(repo.id)}:{cn}",
                    "is_custom_process": True,
                    "is_projection": _is_projection_class_type(c),
                    "projection_type": str(c.get("projection_type") or ""),
                    "display_image_html": display_image_html,
                    "tags": _normalize_node_tags(data),
                    "table_values": tv,
                    "use_standard_commands": bool(std_map.get((repo.id, cn), False)),
                    "repo_uid": repo.config_uid,
                    "dashboard_width": str(c.get("dashboard_width") or c.get("dashboardWidth") or "100"),
                    "dashboard_top": bool(c.get("dashboard_top") or c.get("dashboardTop")),
                })

                if "_sort_string_desc" in data:
                    any_desc = True
                continue

            # data_node
            nodes = _fetch_nodes_for_class(
                repo,
                config_uid=repo.config_uid,
                class_name=cn,
                q=q,
                limit=DEFAULT_LIMIT_PER_CLASS,
                index_name=index_name,
                index_value=index_value,
                parsed_config=parsed,
            )
            for n in nodes:
                data = n.get("_data") or {}
                node_id = n.get("_id") or data.get("_id") or ""

                if isinstance(data, dict) and data.get("_hidden"):
                    continue

                if not _client_user_can_access_node(repo.config_uid, cn, node_id, data if isinstance(data, dict) else {}):
                    continue

                item_tag_ids = remember_tags(data, bool(c.get("show_tag_cloud")))
                if tag_filter and tag_filter not in item_tag_ids:
                    continue

                # cover html for web (priority: display_image_web else cover_image)
                display_image_html = ""
                display_image_html = build_cover_html(cover_render_layout, data, cn, node_id, bool(c.get("show_tag_cloud")), repo_render_cache)
                # try:
                #     if (cover_web_layout or "").strip():
                #         display_image_html = str(render_nodalayout_html(cover_web_layout, data))
                #     else:
                #         display_image_html = str(render_nodalayout_html(cover_layout, data))
                # except Exception:
                #     display_image_html = ""

                # table values (display_image_table spec)
                tv_headers, tv = build_table_values(cover_table_layout, data)
                for h in tv_headers:
                    if h not in table_headers_set:
                        table_headers_set.add(h)
                        table_headers.append(h)

                merged.append({
                    "repo": repo.name,
                    "repo_id": repo.id,
                    "class": cn,
                    "id": node_id,
                    "data": data,
                    "class_key": f"{int(repo.id)}:{cn}",
                    "is_custom_process": False,
                    "is_projection": False,
                    "projection_type": "",
                    "display_image_html": display_image_html,
                    "tags": _normalize_node_tags(data),
                    "table_values": tv,
                    "use_standard_commands": bool(std_map.get((repo.id, cn), False)),
                    "repo_uid": repo.config_uid,
                    "dashboard_width": str(c.get("dashboard_width") or c.get("dashboardWidth") or "100"),
                    "dashboard_top": bool(c.get("dashboard_top") or c.get("dashboardTop")),
                })

                if "_sort_string_desc" in data:
                    any_desc = True

    # Process-local virtual banners are always placed at the very top of Home.
    if is_dashboard:
        try:
            from .ephemeral_banners import list_for as _list_ephemeral_banners
            banner_items = []
            for b in _list_ephemeral_banners([getattr(r, "config_uid", "") for r in repos]):
                kind = str(b.get("kind") or "layout")
                value = b.get("value")
                background = b.get("background")
                if kind == "html":
                    html = str(value or "")
                else:
                    try:
                        html = str(render_nodalayout_html(value, {"_id": "__banner__" + str(b.get("id") or "")}))
                    except Exception:
                        html = str(value or "")
                banner_data = {
                    "_id": "__banner__" + str(b.get("id") or ""),
                    "_virtual_banner": True,
                    "_virtual_banner_kind": kind,
                }
                if background:
                    banner_data["_background"] = str(background)
                banner_items.append({
                    "repo": "", "repo_id": 0, "repo_uid": b.get("config_uid") or "",
                    "class": "__Banner", "id": "__banner__" + str(b.get("id") or ""),
                    "data": banner_data,
                    "class_obj": {}, "is_custom_process": True, "is_projection": False,
                    "projection_type": "", "display_image_html": html, "tags": [],
                    "table_values": {}, "use_standard_commands": False,
                    "dashboard_width": str(b.get("size") or "25"), "dashboard_top": True,
                    "virtual_banner": True, "virtual_banner_kind": kind,
                })
            merged[0:0] = banner_items
        except Exception:
            traceback.print_exc()

    def sort_key(it: Dict[str, Any]) -> str:
        d = it.get("data") or {}
        if "_sort_string_desc" in d:
            return str(d.get("_sort_string_desc") or "")
        if "_sort_string" in d:
            return str(d.get("_sort_string") or "")
        return it.get("id") or ""

    # _semantic_find_ids() and the text/trigram indexes already return ids in
    # relevance order, and _nodes_storage_page() deliberately fetches rows in
    # that order.  The old unconditional sort below sorted the final cards again
    # by node id/_sort_string, completely destroying the model ranking.  Keep
    # ranked-search order; retain the normal section sort only when no ranked
    # search is active.
    selected_filter_kind = str((filter_indexes_map.get(index_name) or {}).get("kind") or "").strip().lower()
    ranked_filter_kinds = {
        "text_index", "text_index_full", "trigram_index",
        "semantic", "semantic_index", "semanic_index",
    }
    ranked_search_active = bool(q) or bool(
        index_name and index_value not in (None, "") and selected_filter_kind in ranked_filter_kinds
    )

    if is_dashboard:
        merged.sort(key=lambda it: (0 if bool(it.get("dashboard_top")) else 1, sort_key(it)))
    elif not ranked_search_active:
        merged.sort(key=sort_key, reverse=any_desc)

    return jsonify(_client_section_json_safe({
        "ok": True,
        "items": merged,
        "count": len(merged),
        "nl_css": DEFAULT_NL_CSS,
        "meta": {
            "classes_ui": classes_ui,
            "table_headers": table_headers,
            "start_menu_cmds_ui": start_menu_cmds_ui,
            "timers_ui": timers_ui,
            "filter_indexes": list(filter_indexes_map.values()),
            "tag_filter": list(tag_filter_map.values()),
            "selected_tag": tag_filter,
            "is_dashboard": bool(is_dashboard),
        }
    }))



@client_bp.route("/api/node/create", methods=["POST"])
@login_required
def api_node_create():
    payload = request.get_json(force=True) or {}
    repo_id = int(payload.get("repo_id") or 0)
    class_name = (payload.get("class_name") or "").strip()
    initial_data = payload.get("data") or payload.get("initial_data") or {}
    if not isinstance(initial_data, dict):
        initial_data = {}
    if not repo_id or not class_name:
        return jsonify({"ok": False, "error": "bad args"}), 400

    repo = models.Repo.query.filter_by(id=repo_id, user_id=current_user.id).first()
    if not repo:
        return jsonify({"ok": False, "error": "repo not found"}), 404

    # deny bulk delete for custom_process classes (they are virtual, not deletable)
    parsed = get_parsed_config(repo, models.db)
    try:
        cmeta = (parsed or {}).get("classes", {}).get(class_name) if isinstance(parsed, dict) else None
        if not isinstance(cmeta, dict):
            return jsonify({"ok": False, "error": "Forbidden"}), 403
        if _is_singleton_class_type(cmeta):
            return jsonify({"ok": False, "error": "singleton process cannot be created here"}), 400
    except Exception:
        return jsonify({"ok": False, "error": "Forbidden"}), 403

    base_url = (repo.base_url or "").strip().rstrip("/")
    current = (request.host_url or "").rstrip("/")

    try:
        if not base_url or base_url == current:
            
            node_id = _node_local_create(repo.config_uid, class_name, initial_data=initial_data)
        else:
            j = _api_post_remote(repo, f"/api/config/{repo.config_uid}/node/{class_name}", json_data=initial_data)
            node_id = None
            if isinstance(j, dict):
                node_id = (j.get("_id") or (j.get("_data") or {}).get("_id"))
            if not node_id:
                return jsonify({"ok": False, "error": "create: no node_id"}), 500

        return jsonify({"ok": True, "node_id": node_id, "config_uid": repo.config_uid})
    except Exception as e:
        return jsonify({"ok": False, 
                        "error": str(e),
                        "message": {"text": f"Handler error: {e}", "level": "error"},
                        }), 200


@client_bp.route("/api/node_delete", methods=["POST"])
@login_required
def api_node_delete():
    payload = request.get_json(force=True) or {}
    repo_id = int(payload.get("repo_id") or 0)
    class_name = (payload.get("class_name") or "").strip()
    node_id = (payload.get("node_id") or "").strip()
    if not repo_id or not class_name or not node_id:
        return jsonify({"ok": False, "error": "bad args"}), 400

    repo = models.Repo.query.filter_by(id=repo_id, user_id=current_user.id).first()
    if not repo:
        return jsonify({"ok": False, "error": "repo not found"}), 404

    # custom_process nodes are virtual and cannot be deleted
    parsed = get_parsed_config(repo, models.db)
    try:
        cmeta = (parsed or {}).get("classes", {}).get(class_name) if isinstance(parsed, dict) else None
        if not isinstance(cmeta, dict):
            return jsonify({"ok": False, "error": "Forbidden"}), 403
        if _is_singleton_class_type(cmeta):
            return jsonify({"ok": False, "error": "singleton process cannot be deleted"}), 400
    except Exception:
        return jsonify({"ok": False, "error": "Forbidden"}), 403
    if _client_repo_is_local(repo) and not _client_user_can_access_node(repo.config_uid, class_name, node_id, {}):
        return jsonify({"ok": False, "error": "Forbidden"}), 403

    base_url = (repo.base_url or "").strip().rstrip("/")
    current = (request.host_url or "").rstrip("/")

    try:
        if not base_url or base_url == current:
            
            _node_local_delete(repo.config_uid, class_name, node_id)
        else:
            _api_delete_remote(repo, f"/api/config/{repo.config_uid}/node/{class_name}/{node_id}")

        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"ok": False, 
                        "error": str(e),
                        "message": {"text": f"Handler error: {e}", "level": "error"},
                        }), 200


@client_bp.route("/api/node/bulk_delete", methods=["POST"])
@login_required
def api_node_bulk_delete():
    payload = request.get_json(force=True) or {}
    repo_id = int(payload.get("repo_id") or 0)
    class_name = (payload.get("class_name") or "").strip()
    ids = payload.get("ids") or []
    if not repo_id or not class_name or not isinstance(ids, list):
        return jsonify({"ok": False, "error": "bad args"}), 400

    repo = models.Repo.query.filter_by(id=repo_id, user_id=current_user.id).first()
    if not repo:
        return jsonify({"ok": False, "error": "repo not found"}), 404

    cfg_uid = repo.config_uid
    parsed = get_parsed_config(repo, models.db) or {}
    cmeta = (parsed.get("classes") or {}).get(class_name) if isinstance(parsed, dict) else None
    if not isinstance(cmeta, dict):
        return jsonify({"ok": False, "error": "Forbidden"}), 403
    if _is_singleton_class_type(cmeta):
        return jsonify({"ok": False, "error": "singleton process cannot be deleted"}), 400
    base_url = (repo.base_url or "").strip().rstrip("/")
    current = (request.host_url or "").rstrip("/")

    deleted = 0
    errors = []

    for node_id in ids:
        node_id = (str(node_id) or "").strip()
        if not node_id:
            continue
        try:
            if not base_url or base_url == current:
                if not _client_user_can_access_node(cfg_uid, class_name, node_id, {}):
                    errors.append({"id": node_id, "error": "Forbidden", "message": {"text": "Forbidden", "level": "error"}})
                    continue
                _node_local_delete(cfg_uid, class_name, node_id)
            else:
                _api_delete_remote(repo, f"/api/config/{cfg_uid}/node/{class_name}/{node_id}")
            deleted += 1
        except Exception as e:
            errors.append({"id": node_id, 
                           "error": str(e),
                           "message": {"text": f"Handler error: {e}", "level": "error"},
                           })

    return jsonify({"ok": True, "deleted": deleted, "errors": errors})



def _client_request_actor_for_external_api():
    """Resolve the user for external client API calls.

    Browser sessions can use the normal Flask-Login session. External clients
    (Android, curl, service integrations) can use HTTP Basic auth with the same
    account that has API access enabled.
    """
    try:
        if getattr(current_user, "is_authenticated", False):
            return current_user
    except Exception:
        pass

    auth = request.authorization
    if auth:
        check_api_auth = getattr(main, "check_api_auth", None)
        try:
            user = check_api_auth(auth.username, auth.password) if callable(check_api_auth) else None
        except Exception:
            user = None
        if user and bool(getattr(user, "can_api", False)):
            return user
        if user:
            abort(403)
    abort(401)


def _get_repo_by_config_uid_for_actor(config_uid: str, actor) -> models.Repo:
    config_uid = str(config_uid or "").strip()
    if not config_uid:
        abort(400)
    repo = models.Repo.query.filter_by(config_uid=config_uid, user_id=int(getattr(actor, "id", 0) or 0)).first()
    if not repo:
        abort(404)
    return repo


def _parse_external_print_form_request(j: Dict[str, Any]) -> Tuple[str, str, Dict[str, Any]]:
    """Parse JSON body for external PrintForm routes.

    Preferred body:
      {"print_form_class": "<config_uid>$<PrintFormClass>", "_data": {...}}

    Also accepts aliases for easier integrations:
      form_class, class_uid, print_form_uid, config_uid + class_name,
      data instead of _data.
    """
    j = j if isinstance(j, dict) else {}
    form_uid = str(
        j.get("print_form_class")
        or j.get("form_class")
        or j.get("class_uid")
        or j.get("print_form_uid")
        or ""
    ).strip()
    config_uid = str(j.get("config_uid") or "").strip()
    class_name = str(j.get("class_name") or j.get("print_class_name") or "").strip()

    if form_uid and "$" in form_uid:
        config_uid, class_name = form_uid.split("$", 1)
        config_uid = config_uid.strip()
        class_name = class_name.strip()
    elif form_uid and not class_name:
        class_name = form_uid

    data = j.get("_data")
    if data is None:
        data = j.get("data")
    if data is None:
        data = {}
    if not isinstance(data, dict):
        abort(400)

    if not config_uid or not class_name:
        abort(400)
    return config_uid, class_name, dict(data)


def _external_print_form_pdf_bytes(config_uid: str, class_name: str, data: Dict[str, Any], actor) -> Tuple[bytes, str, str]:
    repo = _get_repo_by_config_uid_for_actor(config_uid, actor)
    parsed = get_parsed_config(repo, models.db)
    if not parsed:
        abort(404)
    print_cls = ((parsed or {}).get("classes") or {}).get(class_name) or {}
    if not print_cls or not _is_print_form_class_type(print_cls):
        abort(404)

    # External API mode must behave like opening a PrintForm from a normal node:
    # the supplied _data is the original/base document, not the prepared print data.
    # It is injected into _basement_data, then onInputWeb/onInput + listener=onStartForm
    # is executed. The handler fills the final PrintForm _data, and only then the
    # HTML+Jinja template is rendered to PDF.
    base_data = dict(data or {})
    base_class_name = str(
        base_data.get("_class")
        or base_data.get("class_name")
        or base_data.get("_schema_class_name")
        or "ExternalDocument"
    ).strip() or "ExternalDocument"
    base_node_id = str(base_data.get("_id") or base_data.get("id") or f"external_{uuid.uuid4().hex[:12]}").strip()
    base_data.setdefault("_id", base_node_id)
    if base_class_name != "ExternalDocument":
        base_data.setdefault("_class", base_class_name)

    print_node_id = _print_form_node_id(repo.config_uid, class_name, base_class_name, base_node_id)

    _ctx_tokens = _nodes_mod.set_runtime_context(config_uid, parsed, system_user=_client_runtime_system_user_payload())
    try:
        print_data = _execute_print_form_start_handler(
            repo, parsed, print_cls, print_node_id, base_class_name, base_node_id, base_data
        )
        print_html = _render_print_html(repo, print_cls, print_data)
    finally:
        _nodes_mod.reset_runtime_context(_ctx_tokens)

    try:
        from weasyprint import HTML
        pdf_bytes = HTML(string=print_html, base_url=request.host_url).write_pdf()
    except Exception as e:
        raise RuntimeError(f"PDF export error: {e}")

    safe_name = re.sub(r"[^A-Za-z0-9_.-]+", "_", f"{config_uid}_{class_name}_{base_node_id}.pdf")[:180]
    return pdf_bytes, safe_name, print_html


def _send_pdf_to_raw_printer(pdf_bytes: bytes, printer_name: str, printer_port: Optional[int] = None, timeout: int = 20) -> Dict[str, Any]:
    dest = str(printer_name or "").strip()
    if not dest:
        raise ValueError("printer_name is required for raw printing")

    # Direct device path, for example /dev/usb/lp0.
    if dest.startswith("/dev/") or dest.startswith("/tmp/"):
        with open(dest, "wb") as f:
            f.write(pdf_bytes)
        return {"destination": dest, "bytes": len(pdf_bytes), "mode": "device"}

    raw = dest
    if raw.startswith("tcp://"):
        raw = raw[len("tcp://"):]
    if "://" in raw:
        raise ValueError("raw printer_name must be host:port, tcp://host:port, host, or a device path")

    host = raw
    port = int(printer_port or 9100)
    if raw.startswith("[") and "]" in raw:
        # Minimal IPv6 bracket notation: [addr]:9100
        end = raw.find("]")
        host = raw[1:end]
        tail = raw[end + 1:]
        if tail.startswith(":") and tail[1:]:
            port = int(tail[1:])
    elif ":" in raw:
        host, port_s = raw.rsplit(":", 1)
        if port_s.strip():
            port = int(port_s)

    host = host.strip()
    if not host:
        raise ValueError("raw printer host is empty")
    with socket.create_connection((host, port), timeout=timeout) as sock:
        sock.sendall(pdf_bytes)
    return {"destination": f"{host}:{port}", "bytes": len(pdf_bytes), "mode": "socket"}


def _send_pdf_to_cups_printer(pdf_bytes: bytes, printer_name: str, timeout: int = 60) -> Dict[str, Any]:
    printer = str(printer_name or "").strip()
    if not printer:
        raise ValueError("printer_name is required for CUPS printing")
    with tempfile.NamedTemporaryFile(prefix="nodalogic_printform_", suffix=".pdf", delete=True) as tmp:
        tmp.write(pdf_bytes)
        tmp.flush()
        cmd = ["lp", "-d", printer, tmp.name]
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout)
    if proc.returncode != 0:
        err = (proc.stderr or proc.stdout or "CUPS lp failed").strip()
        raise RuntimeError(err)
    return {"destination": printer, "bytes": len(pdf_bytes), "mode": "cups", "lp_output": (proc.stdout or "").strip()}



# ---------- Repository .enod import/export ----------

def _repo_config_dict(repo: models.Repo) -> Dict[str, Any]:
    row = models.RepoConfig.query.filter_by(repo_id=repo.id).first()
    raw = (row.config_json if row else None) or getattr(repo, "config_json", "") or "{}"
    try:
        cfg = json.loads(raw)
        return cfg if isinstance(cfg, dict) else {}
    except Exception:
        return {}


def _repo_cfg_locked_by_ngenie_code(cfg: Dict[str, Any]) -> bool:
    if not isinstance(cfg, dict):
        return False
    return bool(
        cfg.get('ngenie_code_locked')
        or cfg.get('ngenieCodeLocked')
        or cfg.get('nGenieCodeLocked')
    )


def _safe_export_filename(value: str, suffix: str = "") -> str:
    name = re.sub(r"[^A-Za-z0-9_.-]+", "_", str(value or "")).strip("._") or "repo"
    return name + suffix


@client_bp.route("/repos/<int:repo_id>/export_enod")
@login_required
def repo_export_enod(repo_id: int):
    repo = _get_repo_or_404(repo_id)
    cfg = _repo_config_dict(repo)
    if _repo_cfg_locked_by_ngenie_code(cfg):
        flash("This repository configuration is managed by nGenie Code and cannot be exported.", "error")
        return redirect(url_for("client.repos_manage"))
    payload = {
        "NodaLogicFormat": "ENOD",
        "NodaLogicType": "REPO_ENOD",
        "version": 1,
        "exported_at": datetime.now(timezone.utc).isoformat(),
        "repo": {
            "config_uid": repo.config_uid,
            "config_url": repo.config_url,
            "base_url": repo.base_url or "",
            "username": repo.username or "",
            # password is intentionally not exported
            "name": repo.name or "",
            "vendor": repo.vendor or "",
            "version": repo.version or "",
            "display_name": repo.display_name or "",
        },
        "config": cfg,
    }
    data = json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
    filename = _safe_export_filename(repo.display_name or repo.name or repo.config_uid or f"repo_{repo.id}", ".enod")
    return send_file(io.BytesIO(data), mimetype="application/json", as_attachment=True, download_name=filename)


@client_bp.route("/repos/import_enod", methods=["POST"])
@login_required
def repos_import_enod():
    f = request.files.get("enod_file") or request.files.get("file")
    if not f or not f.filename:
        flash("ENOD file is not selected", "error")
        return redirect(url_for("client.repos_manage"))
    try:
        data = f.read().decode("utf-8-sig")
        payload = json.loads(data)
        if not isinstance(payload, dict):
            raise ValueError("root must be JSON object")
        repo_meta = payload.get("repo") if isinstance(payload.get("repo"), dict) else {}
        cfg = payload.get("config") if isinstance(payload.get("config"), dict) else payload
        if not isinstance(cfg, dict):
            raise ValueError("config is missing")
        if _repo_cfg_locked_by_ngenie_code(cfg):
            raise ValueError("nGenie Code configurations cannot be imported into a repository")
        cfg_uid = str(repo_meta.get("config_uid") or cfg.get("uid") or cfg.get("config_uid") or "").strip()
        if not cfg_uid:
            raise ValueError("config uid is missing")
        current = (request.host_url or "").rstrip("/")
        config_url = str(repo_meta.get("config_url") or (current + f"/api/config/{cfg_uid}")).strip()
        base_url = str(repo_meta.get("base_url") or "").strip().rstrip("/")
        if not base_url:
            try:
                parsed_base_url, _parsed_uid, _normalized = parse_config_url(config_url)
            except Exception:
                parsed_base_url = ""
            # Old ENOD exports did not always store base_url.  Keep it empty only
            # for a configuration that actually exists in this Designer DB.
            if not (_client_cfg_by_uid(cfg_uid) is not None and parsed_base_url.rstrip("/") == current):
                base_url = parsed_base_url.rstrip("/")
        repo = models.Repo.query.filter_by(user_id=current_user.id, config_uid=cfg_uid).first()
        if not repo:
            repo = models.Repo(user_id=current_user.id, config_uid=cfg_uid, config_url=config_url)
            models.db.session.add(repo)
            models.db.session.flush()
        repo.config_url = config_url
        repo.base_url = base_url
        repo.username = str(repo_meta.get("username") or repo.username or "")
        # password is not imported/exported by design; keep existing value if any
        cfg_display_name = str(cfg.get("display_name") or cfg.get("name") or "").strip()
        # The configuration JSON is authoritative for its own name.  An ENOD
        # package may contain stale repository metadata from the source server.
        repo.name = cfg_display_name or str(repo_meta.get("name") or repo.name or cfg_uid)
        repo.vendor = str(cfg.get("vendor") or cfg.get("provider") or repo_meta.get("vendor") or repo.vendor or "")
        repo.version = str(cfg.get("version") or repo_meta.get("version") or repo.version or "")
        repo.display_name = cfg_display_name or str(repo_meta.get("display_name") or repo.display_name or repo.name or cfg_uid)
        cfg_json = json.dumps(cfg, ensure_ascii=False)
        repo.config_json = cfg_json
        repo.config_cached_at = datetime.now(timezone.utc)
        row = models.RepoConfig.query.filter_by(repo_id=repo.id).first()
        if not row:
            row = models.RepoConfig(repo_id=repo.id, config_json=cfg_json)
            models.db.session.add(row)
        else:
            row.config_json = cfg_json
            row.updated_at = datetime.now(timezone.utc)
        models.db.session.commit()
        _invalidate_repo_config_mem(repo.id)
        flash("Repository imported from ENOD", "success")
    except Exception as e:
        models.db.session.rollback()
        flash(f"Failed to import ENOD: {e}", "error")
    return redirect(url_for("client.repos_manage"))


# ---------- CodeFrame for local Python handlers ----------

def _require_admin_json() -> Optional[Any]:
    if not _client_is_admin():
        return jsonify({"ok": False, "error": "admin only"}), 403
    return None


def _codeframe_parse_ref(ref: str, fallback_class: str = "") -> Tuple[Optional[str], str]:
    """Parse CodeFrame reference.

    ClassName$method edits a method body inside ClassName.
    method edits a module-level function body in nodes_handlers_server for the current config.
    The fallback_class argument is intentionally no longer used for short refs because
    short refs are now reserved for common server functions.
    """
    raw = str(ref or "").strip()
    if "$" in raw:
        cls, method = raw.split("$", 1)
        cls, method = cls.strip(), method.strip()
    else:
        cls, method = "", raw
    if not method:
        raise ValueError("CodeFrame value must be ClassName$method or module_function")
    if cls and not re.match(r"^[A-Za-z_]\w*$", cls):
        raise ValueError("invalid class name")
    if not re.match(r"^[A-Za-z_]\w*$", method):
        raise ValueError("invalid function or method name")
    return (cls or None), method


def _local_handlers_file(config_uid: str) -> str:
    path = _handlers_file_path(str(config_uid or ""))
    os.makedirs(os.path.dirname(path), exist_ok=True)
    return path


def _local_handlers_code(config_uid: str) -> str:
    path = _local_handlers_file(config_uid)
    if os.path.exists(path):
        try:
            return Path(path).read_text(encoding="utf-8")
        except Exception:
            pass
    Configuration = getattr(main, "Configuration", None)
    if Configuration is not None:
        obj = models.db.session.execute(select(Configuration).where(Configuration.uid == config_uid)).scalar_one_or_none()
        if obj is not None:
            raw = getattr(obj, "nodes_server_handlers", "") or ""
            try:
                return base64.b64decode(raw).decode("utf-8") if raw else ""
            except Exception:
                return raw
    return ""


def _function_body_text(module_code: str, fn_node: ast.AST) -> str:
    lines = (module_code or "").splitlines()
    body = getattr(fn_node, "body", None) or []
    if not body:
        return ""
    start = body[0].lineno
    end = getattr(fn_node, "end_lineno", body[-1].lineno)
    block = lines[start-1:end]
    indents = [len(x) - len(x.lstrip()) for x in block if x.strip()]
    cut = min(indents) if indents else 0
    return "\n".join(x[cut:] if len(x) >= cut else x for x in block).rstrip()


def _extract_method_body(module_code: str, class_name: Optional[str], method_name: str) -> str:
    try:
        tree = ast.parse(module_code or "")
    except Exception:
        return ""
    if not class_name:
        for node in tree.body:
            if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)) and node.name == method_name:
                return _function_body_text(module_code, node)
        return ""
    for node in tree.body:
        if isinstance(node, ast.ClassDef) and node.name == class_name:
            for item in node.body:
                if isinstance(item, (ast.FunctionDef, ast.AsyncFunctionDef)) and item.name == method_name:
                    return _function_body_text(module_code, item)
    return ""


def _class_node_range(tree: ast.Module, class_name: str):
    for node in tree.body:
        if isinstance(node, ast.ClassDef) and node.name == class_name:
            return node
    return None


def _indent_body(body: str, spaces: int = 8) -> str:
    raw = str(body or "").rstrip()
    if not raw:
        raw = "pass"
    pad = " " * spaces
    return "\n".join(pad + line if line.strip() else "" for line in raw.splitlines())


def _upsert_module_function_body(module_code: str, function_name: str, body: str) -> str:
    module_code = module_code or ""
    try:
        tree = ast.parse(module_code or "")
    except SyntaxError as e:
        raise ValueError(f"handlers.py has syntax error before update: {e}")
    lines = module_code.splitlines()
    fn = None
    for node in tree.body:
        if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)) and node.name == function_name:
            fn = node
            break
    body_text = _indent_body(body, 4)
    if fn is not None:
        start = fn.body[0].lineno if fn.body else fn.lineno + 1
        end = getattr(fn, "end_lineno", fn.body[-1].lineno if fn.body else fn.lineno)
        new_lines = lines[:start-1] + body_text.splitlines() + lines[end:]
        out = "\n".join(new_lines).rstrip() + "\n"
    else:
        prefix = module_code.rstrip()
        add = f"\n\ndef {function_name}(input_data=None):\n{body_text}\n"
        out = prefix + add if prefix else add.lstrip()
    try:
        ast.parse(out)
    except SyntaxError as e:
        raise ValueError(f"updated handlers.py has syntax error: {e}")
    return out


def _upsert_method_body(module_code: str, class_name: Optional[str], method_name: str, body: str) -> str:
    if not class_name:
        return _upsert_module_function_body(module_code, method_name, body)
    module_code = module_code or ""
    try:
        tree = ast.parse(module_code or "")
    except SyntaxError as e:
        raise ValueError(f"handlers.py has syntax error before update: {e}")
    lines = module_code.splitlines()
    cls = _class_node_range(tree, class_name)
    body_text = _indent_body(body, 8)
    if cls is None:
        prefix = module_code.rstrip()
        if "from nodes import Node" not in module_code:
            prefix = ("from nodes import Node\n\n" + prefix.lstrip()) if prefix else "from nodes import Node"
        add = f"\n\nclass {class_name}(Node):\n    def __init__(self, id=None, config_uid=None):\n        super().__init__(id, config_uid)\n\n    def {method_name}(self, input_data=None):\n{body_text}\n"
        out = prefix.rstrip() + add
    else:
        method = None
        for item in cls.body:
            if isinstance(item, (ast.FunctionDef, ast.AsyncFunctionDef)) and item.name == method_name:
                method = item
                break
        if method is not None:
            start = method.body[0].lineno if method.body else method.lineno + 1
            end = getattr(method, "end_lineno", method.body[-1].lineno if method.body else method.lineno)
            new_lines = lines[:start-1] + body_text.splitlines() + lines[end:]
            out = "\n".join(new_lines).rstrip() + "\n"
        else:
            insert_at = getattr(cls, "end_lineno", len(lines))
            new_method = f"\n    def {method_name}(self, input_data=None):\n{body_text}"
            new_lines = lines[:insert_at] + new_method.splitlines() + lines[insert_at:]
            out = "\n".join(new_lines).rstrip() + "\n"
    try:
        ast.parse(out)
    except SyntaxError as e:
        raise ValueError(f"updated handlers.py has syntax error: {e}")
    return out


def _save_local_handlers_code(config_uid: str, code: str) -> None:
    path = _local_handlers_file(config_uid)
    Path(path).write_text(code or "", encoding="utf-8")
    Configuration = getattr(main, "Configuration", None)
    if Configuration is not None:
        obj = models.db.session.execute(select(Configuration).where(Configuration.uid == config_uid)).scalar_one_or_none()
        if obj is not None:
            obj.nodes_server_handlers = base64.b64encode((code or "").encode("utf-8")).decode("ascii")
            obj.last_modified = datetime.now(timezone.utc)
            models.db.session.commit()
            sync = getattr(main, "sync_server_methods_from_code", None)
            if callable(sync):
                try:
                    sync(obj, code or "")
                    models.db.session.commit()
                except Exception:
                    models.db.session.rollback()
    SERVER_HANDLERS_MEM.pop(config_uid, None)
    _SERVER_HANDLERS_NS_MEM.pop(config_uid, None)
    for key in list(_SERVER_NODE_CLASS_MEM.keys()):
        if key and key[0] == config_uid:
            _SERVER_NODE_CLASS_MEM.pop(key, None)


@client_bp.route("/api/codeframe/read", methods=["POST"])
@login_required
def api_codeframe_read():
    denied = _require_admin_json()
    if denied is not None:
        return denied
    j = request.get_json(force=True, silent=True) or {}
    repo_id = int(j.get("repo_id") or 0)
    repo = _get_repo_or_404(repo_id)
    if not _is_local_repo(repo):
        return jsonify({"ok": False, "error": "CodeFrame can edit only local Python handlers"}), 400
    try:
        if hasattr(main, "user_can_access_config") and not main.user_can_access_config(current_user, repo.config_uid):
            return jsonify({"ok": False, "error": "forbidden"}), 403
    except Exception:
        pass
    try:
        class_name, method_name = _codeframe_parse_ref(j.get("ref"), j.get("class_name"))
        code = _local_handlers_code(repo.config_uid)
        return jsonify({"ok": True, "class_name": class_name, "method_name": method_name, "code": _extract_method_body(code, class_name, method_name)})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 400


@client_bp.route("/api/codeframe/save", methods=["POST"])
@login_required
def api_codeframe_save():
    denied = _require_admin_json()
    if denied is not None:
        return denied
    j = request.get_json(force=True, silent=True) or {}
    repo_id = int(j.get("repo_id") or 0)
    repo = _get_repo_or_404(repo_id)
    if not _is_local_repo(repo):
        return jsonify({"ok": False, "error": "CodeFrame can edit only local Python handlers"}), 400
    try:
        if hasattr(main, "user_can_access_config") and not main.user_can_access_config(current_user, repo.config_uid):
            return jsonify({"ok": False, "error": "forbidden"}), 403
    except Exception:
        pass
    try:
        class_name, method_name = _codeframe_parse_ref(j.get("ref"), j.get("class_name"))
        module_code = _local_handlers_code(repo.config_uid)
        updated = _upsert_method_body(module_code, class_name, method_name, str(j.get("code") or ""))
        _save_local_handlers_code(repo.config_uid, updated)
        return jsonify({"ok": True, "message": {"text": "CodeFrame saved", "level": "success"}})
    except Exception as e:
        try:
            models.db.session.rollback()
        except Exception:
            pass
        return jsonify({"ok": False, "error": str(e), "message": {"text": str(e), "level": "danger"}}), 200


@client_bp.route("/api/print-form/pdf", methods=["POST"])
def api_external_print_form_pdf():
    actor = _client_request_actor_for_external_api()
    j = request.get_json(force=True, silent=True) or {}
    config_uid, class_name, data = _parse_external_print_form_request(j)
    try:
        pdf_bytes, safe_name, _ = _external_print_form_pdf_bytes(config_uid, class_name, data, actor)
    except RuntimeError as e:
        return Response(str(e), status=500, mimetype="text/plain")

    resp = make_response(pdf_bytes)
    resp.headers["Content-Type"] = "application/pdf"
    resp.headers["Content-Disposition"] = f'attachment; filename="{safe_name}"'
    return resp


@client_bp.route("/api/print-form/print", methods=["POST"])
def api_external_print_form_print():
    actor = _client_request_actor_for_external_api()
    j = request.get_json(force=True, silent=True) or {}
    config_uid, class_name, data = _parse_external_print_form_request(j)

    printer_name = str(j.get("printer_name") or j.get("printer") or "").strip()
    printer_type = str(j.get("printer_type") or "raw").strip().lower()
    if printer_type not in ("raw", "cups"):
        return jsonify({"ok": False, "error": "printer_type must be raw or cups"}), 400
    try:
        printer_port = j.get("printer_port")
        printer_port = int(printer_port) if printer_port not in (None, "") else None
    except Exception:
        return jsonify({"ok": False, "error": "printer_port must be an integer"}), 400

    try:
        pdf_bytes, safe_name, _ = _external_print_form_pdf_bytes(config_uid, class_name, data, actor)
        if printer_type == "cups":
            print_result = _send_pdf_to_cups_printer(pdf_bytes, printer_name)
        else:
            print_result = _send_pdf_to_raw_printer(pdf_bytes, printer_name, printer_port=printer_port)
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

    return jsonify({"ok": True, "filename": safe_name, "printer_type": printer_type, "print": print_result})


@client_bp.route("/print-form/<int:repo_id>/<path:print_class_name>")
@login_required
def print_form_open(repo_id: int, print_class_name: str):
    repo = _get_repo_or_404(repo_id)
    parsed = get_parsed_config(repo, models.db)
    if not parsed:
        abort(404)

    base_class_name = (request.args.get("base_class") or request.args.get("base_class_name") or "").strip()
    base_node_id = (request.args.get("base_node_id") or request.args.get("node_id") or "").strip()
    if not base_class_name or not base_node_id:
        abort(400)

    _ctx_tokens = _nodes_mod.set_runtime_context(repo.config_uid, parsed, system_user=_client_runtime_system_user_payload())
    try:
        print_cls, print_node_id, print_html, print_data = _build_print_form_runtime(
            repo, parsed, print_class_name, base_class_name, base_node_id
        )
    finally:
        _nodes_mod.reset_runtime_context(_ctx_tokens)
    try:
        data_json = json.dumps(print_data if isinstance(print_data, dict) else {}, ensure_ascii=False, indent=2)
    except Exception:
        data_json = "{}"

    pdf_url = url_for(
        "client.print_form_pdf",
        repo_id=repo.id,
        print_class_name=print_class_name,
        base_class=base_class_name,
        base_node_id=base_node_id,
    )

    return render_template(
        "client/node_form.html",
        title=f"{print_class_name} — {base_class_name}/{base_node_id}",
        node_id=print_node_id,
        discussion_node_id=print_node_id,
        class_name=print_class_name,
        repo=repo,
        repo_id=repo.id,
        error="",
        layout_html="",
        print_html=print_html,
        print_pdf_url=pdf_url,
        node_data=print_data,
        data_json=data_json,
        use_standard_commands=False,
        has_onshowweb=False,
        api_event_web=url_for("client.api_node_event_web"),
        api_save_url=url_for("client.api_node_save"),
        api_delete_url=url_for("client.api_node_delete"),
        api_register_url=url_for("client.api_node_register"),
        is_custom_process=False,
        is_projection=False,
        is_print_form=True,
        projection_type="",
        api_projection_kanban_data=url_for("client.api_projection_kanban_data"),
        api_projection_kanban_move=url_for("client.api_projection_kanban_move"),
        api_projection_diagram_data=url_for("client.api_projection_diagram_data"),
        api_projection_diagram_move=url_for("client.api_projection_diagram_move"),
        api_projection_schedule_data=url_for("client.api_projection_schedule_data"),
        api_projection_schedule_move=url_for("client.api_projection_schedule_move"),
        api_projection_gantt_data=url_for("client.api_projection_gantt_data"),
        api_projection_gantt_move=url_for("client.api_projection_gantt_move"),
            api_projection_node_list_data=url_for("client.api_projection_node_list_data"),
            api_projection_html_data=url_for("client.api_projection_html_data"),
            api_codeframe_read=url_for("client.api_codeframe_read"),
            api_codeframe_save=url_for("client.api_codeframe_save"),
            current_user_is_admin=_client_is_admin(),
            show_node_json=_client_show_node_json(),
        show_register_command=False,
        default_room_uid="",
        initial_message=None,
        ui_plugins=[],
        class_obj=print_cls,
        is_raw_node=False,
        print_forms_for_class=[],
    )


@client_bp.route("/print-form/<int:repo_id>/<path:print_class_name>/pdf")
@login_required
def print_form_pdf(repo_id: int, print_class_name: str):
    repo = _get_repo_or_404(repo_id)
    parsed = get_parsed_config(repo, models.db)
    if not parsed:
        abort(404)

    base_class_name = (request.args.get("base_class") or request.args.get("base_class_name") or "").strip()
    base_node_id = (request.args.get("base_node_id") or request.args.get("node_id") or "").strip()
    if not base_class_name or not base_node_id:
        abort(400)

    _ctx_tokens = _nodes_mod.set_runtime_context(repo.config_uid, parsed, system_user=_client_runtime_system_user_payload())
    try:
        _, _, print_html, _ = _build_print_form_runtime(repo, parsed, print_class_name, base_class_name, base_node_id)
    finally:
        _nodes_mod.reset_runtime_context(_ctx_tokens)
    try:
        from weasyprint import HTML
        pdf_bytes = HTML(string=print_html, base_url=request.host_url).write_pdf()
    except Exception as e:
        return Response(f"PDF export error: {escape(str(e))}", status=500, mimetype="text/plain")

    resp = make_response(pdf_bytes)
    resp.headers["Content-Type"] = "application/pdf"
    safe_name = re.sub(r"[^A-Za-z0-9_.-]+", "_", f"{print_class_name}_{base_class_name}_{base_node_id}.pdf")[:180]
    resp.headers["Content-Disposition"] = f'inline; filename="{safe_name}"'
    return resp


@client_bp.route("/node/<path:config_uid>/<path:class_name>/<path:node_id>")
@login_required
def node_form(config_uid: str, class_name: str, node_id: str):
    repo = _get_repo_by_config_uid_or_404(config_uid)
    if not repo:
        abort(404)

    if _client_repo_is_local(repo) and not _client_user_can_access_class(config_uid, class_name):
        abort(403)

    parsed = get_parsed_config(repo, models.db)
    if not parsed:
        abort(404)

    cls = parsed["classes"].get(class_name)
    if not cls:
        abort(404)

    # class-level PlugIn for Web (stored as JSON text in class.plug_in_web)
    def _parse_plugins(s: str):
        s = (s or "").strip()
        if not s:
            return []
        try:
            obj = json.loads(s)
            if isinstance(obj, list):
                return obj
            if isinstance(obj, dict):
                return [obj]
        except Exception:
            return []
        return []

    class_plugins_web = _parse_plugins(cls.get("plug_in_web") or "")

    use_std = bool(cls.get("use_standard_commands"))
    is_custom_process = _is_singleton_class_type(cls)
    is_projection = _is_projection_class_type(cls)
    projection_type = str(cls.get("projection_type") or "").strip()
    has_onshowweb = any((ev.get("event") or "") == "onShowWeb" for ev in (cls.get("events") or []))
    print_forms_for_class = _print_forms_for_class(parsed, class_name) if not _is_print_form_class_type(cls) else []

    def actions_for(event_name: str) -> List[Dict[str, Any]]:
        out = []
        for ev in (cls.get("events") or []):
            if ev.get("event") == event_name:
                out.extend(ev.get("actions") or [])
        return out

    base_url = (repo.base_url or "").strip().rstrip("/")
    current = (request.host_url or "").rstrip("/")

    layout = None
    node_data: Dict[str, Any] = {}
    ui_plugins = None

    try:
        # ---------------- REMOTE ----------------
        if base_url and base_url != current:
            for a in actions_for("onShowWeb"):
                m = (a.get("method") or "").strip()
                if not m:
                    continue
                runtime_method_name = m
                r = _api_post_remote(
                    repo,
                    f"/api/config/{repo.config_uid}/node/{class_name}/{node_id}/{m}",
                    json_data={},
                )
                if isinstance(r, dict) and isinstance(r.get("data"), dict) and ("_ui_layout" in r["data"]):
                    layout = r["data"].get("_ui_layout")

            if is_custom_process:
                # remote custom_process: используем то, что в конфиге
                node_data = (cls.get("_data") or {}).copy()
                _apply_projection_defaults_to_data(cls, node_data, repo.config_uid, class_name, node_id)
                node_data.setdefault("_id", node_id)
                node_data.setdefault("_class", class_name)
            else:
                n = _api_get_remote(repo, f"/api/config/{repo.config_uid}/node/{class_name}/{node_id}")
                node_data = (n.get("_data") or {}) if isinstance(n, dict) else {}

            # Remote mode: we can't execute server-side PlugIn(), but we can still
            # expose class-level plug_in_web to the template.
            if class_plugins_web:
                ui_plugins = class_plugins_web

        # ---------------- LOCAL ----------------
        else:
            node_class = _load_server_node_class(repo.config_uid, class_name)

            
            if is_custom_process:
                try:
                    node = node_class.get(node_id, repo.config_uid)
                except Exception:
                    node = None
                if not node:
                    node = node_class(node_id, repo.config_uid)
            else:
                node = node_class.get(node_id, repo.config_uid)
                if not node:
                    abort(404)

            _nodes_mod.CURRENT_NODE = node
            
            stored_data = {}
            try:
                stored_data = node.get_data() or {}
            except Exception:
                stored_data = {}

            if not isinstance(stored_data, dict):
                stored_data = {}

            if not is_custom_process and _client_repo_is_local(repo) and not _client_user_can_access_node(repo.config_uid, class_name, node_id, stored_data):
                abort(403)

            if is_custom_process:
                
                defaults = (cls.get("_data") or {})
                if not isinstance(defaults, dict):
                    defaults = {}

                merged = stored_data.copy()
                for k, v in defaults.items():
                    merged.setdefault(k, v)

                _apply_projection_defaults_to_data(cls, merged, repo.config_uid, class_name, node_id)
                merged.setdefault("_id", node_id)
                merged.setdefault("_class", class_name)
                node_data = merged

            else:
                node_data = stored_data.copy()

            
            try:
                node._data_cache = node_data.copy()
            except Exception:
                pass

                  
            
            for a in actions_for("onShowWeb"):
                m = (a.get("method") or "").strip()
                if m.lower() == "nodascript":
                    code = a.get("methodText") or a.get("method_text") or a.get("text") or a.get("code") or ""
                    _run_web_nodascript_action(code, node._data_cache, node, repo.config_uid)
                elif m and hasattr(node, m):
                    getattr(node, m)({})
                if not _layout_is_empty(getattr(node, "_ui_layout", None)):
                    layout = node._ui_layout

                post_m = str(a.get("postExecuteMethod") or a.get("post_execute_method") or "").strip().lower()
                if post_m == "nodascript":
                    post_code = a.get("postExecuteMethodText") or a.get("post_execute_method_text") or a.get("postText") or a.get("post_text") or ""
                    _run_web_nodascript_action(post_code, node._data_cache, node, repo.config_uid)

            # Class-level PlugIn for web (like calling self.PlugIn(...))
            if class_plugins_web:
                try:
                    if hasattr(node, "PlugIn"):
                        node.PlugIn(class_plugins_web)
                except Exception:
                    pass

            ui_message = getattr(node, "_ui_message", None)
            try:
                if hasattr(node, "_ui_message"):
                    delattr(node, "_ui_message")   # one-shot
            except Exception:
                pass
            
            ui_plugins = getattr(node, "_ui_plugins", None)
            try:
                if hasattr(node, "_ui_plugins"):
                    delattr(node, "_ui_plugins")   # one-shot
            except Exception:
                pass

            
            try:
                if getattr(node, "_data_cache", None) is not None and isinstance(node._data_cache, dict):
                    node_data = node._data_cache.copy()
                else:
                    node_data = node.get_data() or {}
            except Exception:
                pass
            #finally:
            #    _nodes_mod.CURRENT_NODE = None

            
            if not _layout_is_empty(getattr(node, "_ui_layout", None)):
                layout = node._ui_layout

    except Exception as e:
        editable_data = node_data if isinstance(node_data, dict) else {}
        try:
            data_json = json.dumps(editable_data, ensure_ascii=False, indent=2)
        except Exception:
            data_json = "{}"
        ui_message = str(e)
        return render_template(
            "client/node_form.html",
            title=f"{class_name}/{node_id}",
            node_id=node_id,
            class_name=class_name,
            repo=repo,
            repo_id=repo.id,
            error=str(e),
            layout_html="",
            node_data=node_data,
            data_json=data_json,
            use_standard_commands=use_std,
            has_onshowweb=has_onshowweb,
            api_event_web=url_for("client.api_node_event_web"),
            api_save_url=url_for("client.api_node_save"),
            api_delete_url=url_for("client.api_node_delete"),
            api_register_url=url_for("client.api_node_register"),
            is_custom_process=is_custom_process,
            is_projection=is_projection,
            is_print_form=False,
            print_html="",
            print_pdf_url="",
            print_forms_for_class=locals().get("print_forms_for_class", []),
            projection_type=projection_type,
            api_projection_kanban_data=url_for("client.api_projection_kanban_data"),
            api_projection_kanban_move=url_for("client.api_projection_kanban_move"),
            api_projection_diagram_data=url_for("client.api_projection_diagram_data"),
            api_projection_diagram_move=url_for("client.api_projection_diagram_move"),
            api_projection_schedule_data=url_for("client.api_projection_schedule_data"),
            api_projection_schedule_move=url_for("client.api_projection_schedule_move"),
            api_projection_gantt_data=url_for("client.api_projection_gantt_data"),
            api_projection_gantt_move=url_for("client.api_projection_gantt_move"),
            api_projection_node_list_data=url_for("client.api_projection_node_list_data"),
            api_projection_html_data=url_for("client.api_projection_html_data"),
            api_codeframe_read=url_for("client.api_codeframe_read"),
            api_codeframe_save=url_for("client.api_codeframe_save"),
            current_user_is_admin=_client_is_admin(),
            show_node_json=_client_show_node_json(),
            show_register_command=bool(cls.get("migration_register_command")) and bool(use_std),
            default_room_uid=_resolve_class_default_room_uid(parsed, cls),
            initial_message=ui_message,
            class_obj=cls,
            
        )

    # Default screen layout sources. Empty dynamic/data layouts are not forms:
    # ``[]`` is often only an Android placeholder and must not suppress the
    # dedicated browser layout.
    if _layout_is_empty(layout):
        data_layout = node_data.get("_layout") if isinstance(node_data, dict) else None
        layout = _first_usable_layout(
            data_layout,
            cls.get("init_screen_layout_web"),
            cls.get("init_screen_layout"),
        )

    # Resolve '^layout_id' via CommonLayouts
    layout = resolve_common_layout(parsed, layout)

    if layout is not None and isinstance(node_data, dict):
        _fill_nodeinput_views(repo, parsed, layout, node_data)


    layout_html = ""
    if layout is not None:
        try:
            layout_html = render_nodalayout_html(
                layout,
                node_data if isinstance(node_data, dict) else {},
                assets_base_dir=_userfiles_dir_for_repo(repo),
                context=_nl_context(repo, class_name=class_name, node_id=node_id),
            )
        except Exception:
            layout_html = ""

    editable_data = node_data if isinstance(node_data, dict) else {}
    try:
        data_json = json.dumps(editable_data, ensure_ascii=False, indent=2)
    except Exception:
        data_json = "{}"

    return render_template(
        "client/node_form.html",
        title=f"{class_name}/{node_id}",
        node_id=node_id,
        class_name=class_name,
        repo=repo,
        repo_id=repo.id,
        error="",
        layout_html=layout_html,
        node_data=node_data,
        data_json=data_json,
        use_standard_commands=use_std,
        has_onshowweb=has_onshowweb,
        api_event_web=url_for("client.api_node_event_web"),
        api_save_url=url_for("client.api_node_save"),
        api_delete_url=url_for("client.api_node_delete"),
        api_register_url=url_for("client.api_node_register"),
        is_custom_process=is_custom_process,
        is_projection=is_projection,
        is_print_form=False,
        print_html="",
        print_pdf_url="",
        print_forms_for_class=locals().get("print_forms_for_class", []),
        projection_type=projection_type,
        api_projection_kanban_data=url_for("client.api_projection_kanban_data"),
        api_projection_kanban_move=url_for("client.api_projection_kanban_move"),
        api_projection_diagram_data=url_for("client.api_projection_diagram_data"),
        api_projection_diagram_move=url_for("client.api_projection_diagram_move"),
            api_projection_schedule_data=url_for("client.api_projection_schedule_data"),
            api_projection_schedule_move=url_for("client.api_projection_schedule_move"),
            api_projection_gantt_data=url_for("client.api_projection_gantt_data"),
            api_projection_gantt_move=url_for("client.api_projection_gantt_move"),
            api_projection_node_list_data=url_for("client.api_projection_node_list_data"),
            api_projection_html_data=url_for("client.api_projection_html_data"),
            api_codeframe_read=url_for("client.api_codeframe_read"),
            api_codeframe_save=url_for("client.api_codeframe_save"),
            current_user_is_admin=_client_is_admin(),
            show_node_json=_client_show_node_json(),
        show_register_command=bool(cls.get("migration_register_command")) and bool(use_std),
        default_room_uid=_resolve_class_default_room_uid(parsed, cls),
        initial_message=ui_message,
        ui_plugins=ui_plugins,
        class_obj=cls,
        is_raw_node=False,
    )


def _parse_plugins_json(s: str):
    s = (s or "").strip()
    if not s:
        return []
    try:
        obj = json.loads(s)
        if isinstance(obj, list):
            return obj
        if isinstance(obj, dict):
            return [obj]
    except Exception:
        return []
    return []


@client_bp.route("/raw-node/<path:raw_node_id>")
@login_required
def raw_node_form(raw_node_id: str):
    RawNode = _server_model("RawNode")
    if RawNode is None:
        abort(404)

    obj = RawNode.query.filter_by(node_id=str(raw_node_id or "").strip()).first()
    if not obj:
        abort(404)
    if not _current_user_can_access_raw_node(raw_node_id, obj=obj):
        abort(403)

    payload = _raw_node_payload(obj)
    class_name, payload_node_id, node_data = _raw_node_identity(payload, raw_node_id)
    node_id = payload_node_id or str(raw_node_id or "").strip()
    node_data = dict(node_data or {})
    node_data.setdefault("_id", node_id)
    if class_name:
        node_data.setdefault("_class", class_name)
    node_data.setdefault("_raw_node_id", str(raw_node_id or ""))
    download_url = _raw_node_download_ref(payload, raw_node_id)
    event_download_url = _raw_node_download_url(raw_node_id)
    node_data.setdefault("_download_url", download_url)

    repo, parsed, cls = _resolve_raw_node_class(payload, class_name)
    if repo is None:
        repos = models.Repo.query.filter_by(user_id=_ngenie_effective_user_id()).all()
        repo = repos[0] if repos else None
    if repo is None:
        abort(404)

    cls = cls or {}
    parsed = parsed or get_parsed_config(repo, models.db) or {}
    if isinstance(cls, dict):
        resolved_name = _class_name_from_embedded_class(cls, fallback=class_name, payload=payload)
        if resolved_name:
            class_name = resolved_name
            node_data.setdefault("_class", class_name)
    ui_plugins = _parse_plugins_json(cls.get("plug_in_web") or "") if isinstance(cls, dict) else []
    ui_message = None
    use_std = bool(cls.get("use_standard_commands")) if isinstance(cls, dict) else False
    is_custom_process = _is_singleton_class_type(cls) if isinstance(cls, dict) else False
    is_projection = _is_projection_class_type(cls) if isinstance(cls, dict) else False
    projection_type = str(cls.get("projection_type") or "").strip() if isinstance(cls, dict) else ""
    has_onshowweb = any((ev.get("event") or "") == "onShowWeb" for ev in (cls.get("events") or [])) if isinstance(cls, dict) else False
    layout = None

    # Embedded raw-node classes may contain handlers. Execute onShowWeb through
    # the same download_url/raw-node path used by server-side PythonScript
    # handling, so a raw document behaves like a normal node form instead of a
    # plain JSON viewer.
    if has_onshowweb and _extract_raw_node_class_json(payload):
        try:
            ctx_builder = getattr(main, "resolve_download_url_node_context", None)
            event_runner = getattr(main, "execute_download_url_node_event", None)
            if callable(ctx_builder) and callable(event_runner):
                ctx = ctx_builder(node_id=node_id, fallback_class_name=class_name, download_url=event_download_url)
                event_runner(ctx, "onShowWeb", "", {})
                raw_event_node = ctx.get("node") if isinstance(ctx, dict) else None
                if raw_event_node is not None:
                    try:
                        event_data = raw_event_node.get_data() or {}
                        if isinstance(event_data, dict):
                            node_data = event_data.copy()
                            node_data.setdefault("_id", node_id)
                            node_data.setdefault("_raw_node_id", str(raw_node_id or ""))
                            node_data.setdefault("_download_url", download_url)
                            if class_name:
                                node_data.setdefault("_class", class_name)
                    except Exception:
                        pass
                    if not _layout_is_empty(getattr(raw_event_node, "_ui_layout", None)):
                        layout = getattr(raw_event_node, "_ui_layout", None)
                    if getattr(raw_event_node, "_ui_message", None) is not None:
                        ui_message = getattr(raw_event_node, "_ui_message", None)
                    if getattr(raw_event_node, "_ui_plugins", None) is not None:
                        ui_plugins = getattr(raw_event_node, "_ui_plugins", None)
        except Exception as e:
            ui_message = str(e)

    if _layout_is_empty(layout):
        data_layout = node_data.get("_layout") if isinstance(node_data, dict) else None
        layout = _first_usable_layout(
            data_layout,
            cls.get("init_screen_layout_web") if isinstance(cls, dict) else None,
            cls.get("init_screen_layout") if isinstance(cls, dict) else None,
        )

    layout = resolve_common_layout(parsed, layout)
    if layout is not None and isinstance(node_data, dict):
        try:
            _fill_nodeinput_views(repo, parsed, layout, node_data)
        except Exception:
            pass

    layout_html = ""
    if layout is not None:
        try:
            layout_html = render_nodalayout_html(
                layout,
                node_data if isinstance(node_data, dict) else {},
                assets_base_dir=_userfiles_dir_for_repo(repo),
                context=_nl_context(repo, class_name=class_name, node_id=node_id),
            ) or ""
        except Exception:
            layout_html = ""

    try:
        data_json = json.dumps(node_data if isinstance(node_data, dict) else {}, ensure_ascii=False, indent=2)
    except Exception:
        data_json = "{}"

    return render_template(
        "client/node_form.html",
        title=f"{RAW_NODES_SECTION_NAME} — {class_name or 'raw-node'}/{node_id}",
        node_id=node_id,
        class_name=class_name or "raw-node",
        repo=repo,
        repo_id=repo.id,
        error="",
        layout_html=layout_html,
        node_data=node_data,
        data_json=data_json,
        use_standard_commands=use_std,
        has_onshowweb=has_onshowweb,
        api_event_web=url_for("client.api_node_event_web"),
        api_save_url=url_for("client.api_node_save"),
        api_delete_url=url_for("client.api_node_delete"),
        api_register_url=url_for("client.api_node_register"),
        is_custom_process=is_custom_process,
        is_projection=is_projection,
        is_print_form=False,
        print_html="",
        print_pdf_url="",
        print_forms_for_class=locals().get("print_forms_for_class", []),
        projection_type=projection_type,
        api_projection_kanban_data=url_for("client.api_projection_kanban_data"),
        api_projection_kanban_move=url_for("client.api_projection_kanban_move"),
        api_projection_diagram_data=url_for("client.api_projection_diagram_data"),
        api_projection_diagram_move=url_for("client.api_projection_diagram_move"),
            api_projection_schedule_data=url_for("client.api_projection_schedule_data"),
            api_projection_schedule_move=url_for("client.api_projection_schedule_move"),
            api_projection_gantt_data=url_for("client.api_projection_gantt_data"),
            api_projection_gantt_move=url_for("client.api_projection_gantt_move"),
            api_projection_node_list_data=url_for("client.api_projection_node_list_data"),
            api_projection_html_data=url_for("client.api_projection_html_data"),
            api_codeframe_read=url_for("client.api_codeframe_read"),
            api_codeframe_save=url_for("client.api_codeframe_save"),
            current_user_is_admin=_client_is_admin(),
            show_node_json=_client_show_node_json(),
        show_register_command=False,
        default_room_uid="",
        initial_message=ui_message,
        ui_plugins=ui_plugins,
        class_obj=cls,
        is_raw_node=True,
        discussion_node_id=str(raw_node_id or ""),
    )


@client_bp.route("/api/s3/cached-image", methods=["GET"])
@login_required
def api_client_cached_s3_image():
    image_url = str(request.args.get("url") or "").strip()
    if not image_url:
        return jsonify({"ok": False, "error": "image_url_required"}), 400
    if not _is_cacheable_chat_image_url(image_url):
        return jsonify({"ok": False, "error": "unsupported_image_url"}), 400

    downloader = getattr(main, "_runtime_download_bytes_cached", None)
    if not callable(downloader):
        return jsonify({"ok": False, "error": "runtime_cache_unavailable"}), 500

    try:
        data = downloader(image_url, timeout=20)
    except Exception as e:
        return jsonify({"ok": False, "error": "image_cache_failed", "details": str(e)}), 502

    mimetype = _guess_image_mimetype_from_url(image_url)
    resp = Response(data, mimetype=mimetype)
    resp.headers["Cache-Control"] = "private, max-age=86400"
    resp.headers["X-Content-Type-Options"] = "nosniff"
    return resp


@client_bp.route("/api/node-discussion/by-node/<path:node_id>/messages", methods=["GET"])
@login_required
def api_client_node_discussion_messages(node_id: str):
    try:
        getter = getattr(main, "_get_node_discussion_messages_by_node_id", None)
        localize = getattr(main, "_localize_node_discussion_message_times", None)
        tz_name_fn = getattr(main, "_node_discussion_response_timezone_name", None)
        if callable(getter):
            try:
                messages = getter(node_id, viewer_user=current_user)
            except TypeError:
                messages = getter(node_id)
                messages = [m for m in messages if _message_dict_visible_to_current_user(m)]
        else:
            messages = []
        if callable(localize):
            tz_name = tz_name_fn() if callable(tz_name_fn) else None
            messages = [localize(m, tz_name=tz_name) for m in messages]
        return jsonify({"ok": True, "messages": messages, "count": len(messages)})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e), "messages": []}), 500


@client_bp.route("/api/node-discussion/by-node/<path:node_id>/messages", methods=["POST"])
@login_required
def api_client_post_node_discussion_message(node_id: str):
    data = request.get_json(silent=True) or {}
    if not isinstance(data, dict):
        data = {}

    text = data.get("text")
    if text is None:
        text = data.get("message")
    if text is None:
        text = data.get("body")
    image = data.get("image")
    image_url = data.get("image_url")

    if text in (None, "") and image in (None, "") and image_url in (None, ""):
        return jsonify({"ok": False, "error": "text_or_image_required", "node_id": node_id}), 400

    try:
        # The browser is authenticated with Flask-Login; sender_user from JSON
        # is intentionally ignored to prevent spoofing another user.
        sender_user = str(getattr(current_user, "email", "") or "").strip()
        sender_display_name = str(getattr(current_user, "config_display_name", "") or sender_user).strip()

        if not _current_user_can_access_node_discussion(node_id):
            # Allow the first web message only when the raw node itself belongs
            # to/is visible for this user. Existing ordinary-node discussions
            # are allowed by visible history above.
            if not _current_user_can_access_raw_node(node_id):
                return jsonify({"ok": False, "error": "forbidden", "messages": []}), 403

        targets = []
        if hasattr(main, "_find_node_discussion_targets"):
            targets = main._find_node_discussion_targets(node_id, sender_user=sender_user) or []
        if not targets and hasattr(main, "_create_node_discussion_targets_from_request"):
            targets = main._create_node_discussion_targets_from_request(data, node_id, sender_user) or []
        if not targets and hasattr(main, "_find_node_discussion_targets_from_raw_node"):
            targets = main._find_node_discussion_targets_from_raw_node(node_id, sender_user=sender_user) or []

        if not targets:
            return jsonify({
                "ok": False,
                "error": "node_discussion_target_required",
                "details": "No existing discussion target was found. For the first by-node message provide target_user/user_key/target_key/recipient/to/peer, or members/group_id.",
                "node_id": str(node_id or "").strip(),
                "results": [],
                "messages": [],
            }), 400

        if hasattr(main, "_node_discussion_thread_ref"):
            thread_ref = main._node_discussion_thread_ref(node_id, data.get("thread_ref"))
        else:
            thread_ref = str(data.get("thread_ref") or "")

        message_type = "image" if image not in (None, "") or image_url not in (None, "") else "text"
        payload = {
            "type": message_type,
            "thread_type": "node_discussion",
            "thread_ref": thread_ref,
            "node_id": str(node_id or "").strip(),
            "node_uid": str(node_id or "").strip(),
            "text": text or "",
        }
        if image is not None:
            payload["image"] = image
        if image_url is not None:
            payload["image_url"] = image_url
        if sender_user:
            payload["sender_user"] = sender_user
        if sender_display_name:
            payload["sender_display_name"] = sender_display_name

        extra = data.get("data")
        if isinstance(extra, dict):
            payload.update(extra)
            payload["type"] = message_type
            payload.pop("message_type", None)
            payload["thread_type"] = "node_discussion"
            payload["thread_ref"] = thread_ref
            payload["node_id"] = str(node_id or "").strip()
            payload["node_uid"] = str(node_id or "").strip()
        if sender_user:
            payload["sender_user"] = sender_user
        if sender_display_name:
            payload["sender_display_name"] = sender_display_name

        title = data.get("title") or sender_display_name or "Node discussion"
        body = text or data.get("body") or data.get("message") or "New message"

        results = []
        delivery_ok_count = 0
        accepted_count = 0
        saved_messages = []

        for target in targets:
            target_type = target.get("target_type")
            target_id = target.get("target_id")
            item_payload = dict(payload)

            if target_type == "group":
                item_payload["group_id"] = target_id
                result = main.send_message_to_group_global(target_id, title, body, item_payload, sender_user=sender_user)
            elif target_type == "user":
                item_payload["user_key"] = target_id
                result = main.send_message_to_user_global(target_id, title, body, item_payload, sender_user=sender_user)
            else:
                result = {"ok": False, "error": "unsupported_target_type"}

            client_message_id = result.get("client_message_id") if isinstance(result, dict) else None
            history_msg = None
            if client_message_id and hasattr(main, "NodeDiscussionMessage"):
                try:
                    history_msg = main.NodeDiscussionMessage.query.filter_by(client_message_id=client_message_id).first()
                except Exception:
                    history_msg = None

            if not history_msg and client_message_id and hasattr(main, "_save_node_discussion_history_message"):
                try:
                    history_msg = main._save_node_discussion_history_message(
                        node_id=node_id,
                        client_message_id=client_message_id,
                        sender_user=sender_user,
                        sender_display_name=sender_display_name,
                        target_type=target_type,
                        target_id=target_id,
                        text=text or "",
                        image=image,
                        image_url=image_url,
                        payload=item_payload,
                        delivery_status="accepted" if bool((result or {}).get("ok")) else "queued",
                    )
                except Exception:
                    history_msg = None

            if history_msg and hasattr(main, "_serialize_node_discussion_history_message"):
                accepted_count += 1
                saved_messages.append(main._serialize_node_discussion_history_message(history_msg))

            if bool((result or {}).get("ok")):
                delivery_ok_count += 1

            results.append({
                "target_type": target_type,
                "target_id": target_id,
                "client_message_id": client_message_id,
                "ok": bool(history_msg),
                "delivery_ok": bool((result or {}).get("ok")),
                "history_saved": bool(history_msg),
                "result": result,
            })

        return jsonify({
            "ok": accepted_count > 0,
            "node_id": str(node_id or "").strip(),
            "count": len(saved_messages),
            "accepted_count": accepted_count,
            "delivery_ok_count": delivery_ok_count,
            "messages": saved_messages,
            "results": results,
        }), (200 if accepted_count > 0 else 400)
    except Exception as e:
        return jsonify({"ok": False, "error": str(e), "messages": []}), 500


def _userfiles_root() -> str:
    base = os.path.join(os.path.dirname(__file__), "..", "UserFiles")
    base = os.path.abspath(base)
    try:
        os.makedirs(base, exist_ok=True)
    except Exception:
        pass
    return base


def _userfiles_dir_for_repo(repo) -> str:
    p = os.path.join(_userfiles_root(), str(repo.config_uid))
    try:
        os.makedirs(p, exist_ok=True)
    except Exception:
        pass
    return p


def _safe_filename(name: str) -> str:
    name = (name or "").replace("\\", "/")
    name = name.split("/")[-1]
    name = re.sub(r"[^A-Za-z0-9._\- ]+", "_", name)
    return name.strip()[:180]


def _truthy_form_value(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return value != 0
    if isinstance(value, str):
        return value.strip().lower() in ("1", "true", "yes", "on", "y", "да")
    return bool(value)


def _public_s3_url_for_key(object_key: str) -> str:
    endpoint = str(getattr(main, "S3_ENDPOINT", "") or "").rstrip("/")
    bucket = str(getattr(main, "S3_BUCKET", "") or "").strip("/")
    if not endpoint or not bucket:
        raise RuntimeError("S3 settings are not configured")
    return f"{endpoint}/{bucket}/{str(object_key).lstrip('/')}"


def _upload_userfile_to_s3(file_storage, *, owner_id: str, filename: str, content_type: str) -> str:
    s3_client = getattr(main, "s3", None)
    bucket = getattr(main, "S3_BUCKET", None)
    if s3_client is None or not bucket:
        raise RuntimeError("S3 client is not configured")
    base_name = _safe_filename(filename) or "file.bin"
    _, ext = os.path.splitext(base_name)
    object_key = f"uploads/client_userfiles/{owner_id or 'user'}/{uuid.uuid4().hex}{ext.lower()}"
    extra = {"ContentType": content_type or "application/octet-stream"}
    file_storage.stream.seek(0)
    s3_client.upload_fileobj(file_storage.stream, bucket, object_key, ExtraArgs=extra)
    try:
        invalidate = getattr(main, "_runtime_cache_invalidate", None)
        if callable(invalidate):
            invalidate(_public_s3_url_for_key(object_key))
    except Exception:
        pass
    return _public_s3_url_for_key(object_key)


@client_bp.route("/api/userfiles/<int:repo_id>/list")
@login_required
def api_userfiles_list(repo_id: int):
    repo = _get_repo_or_404(repo_id)
    d = _userfiles_dir_for_repo(repo)
    try:
        items = [f for f in os.listdir(d) if os.path.isfile(os.path.join(d, f))]
        items.sort(key=lambda s: s.lower())
    except Exception:
        items = []
    return jsonify({"ok": True, "files": items})


@client_bp.route("/api/userfiles/<int:repo_id>/upload", methods=["POST"])
@login_required
def api_userfiles_upload(repo_id: int):
    repo = _get_repo_or_404(repo_id)
    d = _userfiles_dir_for_repo(repo)
    if "file" not in request.files:
        return jsonify({"ok": False, "error": "no file"}), 400
    f = request.files["file"]
    if not f or not getattr(f, "filename", ""):
        return jsonify({"ok": False, "error": "empty file"}), 400
    name = _safe_filename(f.filename)
    if not name:
        return jsonify({"ok": False, "error": "bad filename"}), 400

    base, ext = os.path.splitext(name)
    out_name = name
    n = 1
    while os.path.exists(os.path.join(d, out_name)):
        out_name = f"{base}_{n}{ext}"
        n += 1

    upload_s3 = _truthy_form_value(request.form.get("upload_s3"))
    if upload_s3:
        try:
            owner_id = str(getattr(current_user, "id", "") or "user").strip() or "user"
            public_url = _upload_userfile_to_s3(
                f,
                owner_id=owner_id,
                filename=name,
                content_type=getattr(f, "mimetype", None) or "application/octet-stream",
            )
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500
        return jsonify({"ok": True, "filename": public_url, "url": public_url, "s3": True})

    try:
        f.save(os.path.join(d, out_name))
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
    return jsonify({"ok": True, "filename": out_name, "s3": False})


@client_bp.route("/api/userfiles/<int:repo_id>/delete", methods=["POST"])
@login_required
def api_userfiles_delete(repo_id: int):
    repo = _get_repo_or_404(repo_id)
    d = _userfiles_dir_for_repo(repo)
    payload = request.get_json(silent=True) or {}
    raw_filename = str(payload.get("filename") or "").strip()
    s3_key_from_public_url = getattr(main, "_s3_key_from_public_url", None)
    s3_key = ""
    try:
        if callable(s3_key_from_public_url):
            s3_key = s3_key_from_public_url(raw_filename)
    except Exception:
        s3_key = ""
    if s3_key and raw_filename.startswith(("http://", "https://")):
        try:
            s3_client = getattr(main, "s3", None)
            bucket = getattr(main, "S3_BUCKET", None)
            if s3_client is None or not bucket:
                raise RuntimeError("S3 client is not configured")
            s3_client.delete_object(Bucket=bucket, Key=s3_key)
            invalidate = getattr(main, "_runtime_cache_invalidate", None)
            if callable(invalidate):
                invalidate(raw_filename)
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500
        return jsonify({"ok": True, "s3": True})

    name = _safe_filename(raw_filename)
    if not name:
        return jsonify({"ok": False, "error": "no filename"}), 400
    p = os.path.join(d, name)
    try:
        if os.path.isfile(p):
            os.remove(p)
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
    return jsonify({"ok": True, "s3": False})


@client_bp.route("/api/userfiles/<int:repo_id>/raw/<path:filename>")
@login_required
def api_userfiles_raw(repo_id: int, filename: str):
    repo = _get_repo_or_404(repo_id)
    d = _userfiles_dir_for_repo(repo)
    name = _safe_filename(filename)
    if not name:
        abort(404)
    return send_from_directory(d, name, as_attachment=False)



@client_bp.route("/node_r/<int:repo_id>/<path:class_name>/<path:node_id>")
@login_required
def node_form_redirect(repo_id: int, class_name: str, node_id: str):
    repo = _get_repo_or_404(repo_id)
    return redirect(url_for("client.node_form", config_uid=repo.config_uid, class_name=class_name, node_id=node_id))


from typing import Any, Dict

def _coerce_input_value(payload: Dict[str, Any]) -> Any:

    v = payload.get("value")

    
    t = (payload.get("type") or payload.get("input_type") or payload.get("value_type") or "").lower()
    if t in ("checkbox", "switch", "bool", "boolean"):
        return bool(v)

    
    if isinstance(v, bool):
        return v

    # 2) number
    if t in ("number", "int", "integer", "float", "double"):
        if v is None or v == "":
            return None
        try:
            # если "12.3" -> float, если "12" -> int
            if isinstance(v, (int, float)):
                return v
            s = str(v).strip()
            if "." in s or "," in s:
                return float(s.replace(",", "."))
            return int(s)
        except Exception:
            return v

    
    # (тут ничего не приводим)
    return v

import re

def _parse_path_tokens(path: str):
    #
    tokens = []
    cur = ""
    i = 0
    while i < len(path):
        ch = path[i]
        if ch == ".":
            if cur:
                tokens.append(cur); cur = ""
            i += 1
            continue
        if ch == "[":
            if cur:
                tokens.append(cur); cur = ""
            end = path.find("]", i+1)
            if end == -1:
                tokens.append(path[i:])
                return tokens
            inside = path[i+1:end].strip()
            try:
                tokens.append(int(inside))
            except Exception:
                tokens.append(inside)
            i = end + 1
            continue
        cur += ch
        i += 1
    if cur:
        tokens.append(cur)
    return tokens

def _set_by_path(obj: dict, path: str, value):
    if not isinstance(obj, dict):
        return
    tokens = _parse_path_tokens(path or "")
    if not tokens:
        return
    cur = obj
    for k in tokens[:-1]:
        if isinstance(k, int):
            
            return
        if k not in cur or not isinstance(cur[k], (dict, list)):
            
            nxt = tokens[tokens.index(k)+1]  
            cur[k] = [] if isinstance(nxt, int) else {}
        cur = cur[k]
        if isinstance(cur, list):
            return
    last = tokens[-1]
    if isinstance(last, int):
        return
    cur[last] = value


def _apply_web_payload_to_node_data(node, payload: dict):
    if not isinstance(payload, dict):
        return

    base = None
    if getattr(node, "_data_cache", None) is not None and isinstance(node._data_cache, dict):
        base = node._data_cache
    elif getattr(node, "_data", None) is not None and isinstance(node._data, dict):
        base = node._data
    else:
        base = {}
        try:
            node._data_cache = base
        except Exception:
            pass

    listener = payload.get("listener") or payload.get("id")
    if listener is not None:
        base["listener"] = str(listener)

    el_id = payload.get("id")
    if el_id:
        if "value" in payload:
            base[str(el_id)] = payload.get("value")

        if "date_iso" in payload:
            base["_d" + str(el_id)] = payload.get("date_iso")

    try:
        p = payload.get("path")
        if isinstance(p, str) and p.strip() and ("value" in payload):
            _set_by_path(base, p.strip(), payload.get("value"))
    except Exception:
        pass

    # ВАЖНО: дополнительные значения от UI, например <field>_view
    # DatasetField / DatasetInput могут присылать:
    # "extra": [{"path": "customer_view", "value": "ООО Ромашка"}]
    extra = payload.get("extra")
    if isinstance(extra, list):
        for item in extra:
            if not isinstance(item, dict):
                continue
            try:
                ep = item.get("path")
                if isinstance(ep, str) and ep.strip() and ("value" in item):
                    _set_by_path(base, ep.strip(), item.get("value"))
            except Exception:
                pass

    fd = payload.get("full_data")
    if isinstance(fd, dict):
        try:
            for k, v in fd.items():
                base[k] = v
        except Exception:
            pass

    dv = payload.get("dialog_values")
    if isinstance(dv, dict):
        for p, v in dv.items():
            if isinstance(p, str) and p.strip():
                _set_by_path(base, p.strip(), v)

    passthrough_keys = [
        "row", "col", "row_id", "col_id",
        "selected", "selected_ids",
        "page", "page_size",
        "sort", "filter",
    ]
    for k in passthrough_keys:
        if k in payload:
            base[k] = payload[k]

    try:
        if getattr(node, "_data", None) is None or not isinstance(node._data, dict):
            node._data = {}
        node._data.update(base)
    except Exception:
        pass


def _raw_node_listener_from_payload(payload: Dict[str, Any]) -> str:
    try:
        if isinstance(payload, dict):
            return str(payload.get("listener") or payload.get("id") or "").strip()
    except Exception:
        pass
    return ""


def _raw_node_has_event_action(class_obj: Dict[str, Any], event_name: str, listener: str = "") -> bool:
    if not isinstance(class_obj, dict):
        return False
    for ev in (class_obj.get("events") or class_obj.get("Events") or []):
        if not isinstance(ev, dict):
            continue
        if ev.get("event") != event_name:
            continue
        ev_listener = str(ev.get("listener") or "").strip()
        if listener:
            if ev_listener and ev_listener != listener:
                continue
        else:
            if ev_listener:
                continue
        actions = ev.get("actions") or ev.get("Actions") or []
        return bool(actions)
    return False


def _raw_node_runner_listener(class_obj: Dict[str, Any], event_name: str, listener: str = "") -> str:
    """Match normal web-event listener semantics for app.py's stricter runner."""
    listener = str(listener or "").strip()
    if not isinstance(class_obj, dict):
        return listener
    if listener:
        for ev in (class_obj.get("events") or class_obj.get("Events") or []):
            if not isinstance(ev, dict) or ev.get("event") != event_name:
                continue
            if str(ev.get("listener") or "").strip() == listener and (ev.get("actions") or ev.get("Actions") or []):
                return listener
        for ev in (class_obj.get("events") or class_obj.get("Events") or []):
            if not isinstance(ev, dict) or ev.get("event") != event_name:
                continue
            if not str(ev.get("listener") or "").strip() and (ev.get("actions") or ev.get("Actions") or []):
                return ""
    return listener



def _debug_break_view_payload(repo: models.Repo, payload: Dict[str, Any]) -> Dict[str, Any]:
    if not isinstance(payload, dict):
        payload = {}
    items = payload.get("items") or payload.get("nodes") or []
    title = str(payload.get("title") or "Debug nodes")
    objects = []
    for uid in _normalize_projection_object_ids(items):
        obj = _projection_object_payload(repo, "__debug_break__", uid)
        if obj:
            objects.append(obj)
    return {"title": title, "objects": objects}


def _render_raw_node_layout_response(repo, parsed: Dict[str, Any], class_name: str, node_id: str, layout: Any, node_data: Dict[str, Any]) -> str:
    layout = resolve_common_layout(parsed, layout)
    if layout is None:
        return ""
    try:
        if isinstance(node_data, dict):
            _fill_nodeinput_views(repo, parsed, layout, node_data)
    except Exception:
        pass
    try:
        return render_nodalayout_html(
            layout,
            node_data if isinstance(node_data, dict) else {},
            assets_base_dir=_userfiles_dir_for_repo(repo),
            context=_nl_context(repo, class_name=class_name, node_id=node_id),
        ) or ""
    except Exception:
        return ""


def _api_raw_node_event_web(j: Dict[str, Any]):
    """Handle UI events for raw-nodes that carry embedded _class JSON.

    Returns None when the raw-node only references a normal server class by
    string; in that case the existing normal-node event path remains the best
    match.
    """
    raw_node_id = str(j.get("raw_node_id") or j.get("raw_id") or "").strip()
    if not raw_node_id:
        return None

    RawNode = _server_model("RawNode")
    if RawNode is None:
        return jsonify({"ok": False, "error": "raw_node_model_unavailable"}), 404

    obj = RawNode.query.filter_by(node_id=raw_node_id).first()
    if not obj:
        return jsonify({"ok": False, "error": "raw_node_not_found"}), 404
    if not _current_user_can_access_raw_node(raw_node_id, obj=obj):
        return jsonify({"ok": False, "error": "forbidden"}), 403

    payload_json = _raw_node_payload(obj)
    embedded_class = _extract_raw_node_class_json(payload_json)
    if not embedded_class:
        return None

    event = str(j.get("event") or "").strip()
    payload = j.get("payload") or {}
    if not isinstance(payload, dict):
        payload = {}

    class_name, payload_node_id, node_data = _raw_node_identity(payload_json, raw_node_id)
    class_name = _class_name_from_embedded_class(embedded_class, fallback=class_name, payload=payload_json)
    node_id = payload_node_id or str(j.get("node_id") or raw_node_id).strip()
    download_url = _raw_node_download_ref(payload_json, raw_node_id)
    event_download_url = _raw_node_download_url(raw_node_id)

    repo, parsed, cls = _resolve_raw_node_class(payload_json, class_name)
    if repo is None:
        repos = models.Repo.query.filter_by(user_id=_ngenie_effective_user_id()).order_by(models.Repo.id.asc()).all()
        repo = repos[0] if repos else None
    if repo is None:
        return jsonify({"ok": False, "error": "repo_not_found"}), 404
    parsed = parsed or get_parsed_config(repo, models.db) or {}
    cls = cls or embedded_class

    listener = _raw_node_listener_from_payload(payload)
    if not _raw_node_has_event_action(cls, event, listener):
        return jsonify({
            "ok": True,
            "noop": True,
            "handled": {"class_name": class_name, "node_id": node_id},
            "node_data": {},
            "patches": [],
        })

    node_data = dict(node_data or {})
    node_data.setdefault("_id", node_id)
    node_data.setdefault("_class", class_name)
    node_data.setdefault("_raw_node_id", raw_node_id)
    node_data.setdefault("_download_url", download_url)

    try:
        ctx_builder = getattr(main, "resolve_download_url_node_context", None)
        event_runner = getattr(main, "execute_download_url_node_event", None)
        if not callable(ctx_builder) or not callable(event_runner):
            return jsonify({"ok": False, "error": "raw_node_event_runtime_unavailable"}), 500

        ctx = ctx_builder(node_id=node_id, fallback_class_name=class_name, download_url=event_download_url)
        raw_event_node = ctx.get("node") if isinstance(ctx, dict) else None
        if raw_event_node is not None:
            try:
                if getattr(raw_event_node, "_data", None) is None or not isinstance(raw_event_node._data, dict):
                    raw_event_node._data = {}
                raw_event_node._data.update(node_data)
            except Exception:
                pass
            _apply_web_payload_to_node_data(raw_event_node, payload)

        event_runner(ctx, event, _raw_node_runner_listener(cls, event, listener), payload)

        if raw_event_node is not None:
            try:
                updated = raw_event_node.get_data() or {}
                if isinstance(updated, dict):
                    node_data = updated.copy()
            except Exception:
                pass

        node_data.setdefault("_id", node_id)
        node_data.setdefault("_class", class_name)
        node_data.setdefault("_raw_node_id", raw_node_id)
        node_data.setdefault("_download_url", download_url)

        new_layout = getattr(raw_event_node, "_ui_layout", None) if raw_event_node is not None else None
        ui_message = getattr(raw_event_node, "_ui_message", None) if raw_event_node is not None else None
        ui_dialog = getattr(raw_event_node, "_ui_dialog", None) if raw_event_node is not None else None
        ui_open = getattr(raw_event_node, "_ui_open", None) if raw_event_node is not None else None
        ui_close = getattr(raw_event_node, "_ui_close", None) if raw_event_node is not None else None
        ui_run_projection = getattr(raw_event_node, "_ui_run_projection", None) if raw_event_node is not None else None
        ui_plugins = getattr(raw_event_node, "_ui_plugins", None) if raw_event_node is not None else None

        if new_layout is not None:
            resp: Dict[str, Any] = {
                "ok": True,
                "layout_html": _render_raw_node_layout_response(repo, parsed, class_name, node_id, new_layout, node_data),
                "node_data": node_data,
            }
        else:
            resp = {"ok": True, "node_data": node_data}

        resp["handled"] = {"class_name": class_name, "node_id": node_id}
        resp["patches"] = []

        if ui_plugins is not None:
            resp["plugins"] = ui_plugins
        if ui_message is not None:
            resp["message"] = ui_message
        if ui_open is not None:
            resp["open"] = ui_open
        if ui_close:
            resp["close"] = True
        if ui_run_projection:
            resp["run_projection"] = True

        if isinstance(ui_dialog, dict):
            layout_html = ""
            if ui_dialog.get("layout") is not None:
                layout_html = _render_raw_node_layout_response(repo, parsed, class_name, node_id, ui_dialog.get("layout"), node_data)
            resp["dialog"] = {
                "id": ui_dialog.get("id") or "dialog",
                "title": ui_dialog.get("title") or "",
                "positive": ui_dialog.get("positive") or "OK",
                "negative": ui_dialog.get("negative") or "Cancel",
                "layout_html": layout_html,
                "html": ui_dialog.get("html") or "",
            }

        return jsonify(resp)
    except getattr(_nodes_mod, "UiBreak", Exception) as e:
        try:
            return jsonify({"ok": True, "break_view": _debug_break_view_payload(repo, getattr(e, "payload", {}) or {})})
        except Exception as be:
            return jsonify({"ok": False, "error": str(be), "message": {"text": str(be), "level": "error"}}), 200
    except Exception as e:
        return jsonify({
            "ok": False,
            "error": str(e),
            "message": {"text": f"Raw-node handler error: {e}", "level": "error"},
        }), 200



def _run_web_nodascript_action(code: Any, data_root: Dict[str, Any], host: Any, config_uid: str = "") -> Any:
    """Execute one NodaScript action as a web-backend fallback.

    The browser remains the preferred runtime for web NodaScript. This helper is
    used when the browser engine is unavailable/failed and for CommonEvents,
    which never had a browser-side NodaScript pass.
    """
    from .nodascript import NodaScriptEngine

    text = str(code or "").strip()
    if not text:
        raise ValueError("NodaScript action has no code")
    if not isinstance(data_root, dict):
        data_root = {}

    engine = NodaScriptEngine()

    def _message(text_value="", level="info"):
        if host is not None and callable(getattr(host, "Message", None)):
            return host.Message(text_value, level)
        return None

    def _dialog(*args, **kwargs):
        if host is None:
            return None
        if args and isinstance(args[0], dict):
            spec = dict(args[0])
            host._ui_dialog = {
                "id": str(spec.get("id") or "dialog"),
                "title": str(spec.get("title") or ""),
                "positive": str(spec.get("positive") or "OK"),
                "negative": str(spec.get("negative") or "Cancel"),
                "layout": spec.get("layout"),
                "html": str(spec.get("html") or ""),
            }
            return True
        fn = getattr(host, "Dialog", None)
        return fn(*args, **kwargs) if callable(fn) else None

    def _show(layout):
        fn = getattr(host, "Show", None)
        return fn(layout) if callable(fn) else None

    def _close():
        fn = getattr(host, "CloseNode", None)
        return fn() if callable(fn) else None

    def _run_projection():
        fn = getattr(host, "RunProjection", None)
        return fn() if callable(fn) else _nodes_mod.RunProjection()

    for name in ("message", "Message", "Сообщение", "Сообщить"):
        engine.register(name, _message)
    for name in ("dialog", "Dialog", "Диалог"):
        engine.register(name, _dialog)
    for name in ("show", "Show", "Показать"):
        engine.register(name, _show)
    for name in ("close", "CloseNode", "Закрыть"):
        engine.register(name, _close)
    for name in ("RunProjection", "runprojection", "СформироватьПроекцию"):
        engine.register(name, _run_projection)

    prev_current = getattr(_nodes_mod, "CURRENT_NODE", None)
    token = None
    try:
        if host is not None:
            try:
                host._config_uid = str(config_uid or "")
            except Exception:
                pass
        setattr(_nodes_mod, "CURRENT_NODE", host)
        try:
            token = _nodes_mod.CURRENT_CONFIG_UID.set(str(config_uid or "") or None)
        except Exception:
            token = None
        return engine.execute(text, data_root)
    finally:
        try:
            if token is not None:
                _nodes_mod.CURRENT_CONFIG_UID.reset(token)
        except Exception:
            pass
        setattr(_nodes_mod, "CURRENT_NODE", prev_current)


@client_bp.route("/api/node/event_job/<job_id>", methods=["GET"])
@login_required
def api_node_event_job_status(job_id):
    owner = "web:" + str(current_user.get_id() or "")
    job = _background_jobs.get(job_id, owner=owner)
    if job is None:
        return jsonify({"ok": False, "error": "job not found"}), 404
    return jsonify({"ok": True, "job": job})


@client_bp.route("/api/node/event_job/<job_id>/cancel", methods=["POST"])
@login_required
def api_node_event_job_cancel(job_id):
    owner = "web:" + str(current_user.get_id() or "")
    ok = _background_jobs.request_cancel(job_id, owner=owner)
    return jsonify({"ok": bool(ok)})


@client_bp.route("/api/node/event_web", methods=["POST"])
@login_required
def api_node_event_web():
    j = request.get_json(force=True) or {}

    if j.get("is_raw_node") or j.get("raw_node_id"):
        raw_resp = _api_raw_node_event_web(j)
        if raw_resp is not None:
            return raw_resp

    repo_id = int(j.get("repo_id") or 0)

    
    class_name = str(j.get("class_name") or "").strip()
    node_id = str(j.get("node_id") or "").strip()

    
    target_class_name = str(j.get("target_class_name") or "").strip()
    target_node_id = str(j.get("target_node_id") or "").strip()

    event = str(j.get("event") or "").strip()
    payload = j.get("payload") or {}

    if not repo_id or not class_name or not node_id or not event:
        return jsonify({"ok": False, "error": "bad args"}), 400

    repo = _get_repo_or_404(repo_id)

   
    eff_class = target_class_name or class_name
    eff_id = target_node_id or node_id

    parsed = get_parsed_config(repo, models.db) or {}
    _ctx_tokens = _nodes_mod.set_runtime_context(repo.config_uid, parsed, system_user=_client_runtime_system_user_payload())

    @after_this_request
    def _reset_ctx(resp):
        _nodes_mod.reset_runtime_context(_ctx_tokens)
        return resp

    eff_cls_cfg = (parsed.get("classes") or {}).get(eff_class) or {}
    is_custom_process = _is_singleton_class_type(eff_cls_cfg)

    # listener matching
    listener = ""
    try:
        if isinstance(payload, dict):
            listener = str(payload.get("listener") or payload.get("id") or "").strip()
    except Exception:
        listener = ""

    actions: list[dict] = []
    for ev in (eff_cls_cfg.get("events") or []):
        if ev.get("event") != event:
            continue

        ev_listener = str(ev.get("listener") or "").strip()
        
        if listener:
            if ev_listener and ev_listener != listener:
                continue
        else:
            if ev_listener:
                continue

        actions.extend(ev.get("actions") or [])

    if not actions:
        return jsonify({
            "ok": True,
            "noop": True,
            "handled": {"class_name": eff_class, "node_id": eff_id},
            "node_data": {},
            "patches": []
        })

    # runprogress and runasync are real background modes. The first request only
    # queues the event; a worker executes the exact same endpoint once with an
    # internal marker, preserving all existing handler/rendering semantics.
    internal_job_execution = bool(j.get("__event_job_execute"))
    action_modes = []
    for action_item in actions:
        method_name = str((action_item or {}).get("method") or "").strip().lower()
        if method_name == "nodascript":
            continue
        mode_name = str((action_item or {}).get("action") or "run").strip().lower()
        if mode_name in ("runprogress", "runasync"):
            action_modes.append(mode_name)
    execution_mode = "runprogress" if "runprogress" in action_modes else ("runasync" if "runasync" in action_modes else "run")

    # Reports/projections are potentially expensive and must never occupy the
    # request worker by default.  Existing explicit runasync/runprogress modes
    # still win; only the standard onRunProjection action is upgraded from run.
    if execution_mode == "run" and _is_projection_class_type(eff_cls_cfg):
        projection_methods = {
            str((action_item or {}).get("method") or "").strip().lower()
            for action_item in actions
            if str((action_item or {}).get("method") or "").strip()
        }
        if str(listener or "").strip().lower() == "onrunprojection" or "onrunprojection" in projection_methods:
            execution_mode = "runprogress"

    if not internal_job_execution and execution_mode in ("runprogress", "runasync"):
        app_obj = current_app._get_current_object()
        user_id = str(current_user.get_id() or "")
        request_body = dict(j)
        request_body["__event_job_execute"] = True

        def _execute_event_job():
            with app_obj.test_client() as test_client:
                with test_client.session_transaction() as session_data:
                    session_data["_user_id"] = user_id
                    session_data["_fresh"] = True
                response = test_client.post(
                    "/client/api/node/event_web",
                    json=request_body,
                    headers={"X-Noda-Background-Job": "1"},
                )
                result = response.get_json(silent=True)
                if result is None:
                    result = {"ok": False, "error": response.get_data(as_text=True)[:2000]}
                if response.status_code >= 400:
                    raise RuntimeError(str(result.get("error") or f"HTTP {response.status_code}"))
                return result

        job = _background_jobs.submit(
            _execute_event_job,
            owner="web:" + user_id,
            mode=execution_mode,
            title=str(event or "Operation"),
        )
        job_id = str(job.get("id") or "")
        return jsonify({
            "ok": True,
            "background": True,
            "mode": execution_mode,
            "job": {
                **job,
                "status_url": f"/client/api/node/event_job/{job_id}",
                "cancel_url": f"/client/api/node/event_job/{job_id}/cancel",
            },
        }), 202

    base_url = (repo.base_url or "").strip().rstrip("/")
    current = (request.host_url or "").rstrip("/")

    new_layout = None
    node_data: Dict[str, Any] = {}
    ui_message = None
    ui_dialog = None
    ui_open = None
    ui_close = None
    ui_run_projection = None
    patches: list[dict] = []
    runtime_method_name = ""

    try:
        # -------- REMOTE --------
        if base_url and base_url != current:
            for a in actions:
                m = (a.get("method") or "").strip()
                if not m:
                    continue
                if m.lower() == "nodascript":
                    # Normally already executed in the browser. A failed browser
                    # fallback cannot safely mutate a remote repository here.
                    if bool(payload.get("__nodascript_fallback")):
                        ui_message = (ui_message or []) + [{"text": "NodaScript browser fallback is available only for local configurations", "level": "error"}]
                    continue

                r = _api_post_remote(
                    repo,
                    f"/api/config/{repo.config_uid}/node/{eff_class}/{eff_id}/{m}",
                    json_data=payload,
                )

                if isinstance(r, dict) and isinstance(r.get("data"), dict):
                    data = r["data"]
                    if "_ui_layout" in data:
                        new_layout = data.get("_ui_layout")
                    if "_ui_message" in data:
                        ui_message = data.get("_ui_message")
                    if "_ui_dialog" in data:
                        ui_dialog = data.get("_ui_dialog")
                    if "_ui_open" in data:
                        ui_open = data.get("_ui_open")
                    if "_ui_close" in data:
                        ui_close = data.get("_ui_close")
                    if "_ui_run_projection" in data:
                        ui_run_projection = data.get("_ui_run_projection")

            
            node_data = {}

        # -------- LOCAL --------
        else:
            node_class = _load_server_node_class(repo.config_uid, eff_class)

            if is_custom_process:
                defaults = (eff_cls_cfg.get("_data") or {})
                if not isinstance(defaults, dict):
                    defaults = {}
                try:
                    node = node_class.get(eff_id, repo.config_uid)
                except Exception:
                    node = None
                if not node:
                    node = node_class(eff_id, repo.config_uid)

                stored_data = {}
                try:
                    stored_data = node.get_data() or {}
                except Exception:
                    stored_data = {}
                if not isinstance(stored_data, dict):
                    stored_data = {}

                node_data = stored_data.copy()
                for k, v in defaults.items():
                    node_data.setdefault(k, v)
                _apply_projection_defaults_to_data(eff_cls_cfg, node_data, repo.config_uid, eff_class, eff_id)
                node_data.setdefault("_id", eff_id)
                node_data.setdefault("_class", eff_class)

                try:
                    node._data_cache = node_data.copy()
                except Exception:
                    pass

            else:
                node = node_class.get(eff_id, repo.config_uid)
                if not node:
                    return jsonify({"ok": False, "error": "node not found"}), 404

            
            try:
                node._schema_class_name = eff_class
            except Exception:
                pass

            
            try:
                if getattr(node, "_data_cache", None) is None:
                    full = node.get_data() or {}
                    node._data_cache = dict(full) if isinstance(full, dict) else {}
            except Exception:
                if getattr(node, "_data_cache", None) is None:
                    node._data_cache = {}

            
            _apply_web_payload_to_node_data(node, payload if isinstance(payload, dict) else {})

            # call methods. NodaScript normally runs in the browser; execute it
            # here only when the browser marked a fallback (or post-only fallback).
            prev_current = getattr(_nodes_mod, "CURRENT_NODE", None)
            setattr(_nodes_mod, "CURRENT_NODE", node)
            try:
                ns_fallback = bool(payload.get("__nodascript_fallback"))
                ns_post_fallback = bool(payload.get("__nodascript_post_fallback"))
                ns_only = bool(payload.get("__nodascript_only"))
                for a in actions:
                    m = (a.get("method") or "").strip()
                    runtime_method_name = m or runtime_method_name
                    if m.lower() == "nodascript":
                        if ns_fallback and not ns_post_fallback:
                            code = a.get("methodText") or a.get("method_text") or a.get("text") or a.get("code") or ""
                            _run_web_nodascript_action(code, node._data_cache, node, repo.config_uid)
                        continue
                    if not ns_only and m and hasattr(node, m):
                        getattr(node, m)(payload)
                        if getattr(node, "_ui_layout", None) is not None:
                            new_layout = node._ui_layout

                    post_m = str(a.get("postExecuteMethod") or a.get("post_execute_method") or "").strip().lower()
                    if ns_post_fallback and post_m == "nodascript":
                        post_code = a.get("postExecuteMethodText") or a.get("post_execute_method_text") or a.get("postText") or a.get("post_text") or ""
                        _run_web_nodascript_action(post_code, node._data_cache, node, repo.config_uid)
            finally:
                setattr(_nodes_mod, "CURRENT_NODE", prev_current)

            ui_message = getattr(node, "_ui_message", None)
            ui_dialog = getattr(node, "_ui_dialog", None)
            ui_open = getattr(node, "_ui_open", None)
            ui_close = getattr(node, "_ui_close", None)
            ui_run_projection = getattr(node, "_ui_run_projection", None)

            ui_plugins = getattr(node, "_ui_plugins", None)
            try:
                if hasattr(node, "_ui_plugins"):
                    delattr(node, "_ui_plugins")   # one-shot
            except Exception:
                pass


            # refresh data for @vars
            try:
                if getattr(node, "_data_cache", None) is not None:
                    node_data = node._data_cache or {}
                else:
                    node_data = node.get_data() or {}
            except Exception:
                node_data = node_data or {}

            # clear one-shot ui fields
            try:
                for k in ("_ui_message", "_ui_dialog", "_ui_layout", "_ui_open", "_ui_close", "_ui_run_projection"):
                    if hasattr(node, k):
                        delattr(node, k)
            except Exception:
                pass

        
        if target_class_name and target_node_id:
            try:
                patches.append({
                    "type": "cover",
                    "class_name": eff_class,
                    "node_id": eff_id,
                    "html": _node_cover_html(repo, eff_class, eff_id),
                })
            except Exception:
                patches.append({
                    "type": "cover",
                    "class_name": eff_class,
                    "node_id": eff_id,
                    "html": "",
                })

        # dialog render
        ui_dialog_payload = None
        if ui_dialog is not None and isinstance(ui_dialog, dict):
            layout_html = ""
            if ui_dialog.get("layout") is not None:
                try:
                    dialog_layout = resolve_common_layout(parsed, ui_dialog.get("layout"))
                    dialog_data = node_data if isinstance(node_data, dict) else {}
                    if dialog_layout is not None and isinstance(dialog_data, dict):
                        _fill_nodeinput_views(repo, parsed, dialog_layout, dialog_data)
                    layout_html = render_nodalayout_html(
                        dialog_layout,
                        dialog_data,
                        assets_base_dir=_userfiles_dir_for_repo(repo),
                        context=_nl_context(repo, class_name=eff_class, node_id=eff_id),
                    )
                except Exception:
                    layout_html = ""
            ui_dialog_payload = {
                "id": ui_dialog.get("id") or "dialog",
                "title": ui_dialog.get("title") or "",
                "positive": ui_dialog.get("positive") or "OK",
                "negative": ui_dialog.get("negative") or "Cancel",
                "layout_html": layout_html,
                "html": ui_dialog.get("html") or "",
            }

        # layout render
        resp: Dict[str, Any]
        if new_layout is not None:
            try:
                runtime_layout = resolve_common_layout(parsed, new_layout)
                data_for_layout = node_data if isinstance(node_data, dict) else {}
                if runtime_layout is not None and isinstance(data_for_layout, dict):
                    _fill_nodeinput_views(repo, parsed, runtime_layout, data_for_layout)
                layout_html = render_nodalayout_html(
                    runtime_layout,
                    data_for_layout,
                    assets_base_dir=_userfiles_dir_for_repo(repo),
                    context=_nl_context(repo, class_name=eff_class, node_id=eff_id),
                )
            except Exception:
                layout_html = ""
            resp = {"ok": True, "layout_html": layout_html, "node_data": node_data}
        else:
            resp = {"ok": True, "node_data": node_data}

        resp["handled"] = {"class_name": eff_class, "node_id": eff_id}
        resp["patches"] = patches

        if ui_plugins is not None:
            resp["plugins"] = ui_plugins


        if ui_message is not None:
            resp["message"] = ui_message
        if ui_dialog_payload is not None:
            resp["dialog"] = ui_dialog_payload
        if ui_open is not None:
            resp["open"] = ui_open
        if ui_close:
            resp["close"] = True
        if ui_run_projection:
            resp["run_projection"] = True

        return jsonify(resp)

    except getattr(_nodes_mod, "UiBreak", Exception) as e:
        try:
            return jsonify({"ok": True, "break_view": _debug_break_view_payload(repo, getattr(e, "payload", {}) or {})}), 200
        except Exception as be:
            return jsonify({"ok": False, "error": str(be), "message": {"text": str(be), "level": "error"}}), 200

    except _nodes_mod.AcceptRejected as e:


        payload = getattr(e, 'payload', None) or {}


        msg = payload.get('message')


        if not isinstance(msg, dict):


            msg = {'text': payload.get('error') or 'Save rejected', 'level': 'error'}


        return jsonify({'ok': False, 'error': payload.get('error') or 'rejected', 'message': msg}), 200


    except Exception as e:
        try:
            import traceback as _tb
            from solutions.runtime_errors import record_runtime_error
            record_runtime_error(
                repo.config_uid, exception=e, traceback_text=_tb.format_exc(), source="web_event",
                class_name=eff_class, node_id=eff_id, event_name=str(event or ""),
                method_name=runtime_method_name, input_data=payload,
                context={
                    "listener": str(listener or ""),
                    "request_path": request.path,
                    "action_modes": action_modes,
                    "background": bool(internal_job_execution),
                },
            )
        except Exception:
            pass
        return jsonify({
            "ok": False,
            "error": str(e),
            "message": {"text": f"Handler error: {e}", "level": "error"},
        }), 200


class _UiHost:
    """
    UI host for CommonEvents (no real node).
    nodes.py free-functions (message/Dialog/CloseNode) use CURRENT_NODE, so we provide the same API as Node.
    """

    def __init__(self):
        # keep the same shapes as in nodes.py Node.Message/Node.Dialog/Node.Show
        self._ui_message = None   # list[{"text":..,"level":..}]
        self._ui_dialog = None    # dict with keys: id,title,positive,negative,layout,html
        self._ui_layout = None    # layout spec (will be rendered by render_nodalayout_html)
        self._ui_open = None
        self._ui_close = None

    def Message(self, text: str, level: str = "info"):
        try:
            msgs = getattr(self, "_ui_message", None)
            if not isinstance(msgs, list):
                msgs = []
            msgs.append({"text": str(text), "level": str(level or "info")})
            self._ui_message = msgs
        except Exception:
            pass

    def Dialog(self, dialog_id: str, title: str = "", *, positive: str = "OK", negative: str = "Cancel", layout=None, html: str = ""):
        self._ui_dialog = {
            "id": str(dialog_id or "dialog"),
            "title": str(title or ""),
            "positive": str(positive or "OK"),
            "negative": str(negative or "Cancel"),
            "layout": layout,
            "html": html,
        }

    def Show(self, layout):
        self._ui_layout = layout
        return True

    def CloseNode(self):
        self._ui_close = True
        return True

    def RunProjection(self):
        self._ui_run_projection = True
        return True



@client_bp.route("/api/common/event_web", methods=["POST"])
@login_required
def api_common_event_web():
    """
    Execute CommonEvents (Configuration.config_events) in web client context.

    Works like onInputWeb but:
      - does NOT call node methods
      - calls python functions from handlers module:
            def <method>(input_data)
      - injects `nodes.CURRENT_NODE = _UiHost()` so message()/Dialog() works

    Payload rules:
      - listener matching
      - expected payload is dict
    """
    j = request.get_json(force=True) or {}
    repo_id = int(j.get("repo_id") or 0)
    event = str(j.get("event") or "")
    payload = j.get("payload") or {}

    if not repo_id or not event:
        return jsonify({"ok": False, "error": "bad args"}), 400
    if not isinstance(payload, dict):
        payload = {}

    repo = _get_repo_or_404(repo_id)
    parsed = get_parsed_config(repo, models.db) or {}
    cfg = (parsed.get("cfg") or {}) if isinstance(parsed, dict) else {}

    common_events = cfg.get("CommonEvents") or []

    # listener matching exactly like in api_node_event_web
    listener = ""
    try:
        listener = str(payload.get("listener") or payload.get("id") or "").strip()
    except Exception:
        listener = ""

    actions = []
    for ev in (common_events or []):
        if (ev.get("event") or "") != event:
            continue
        ev_listener = str(ev.get("listener") or "").strip()

        if listener:
            # allow exact + global(empty)
            if ev_listener and ev_listener != listener:
                continue
        else:
            # only global(empty)
            if ev_listener:
                continue

        actions.extend(ev.get("actions") or [])

    if not actions:
        return jsonify({"ok": True, "noop": True})

    base_url = (repo.base_url or "").strip().rstrip("/")
    current = (request.host_url or "").rstrip("/")

    ui_message = None
    ui_dialog_payload = None
    layout_html = None
    ui_open = None
    ui_close = None
    ui_run_projection = None

    try:
        
        if base_url and base_url != current:
            return jsonify({
                "ok": False,
                "error": "CommonEvents are local-only (no remote endpoint implemented)",
                "message": [{"text": "CommonEvents: remote call not supported", "level": "warning"}],
            }), 200

        # -------- LOCAL --------
        ns = _load_server_handlers_ns(repo.config_uid)

        ui = _UiHost()

        # inject CURRENT_NODE so nodes.message()/nodes.Dialog() works
        prev_current = getattr(_nodes_mod, "CURRENT_NODE", None)
        setattr(_nodes_mod, "CURRENT_NODE", ui)

        try:
            for a in actions:
                m = (a.get("method") or "").strip()
                if not m:
                    continue
                if m.lower() == "nodascript":
                    code = a.get("methodText") or a.get("method_text") or a.get("text") or a.get("code") or ""
                    _run_web_nodascript_action(code, payload, ui, repo.config_uid)
                else:
                    fn = ns.get(m)
                    if not callable(fn):
                        raise ValueError(f"CommonEvent function '{m}' not found/callable in handlers")
                    fn(payload)  # function(input_data)

                post_m = str(a.get("postExecuteMethod") or a.get("post_execute_method") or "").strip().lower()
                if post_m == "nodascript":
                    post_code = a.get("postExecuteMethodText") or a.get("post_execute_method_text") or a.get("postText") or a.get("post_text") or ""
                    _run_web_nodascript_action(post_code, payload, ui, repo.config_uid)

        
        finally:
            setattr(_nodes_mod, "CURRENT_NODE", prev_current)

        # collect ui fields
        ui_message = getattr(ui, "_ui_message", None)
        ui_dialog = getattr(ui, "_ui_dialog", None)
        ui_open = getattr(ui, "_ui_open", None)
        ui_close = getattr(ui, "_ui_close", None)
        ui_run_projection = getattr(ui, "_ui_run_projection", None)

        # dialog render (same as in api_node_event_web, but data for vars is just payload)
        if ui_dialog is not None and isinstance(ui_dialog, dict):
            dlg_layout_html = ""
            if ui_dialog.get("layout") is not None:
                try:
                    dlg_layout_html = render_nodalayout_html(
                        ui_dialog.get("layout"),
                        payload if isinstance(payload, dict) else {},
                        assets_base_dir=_userfiles_dir_for_repo(repo),
                        context=_nl_context(repo, class_name="", node_id="")
                    )
                except Exception:
                    dlg_layout_html = ""

            ui_dialog_payload = {
                "id": ui_dialog.get("id") or "dialog",
                "title": ui_dialog.get("title") or "",
                "positive": ui_dialog.get("positive") or "OK",
                "negative": ui_dialog.get("negative") or "Cancel",
                "layout_html": dlg_layout_html,
                "html": ui_dialog.get("html") or "",
            }

        # layout render (Show(layout))
        if getattr(ui, "_ui_layout", None) is not None:
            try:
                layout_html = render_nodalayout_html(
                    getattr(ui, "_ui_layout"),
                    payload if isinstance(payload, dict) else {},
                    assets_base_dir=_userfiles_dir_for_repo(repo),
                    context=_nl_context(repo, class_name="", node_id="")
                )
            except Exception:
                layout_html = ""

        resp = {"ok": True}
        if ui_message is not None:
            resp["message"] = ui_message
        if ui_dialog_payload is not None:
            resp["dialog"] = ui_dialog_payload
        if layout_html is not None:
            resp["layout_html"] = layout_html
        if ui_open is not None:
            resp["open"] = ui_open
        if ui_close:
            resp["close"] = True
        if ui_run_projection:
            resp["run_projection"] = True

        return jsonify(resp)

    except Exception as e:
        return jsonify({
            "ok": False,
            "error": str(e),
            "message": [{"text": f"CommonEvent error: {e}", "level": "error"}],
        }), 200



TIMER_MIN_PERIOD_SECONDS = 900
SERVER_TIMER_SCAN_SECONDS = 30
_SERVER_TIMER_STATE: Dict[str, Dict[str, Any]] = {}
_SERVER_TIMER_LOCK = threading.RLock()
_SERVER_TIMER_STARTED = False
_SERVER_TIMER_STOP = None


def _timer_bool(value, default=False) -> bool:
    if value is None:
        return default
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return value != 0
    value = str(value).strip().lower()
    if value in {'1', 'true', 'on', 'yes'}:
        return True
    if value in {'0', 'false', 'off', 'no'}:
        return False
    return default


def _timer_runtime(timer_cfg: Dict[str, Any]) -> str:
    if not isinstance(timer_cfg, dict):
        return 'server'
    raw = str(timer_cfg.get('runtime') or timer_cfg.get('run_on') or '').strip().lower()
    if raw in {'client', 'android', 'mobile', 'клиент', 'андроид'}:
        return 'client'
    if raw in {'server', 'backend', 'сервер'}:
        return 'server'
    # Backward compatibility: timers created before the Server/Client switch
    # were intended by design to be server timers.
    return 'server'


def _timer_period_seconds(timer_cfg: Dict[str, Any]) -> int:
    try:
        period = int(float((timer_cfg or {}).get('period_seconds') or (timer_cfg or {}).get('period') or 0))
    except Exception:
        period = 0
    runtime = _timer_runtime(timer_cfg)
    worker = _timer_bool((timer_cfg or {}).get('worker'), False)
    min_period = TIMER_MIN_PERIOD_SECONDS if runtime == 'server' or worker else 1
    return max(min_period, period)


def _timer_id(timer_cfg: Dict[str, Any]) -> str:
    if not isinstance(timer_cfg, dict):
        return ''
    return str(timer_cfg.get('timer_id') or timer_cfg.get('id') or '').strip()


def _timer_actions(timer_cfg: Dict[str, Any]) -> List[Dict[str, Any]]:
    actions = (timer_cfg or {}).get('actions') or (timer_cfg or {}).get('Actions') or []
    return [a for a in actions if isinstance(a, dict)]


def _timer_signature(timer_cfg: Dict[str, Any]) -> str:
    try:
        stable = {
            'timer_id': _timer_id(timer_cfg),
            'period_seconds': _timer_period_seconds(timer_cfg),
            'runtime': _timer_runtime(timer_cfg),
            'active': _timer_bool((timer_cfg or {}).get('active'), True),
            'worker': _timer_bool((timer_cfg or {}).get('worker'), False),
            'actions': _timer_actions(timer_cfg),
        }
        return hashlib.sha256(json.dumps(stable, ensure_ascii=False, sort_keys=True).encode('utf-8')).hexdigest()
    except Exception:
        return str(time.time())


def _execute_timer_actions_for_repo(repo, timer_cfg: Dict[str, Any], payload: Optional[Dict[str, Any]] = None, *, include_ui: bool = False, parsed_config: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    """Execute timer actions exactly like CommonEvents, without requiring a browser.

    Browser/client calls pass include_ui=True so UI helper results can be rendered
    back to the current page. Server scheduler calls include_ui=False; message(),
    Dialog(), Show(), CloseNode() are still safely captured but only logged/returned.
    """
    timer_id = _timer_id(timer_cfg)
    if not timer_id:
        return {"ok": False, "error": "Timer ID is empty"}
    if not _timer_bool(timer_cfg.get("active"), True):
        return {"ok": True, "noop": True, "inactive": True}

    actions = _timer_actions(timer_cfg)
    if not actions:
        return {"ok": True, "noop": True}

    timer_payload = dict(payload or {}) if isinstance(payload, dict) else {}
    timer_payload["_timer_id"] = timer_id
    timer_payload["_timer_runtime"] = _timer_runtime(timer_cfg)
    timer_payload["_timer_period_seconds"] = _timer_period_seconds(timer_cfg)
    data_payload = timer_payload.get("_data")
    if not isinstance(data_payload, dict):
        data_payload = {}
    data_payload["_timer_id"] = timer_id
    data_payload["_timer_runtime"] = _timer_runtime(timer_cfg)
    data_payload["_timer_period_seconds"] = _timer_period_seconds(timer_cfg)
    timer_payload["_data"] = data_payload

    if parsed_config is None:
        try:
            parsed_config = get_parsed_config(repo, models.db) or {}
        except Exception:
            parsed_config = {}

    config_uid = str(getattr(repo, "config_uid", "") or "").strip()
    try:
        cfg_json = parsed_config.get("cfg") if isinstance(parsed_config, dict) else None
        if not config_uid and isinstance(cfg_json, dict):
            config_uid = str(cfg_json.get("uid") or "").strip()
    except Exception:
        pass

    ns = _load_server_handlers_ns(config_uid, parsed_config)
    ui = _UiHost()
    prev_current = getattr(_nodes_mod, "CURRENT_NODE", None)
    ctx_tokens = None
    try:
        ctx_tokens = _nodes_mod.set_runtime_context(
            config_uid,
            parsed_config,
            system_user={"id": int(getattr(repo, "user_id", 0) or 0)},
        )
    except Exception:
        ctx_tokens = None
    setattr(_nodes_mod, "CURRENT_NODE", ui)

    executed = []
    try:
        for a in actions:
            m = (a.get("method") or "").strip()
            if not m:
                continue
            try:
                if m.lower() == "nodascript":
                    code = a.get("methodText") or a.get("method_text") or a.get("text") or a.get("code") or ""
                    _run_web_nodascript_action(code, timer_payload, ui, config_uid)
                else:
                    fn = ns.get(m)
                    if not callable(fn):
                        raise ValueError(f"Timer function '{m}' not found/callable in handlers")
                    fn(timer_payload)
                executed.append(m)

                post_m = str(a.get("postExecuteMethod") or a.get("post_execute_method") or "").strip().lower()
                if post_m == "nodascript":
                    post_code = a.get("postExecuteMethodText") or a.get("post_execute_method_text") or a.get("postText") or a.get("post_text") or ""
                    _run_web_nodascript_action(post_code, timer_payload, ui, config_uid)
            except Exception as e:
                print(f"Timer {timer_id} handler {m} error: {e}")
                try:
                    traceback.print_exc()
                except Exception:
                    pass
    finally:
        setattr(_nodes_mod, "CURRENT_NODE", prev_current)
        if ctx_tokens is not None:
            try:
                _nodes_mod.reset_runtime_context(ctx_tokens)
            except Exception:
                pass

    resp: Dict[str, Any] = {"ok": True, "timer_id": timer_id, "executed": executed}

    ui_message = getattr(ui, "_ui_message", None)
    ui_dialog = getattr(ui, "_ui_dialog", None)
    ui_open = getattr(ui, "_ui_open", None)
    ui_close = getattr(ui, "_ui_close", None)
    ui_run_projection = getattr(ui, "_ui_run_projection", None)

    if ui_message is not None:
        resp["message"] = ui_message

    if include_ui:
        ui_dialog_payload = None
        layout_html = None

        if ui_dialog is not None and isinstance(ui_dialog, dict):
            dlg_layout_html = ""
            if ui_dialog.get("layout") is not None:
                try:
                    dlg_layout_html = render_nodalayout_html(
                        ui_dialog.get("layout"),
                        timer_payload,
                        assets_base_dir=_userfiles_dir_for_repo(repo),
                        context=_nl_context(repo, class_name="", node_id="")
                    )
                except Exception:
                    dlg_layout_html = ""

            ui_dialog_payload = {
                "id": ui_dialog.get("id") or "dialog",
                "title": ui_dialog.get("title") or "",
                "positive": ui_dialog.get("positive") or "OK",
                "negative": ui_dialog.get("negative") or "Cancel",
                "layout_html": dlg_layout_html,
                "html": ui_dialog.get("html") or "",
            }

        if getattr(ui, "_ui_layout", None) is not None:
            try:
                layout_html = render_nodalayout_html(
                    getattr(ui, "_ui_layout"),
                    timer_payload,
                    assets_base_dir=_userfiles_dir_for_repo(repo),
                    context=_nl_context(repo, class_name="", node_id="")
                )
            except Exception:
                layout_html = ""

        if ui_dialog_payload is not None:
            resp["dialog"] = ui_dialog_payload
        if layout_html is not None:
            resp["layout_html"] = layout_html
        if ui_open is not None:
            resp["open"] = ui_open
        if ui_close:
            resp["close"] = True
        if ui_run_projection:
            resp["run_projection"] = True
    else:
        if ui_dialog is not None:
            resp["dialog"] = True
        if getattr(ui, "_ui_layout", None) is not None:
            resp["layout"] = True
        if ui_open is not None:
            resp["open"] = ui_open
        if ui_close:
            resp["close"] = True
        if ui_run_projection:
            resp["run_projection"] = True

    return resp


@client_bp.route("/api/timer/event_web", methods=["POST"])
@login_required
def api_timer_event_web():
    """Compatibility endpoint.

    runtime=client is the Android runtime, not the browser.  Web pages never
    execute configuration timers; server timers are handled by the scheduler.
    """
    j = request.get_json(force=True) or {}
    repo_id = int(j.get("repo_id") or 0)
    timer_id = str(j.get("timer_id") or j.get("id") or "").strip()
    if not repo_id or not timer_id:
        return jsonify({"ok": False, "error": "bad args"}), 400

    repo = _get_repo_or_404(repo_id)
    parsed = get_parsed_config(repo, models.db) or {}
    cfg = (parsed.get("cfg") or {}) if isinstance(parsed, dict) else {}
    timer_cfg = next((
        t for t in (cfg.get("Timers") or cfg.get("timers") or [])
        if isinstance(t, dict) and _timer_id(t) == timer_id
    ), None)
    if not timer_cfg:
        return jsonify({"ok": False, "error": "Timer not found"}), 404
    if _timer_runtime(timer_cfg) == "client":
        return jsonify({"ok": True, "noop": True, "android_timer": True})
    return jsonify({"ok": True, "noop": True, "server_timer": True})




def _refresh_local_repo_config_cache(repo) -> None:
    """Refresh repository cache only when the config exists in this Designer DB.

    A client repository can contain a cached configuration that was added by a
    public link or copied from another server. In that case repo.base_url may be
    empty because node API access is configured separately, but the configuration
    UID is not present in the local Designer DB. Server timers must still use the
    cached RepoConfig snapshot and should not spam logs with "Configuration ...
    not found in DB".
    """
    try:
        if (getattr(repo, "base_url", "") or "").strip():
            return

        config_uid = str(getattr(repo, "config_uid", "") or "").strip()
        if not config_uid:
            return

        Configuration = getattr(main, "Configuration", None)
        if Configuration is None:
            return

        # Compare the cheap SQL revision first. The previous implementation
        # rebuilt the complete configuration (including every method/event/action)
        # for every repository on every scheduler pass, even when nothing changed.
        cfg_meta = models.db.session.execute(
            select(Configuration.id, Configuration.last_modified).where(Configuration.uid == config_uid)
        ).first()
        if not cfg_meta:
            return
        db_modified = cfg_meta[1]
        db_stamp = db_modified.isoformat() if hasattr(db_modified, "isoformat") else str(db_modified or "")

        row = models.RepoConfig.query.filter_by(repo_id=repo.id).first()
        old_cfg = None
        if row and row.config_json:
            try:
                old_cfg = json.loads(row.config_json)
            except Exception:
                old_cfg = None
        if isinstance(old_cfg, dict) and str(old_cfg.get("last_modified") or "") == db_stamp:
            return

        actor = _client_repo_actor(repo)
        cfg = fetch_config_from_local_db(config_uid, user=actor)
        if _repo_cfg_locked_by_ngenie_code(cfg):
            return

        # request.host_url is unavailable in the scheduler. Preserve the
        # previously stored URL so this cosmetic difference cannot trigger an
        # endless cache rewrite/invalidation loop.
        if isinstance(old_cfg, dict) and old_cfg.get("url"):
            cfg["url"] = old_cfg.get("url")

        cfg_json = json.dumps(cfg, ensure_ascii=False)
        changed = old_cfg != cfg
        if not row:
            row = models.RepoConfig(repo_id=repo.id, config_json=cfg_json)
            models.db.session.add(row)
            changed = True
        elif changed:
            row.config_json = cfg_json

        if changed:
            row.updated_at = datetime.now(timezone.utc)
            try:
                repo.config_json = cfg_json
                repo.config_cached_at = row.updated_at
            except Exception:
                pass
            models.db.session.commit()
            _invalidate_repo_config_mem(repo.id)
    except Exception as e:
        try:
            models.db.session.rollback()
        except Exception:
            pass
        print(f"Server timer: local repo cache refresh skipped for repo={getattr(repo, 'id', '?')}: {e}")


def _iter_server_timer_candidates() -> List[Tuple[Any, Dict[str, Any], Dict[str, Any]]]:
    out: List[Tuple[Any, Dict[str, Any], Dict[str, Any]]] = []
    repos = models.Repo.query.order_by(models.Repo.id.asc()).all()
    for repo in repos:
        _refresh_local_repo_config_cache(repo)
        try:
            parsed = get_parsed_config(repo, models.db, user=_client_repo_actor(repo)) or {}
            cfg = (parsed.get("cfg") or {}) if isinstance(parsed, dict) else {}
        except Exception as e:
            print(f"Server timer: cannot read repo {getattr(repo, 'id', '?')}: {e}")
            continue

        for timer_cfg in (cfg.get("Timers") or cfg.get("timers") or []):
            if not isinstance(timer_cfg, dict):
                continue
            if not _timer_bool(timer_cfg.get("active"), True):
                continue
            if _timer_runtime(timer_cfg) != "server":
                continue
            if not _timer_id(timer_cfg):
                continue
            out.append((repo, timer_cfg, parsed))
    return out


def _server_timer_scheduler_tick() -> None:
    now = time.time()
    active_keys = set()

    for repo, timer_cfg, parsed in _iter_server_timer_candidates():
        timer_id = _timer_id(timer_cfg)
        key = f"{int(getattr(repo, 'id', 0) or 0)}:{timer_id}"
        active_keys.add(key)
        period = _timer_period_seconds(timer_cfg)
        sig = _timer_signature(timer_cfg)

        with _SERVER_TIMER_LOCK:
            st = _SERVER_TIMER_STATE.get(key)
            if not st:
                # New active server timer: fire on the next scheduler pass immediately,
                # then continue by period. The period is still normalized to at least
                # 15 minutes for Server/Worker timers.
                _SERVER_TIMER_STATE[key] = {
                    "next_at": now,
                    "period": period,
                    "signature": sig,
                    "running": True,
                    "last_error": "",
                }
            elif st.get("signature") != sig:
                # Timer config/action changed: apply it right away.
                st.update({
                    "next_at": now,
                    "period": period,
                    "signature": sig,
                    "running": True,
                    "last_error": "",
                })
            else:
                if st.get("running") or now < float(st.get("next_at") or 0):
                    continue
                st["running"] = True

        started_at = time.time()
        try:
            print(f"Server timer fire: repo={getattr(repo, 'id', '?')} config={getattr(repo, 'config_uid', '')} timer={timer_id}")
            resp = _execute_timer_actions_for_repo(
                repo,
                timer_cfg,
                {
                    "_timer_id": timer_id,
                    "_server_timer": True,
                    "repo_id": int(getattr(repo, "id", 0) or 0),
                    "config_uid": str(getattr(repo, "config_uid", "") or ""),
                },
                include_ui=False,
                parsed_config=parsed,
            )
            if resp.get("message"):
                print(f"Server timer message {timer_id}: {resp.get('message')}")
        except Exception as e:
            print(f"Server timer error: repo={getattr(repo, 'id', '?')} timer={timer_id}: {e}")
            try:
                traceback.print_exc()
            except Exception:
                pass
            with _SERVER_TIMER_LOCK:
                if key in _SERVER_TIMER_STATE:
                    _SERVER_TIMER_STATE[key]["last_error"] = str(e)
        finally:
            with _SERVER_TIMER_LOCK:
                st = _SERVER_TIMER_STATE.setdefault(key, {})
                st["running"] = False
                st["last_fired_at"] = started_at
                # Schedule from the start time to avoid drift, but never fire immediately
                # in a tight loop if the handler ran longer than the period.
                st["next_at"] = max(started_at + period, time.time() + 1)
                st["period"] = period
                st["signature"] = sig

    with _SERVER_TIMER_LOCK:
        for key in list(_SERVER_TIMER_STATE.keys()):
            if key not in active_keys:
                _SERVER_TIMER_STATE.pop(key, None)


def _server_timer_scheduler_loop(app_obj):
    print("Server timer scheduler started")
    while True:
        stop_event = globals().get("_SERVER_TIMER_STOP")
        if stop_event is not None and stop_event.is_set():
            break
        try:
            with app_obj.app_context():
                _server_timer_scheduler_tick()
        except Exception as e:
            print(f"Server timer scheduler tick error: {e}")
            try:
                traceback.print_exc()
            except Exception:
                pass
        stop_event = globals().get("_SERVER_TIMER_STOP")
        if stop_event is not None and stop_event.wait(SERVER_TIMER_SCAN_SECONDS):
            break


def start_server_timer_scheduler(app_obj) -> bool:
    """Start one background scheduler thread for Server timers."""
    global _SERVER_TIMER_STARTED, _SERVER_TIMER_STOP
    if _SERVER_TIMER_STARTED:
        return False
    _SERVER_TIMER_STOP = threading.Event()
    t = threading.Thread(
        target=_server_timer_scheduler_loop,
        args=(app_obj,),
        name="noda-server-timer-scheduler",
        daemon=True,
    )
    t.start()
    _SERVER_TIMER_STARTED = True
    return True


@client_bp.route("/api/class/event_web", methods=["POST"])
@login_required
def api_class_event_web():
    j = request.get_json(force=True) or {}
    repo_id = int(j.get("repo_id") or 0)
    class_name = str(j.get("class_name") or "")
    event = str(j.get("event") or "")
    payload = j.get("payload") or {}

    if not repo_id or not class_name or not event:
        return jsonify({"ok": False, "error": "bad args"}), 400
    if not isinstance(payload, dict):
        payload = {}

    repo = _get_repo_or_404(repo_id)
    parsed = get_parsed_config(repo, models.db) or {}
    _ctx_tokens = _nodes_mod.set_runtime_context(repo.config_uid, parsed, system_user=_client_runtime_system_user_payload())

    @after_this_request
    def _reset_ctx(resp):
        _nodes_mod.reset_runtime_context(_ctx_tokens)
        return resp

    cls = (parsed.get("classes") or {}).get(class_name) or {}

    # listener matching 
    listener = ""
    try:
        listener = str(payload.get("listener") or payload.get("id") or "").strip()
    except Exception:
        listener = ""

    actions = []
    for ev in (cls.get("events") or []):
        if ev.get("event") != event:
            continue
        ev_listener = str(ev.get("listener") or "").strip()

        if listener:
            if ev_listener and ev_listener != listener:
                continue
        else:
            if ev_listener:
                continue

        actions.extend(ev.get("actions") or [])

    if not actions:
        return jsonify({"ok": True, "noop": True})

    base_url = (repo.base_url or "").strip().rstrip("/")
    current = (request.host_url or "").rstrip("/")

    
    if base_url and base_url != current:
        return jsonify({
            "ok": False,
            "error": "class event remote call not implemented",
            "message": [{"text": "ClassCommandWeb: remote call not supported", "level": "warning"}],
        }), 200

    try:
        node_class = _load_server_node_class(repo.config_uid, class_name)

        
        ui = _UiHost()
        prev_current = getattr(_nodes_mod, "CURRENT_NODE", None)
        setattr(_nodes_mod, "CURRENT_NODE", ui)
        try:
            for a in actions:
                m = (a.get("method") or "").strip()
                if not m:
                    continue

                fn = getattr(node_class, m, None)
                if not callable(fn):
                    raise ValueError(f"Class handler '{m}' not found on class '{class_name}'")


                fn(payload)

        finally:
            setattr(_nodes_mod, "CURRENT_NODE", prev_current)

        resp = {"ok": True}

        ui_message = getattr(ui, "_ui_message", None)
        ui_dialog = getattr(ui, "_ui_dialog", None)

        if ui_message is not None:
            resp["message"] = ui_message
        if ui_dialog is not None:
            resp["dialog"] = ui_dialog

        return jsonify(resp)

    except _nodes_mod.AcceptRejected as e:


        payload = getattr(e, 'payload', None) or {}


        msg = payload.get('message')


        if not isinstance(msg, dict):


            msg = {'text': payload.get('error') or 'Save rejected', 'level': 'error'}


        return jsonify({'ok': False, 'error': payload.get('error') or 'rejected', 'message': msg}), 200


    except Exception as e:
        return jsonify({
            "ok": False,
            "error": str(e),
            "message": [{"text": f"class-event error: {e}", "level": "error"}],
        }), 200


@client_bp.route("/api/class_nodes")
@login_required
def api_class_nodes():
    repo_id = int(request.args.get("repo_id") or 0)
    class_name = (request.args.get("class_name") or "").strip()
    q = (request.args.get("q") or "").strip()
    limit = int(request.args.get("limit") or 50)

    if not class_name:
        return jsonify({"ok": False, "error": "bad args"}), 400

    # NodeInput keeps its normal class-based API.  The only exception is the
    # reserved _User class, whose nodes belong to the owner's hidden _System
    # configuration rather than to the current business repository.
    if class_name == "_User":
        return jsonify({
            "ok": True,
            "items": _client_system_user_picker_items(q=q, limit=limit),
        })

    if repo_id:
        repos = [models.Repo.query.filter_by(id=repo_id, user_id=current_user.id).first()]
    else:
        repos = models.Repo.query.filter_by(user_id=_ngenie_effective_user_id()).all()
    repos = [r for r in repos if r]
    if not repos:
        return jsonify({"ok": False, "error": "repo not found"}), 404

    items = []
    seen = set()
    for repo in repos:
        try:
            nodes = _fetch_nodes_for_class(repo, config_uid=repo.config_uid, class_name=class_name, q=q, limit=limit) or []
        except Exception:
            nodes = []

        parsed_cfg = get_parsed_config(repo, models.db)
        for n in nodes:
            data = n.get("_data") or {}
            nid = n.get("_id") or data.get("_id") or ""
            if not nid:
                continue
            uid = _nodes_mod.normalize_own_uid(repo.config_uid, class_name, str(nid))
            if not uid or uid in seen:
                continue
            seen.add(uid)
            view = _render_class_record_view(parsed_cfg, class_name, str(nid), data)
            cover_html = ""
            try:
                cover_html = _node_cover_html(repo, class_name, str(nid), mode="table")
            except Exception:
                cover_html = ""
            items.append({
                "uid": uid,
                "_id": str(nid),
                "_class": class_name,
                "_view": str(view),
                "cover_html": cover_html,
                "data": data,
                "repo_id": repo.id,
                "repo_uid": repo.config_uid,
            })
            if len(items) >= limit:
                return jsonify({"ok": True, "items": items})

    return jsonify({"ok": True, "items": items})

@client_bp.route("/api/dataset_items")
@login_required
def api_dataset_items():
    repo_id = int(request.args.get("repo_id") or 0)
    ds_name = (request.args.get("dataset") or "").strip()
    q = (request.args.get("q") or "").strip().lower()
    limit = int(request.args.get("limit") or 100)

    if not ds_name:
        return jsonify({"ok": False, "error": "bad args"}), 400

    if repo_id:
        repos = [models.Repo.query.filter_by(id=repo_id, user_id=current_user.id).first()]
    else:
        repos = models.Repo.query.filter_by(user_id=_ngenie_effective_user_id()).all()
    repos = [r for r in repos if r]
    if not repos:
        return jsonify({"ok": False, "error": "repo not found"}), 404

    ds = None
    repo = None
    cfg = None
    for candidate_repo in repos:
        cfg_uid = (candidate_repo.config_uid or "").strip()
        if not cfg_uid:
            continue
        candidate_cfg = main.Configuration.query.filter_by(uid=cfg_uid).first()
        if not candidate_cfg:
            continue
        candidate_ds = main.Dataset.query.filter_by(config_id=candidate_cfg.id, name=ds_name).first()
        if candidate_ds:
            repo = candidate_repo
            cfg = candidate_cfg
            ds = candidate_ds
            break

    if not repo or not cfg or not ds:
        return jsonify({"ok": False, "error": "dataset not found"}), 404

    # helper: render view_template like "Item: @name (@code)"
    tmpl_re = re.compile(r"\{([\w.]+)\}", re.UNICODE)
    def render_view(data: dict) -> str:
        if isinstance(data, dict):
            v = data.get("_view")
            if isinstance(v, str) and v.strip():
                return v.strip()

        tpl = (ds.view_template or "").strip()
        if tpl and isinstance(data, dict):
            def repl(m: re.Match) -> str:
                k = m.group(1)
                v = data.get(k)
                return "" if v is None else str(v)
            s = tmpl_re.sub(repl, tpl).strip()
            if s:
                return s
        return ""

    # load items
    items_q = main.DatasetItem.query.filter_by(dataset_id=ds.id)

    
    hard_limit = max(1, min(limit, 500))
    items = items_q.limit(hard_limit * 3).all() if q else items_q.limit(hard_limit).all()

    out = []
    pos = 0
    for it in items:
        data = it.data or {}
        if not isinstance(data, dict):
            data = {}

        item_id = (it.item_id or "").strip()
        if not item_id:
            continue

        view = render_view(data) or item_id

        # search in: item_id + view + top-level string fields
        if q:
            hay = [item_id, view]
            for _, v in data.items():
                if isinstance(v, str):
                    hay.append(v)
            if q not in (" ".join(hay).lower()):
                continue

        pos += 1
        out.append({
            "key": f"{ds_name}${item_id}",
            "_id": item_id,
            "_view": view,
            "data": data,
            "position": pos,
        })

        if len(out) >= hard_limit:
            break

    return jsonify({"ok": True, "items": out})